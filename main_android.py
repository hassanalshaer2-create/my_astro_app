# 1. المكتبات الأساسية للنظام وواجهة المستخدم
import flet as ft
import hashlib
import os
import platform
import math
import uuid
import time
from datetime import datetime, timezone

# 2. المكتبات الفلكية والجغرافية (تأكد من وجودها في requirements.txt)
import swisseph as swe
from ummalqura.hijri_date import HijriDate
from geopy.geocoders import Nominatim

# 3. مكتبات معالجة اللغة العربية وتنسيق النصوص
import arabic_reshaper
from bidi.algorithm import get_display

try:
    from docx import Document
except ImportError:
    # هذا السطر يحمي التطبيق من الانهيار إذا لم يجد المكتبة
    Document = None 
def fix_ar(text):
    if not text: return ""
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)
# ---------------------------------------------------------
# إعدادات الحماية والتشغيل (المتغيرات العامة)
# ---------------------------------------------------------
SECRET_SALT = "MEEQAT_PRO_2026_ULTIMATE"
TRIAL_DAYS = 3
# ملاحظة للأندرويد: لا تستخدم os.getcwd() كمسار حفظ نهائي، اعتمد على page.client_storage
# بدلاً من المسار القديم، نستخدم هذا المسار المتوافق مع الموبايل والويندوز معاً
if platform.system() == "Android":
    # هذا المسار يضمن لك صلاحية الكتابة داخل أندرويد
    data_dir = os.environ.get("PYTHONHOME", "/data/user/0/com.your.app/files") 
    DATA_FILE = os.path.join(data_dir, ".hassan_v15_final.dat")
PLANET_MEANINGS = {
    "الشمس": "جوهر الذات، الإرادة الحيوية، والقدرة على القيادة والتألق،",
    "القمر": "الاحتياجات العاطفية، العقل الباطن، وكيفية الشعور بالأمان،",
    "عطارد": "الذكاء التحليلي، أسلوب التواصل، وقدرة العقل على معالجة البيانات،",
    "الزهرة": "التناغم الاجتماعي، القيم المادية، وطريقة التعبير عن الحب والجمال،",
    "المريخ": "طاقة المبادرة، الشجاعة، وكيفية توجيه القوة والاندفاع نحو الأهداف،",
    "المشتري": "فرص التوسع، الحظ، الفلسفة العالية وقدرة الشخص على النمو،",
    "زحل": "هيكلية الانضباط، الدروس الكرمية، والمسؤوليات التي تتطلب صبراً،",
    "أورانوس": "شرارة العبقرية، التغيير المفاجئ، والميل نحو التحرر والابتكار،",
    "نبتون": "الخيال الروحي، الحدس العميق، والقدرة على التماهي مع الفنون والروحانيات،",
    "بلوتو": "قوة التحول الجذري، تدمير القديم لبناء الجديد، والقدرة على النهوض،",
    "العقدة الشمالية": "بوصلة الروح والغاية القصوى التي يجب على الإنسان تعلمها في حياته،",
    "العقدة الجنوبية": "الرصيد الفطري والمهارات التي ولدت بها ولكنها قد تعيقك إذا تمسكت بها،"
}
# تأثير البروج (الصبغة أو الأسلوب)
SIGN_FLAVORS = {
    "الحمل": "بأسلوب ناري، مبادر، ويميل للاندفاع والشجاعة المطلقة.",
    "الثور": "بأسلوب ترابي ثابت، يميل للواقعية، الصبر، والبحث عن الاستقرار المادي.",
    "الجوزاء": "بأسلوب هوائي مرن، فضولي، ويميل للتواصل المتعدد والذكاء الاجتماعي.",
    "السرطان": "بأسلوب مائي عاطفي، يميل للحماية، الرعاية، والارتباط العميق بالجذور.",
    "الأسد": "بأسلوب ناري قيادي، يميل للظهور، الثقة بالنفس، وطلب التقدير والتميز.",
    "العذراء": "بأسلوب ترابي تحليلي، يميل للدقة، التنظيم، والبحث عن الكمال والخدمة.",
    "الميزان": "بأسلوب هوائي دبلوماسي، يميل للتوازن، العدل، والبحث عن الجمال والشراكة.",
    "العقرب": "بأسلوب مائي مكثف، يميل للعمق، التحول، والقوة الخفية والإرادة الصلبة.",
    "القوس": "بأسلوب ناري استكشافي، يميل للتفاؤل، الفلسفة، وحب الحرية والمعرفة العالية.",
    "الجدي": "بأسلوب ترابي طموح، يميل للجدية، الانضباط، وبناء المكانة الاجتماعية المرموقة.",
    "الدلو": "بأسلوب هوائي ابتكاري، يميل للتحرر، التفكير الإنساني، والتمرد على القوالب التقليدية.",
    "الحوت": "بأسلوب مائي روحاني، يميل للخيال، التضحية، والتماهي العميق مع المشاعر والكون."
}
# دلالات البيوت (مجال التجربة الحياتية)
HOUSE_FIELDS = {
    "1": "في بناء الهوية الشخصية، المظهر الخارجي، وكيفية مواجهة العالم لأول مرة.",
    "2": "في ميدان الموارد المالية، القيم الذاتية، والممتلكات والقدرة على الكسب.",
    "3": "في مجال التواصل اليومي، الإخوة، الدراسات الأولية، والذكاء العملي والأسفار القصيرة.",
    "4": "في أعماق الحياة الخاصة، المنزل، الجذور العائلية، والارتباط بالأرض والوالدين.",
    "5": "في عالم الإبداع، التعبير عن الذات، المواهب الفنية، العواطف، والأطفال.",
    "6": "في مجال العمل الروتيني، الصحة الجسدية، تطوير المهارات، والخدمة والواجبات.",
    "7": "في ميدان الشراكات الرسمية، الزواج، العلاقات المتوازنة، والأعداء الظاهرين.",
    "8": "في مناطق التحول الجذري، المواريث، المال المشترك، الأسرار، والعلوم الباطنية.",
    "9": "في مجال الوعي العالي، السفر البعيد، التعليم الجامعي، الفلسفة، والقانون.",
    "10": "في قمة الطموح المهني، السمعة العامة، السلطة، والنجاح في العالم الخارجي.",
    "11": "في دائرة الأصدقاء، المجموعات، الآمال الكبيرة، والمشاريع الجماعية الإنسانية.",
    "12": "في عالم الخفاء، الروحانيات، العزلة، والتعامل مع الأعداء المخفيين والماضي الكرمي."
}
ASTRO_INTERPRETATIONS = {
    "fortune_part": {
        "الشمس": "وفرة مالية عبر المناصب والظهور الاجتماعي.",
        "القمر": "رزق يأتي من التجارة، السفر، أو الدعم العائلي.",
        "المريخ": "مكاسب تأتي من الجهد البدني، الشجاعة، أو المبادرة.",
        # أضف باقي الكواكب...
    },
    "firdar_ruler": {
        "الشمس": "فترة للبروز، تحقيق الذات، وريادة الأعمال.",
        "الزهرة": "فترة ذهبية للزواج، الفنون، وتحسين العلاقات الاجتماعية.",
        "زحل": "فترة تتطلب الصبر، العمل الجاد، وبناء الأسس المتينة.",
        # أضف باقي الكواكب...
    }
}
PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]

PLANETARY_HOURS_INFO = {
    "الشمس": "ساعة للظهور، طلب المناصب، التعامل مع المسؤولين، والبدايات التي تحتاج شهرة.",
    "الزهرة": "ساعة للجمال، الزواج، التصالح، الفنون، وشراء الأشياء الثمينة.",
    "عطارد": "ساعة للتجارة، الكتابة، الدراسة، المراسلات، والعمليات الحسابية والذكاء.",
    "القمر": "ساعة للتغيرات، التنقل، التعامل مع الجمهور، الأمور العائلية، والعواطف.",
    "زحل": "ساعة للعمل الشاق، البناء، العقارات، التعامل مع كبار السن، وتتطلب الصبر.",
    "المشتري": "ساعة الحظ، التوسع، طلب العلم، الأمور الدينية، والنمو المالي.",
    "المريخ": "ساعة القوة، الرياضة، الجراحة، المواجهة، ويُنصح فيها بالحذر من الاندفاع."
}
# الترتيب الكلداني للكواكب
PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]
# حكام الأيام (يبدأ اليوم الفلكي من الشروق)
DAY_RULERS = {
    "Sunday": "الشمس", "Monday": "القمر", "Tuesday": "المريخ", 
    "Wednesday": "عطارد", "Thursday": "المشتري", "Friday": "الزهرة", "Saturday": "زحل"
}

PLANETARY_WARNINGS_DB = {
    # --- التراجعات (Retrogrades) ---
    "retro_mercury": "⚠️ تراجع عطارد: فساد في الفكر والرسائل، تعطل الأسفار القريبة، وخلل في البيع والشراء.",
    "retro_venus": "⚠️ تراجع الزهرة: فساد في النكاح، نفور بين المحبين، ومراجعة في أمور الزينة والمال.",
    "retro_mars": "⚠️ تراجع المريخ: حدة في الغضب، تعطل المبادرات، ووهن في العزيمة البدنية.",
    "retro_jupiter": "⚠️ تراجع المشتري: نقص في البركة، ضيق في الرزق الواسع، ومراجعة للقضايا القانونية.",
    "retro_saturn": "⚠️ تراجع زحل: ثقل في المسؤوليات، تأخير في البناء والعقارات، وعودة ديون قديمة.",

    # --- القواطع والاقترانات الصعبة ---
    "malefic_conjunction": "🚨 اقتران النحسين (زحل والمريخ): اجتماع 'القاطع' و'المانع'. شدة وضيق تتطلب صبراً عظيماً.",
    "saturn_vital": "⚠️ حصار زحل لكوكب حيوي: تقييد، حزن، أو تأخير قسري يطال شؤون {planet}.",
    "mars_vital": "⚠️ احتكاك المريخ بكوكب حيوي: تهور، اندفاع، أو خطر التهاب ونزاع يخص {planet}.",
    "mars_opposition": "⚡ مقابلة المريخ: مواجهة مباشرة مع الأعداء، أو تهور يؤدي لنتائج عكسية.",

    # --- الحالات الفنية (احتراق، خلو مسار، درجات حرجة) ---
    "combust": "☄️ احتراق {planet}: الكوكب في 'حيز الشمس'؛ ضعفت قوته واختفت دلالته تحت الشعاع.",
    "void_of_course": "🌑 القمر خالي المسار: سعي بلا طائل. تجنب البدء بأي أمر ذي بال حتى ينتقل القمر.",
    "critical_degree": "📐 درجة حرجة (0 أو 29): الكوكب في حالة 'المنقلب'؛ نهاية مرحلة وبداية تحول مفاجئ.",
    "saturn_return": "🪐 عودة زحل: ساعة المحاسبة الكبرى. اختبار للنضج وإعادة هيكلة أساسات الحياة.",

    # --- النجوم الثابتة (Fixed Stars) ---
    "star_algol": "💀 اقتران رأس الغول (Algol): أشر النجوم؛ خطر فقدان السيطرة، أو قرارات مدمرة للعقل.",
    "star_aldebaran": "🌟 اقتران الدبران: نجاح عظيم محفوف بمخاطر السقوط الأخلاقي أو النزاعات القانونية.",
    "star_regulus": "🦁 اقتران قلب الأسد (Regulus): ارتقاء للمناصب العالية، مع خطر الانتقام من الأنداد.",
    "star_fomalhaut": "🌟 اقتران فم الحوت: نجاح في العلوم الروحية والفنون، مع خطر الوقوع في الأوهام.",
    "star_deneb_kaitos": "🌊 اقتران ذنب قيطس: دلالة على العزلة الإجبارية، التعب الجسدي، أو الحزن المفاجئ.",
    "star_antares": "🌟 اقتران قلب العقرب (Antares): تهور قتالي، نجاح سريع متبوع بنهاية مفاجئة.",

    # --- قواطع كتب السابقين ---
    "besieged": "🔒 كوكب محاصر بين النحسين: الكوكب {planet} في ضيق شديد، دلالته معطلة تماماً.",
    "hayz_violation": "🚫 فساد الهيئة (خروج عن الحيز): الكوكب يعمل عكس طبيعته، مما يسبب نتائج غير متوقعة."
}
FIRDAR_INTERPRETATIONS = {
    "الشمس": {
        "title": "فترة الظهور والسيادة (10 سنوات)",
        "desc": "مرحلة تتركز حول تحقيق الذات، الطموح المهني، والتعامل مع أصحاب السلطة. فترة جيدة للبروز الاجتماعي."
    },
    "الزهرة": {
        "title": "فترة الجمال والمال (8 سنوات)",
        "desc": "مرحلة ذهبية للعلاقات العاطفية، الزواج، الفنون، وتحسن الدخل المادي والرفاهية."
    },
    "عطارد": {
        "title": "فترة العلم والتجارة (13 سنة)",
        "desc": "نشاط ذهني مكثف، كثرة الأسفار، نجاح في الدراسة أو التجارة والكتابة وتوسيع شبكة التواصل."
    },
    "القمر": {
        "title": "فترة المتغيرات والعائلة (9 سنوات)",
        "desc": "تركز على الاستقرار المنزلي، المشاعر، والتغيرات في السكن أو الحالة العائلية."
    },
    "زحل": {
        "title": "فترة الحصاد والمسؤولية (11 سنة)",
        "desc": "مرحلة تتطلب الصبر والعمل الجاد. بناء أسس طويلة الأمد أو مواجهة تحديات تتطلب نضجاً."
    },
    "المشتري": {
        "title": "فترة الوفرة والحظ (12 سنة)",
        "desc": "توسع في الرزق، تفاؤل، نجاحات قانونية أو دينية، وفترة مناسبة للسفر البعيد والنمو."
    },
    "المريخ": {
        "title": "فترة القوة والصراع (7 سنوات)",
        "desc": "طاقة عالية، اندفاع، مبادرات عملية شجاعة، لكنها تتطلب الحذر من الخلافات أو التهور الجسدي."
    },
    "الرأس": {
        "title": "فترة الصعود المفاجئ (3 سنوات)",
        "desc": "تغييرات قدرية، طموحات لا محدودة، وغالباً ما تشهد قفزة نوعية في مكانة الشخص."
    },
    "الذنب": {
        "title": "فترة التطهير والروحانية (سنتان)",
        "desc": "فترة زهد أو مراجعة للماضي، تخلص من علاقات أو أشياء لم تعد تخدمك، والتركيز على الباطن."
    }
}
LOTS_INTERPRETATIONS = {
    "fortune": "معدن الرزق، الحظ الجسدي، والمال الحلال وصحة البدن.",
    "spirit": "معدن القوة الروحية، النوايا، الطموح، وما يخطط له العقل.",
    "love": "سهم الزهرة: يتعلق بالعواطف، الصداقات، والمتعة والحياة الاجتماعية.",
    "necessity": "سهم الشدة: يتعلق بالقيود، المتاعب، وما يُفرض على الإنسان قسراً.",
    "victory": "سهم النصر: يتعلق بالنجاح في المنازعات، القوة، والوصول للأهداف.",
    "marriage": "سهم الزواج: يحدد جودة الشراكة العاطفية واستقرار الحياة الزوجية.",
    "enemies": "سهم الأعداء: يكشف عن الخصوم المخفيين والمكائد التي قد تحاك."
}
def check_comprehensive_warnings(data):
    alerts = []
    
    # تأكد من وجود قاموس التحذيرات لتجنب KeyError
    # يفترض أن PLANETARY_WARNINGS_DB معرفة في مكان ما في الكود
    
    # 1. التراجعات
    retro_map = {'عطارد': 'mercury', 'الزهرة': 'venus', 'المريخ': 'mars', 'المشتري': 'jupiter', 'زحل': 'saturn'}
    for ar_p, en_p in retro_map.items():
        if data.get(f'is_{en_p}_retro'):
            alerts.append(f"⚠️ تراجع {ar_p}: " + PLANETARY_WARNINGS_DB.get(f"retro_{en_p}", "تنبيه تراجع"))

    # 2. النحسين (زحل والمريخ) - اقتران النحوس
    if data.get('mars_saturn_dist', 100) < 6:
        alerts.append("🔴 " + PLANETARY_WARNINGS_DB.get("malefic_conjunction", "اقتران المريخ وزحل"))

    # 3. الاحتراق (تعديل لتجنب احتراق الشمس بنفسها)
    sun_distances = data.get('sun_distances', {})
    for p, dist in sun_distances.items():
        if p != "الشمس" and dist < 8.5:
            msg = PLANETARY_WARNINGS_DB.get("combust", "احتراق كوكب {planet}")
            alerts.append(f"🔥 {msg.format(planet=p)}")

    # 4. النجوم الثابتة (إضافة أورب دقيق)
    for star in ['algol', 'aldebaran', 'regulus', 'fomalhaut', 'deneb_kaitos', 'antares']:
        if data.get(f'conjunct_{star}'):
            alerts.append("✨ " + PLANETARY_WARNINGS_DB.get(f"star_{star}", f"اقتران مع نجم {star}"))

    # 5. خلو المسار والدرجات الحرجة
    if data.get('is_moon_voc'):
        alerts.append("🚫 " + PLANETARY_WARNINGS_DB.get("void_of_course", "القمر في خلو مسار"))
    
    # الدرجات الحرجة (أول وآخر درجة في البرج)
    for p, deg in data.get('positions', {}).items():
        pos_in_sign = deg % 30
        if pos_in_sign < 1.0 or pos_in_sign > 29.0:
            msg = PLANETARY_WARNINGS_DB.get('critical_degree', "درجة حرجة")
            alerts.append(f"📍 {p}: {msg}")

    # تحويل النص النهائي للعربية الصحيحة (Bidi) قبل العرض
    final_report = "\n".join(alerts) if alerts else "✅ الهيئة متزنة ولا توجد قواطع سلبية حالياً."
    return final_report
def open_astro_logic(e):
    page.clean()
    
    # محاكاة بيانات (سيتم استبدالها لاحقاً بحسابات swisseph)
    my_data = {
        "is_mercury_retro": True,
        "sun_distances": {"الزهرة": 5.2},
        "positions": {"القمر": 29.5},
        "is_moon_voc": False
    }
    
    result_text = check_comprehensive_warnings(my_data)
    
    # تحويل النص ليدعم العرض الصحيح (Bidi)
    # ملاحظة: Flet الحديث يدعم العربية ولكن bidi تضمن التنسيق في القوائم
    display_text = get_display(arabic_reshaper.reshape(result_text))

    page.add(
        ft.AppBar(
            title=ft.Text(fx("نتائج فحص الهيئة"), color="white"),
            bgcolor=ft.colors.BLUE_GREY_900
        ),
        ft.Container(
            content=ft.Column([
                ft.Text(
                    display_text, 
                    rtl=True, 
                    size=18,
                    color=ft.colors.WHITE,
                    text_align=ft.TextAlign.RIGHT
                ),
                ft.ElevatedButton(
                    text=fx("العودة للرئيسية"), 
                    on_click=lambda _: main_menu(page) # افترضت وجود دالة رئيسية
                )
            ], scroll=ft.ScrollMode.ADAPTIVE),
            padding=20,
            expand=True,
            bgcolor=ft.colors.BLACK87
        )
    )
    page.update()

def fx(text):
    if not text: return ""
    # إعادة تشكيل الحروف العربية لتظهر متصلة وصحيحة
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)

def get_hwid(page: ft.Page):
    # 1. التحقق من وجود البصمة المخزنة
    stored_hwid = page.client_storage.get("user_hwid")
    
    if stored_hwid is None:
        # 2. إنشاء بصمة تعتمد على خصائص الجهاز الفعلية (إن وجدت) + UUID
        # نستخدم uuid4 مع Salt لتعقيد البصمة
        raw_id = f"{platform.processor()}_{platform.node()}_{uuid.uuid4()}"
        
        # 3. تشفير البصمة باستخدام SHA-256 لجعلها ثابتة الطول (مثلاً 16 حرفاً)
        hashed_id = hashlib.sha256((raw_id + SECRET_SALT).encode()).hexdigest().upper()
        short_id = hashed_id[:16] # نأخذ أول 16 حرف لتسهيل التعامل معها
        
        # 4. حفظ البصمة في ذاكرة الجهاز الدائمة
        page.client_storage.set("user_hwid", short_id)
        return short_id
    
    return stored_hwid
def generate_activation_key(hwid):
    """توليد كود التفعيل بناءً على بصمة الجهاز"""
    raw_key = hwid + SECRET_SALT
    return hashlib.sha256(raw_key.encode()).hexdigest()[:12].upper()

def get_trial_info(page: ft.Page):
    """حساب الأيام المتبقية باستخدام ذاكرة التطبيق (أضمن للأندرويد)"""
    today = datetime.now().date()
    
    # التحقق من تاريخ أول تشغيل
    if not page.client_storage.contains_key("first_run_date"):
        page.client_storage.set("first_run_date", today.strftime("%Y-%m-%d"))
        return TRIAL_DAYS
    
    try:
        first_date_str = page.client_storage.get("first_run_date")
        first_date = datetime.strptime(first_date_str, "%Y-%m-%d").date()
        elapsed = (today - first_date).days
        return max(0, TRIAL_DAYS - elapsed)
    except:
        return 0

def show_activation_dialog(page: ft.Page, hwid):
    page.clean()
    code_input = ft.TextField(label=fx("أدخل كود التفعيل"), password=True, can_reveal_password=True)
    
    def verify_click(e):
        correct_key = generate_activation_key(hwid)
        if code_input.value == correct_key:
            page.client_storage.set("is_activated", True)
            page.snack_bar = ft.SnackBar(ft.Text(fx("تم التفعيل! أعد تشغيل التطبيق")))
            page.snack_bar.open = True
            page.update()
            time.sleep(2)
            main_dashboard(page)
        else:
            code_input.error_text = fx("الكود غير صحيح!")
            page.update()

    page.add(
        ft.Text(fx("انتهت الفترة التجريبية"), size=25, color="red"),
        ft.Text(fx(f"معرف الجهاز: {hwid}"), selectable=True),
        code_input,
        ft.ElevatedButton(text=fx("تفعيل الآن"), on_click=verify_click)
    )
def show_splash(page: ft.Page):
    page.clean()
    page.bgcolor = "black"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    
    logo = ft.Text("✧ ASC ✧", size=70, weight="bold", color="#FFD700")
    title = ft.Text("MEEQAT ULTIMATE PRO", size=30, weight="bold", color="#39FF14")
    progress = ft.ProgressBar(width=400, color="#FFD700", value=0)
    
    page.add(logo, title, ft.Text("By Hassan Al-Shaer", color="#BB86FC"), progress)
    
    # حلقة التحميل
    for i in range(1, 11):
        progress.value = i / 10
        page.update()
        time.sleep(0.15)
    
    # الانتقال لنظام التفعيل الذي صممته أنت
    run_auth_system(page) 
def run_auth_system(page: ft.Page):
    hwid = get_hwid(page)
    valid_key = generate_activation_key(hwid)
    days_left = get_trial_info(page)

    page.clean()
    page.bgcolor = ft.colors.BLACK
    
    # حقل النص
    entry = ft.TextField(
        label="أدخل كود التفعيل", 
        width=300, 
        password=True, 
        can_reveal_password=True,
        text_align=ft.TextAlign.CENTER
    )

    # دالة اللصق للأندرويد
def do_paste(e):
    res = page.get_clipboard()
    if res:
        entry.value = res
        page.update()

    def on_verify(e):
        if entry.value.upper().strip() == valid_key:
            page.client_storage.set("license_key", valid_key)
            open_astro_logic(page) # دالة فتح البرنامج
        else:
            page.snack_bar = ft.SnackBar(ft.Text("كود تفعيل خاطئ"))
            page.snack_bar.open = True
            page.update()

    def on_trial(e):
        if days_left > 0:
            open_astro_logic(page)
        else:
            page.snack_bar = ft.SnackBar(ft.Text("انتهت الفترة التجريبية"))
            page.snack_bar.open = True
            page.update()

    # بناء الواجهة
    page.add(
        ft.Text("تفعيل الميقاتي برو", size=25, weight="bold", color="white"),
        ft.Text(f"ID: {hwid}", color="gray", size=12),
        ft.Row([
            entry, 
            ft.IconButton(icon=ft.icons.PASTE, on_click=do_paste, tooltip="لصق الكود")
        ], alignment=ft.MainAxisAlignment.CENTER),
        ft.ElevatedButton("تفعيل النسخة الأصلية", on_click=on_verify, bgcolor="green", color="white"),
        ft.TextButton(f"الدخول كنسخة تجريبية ({days_left} يوم)", on_click=on_trial),
        ft.Text(f"الأيام المتبقية: {days_left}", color="orange")
    )
    page.update()
def main(page: ft.Page):
    # إعدادات الصفحة الأساسية
    page.window_width = 400
    page.window_height = 800
    # ابدأ بشاشة التحميل
    show_splash(page)

ft.app(target=main)
# ==========================================
# 3. محرك البرنامج الرئيسي (Master Engine)
# ==========================================
class HassanAstroMobile:
    def __init__(self, page: ft.Page, mode):
        self.page = page
        
        # 1. إعدادات العنوان (المتوافقة مع أندرويد)
        version_type = "الكاملة" if mode == "full" else "التجريبية"
        self.page.title = f"Hassan Astro Pro - {version_type}"
        self.page.bgcolor = "#ffffff"
        self.page.rtl = True # دعم العربية
        
        # 2. تعريف المتغيرات (بدل tk.StringVar نستخدم متغيرات عادية)
        self.house_system = "Placidus"
 # 1. أرباب البيوت
        self.domiciles = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", 
                          "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
        
        # 2. أرباب الشرف
        self.exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}

        # 3. أرباب المثلثات (نهار/ليل)
        self.triplicities = {
            "ناري":  {"day": "الشمس", "night": "المشتري"},
            "ترابي": {"day": "الزهرة", "night": "القمر"},
            "هوائي": {"day": "زحل", "night": "عطارد"},
            "مائي":  {"day": "الزهرة", "night": "المريخ"}
        }

        # 4. الحدود المصرية (Egyptian Terms) - الدقة القصوى
        # كل قائمة تحتوي على (الدرجة القصوى للحد، الكوكب صاحب الحد)
        self.egyptian_terms = {
            0:  [(6, "المشتري"), (12, "الزهرة"), (20, "عطارد"), (25, "المريخ"), (30, "زحل")],  # الحمل
            1:  [(8, "الزهرة"), (14, "عطارد"), (22, "المشتري"), (27, "زحل"), (30, "المريخ")],  # الثور
            2:  [(6, "عطارد"), (12, "المشتري"), (17, "الزهرة"), (24, "المريخ"), (30, "زحل")],  # الجوزاء
            3:  [(7, "المريخ"), (13, "الزهرة"), (19, "عطارد"), (26, "المشتري"), (30, "زحل")],  # السرطان
            4:  [(6, "المشتري"), (11, "الزهرة"), (18, "زحل"), (24, "عطارد"), (30, "المريخ")],  # الأسد
            5:  [(7, "عطارد"), (17, "الزهرة"), (21, "المشتري"), (28, "المريخ"), (30, "زحل")],  # العذراء
            6:  [(6, "زحل"), (14, "الزهرة"), (21, "المشتري"), (28, "عطارد"), (30, "المريخ")],  # الميزان
            7:  [(7, "المريخ"), (11, "الزهرة"), (19, "عطارد"), (24, "المشتري"), (30, "زحل")],  # العقرب
            8:  [(12, "المشتري"), (17, "الزهرة"), (21, "عطارد"), (26, "زحل"), (30, "المريخ")],  # القوس
            9:  [(7, "عطارد"), (14, "المشتري"), (22, "الزهرة"), (26, "زحل"), (30, "المريخ")],  # الجدي
            10: [(7, "عطارد"), (13, "الزهرة"), (20, "المشتري"), (25, "المريخ"), (30, "زحل")],  # الدلو
            11: [(12, "الزهرة"), (16, "المشتري"), (19, "عطارد"), (28, "المريخ"), (30, "زحل")]   # الحوت
        }        
        self.z_syms = ["♈", "♉", "♊", "♋", "♌", "♍", "♎", "♏", "♐", "♑", "♒", "♓"]
        self.z_names = ["الحمل ♈", "الثور ♉", "الجوزاء ♊", "السرطان ♋", "الأسد ♌", "العذراء ♍", "الميزان ♎", "العقرب ♏", "القوس ♐", "الجدي ♑", "الدلو ♒", "الحوت ♓"]

  # --- الإضافة الجديدة: إعدادات محرك تفسير الأحلام الفلسفي ---
        # حساب اليوم القمري بدقة باستخدام مكتبة أم القرى
        try:
            now = datetime.now()
            hijri_now = HijriDate(now.year, now.month, now.day, gr=True)
            self.hijri_day = hijri_now.day
            self.hijri_month_name = hijri_now.month_name() # إضافة اسم الشهر لجمالية الواجهة
        except Exception as e:
            # في حال حدوث خطأ في التاريخ، نستخدم الحساب التقريبي كخطة بديلة
            import math
            diff = datetime.now() - datetime(2000, 1, 6)
            self.hijri_day = math.floor(((diff.days / 29.5305) - math.floor(diff.days / 29.5305)) * 29.53) + 1
            print(f"Fallback to math-based lunar day: {e}")        
        # تعريف فلسفة ابن عربي والكلدانيين كمرجع ثابت
        self.dream_philosophy = {
            "ibn_arabi": "الرؤيا هي اتصال الروح بعالم المثال (الخيال المنفصل).",
            "chaldean": "الأحلام هي انعكاس لنفوذ الكواكب السبعة على النفس البشرية."
        }
        # -------------------------------------------------------

        self.build_flet_ui() # دالة سنقوم بتعريفها لبناء واجهة Flet
    def build_flet_ui(self):
        # تنظيف الصفحة وبناء الواجهة الأساسية
        self.page.clean()
        
        # حاوية النتائج (التي ستستقبل الـ 104 أسطر)
        self.report_container = ft.Column(
            controls=[], 
            scroll=ft.ScrollMode.ALWAYS,
            expand=True
        )

        self.page.add(
            ft.AppBar(
                title=ft.Text("الميقاتي الفلكي - Master Edition"), 
                bgcolor=ft.colors.SURFACE_VARIANT,
                actions=[
                    ft.IconButton(ft.icons.COPY, on_click=self.copy_report) # زر لنسخ التقرير كاملاً
                ]
            ),
            ft.Text(f"📅 اليوم القمري: {self.hijri_day}", size=18, weight="bold", rtl=True),
            ft.Divider(),
            self.report_container,
            ft.FloatingActionButton(
                icon=ft.icons.PLAY_ARROW, 
                text="بدء الحسابات الملحمية",
                on_click=lambda _: self.start_calculations()
            )
        )

    # هذه الدالة يجب أن تكون خارج build_flet_ui ولكن داخل الـ Class
    def log_to_flet(self, message, color="black"):
        """البديل الرسمي لـ txt.insert - متوافق مع أندرويد"""
        self.report_container.controls.append(
            ft.Text(value=message, color=color, size=16, rtl=True, selectable=True)
        )
        self.page.update()
# الدالة هنا (داخل الكلاس)
    def copy_report(self, e):
        """دالة لنسخ كافة النتائج المسجلة في الحاوية"""
        full_text = "\n".join([c.value for c in self.report_container.controls if isinstance(c, ft.Text)])
        self.page.set_clipboard(full_text)
        
        self.page.snack_bar = ft.SnackBar(
            content=ft.Text(fx("تم نسخ التقرير الفلكي بنجاح!")),
        )
        self.page.snack_bar.open = True
        self.page.update()
    def get_egyptian_bounds(self, sign_idx, deg): 
        """استخراج صاحب الحد بناءً على البرج والدرجة"""
        terms = {
            0: [(6, "المشتري"), (12, "الزهرة"), (20, "عطارد"), (25, "المريخ"), (30, "زحل")],
            1: [(8, "الزهرة"), (14, "عطارد"), (22, "المشتري"), (27, "زحل"), (30, "المريخ")],
            2: [(6, "عطارد"), (12, "المشتري"), (17, "الزهرة"), (24, "المريخ"), (30, "زحل")],
            3: [(7, "المريخ"), (13, "الزهرة"), (19, "عطارد"), (26, "المشتري"), (30, "زحل")],
            4: [(6, "المشتري"), (11, "الزهرة"), (18, "زحل"), (24, "عطارد"), (30, "المريخ")],
            5: [(7, "عطارد"), (17, "الزهرة"), (21, "المشتري"), (28, "المريخ"), (30, "زحل")],
            6: [(6, "زحل"), (14, "الزهرة"), (21, "المشتري"), (28, "عطارد"), (30, "المريخ")],
            7: [(7, "المريخ"), (11, "الزهرة"), (19, "عطارد"), (24, "المشتري"), (30, "زحل")],
            8: [(12, "المشتري"), (17, "الزهرة"), (21, "عطارد"), (26, "زحل"), (30, "المريخ")],
            9: [(7, "عطارد"), (14, "المشتري"), (22, "الزهرة"), (26, "زحل"), (30, "المريخ")],
            10:[(7, "عطارد"), (13, "الزهرة"), (20, "المشتري"), (25, "المريخ"), (30, "زحل")],
            11:[(12, "الزهرة"), (16, "المشتري"), (19, "عطارد"), (28, "المريخ"), (30, "زحل")]
        }
        
        # البحث عن الكوكب الذي يقع الدرجة ضمن حده
        for limit, planet in terms.get(sign_idx, []):
            if deg < limit: # تم تعديلها لـ أصغر من لضمان دقة التقسيم
                return planet
        return "زحل" # كوكب افتراضي في حال الخطأ
    def calculate_real_almuten(self, target_deg, jd, is_day):
        """حساب المبتز (المستولي) بناءً على الحظوظ الخمسة"""
        domiciles = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
        exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}
        triplicities = {
            "Fire": ["الشمس", "المشتري"], "Earth": ["الزهرة", "القمر"],
            "Air": ["زحل", "عطارد"], "Water": ["الزهرة", "المريخ"]
        }
        
        sign_idx = int(target_deg / 30) % 12
        deg_in_sign = target_deg % 30
        scores = {k: 0 for k in ["الشمس", "القمر", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل"]}
        
        # 1. رب البيت (5 نقاط)
        scores[domiciles[sign_idx]] += 5
        
        # 2. رب الشرف (4 نقاط)
        if sign_idx in exaltations: 
            scores[exaltations[sign_idx]] += 4
            
        # 3. رب المثلثة (3 نقاط)
        element_idx = [0, 1, 2, 3][sign_idx % 4] # 0:Fire, 1:Earth...
        element_name = ["Fire", "Earth", "Air", "Water"][element_idx]
        scores[triplicities[element_name][0 if is_day else 1]] += 3
        
        # 4. رب الحد (نقطتان) - تأكد من تطابق اسم الدالة هنا مع ما عرفناه سابقاً
        scores[self.get_egyptian_bounds(sign_idx, deg_in_sign)] += 2
        
        # 5. رب الوجه (نقطة واحدة)
        faces = ["المريخ", "الشمس", "الزهرة", "عطارد", "القمر", "زحل", "المشتري"]
        scores[faces[int(target_deg / 10) % 7]] += 1

        # تحديد الفائز (المبتز)
        winner = max(scores, key=scores.get)
        
        # خريطة الثوابت لمكتبة Swiss Ephemeris
        p_map = {
            "الشمس": swe.SUN, "القمر": swe.MOON, "عطارد": swe.MERCURY, 
            "الزهرة": swe.VENUS, "المريخ": swe.MARS, "المشتري": swe.JUPITER, "زحل": swe.SATURN
        }
        
        # حساب موقع الكوكب المبتز في اللحظة المطلوبة
        res = swe.calc_ut(jd, p_map[winner])
        return winner, res[0] # إرجاع اسم الكوكب ودرجته الفلكية
    def draw_astro_wheel(self, jd, title):
        """
        تقوم بحساب المواقع الفلكية وعرضها في واجهة Flet وجدول البيانات.
        """
        # 1. تنظيف الواجهة والجدول قبل البدء
        try:
            self.tree.delete(*self.tree.get_children())
        except:
            pass # في حال لم يكن الجدول متاحاً بعد

        # 2. حساب نظام البيوت والأوتاد
        # التأكد من جلب النظام المختار أو الافتراضي 'P' (Placidus)
        selected_system = self.systems_map.get(self.house_system.get(), b'P')
        cusps, ascmc = swe.houses(jd, self.lat, self.lon, selected_system)
        
        asc = float(ascmc[0])
        mc = float(ascmc[1])
        
        # 3. تعريف بيانات الكواكب (الأساسية والإضافية)
        p_data = {
            0: ("الشمس", "☉", "#f59e0b"), 1: ("القمر", "☽", "#3b82f6"), 
            2: ("عطارد", "☿", "#8b5cf6"), 3: ("الزهرة", "♀", "#ec4899"), 
            4: ("المريخ", "♂", "#dc2626"), 5: ("المشتري", "♃", "#d97706"), 
            6: ("زحل", "♄", "#475569"), 7: ("أورانس", "♅", "#06b6d4"), 
            8: ("نبتون", "♆", "#6366f1"), 9: ("بلوتو", "♇", "#a855f7"), 
            10: ("الرأس", "☊", "#059669"), 12: ("ليلث", "⚸", "#000000"),
            13: ("كايرون", "Chiron", "#5b21b6")
        }

        # 4. عرض الأوتاد الرئيسية
        self.log_to_flet(f"✨ {title}", color="blue", size=20)
        self.log_to_flet("📐 الأوتاد الرئيسية:", color="red", weight="bold")
        
        for deg, (n_ar, n_en, clr) in {asc: ("الطالع", "ASC", "#ef4444"), mc: ("وسط السماء", "MC", "#10b981")}.items():
            d = int(deg % 30)
            m = int((deg - int(deg)) * 60)
            sign_idx = int(deg / 30)
            self.log_to_flet(f"{n_en} {n_ar}: {self.z_names[sign_idx]} {d}°{m:02d}'", color=clr)
            self.tree.insert("", "end", values=(n_en, self.z_names[sign_idx], f"{d}°{m:02d}'", "-"), tags=(n_en,))

        # 5. حساب وعرض الكواكب
        self.log_to_flet("\n🪐 وضعية الكواكب:", color="blue", weight="bold")
        
        for pid, (n, s, c) in p_data.items():
            res, _ = swe.calc_ut(jd, pid)
            pos = float(res[0])
            h = self.get_house(pos, cusps) # استدعاء ميثود حساب البيت
            
            sign_idx = int(pos / 30)
            deg = int(pos % 30)
            mnt = int((pos - int(pos)) * 60)
            
            # الإضافة لـ Flet
            self.log_to_flet(f"{s} {n}: {self.z_names[sign_idx]} {deg}°{mnt:02d}' - البيت ({h})", color=c)
            
            # الإضافة للجدول
            self.tree.insert("", "end", values=(f"{s} {n}", self.z_names[sign_idx], f"{deg}°{mnt:02d}'", h), tags=(n,))
            self.tree.tag_configure(n, foreground=c)

            # إضافة الذنب (العقدة الجنوبية) تلقائياً
            if pid == 10:
                s_node_pos = (pos + 180) % 360
                h_sn = self.get_house(s_node_pos, cusps)
                self.log_to_flet(f"☋ الذنب: {self.z_names[int(s_node_pos/30)]} {int(s_node_pos%30)}° - البيت ({h_sn})", color="#059669")

        # 6. حساب السهام السرية (السعادة والغيب)
        sun_pos_raw, _ = swe.calc_ut(jd, 0)
        moon_res, _ = swe.calc_ut(jd, 1)
        sun, moon = float(sun_pos_raw[0]), float(moon_res[0])
        
        # ميلاد نهاري أم ليلي
        is_day = not (180 <= (sun - asc) % 360 <= 360) 
        fortuna = (asc + moon - sun) % 360 if is_day else (asc + sun - moon) % 360
        spirit = (asc + sun - moon) % 360 if is_day else (asc + moon - sun) % 360

        self.log_to_flet("\n💎 السهام السرية:", color="orange", weight="bold")
        for s_pos, s_name, s_sym in [(fortuna, "سهم السعادة", "⊗"), (spirit, "سهم الغيب", "❂")]:
            h_s = self.get_house(s_pos, cusps)
            sign_s = self.z_names[int(s_pos/30)]
            self.log_to_flet(f"{s_sym} {s_name}: {sign_s} {int(s_pos%30)}° - البيت ({h_s})", color="#f97316")

        # 7. عرض حدود البيوت
        self.log_to_flet("\n🏠 حدود البيوت الفلكية:", color="green", weight="bold")
        for i, cusp_val in enumerate(cusps):
            sign_idx = int(cusp_val / 30)
            deg = int(cusp_val % 30)
            self.log_to_flet(f"البيت {i+1}: {self.z_names[sign_idx]} {deg}°")
    def draw_now(self):
        # الحصول على الوقت الحالي العالمي (UTC)
        n = datetime.now(timezone.utc)
        # تحويل الوقت الحالي إلى التاريخ اليولياني بدقة الثواني
        jd = swe.julday(n.year, n.month, n.day, n.hour + n.minute/60.0 + n.second/3600.0)
        
        # استدعاء دالة الرسم والعرض
        self.draw_astro_wheel(jd, "خريطة اللحظة (Time Now)")

    def draw_2026(self):
        # إعداد التاريخ المحدد (دخول شمس 2026 أو أي حدث تريده)
        # 20 مارس 2026 الساعة 09:24
        jd = swe.julday(2026, 3, 20, 9 + 24/60.0)
        
        # 1. رسم الخريطة الفلكية وعرض البيانات الأساسية
        self.draw_astro_wheel(jd, "خريطة طالع سنة 2026")
        
        # 2. تحليل الهيئة العالمية (Mundane Analysis) 
        # تأكد أن دالة show_mundane_analysis معرفة لديك
        if hasattr(self, 'show_mundane_analysis'):
            self.show_mundane_analysis(jd)
def setup_ui(self):
    # إعدادات عامة للصفحة
    self.page.padding = 10
    self.page.spacing = 10
    self.page.rtl = True  # دعم اللغة العربية بالكامل
    self.page.theme_mode = ft.ThemeMode.LIGHT

    # 1. القائمة اليمين (جدول الكواكب) - DataTable
    self.tree = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("الجرم", weight="bold")),
            ft.DataColumn(ft.Text("البرج", weight="bold")),
            ft.DataColumn(ft.Text("الدرجة", weight="bold")),
            ft.DataColumn(ft.Text("بيت", weight="bold")),
        ],
        rows=[],  # يتم تعبئتها ديناميكياً بواسطة دالة draw_astro_wheel
        column_spacing=15,
        heading_row_height=40,
        data_row_min_height=35,
    )

    left_sidebar = ft.Container(
        content=ft.Column([
            ft.Text("📊 بيانات الهيئة", size=16, weight="bold", color="#1e293b"),
            self.tree
        ], scroll=ft.ScrollMode.ADAPTIVE),
        bgcolor="#ffffff",
        border=ft.border.all(1, "#cbd5e1"),
        border_radius=8,
        padding=10,
        expand=2,
    )

    # 2. القائمة اليسرى (أدوات التحكم)
    def section_title(text, color="#e2e8f0"):
        return ft.Container(
            content=ft.Text(text, weight="bold", size=14, color="#1e293b"),
            bgcolor=color,
            padding=5,
            alignment=ft.alignment.center,
            width=float("inf"),
            border_radius=5
        )

    # حقول إدخال الموقع والبيوت
    self.city_entry = ft.TextField(label="المدينة", value="Tartus", text_align="center", dense=True)
    self.house_menu = ft.Dropdown(
        label="نظام البيوت",
        options=[ft.dropdown.Option(k) for k in self.systems_map.keys()],
        value="Placidus",
        on_change=lambda e: self.draw_now(),
        dense=True
    )
    
    self.coords_lbl = ft.Text("عرض: 00° 00' | طول: 00° 00'", size=12, text_align="center", color="#64748b")

    # حقول إدخال الوقت والتاريخ (مع تفعيل التنقل التلقائي)
    self.day_ent = ft.TextField(label="يوم", width=60, text_align="center")
    self.month_ent = ft.TextField(label="شهر", width=60, text_align="center")
    self.year_ent = ft.TextField(label="سنة", width=80, text_align="center")
    self.hour_ent = ft.TextField(label="س", width=60, text_align="center")
    self.min_ent = ft.TextField(label="د", width=60, text_align="center")

    # ربط الحقول بوظيفة auto_skip
    self.day_ent.on_change = lambda e: self.auto_skip(e, self.month_ent, 2)
    self.month_ent.on_change = lambda e: self.auto_skip(e, self.year_ent, 2)
    self.year_ent.on_change = lambda e: self.auto_skip(e, self.hour_ent, 4)
    self.hour_ent.on_change = lambda e: self.auto_skip(e, self.min_ent, 2)

    # أزرار الروحانيات والأدوات
    spiritual_btns = ft.Column([
        ft.ElevatedButton("🌙 منازل القمر", bgcolor="#8b5cf6", color="white", on_click=self.show_moon_mansion, width=250),
        ft.ElevatedButton("💭 تفسير الأحلام", bgcolor="#4b5563", color="white", on_click=self.show_dream_interpreter, width=250),
        ft.ElevatedButton("✂️ قص الشعر", bgcolor="#ec4899", color="white", on_click=self.hair_cut_calendar, width=250),
        ft.ElevatedButton("💅 قص الأظافر", bgcolor="#be185d", color="white", on_click=self.nail_cut_calendar, width=250),
        ft.ElevatedButton("⏳ ساعات الكواكب", bgcolor="#0f172a", color="white", on_click=self.show_planetary_hours, width=250),
    ], horizontal_alignment="center", spacing=5)

    right_sidebar_content = ft.Column(
        [
            ft.Text("الميقاتي الفلكي - حسان الشاعر", color="#a11616", size=18, weight="bold", text_align="center"),
            section_title("📍 الموقع والبيوت"),
            self.city_entry,
            self.house_menu,
            ft.ElevatedButton("🌍 جلب الإحداثيات", bgcolor="#10b981", color="white", on_click=self.search_city_logic, width=250),
            self.coords_lbl,
            
            section_title("🔮 المحرك الرئيسي"),
            ft.ElevatedButton("⚖️ تصحيح الطالع", bgcolor="#7c3aed", color="white", on_click=self.show_animodar_correction, width=250),
            ft.ElevatedButton("🕒 الوقت الآن", bgcolor="#2563eb", color="white", on_click=lambda _: self.draw_now(), width=250),
            ft.ElevatedButton("🌍 طالع سنة العالم", bgcolor="#dc2626", color="white", on_click=lambda _: self.draw_2026(), width=250),
            
            section_title("👶 بيانات الميلاد"),
            ft.Row([self.day_ent, self.month_ent, self.year_ent], alignment="center", spacing=5),
            ft.Row([self.hour_ent, self.min_ent], alignment="center", spacing=5),
            ft.ElevatedButton("👶 تحليل الولادة", bgcolor="#059669", color="white", on_click=self.analyze_birth_chart, width=250),
            
            section_title("🌙 الروحانيات والأدوات"),
            spiritual_btns,
            
            ft.ElevatedButton("حول البرنامج", bgcolor="#1f538d", color="white", on_click=self.show_about, width=150),
        ],
        scroll=ft.ScrollMode.ALWAYS,
        horizontal_alignment="center",
        spacing=10
    )

    right_sidebar = ft.Container(
        content=right_sidebar_content,
        width=300,
        bgcolor="#f8fafc",
        border=ft.border.all(1, "#cbd5e1"),
        border_radius=10,
        padding=15
    )

    # 3. منطقة الخريطة المركزية (Canvas/Log)
    self.results_list = ft.Column(scroll=ft.ScrollMode.ALWAYS, expand=True)
    self.chart_container = ft.Container(
        content=self.results_list,
        expand=3,
        bgcolor="#ffffff",
        border=ft.border.all(1, "#cbd5e1"),
        border_radius=8,
        padding=10,
        alignment=ft.alignment.top_right
    )

    # تجميع الواجهة النهائية (استخدام Row للتطبيقات المكتبية أو ResponsiveRow للجوال)
    self.page.add(
        ft.Row(
            [
                left_sidebar,
                self.chart_container,
                right_sidebar,
            ],
            expand=True,
            spacing=10
        )
    )

def auto_skip(self, e, next_control, limit):
    if len(e.control.value) >= limit:
        next_control.focus()
        self.page.update()
def show_planetary_hours(self):
    try:
        # 1. جلب البيانات
        lat = getattr(self, 'lat', 34.88)
        lon = getattr(self, 'lon', 35.88)
        city_name = self.city_entry.value if hasattr(self, 'city_entry') else "غير محدد"
        now = datetime.now()
        
        # 2. حساب الحاكم (الترتيب الكلداني)
        # weekday() في بايثون: الاثنين=0 ... الأحد=6
        # نحتاج تحويلها لتناسب ترتيب: الأحد (الشمس)، الاثنين (القمر)...
        weekday = now.weekday() 
        # خريطة لربط اليوم ببداية أول ساعة (شروق الشمس)
        # 0:الاثنين(القمر)، 1:الثلاثاء(المريخ)، 2:الأربعاء(عطارد)، 3:الخميس(المشتري)، 4:الجمعة(الزهرة)، 5:السبت(زحل)، 6:الأحد(الشمس)
        day_rulers_start = {6: 3, 0: 6, 1: 2, 2: 5, 3: 1, 4: 4, 5: 0}
        start_idx = day_rulers_start.get(weekday, 0)
        
        # ملاحظة: للحصول على دقة 100% يفضل حساب الساعات بناءً على الشروق (Sunset/Sunrise)
        # لكن برمجياً سنعتمد الآن على الساعة الحالية:
        ruler = PLANETARY_ORDER[(start_idx + now.hour) % 7]

        # 3. تجهيز نص التحليل
        analysis_text = f"📍 الموقع: {city_name}\n🕒 الوقت: {now.strftime('%H:%M')}\n\n"
        analysis_text += f"📜 التفسير الموسع:\n{PLANETARY_HOURS_INFO.get(ruler)}\n\n"
        analysis_text += f"💡 نصيحة الساعة:\nيفضل ممارسة الأنشطة المرتبطة بطبيعة كوكب {ruler}."

        # 4. عرض النافذة المنبثقة
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            modal=False,
            title=ft.Text(f"⏳ حاكم الساعة: {ruler}", size=22, weight="bold", color="#f59e0b", rtl=True),
            content=ft.Container(
                content=ft.Text(analysis_text, size=16, rtl=True, text_align="right"),
                padding=10,
                width=400,
            ),
            actions=[
                ft.TextButton("نسخ", icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(analysis_text)),
                ft.TextButton("إغلاق", on_click=close_dlg),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

    except Exception as e:
        self.page.snack_bar = ft.SnackBar(ft.Text(f"حدث خلل: {str(e)}"))
        self.page.snack_bar.open = True
        self.page.update()
def show_animodar_correction(self):
        try:
            def clean_float(val):
                if isinstance(val, (list, tuple)): 
                    return float(val[0])
                return float(val)

            # 1. جلب البيانات من واجهة Flet
            lat = float(getattr(self, 'lat', 34.88))
            lon = float(getattr(self, 'lon', 35.88))
            
            try:
                # استخدام .value لجلب النصوص من الحقول وتحويلها
                d = int(self.day_ent.value or 1)
                m = int(self.month_ent.value or 1)
                y = int(self.year_ent.value or 2000)
                h = int(self.hour_ent.value or 12)
                mn = int(self.min_ent.value or 0)
                # استخدام توقيت افتراضي +3 إذا لم يوجد حقل GMT
                offset = float(getattr(self, 'gmt_ent', type('', (), {'value': '3.0'})()).value or 3.0)
            except ValueError:
                raise Exception("يرجى التأكد من إدخال أرقام صحيحة في خانات الميلاد")

            # حساب التاريخ اليولياني (وقت عالمي UT)
            jd_birth = swe.julday(y, m, d, (h + mn/60.0) - offset)
            
            # 2. حساب الطالع
            res_houses = swe.houses_ex(jd_birth, lat, lon, b'P')
            abs_asc = clean_float(res_houses[0][0]) # التصحيح: الكوسب الأول هو الطالع

            # 3. حساب درجة السيجي (اجتماع أو استقبال)
            sun_pos = clean_float(swe.calc_ut(jd_birth, swe.SUN)[0])
            moon_pos = clean_float(swe.calc_ut(jd_birth, swe.MOON)[0])
            
            # تحديد الكوكب الهدف (الشمس للنهاري والقمر لليلي كقاعدة عامة في الأنمودار)
            is_day = not (180 <= (sun_pos - abs_asc) % 360 <= 360)
            target_deg = sun_pos if is_day else moon_pos

            # 4. حساب المبتز على درجة السيجي (تبسيط للمثال)
            p_names = {0: "الشمس", 1: "القمر", 2: "عطارد", 3: "الزهرة", 4: "المريخ", 5: "المشتري", 6: "زحل"}
            domiciles = [4, 2, 2, 1, 0, 2, 3, 4, 5, 6, 6, 5] 
            
            sign_idx = int(target_deg / 30) % 12
            winner_id = domiciles[sign_idx]
            winner_name = p_names[winner_id]
            abs_winner = clean_float(swe.calc_ut(jd_birth, winner_id)[0])

            # 5. تحضير الأسماء للعرض
            z_list = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            s_sign = z_list[int(abs_winner / 30) % 12]
            a_sign = z_list[int(abs_asc / 30) % 12]
            
            # 6. استدعاء دالة العرض التفصيلية (تأكد من تحويلها لـ Flet أيضاً)
            if hasattr(self, 'display_detailed_animodar'):
                self.display_detailed_animodar(
                    y, m, d, h, mn, 
                    s_sign, int(abs_winner % 30), 
                    a_sign, int(abs_asc % 30), 
                    int(abs_winner % 30), 
                    winner_name
                )
            else:
                # عرض سريع في حال عدم وجود دالة التفاصيل
                self.log_to_flet(f"✅ التصحيح: يجب أن تكون درجة الطالع قريبة من {int(abs_winner % 30)}° في برج {s_sign}", color="purple")

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"⚠️ خطأ في التصحيح: {str(e)}", rtl=True))
            self.page.snack_bar.open = True
            self.page.update()
def display_detailed_animodar(self, y, m, d, h, mn, s_sign, s_deg, a_sign, a_curr, a_corr, planet_name):
        # 1. بناء نص التقرير
        full_text = f"""تقرير تصحيح وقت الولادة (النمودار التقليدي المتقدم)
تاريخ الميلاد: {y}/{m}/{d} | الوقت المحلي: {h:02d}:{mn:02d}

۞ البعد الفلسفي للنمودار:
يقوم تصحيح الطالع على مبدأ 'التناظر الكوني' حيث ترتبط درجة الطالع بدرجة الكوكب المهيمن (المبتز) على درجة اجتماع أو استقبال النيرين قبل الولادة.

النتائج النهائية:
• الكوكب المبتز: {planet_name}
• درجة الكوكب: {s_deg}° في برج {s_sign}
• الطالع الحالي: {a_curr}° في برج {a_sign}
• الدرجة المصححة الواجب ضبطها: [{a_corr}°] تماماً.

المطور: الفلكي البرمجي حسان الشاعر"""

        # 2. دالة حفظ ملف Word (تحتاج مكتبة python-docx)
        def save_to_word(e):
            try:
                from docx import Document
                import os
                
                doc = Document()
                doc.add_heading('تقرير تصحيح الطالع - النمودار', 0)
                doc.add_paragraph(full_text)
                
                file_name = f"Animodar_{y}_{m}_{d}.docx"
                # تحديد مسار الحفظ (يتوافق مع أندرويد عبر Flet)
                save_dir = "/storage/emulated/0/Download" if os.path.exists("/storage/emulated/0/Download") else os.getcwd()
                save_path = os.path.join(save_dir, file_name)
                
                doc.save(save_path)
                
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"✅ تم الحفظ في التحميلات: {file_name}"), bgcolor="green"))
            except Exception as ex:
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"❌ خطأ في الحفظ: {str(ex)}"), bgcolor="red"))
            self.page.update()

        # 3. إغلاق النافذة
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # 4. واجهة عرض التقرير (ثيم داكن متناسق)
        report_column = ft.Column(
            scroll=ft.ScrollMode.ALWAYS,
            expand=True,
            controls=[
                ft.Text("التقرير الملكي للنمودار", size=22, weight="bold", color="#38bdf8", text_align="center"),
                ft.Divider(color="#38bdf8"),
                ft.Text(full_text, size=16, color=ft.colors.ON_SURFACE_VARIANT, rtl=True),
            ]
        )

        dlg = ft.AlertDialog(
            modal=True,
            title=ft.Text("نتائج التصحيح الفلكي", rtl=True),
            content=ft.Container(content=report_column, width=500, height=600, padding=10),
            actions=[
                ft.ElevatedButton("💾 حفظ Word", icon=ft.icons.SAVE, on_click=save_to_word, bgcolor="#10b981", color="white"),
                ft.TextButton("إغلاق", on_click=close_dlg),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

def auto_skip(self, e, next_widget, limit):
        """وظيفة التنقل التلقائي بين حقول الإدخال"""
        if len(e.control.value) >= limit:
            if next_widget:
                next_widget.focus()
                self.page.update()
def get_accurate_planets(self, year, month, day, hour_decimal, lat, lon):
        try:
            # 1. تصحيح التوقيت العالمي (بافتراض فرق توقيت 3 ساعات)
            # يفضل لاحقاً جلب الـ offset من حقل gmt_ent إذا كان متاحاً
            utc_hour = hour_decimal - 3 
            jd = swe.julday(year, month, day, utc_hour)
            
            planets_map = {
                "الشمس": swe.SUN, "القمر": swe.MOON, "عطارد": swe.MERCURY,
                "الزهرة": swe.VENUS, "المريخ": swe.MARS, "المشتري": swe.JUPITER,
                "زحل": swe.SATURN, "الرأس": swe.MEAN_NODE
            }
            
            results = {}
            signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                     "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

            # 2. حساب مواقع الكواكب
            for name, p_id in planets_map.items():
                res, _ = swe.calc_ut(jd, p_id)
                lon_pos = float(res[0]) # التأكد من أنه رقم عشري
                results[name] = {
                    "sign": signs[int(lon_pos / 30) % 12], 
                    "full_deg": lon_pos,
                    "deg_in_sign": lon_pos % 30
                }
                
            # 3. حساب البيوت والطالع (نظام بلازيدوس P)
            # houses_ex تعيد (cusps, ascmc)
            cusps, ascmc = swe.houses_ex(jd, lat, lon, b'P')
            
            results["الطالع"] = float(ascmc[0])
            results["وسط السماء"] = float(ascmc[1])
            results["البيوت"] = [float(c) for c in cusps]
            
            return results

        except Exception as e:
            # عرض الخطأ لمستخدم أندرويد عبر Flet SnackBar
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(
                    ft.Text(f"خطأ في الحساب الفلكي: {str(e)}", rtl=True),
                    bgcolor=ft.colors.RED_ACCENT
                )
                self.page.snack_bar.open = True
                self.page.update()
            return None
def calculate_almuten(self, planets_data):
        """
        حساب الكوكب المبتز (المسيطر) بناءً على نقاط القوة في الأبراج.
        """
        # 1. تعريف أرباب البيوت (بالترتيب من الحمل إلى الحوت)
        domiciles = [
            "المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", 
            "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"
        ]
        
        # 2. تعريف أرباب الشرف (البرج: الكوكب)
        # الحمل: شمس، الثور: قمر، السرطان: مشتري، العذراء: عطارد، الميزان: زحل، الجدي: مريخ، الحوت: زهرة
        exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}

        # تهيئة سجل النقاط للكواكب السبعة فقط
        scores = {k: 0 for k in ["الشمس", "القمر", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل"]}
        
        try:
            if not planets_data:
                return "لا توجد بيانات", 0

            for p_name in scores.keys():
                p_info = planets_data.get(p_name)
                if not p_info: 
                    continue
                
                pos = p_info["full_deg"]
                sign_idx = int(pos / 30) % 12
                
                # أ- قوة بيت الكوكب (5 نقاط)
                # إذا كان الكوكب يتواجد في برج هو ربه الأصلي
                if domiciles[sign_idx] == p_name:
                    scores[p_name] += 5
                    
                # ب- قوة الشرف (4 نقاط)
                # إذا كان الكوكب يتواجد في برج شرفه
                if exaltations.get(sign_idx) == p_name:
                    scores[p_name] += 4
                
                # إمكانية التطوير المستقبلي:
                # المثلثات (+3)، الحدود (+2)، الوجوه (+1)

            # 3. تحديد الكوكب صاحب أعلى نقاط
            winner = max(scores, key=scores.get)
            
            # إذا تعادلت النقاط عند الصفر
            if scores[winner] == 0:
                return "غير محدد", 0
                
            return winner, scores[winner]
            
        except Exception as e:
            # تنبيه المستخدم في حال حدوث خطأ أثناء المعالجة
            if hasattr(self, 'page'):
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ في حساب المبتز: {e}")))
            return "خطأ", 0
def analyze_birth_chart(self):
        try:
            import swisseph as swe
            from datetime import datetime

            # 1. جلب المدخلات من واجهة Flet
            def g_v(attr, default):
                element = getattr(self, attr, None)
                return element.value if element and element.value else default

            try:
                y, m, d = int(g_v('year_ent', 1990)), int(g_v('month_ent', 1)), int(g_v('day_ent', 1))
                h, mi = int(g_v('hour_ent', 12)), int(g_v('min_ent', 0))
                offset = float(g_v('gmt_ent', 3.0))
            except ValueError:
                raise Exception("يرجى إدخال أرقام صحيحة في تاريخ ووقت الميلاد")

            lat = float(getattr(self, 'lat', 34.88))
            lon = float(getattr(self, 'lon', 35.88))
            
            # تصحيح التوقيت العالمي (UT)
            ut_h = h - offset + (mi / 60.0)
            jd = swe.julday(y, m, d, ut_h)

            # 2. حساب البيوت والطالع (بلاسيدوس)
            cusps_raw, ascmc_raw = swe.houses(jd, lat, lon, b'P')
            asc_raw = float(ascmc_raw[0])
            c = [float(cusps_raw[i]) for i in range(1, 13)]

            def find_house(deg):
                for i in range(12):
                    s, e = c[i], c[(i + 1) % 12]
                    if s < e:
                        if s <= deg < e: return str(i + 1)
                    else: # معالجة نقطة الصفر (الحمل)
                        if s <= deg or deg < e: return str(i + 1)
                return "12"

            # 3. حساب الكواكب والسرعات
            zodiacs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                       "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            
            p_map = [("الشمس",swe.SUN,"sun"), ("القمر",swe.MOON,"moon"), ("عطارد",swe.MERCURY,"mercury"), 
                     ("الزهرة",swe.VENUS,"venus"), ("المريخ",swe.MARS,"mars"), ("المشتري",swe.JUPITER,"jupiter"), 
                     ("زحل",swe.SATURN,"saturn"), ("أورانوس",swe.URANUS,"uranus"), ("نبتون",swe.NEPTUNE,"neptune"), 
                     ("بلوتو",swe.PLUTO,"pluto"), ("الرأس",swe.MEAN_NODE,"nn")]

            results = {}
            raw_d = {}
            raw_speed = {}

            for ar, code, en in p_map:
                res, _ = swe.calc_ut(jd, code)
                deg = float(res[0])
                speed = float(res[3]) 
                raw_d[en] = deg
                raw_speed[en] = speed
                results[f"{en}_sign"] = zodiacs[int(deg / 30)]
                results[f"{en}_house"] = find_house(deg)

            # 4. السهام والليل/النهار
            s_h = int(results.get('sun_house', 1))
            is_day_time = 7 <= s_h <= 12
            
            if is_day_time:
                fort = (asc_raw + raw_d['moon'] - raw_d['sun']) % 360
                spir = (asc_raw + raw_d['sun'] - raw_d['moon']) % 360
            else:
                fort = (asc_raw + raw_d['sun'] - raw_d['moon']) % 360
                spir = (asc_raw + raw_d['moon'] - raw_d['sun']) % 360

            results['fortune_part'] = f"{zodiacs[int(fort/30)]} ({int(fort%30)}°)"
            results['spirit_part'] = f"{zodiacs[int(spir/30)]} ({int(spir%30)}°)"

            # 5. محرك القواطع والتحذيرات
            warnings_list = []
            for ar, code, en in p_map:
                if raw_speed.get(en, 0) < 0:
                    warnings_list.append(f"⚠️ كوكب {ar} متراجع (Retrograde)")
                if en != "sun":
                    dist_to_sun = abs(raw_d[en] - raw_d['sun'])
                    if dist_to_sun < 8 or dist_to_sun > 352:
                        warnings_list.append(f"🔥 كوكب {ar} محترق (Combust) لقربه من الشمس")
            
            if abs(raw_d['saturn'] - raw_d['mars']) < 5:
                warnings_list.append("💀 تحذير: اقتران النحسين (زحل والمريخ) في الهيئة")
            
            results['warnings'] = "\n".join(warnings_list) if warnings_list else "لا يوجد قواطع فلكية سلبية"

            # 6. تحليل النجوم الثابتة
            fixed_stars_db = [
                {"name": "الثريا (Alcyone)", "pos": 60.1, "effect": "الجاذبية، الشهرة، والنجاح الفني."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "effect": "شجاعة قيادية ونجاح مادي."},
                {"name": "رأس التوأم (Pollux)", "pos": 113.3, "effect": "الجرأة وقوة البيان."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "effect": "الجاه العظيم والسلطة."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "effect": "الثروة والنجاح العلمي."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "effect": "التحولات الكبرى وإعادة البناء."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "effect": "الإبداع الروحاني والتميز الفني."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "effect": "الشهرة العالمية والتميز الروحاني."}
            ]
            
            stars_analysis = []
            my_planets_for_stars = {**raw_d, "Ascendant": asc_raw}
            p_ar_names = {"sun":"الشمس", "moon":"القمر", "mercury":"عطارد", "venus":"الزهرة", "mars":"المريخ", "jupiter":"المشتري", "saturn":"زحل", "Ascendant":"الطالع"}

            for star in fixed_stars_db:
                for p_en, p_deg in my_planets_for_stars.items():
                    diff = abs(p_deg - star['pos'])
                    if diff <= 1.5 or diff > 358.5:
                        p_ar = p_ar_names.get(p_en, p_en)
                        stars_analysis.append(f"⭐ يقترن {p_ar} بنجم {star['name']}.\n   💡 التأثير: {star['effect']}")

            results['stars_report'] = "\n\n".join(stars_analysis) if stars_analysis else "لا توجد اقترانات مباشرة بالنجوم العظمى."

            # 7. تخزين النتائج النهائية واستدعاء التقرير
            self.raw_d = raw_d
            self.asc_raw = asc_raw
            self.is_analyzed = True
            
            if hasattr(self, 'display_comprehensive_report'):
                self.display_comprehensive_report(results)
            else:
                self.log_to_flet("تم التحليل بنجاح، يرجى مراجعة النتائج.", color="green")

        except Exception as e:
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ في التحليل: {str(e)}", rtl=True))
                self.page.snack_bar.open = True
                self.page.update()
def get_house(self, planet_deg, houses):
        """تحديد البيت الذي يقع فيه الكوكب بناءً على حدود البيوت (Cusps)"""
        planet_deg %= 360
        # نستخدم 12 بيتاً؛ houses تبدأ من الفهرس 1 للبيت الأول في swe.houses
        for i in range(1, 13):
            s = houses[i]
            e = houses[i + 1] if i < 12 else houses[1]
            
            if s < e:
                if s <= planet_deg < e:
                    return str(i)
            else:
                # حالة عبور نقطة الصفر (نهاية الحوت وبداية الحمل)
                if s <= planet_deg or planet_deg < e:
                    return str(i)
        return "1"

def get_sign_name(self, deg):
        """تحويل الدرجة المطلقة إلى اسم البرج المقابل لها"""
        signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                 "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
        index = int((deg % 360) / 30)
        return signs[index]

def run_calculation(self, jd, lat, lon):
        """المحرك الحسابي الرئيسي لاستخراج بيانات الهيئة"""
        data = {}
        zodiac_names = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                        "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

        # 1. حساب نظام البيوت (نظام بلاسيدوس P)
        # houses: مصفوفة بـ 13 عنصراً (العنصر 0 مهمل، 1-12 هي البيوت)
        # ascmc: مصفوفة بـ 10 عناصر (0: الطالع، 1: وسط السماء)
        houses, ascmc = swe.houses(jd, lat, lon, b'P') 
        ascendant = float(ascmc[0])
        data['ascendant_deg'] = ascendant
        data['ascendant_sign'] = zodiac_names[int(ascendant / 30)]
        data['mc_deg'] = float(ascmc[1])

        # 2. حساب الكواكب الأساسية
        planets_to_calc = {
            'sun': swe.SUN, 'moon': swe.MOON, 'mercury': swe.MERCURY,
            'venus': swe.VENUS, 'mars': swe.MARS, 'jupiter': swe.JUPITER,
            'saturn': swe.SATURN, 'uranus': swe.URANUS, 'neptune': swe.NEPTUNE,
            'pluto': swe.PLUTO
        }

        for name, code in planets_to_calc.items():
            res, _ = swe.calc_ut(jd, code)
            p_deg = float(res[0])
            data[f'{name}_deg'] = p_deg
            data[f'{name}_sign'] = zodiac_names[int(p_deg / 30)]
            # استخدام دالة get_house الموحدة
            data[f'{name}_house'] = self.get_house(p_deg, houses)

        # 3. العقد القمرية (الرأس)
        res_node, _ = swe.calc_ut(jd, swe.MEAN_NODE)
        nn_deg = float(res_node[0])
        data['north_node_deg'] = nn_deg
        data['north_node_sign'] = zodiac_names[int(nn_deg / 30)]
        data['north_node_house'] = self.get_house(nn_deg, houses)
        
        # 4. حساب السهام الفلكية (بناءً على الطالع المحسوب)
        # ملاحظة: هذه الحسابات تفترض الصيغة النهارية، يفضل ربطها بـ is_day_time مستقبلاً
        fortune_deg = (ascendant + data['moon_deg'] - data['sun_deg']) % 360
        spirit_deg = (ascendant + data['sun_deg'] - data['moon_deg']) % 360
        
        data['fortune_deg'] = fortune_deg
        data['spirit_deg'] = spirit_deg
        data['fortune_part'] = f"{zodiac_names[int(fortune_deg / 30)]} ({int(fortune_deg % 30)}°)"
        data['spirit_part'] = f"{zodiac_names[int(spirit_deg / 30)]} ({int(spirit_deg % 30)}°)"

        return data
def get_hijri_date_string(self):
        """جلب التاريخ الهجري الحالي باستخدام تقويم أم القرى"""
        try:
            from ummalqura.hijri_date import HijriDate
            now = HijriDate.today()
            return f"{now.day} / {now.month} / {now.year} هـ"
        except ImportError:
            return "مكتبة ummalqura غير مثبتة"
        except Exception as e:
            return f"خطأ: {str(e)}"

def get_daily_wisdom(self, day):
        """قاعدة بيانات تفسير رؤى أيام الشهر القمري"""
        # الأيام مبنية على رؤية الصادق وابن سيرين للتراث الروحاني
        wisdom_db = {
            1: {"sadiq": "رؤيا صادقة جداً وتتحقق سريعاً.", "sirin": "بشارة خير في أول الشهر.", "philosophy": "بداية دورة جديدة في عالم المثال."},
            2: {"sadiq": "تتأخر قليلاً لكنها خير.", "sirin": "تدل على نيل رزق أو فائدة.", "philosophy": "انعكاس للرغبات الكامنة."},
            15: {"sadiq": "رؤيا وسطى بين الحق والباطل.", "sirin": "تدل على بلوغ منتصف الأمر.", "philosophy": "اكتمال الوعي القمري."},
            # يمكن إضافة بقية الأيام حتى 30 بنفس التنسيق
        }
        return wisdom_db.get(day, {
            "sadiq": "تختلف حسب حال الرائي ووقت الرؤيا.", 
            "sirin": "تتطلب تأويلاً دقيقاً من خبير.", 
            "philosophy": "انعكاس لمرآة النفس الفلكية."
        })

def show_about(self, e=None):
        """نافذة معلومات النظام المتوافقة مع أندرويد"""
        about_text = (
            "🚀 الميقاتي الفلكي - النسخة الاحترافية\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "تطوير الخبير الفلكي: حسان الشاعر\n"
            "إصدار: 2026 PRO Mobile (Flet Version)\n\n"
            "نظام فلكي متكامل يجمع بين المحرك السويسري العالمي\n"
            "والاستنباطات الفلسفية التراثية، مصمم ليعمل\n"
            "بكفاءة عالية على الهواتف الذكية."
        )

        def close_dlg(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            title=ft.Text("حول النظام", weight="bold", rtl=True),
            content=ft.Container(
                content=ft.Text(about_text, rtl=True, size=15),
                padding=10
            ),
            actions=[
                ft.TextButton("إغلاق", on_click=close_dlg),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
def show_dream_interpreter(self):
        try: 
            # 1. حساب اليوم القمري (المحرك الفلكي الدقيق)
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, 12)
            res_m, _ = swe.calc_ut(jd, swe.MOON) # موقع القمر
            res_s, _ = swe.calc_ut(jd, swe.SUN)  # موقع الشمس
            
            # حساب الفرق الزاوي لتعيين اليوم القمري (دورة 29.5 يوم تقريباً)
            moon_pos = float(res_m[0])
            sun_pos = float(res_s[0])
            h_day = int(((moon_pos - sun_pos) % 360 / 12.19) + 1)
            
            # جلب حكمة اليوم القمري
            wisdom = self.get_daily_wisdom(h_day)  
            
            # 2. قاموس الرموز            
            # 2. قاموس الرموز الماروني (العربي، الفارسي، اليوناني)
            symbols_dictionary = {
                "نبي": "🌙 العرب: عز وشرف وبشارة.\n🏹 الفرس: استقامة دين وظفر.\n🏛️ اليونان: اتصال بالعقل الكلي والحكمة.",
                "ملك": "🌙 العرب: نصر وعز وقضاء حاجة.\n🏹 الفرس: نيل منصب رفيع.\n🏛️ اليونان: سيطرة على النفس والأنا العليا.",
                "شمس": "🌙 العرب: ولي الأمر أو الوالد.\n🏹 الفرس: الحاكم الأعظم.\n🏛️ اليونان: الحقيقة والوعي المطلق.",
                "ماء": "🌙 العرب: حياة طيبة ورزق.\n🏹 الفرس: استقرار ونمو تجاري.\n🏛️ اليونان: تطهير وشفاء جسدي.",
                "نار": "🌙 العرب: فتنة أو هداية.\n🏹 الفرس: منصب نوراني ورفعة.\n🏛️ اليونان: تحول وإبداع أو غضب.",
                "ذهب": "🌙 العرب: همّ للرجل وفرح للنساء.\n🏹 الفرس: ميراث أو ولاية.\n🏛️ اليونان: خلود وفكر نقي.",
                "فضة": "🌙 العرب: مال وراحة.\n🏹 الفرس: زواج أو ولادة.\n🏛️ اليونان: حدس وأنوار باطنية.",
                "ثعبان": "🌙 العرب: عدو من الأهل.\n🏹 الفرس: عدو ذو مال.\n🏛️ اليونان: شفاء وحكمة الأرض.",
                "أسد": "🌙 العرب: عدو جبار.\n🏹 الفرس: قوة ونفوذ.\n🏛️ اليونان: شجاعة وسيادة.",
                "طيران": "🌙 العرب: سفر ورفعة.\n🏹 الفرس: غنى سريع.\n🏛️ اليونان: تحرر من قيود المادة.",
                "موت": "🌙 العرب: عمر طويل وتوبة.\n🏹 الفرس: بدء حياة جديدة.\n🏛️ اليونان: تحول الوعي.",
                "بحر": "🌙 العرب: الدنيا وتقلباتها.\n🏹 الفرس: ملك لا يدرك غوره.\n🏛️ اليونان: اللاوعي العميق.",
                "خبز": "🌙 العرب: الإسلام والرزق.\n🏹 الفرس: أمن من الفقر.\n🏛️ اليونان: نمو روحي.",
                "لبن": "🌙 العرب: فطرة ورزق حلال.\n🏹 الفرس: علم وحكمة.\n🏛️ اليونان: تغذية النفس.",
                "عسل": "🌙 العرب: حلاوة الإيمان.\n🏹 الفرس: مال موروث.\n🏛️ اليونان: رحيق الحكمة.",
                "سيف": "🌙 العرب: ولد أو حجة قوية.\n🏹 الفرس: نصر وقوة.\n🏛️ اليونان: إرادة قاطعة.",
                "خاتم": "🌙 العرب: أمان أو منصب أو زواج.\n🏹 الفرس: إتمام أمر عظيم.\n🏛️ اليونان: عهد أبدي ودائرة الكمال.",
                "شجر": "🌙 العرب: رجال أو أعمال.\n🏹 الفرس: طول عمر وبركة.\n🏛️ اليونان: نمو الشخصية والجذور.",
                "كعبة": "🌙 العرب: حج وأمن وبشارة.\n🏹 الفرس: استقامة وقبلة المراد.\n🏛️ اليونان: مركز الوجود والاتزان.",
                "ضحك": "🌙 العرب: حزن (إلا بشارة بمولود).\n🏹 الفرس: خبر طارئ.\n🏛️ اليونان: إفراغ توتر العقل الباطن.",
                "بكاء": "🌙 العرب: فرج ونفاذ ضيق.\n🏹 الفرس: زوال مرض أو كرب.\n🏛️ اليونان: تطهير عاطفي (Catharsis)."
            }
            # 2. القاموس الموسع الشامل (موسوعة ابن سيرين والصادق المحدثة 2026)
            symbols_dictionary = {
                # --- الرموز الروحية والعلوية ---
                "نبي": "ابن سيرين: عز وشرف في الدارين ورؤياهم بشارة. الصادق: هداية ونيل أمانة وخير مستدام.",
                "ملك": "ابن سيرين: نصر على الأعداء ونيل سلطة وجاه. الصادق: قضاء حاجة وبشارة بالظفر.",
                "كعبة": "ابن سيرين: أمن وبشارة واجتماع بولي الأمر. الصادق: إيمان، إسلام، أمن، أو حج قريب.",
                "قرآن": "ابن سيرين: حكمة وميراث وعلم نافع. الصادق: انتشار علم ورزق طيب وتطهير للروح.",
                "محراب": "ابن سيرين: حصول ولد صالح أو إمام عادل. الصادق: امرأة صالحة أو رزق من جهة الدين.",
                "أذان": "ابن سيرين: حج لمن أذن في أشهره، أو نيل منصب. الصادق: دعوة للخير أو رفعة ذكر.",

                # --- الطبيعة والعناصر الكونيه ---
                "ماء": "ابن سيرين: حياة طيبة. الصادق: فتنة إن كان كدراً، وعلم ورزق إن كان صافياً.",
                "نار": "ابن سيرين: فتنة أو سلطان أو بشارة ورزق. الصادق: عذاب أو طريق هداية مستقيم.",
                "بحر": "ابن سيرين: سلطان مهاب. الصادق: ملك، رئيس، عالم، علم، مال، أو شغل عظيم.",
                "مطر": "ابن سيرين: رحمة عامة ما لم يكن مخرباً. الصادق: فرج من هم واستجابة دعاء.",
                "شمس": "ابن سيرين: ولي الأمر أو الوالد أو الملك. الصادق: نور في الدين وجاه في الدنيا.",
                "قمر": "ابن سيرين: وزير الملك أو الزوجة أو العالم. الصادق: زيادة في العلم أو نيل منصب.",
                "تراب": "ابن سيرين: مال ورزق بعد تعب. الصادق: يؤول بالدنيا أو الموت أو الفقر والحاجة.",

                # --- الحيوانات والطيور والحشرات ---
                "أسد": "ابن سيرين: عدو جبار أو سلطان جائر. الصادق: شدة وضيق من جهة صاحب نفوذ.",
                "خيل": "ابن سيرين: عز وشرف ومرتبة عالية. الصادق: سفر مبارك أو نيل سيادة على قومه.",
                "جمل": "ابن سيرين: حزن أو رجل جاهل أو سفر طويل. الصادق: نفوذ وقوة وصبر على المحن.",
                "ثعبان": "ابن سيرين: عدو يكاتم العداوة. الصادق: امرأة، ولد، عدو ذو مال، أو كنوز مخفية.",
                "عقرب": "ابن سيرين: عدو ضعيف يغتاب. الصادق: نكد وهمّ يسببه شخص نمام قريب.",
                "نحل": "ابن سيرين: غنيمة بلا تعب وعمل صالح. الصادق: عسكر، شفاء من مرض، أو رزق مبارك.",
                "غراب": "ابن سيرين: رجل فاسق أو سفر بعيد. الصادق: يؤول بالرجل الكاذب أو الحزن.",
                "حوت": "ابن سيرين: وزير الملك أو ضيق يتبعه فرج (قصة يونس). الصادق: رزق من البحر أو علم.",

                # --- الثمار والأطعمة ---
                "خبز": "ابن سيرين: الإسلام والعمر والمال الحلال. الصادق: عيش طيب وعلم نافع.",
                "عسل": "ابن سيرين: ميراث أو مال مجموع. الصادق: حلاوة الإيمان والقرآن والشفاء.",
                "لبن": "ابن سيرين: فطرة الإسلام والرزق الفطري. الصادق: علم واسع وحكمة.",
                "تفاح": "ابن سيرين: همّة الرجل في عمله. الصادق: نيل بشارة ومنفعة من صديق.",
                "عنب": "ابن سيرين: رزق واسع، والأسود همّ. الصادق: مال مجموع أو منفعة دائمة.",
                "تمر": "ابن سيرين: مطر أو رزق حلال لا يدوم. الصادق: حلاوة الإيمان ونيل مقصود.",
                "لحم": "ابن سيرين: المطبوخ مال، والنيء وجع ومرض. الصادق: غنيمة أو غيبة (حسب السياق).",

                # --- المعادن والكنوز والأدوات ---
                "ذهب": "ابن سيرين: همّ وغرامة للرجل. الصادق: فرح للنساء، وللرجل ذهاب منصب.",
                "فضة": "ابن سيرين: مال مجموع وراحة بال. الصادق: امرأة جميلة، ولد، أو خير كثير.",
                "جوهر": "ابن سيرين: علم وجمال. الصادق: يؤول بالمرأة الحسناء أو التجارة الرابحة.",
                "كرسي": "ابن سيرين: رفعة وشأن. الصادق: عز، شرف، ولاية، نصر، أو سعادة.",
                "مفتاح": "ابن سيرين: نصر وفتح أبواب الرزق. الصادق: قضاء حاجة، حج، ومال ميسر.",
                "ثوب": "ابن سيرين: ستر ودين ووقار. الصادق: عز وجاه وحال مستورة.",
                "سيف": "ابن سيرين: ولد أو سلطان أو كلام. الصادق: نصر على الأعداء أو حجة قوية.",

                # --- أحوال الإنسان وأفعاله ---
                "موت": "ابن سيرين: طول عمر أو ندم وتوبة. الصادق: انقطاع أمر أو زواج للأعزب.",
                "طيران": "ابن سيرين: سفر وتغير حال وطلب رفعة. الصادق: استجابة دعاء وعلو مرتبة.",
                "بكاء": "ابن سيرين: فرج وسرور (بدون عويل). الصادق: نجاة من كرب وفرح عاجل.",
                "ضحك": "ابن سيرين: حزن وندم. الصادق: بشارة بقدوم ولد أو خبر سار.",
                "زواج": "ابن سيرين: عناية إلهية أو قيد وهمّ. الصادق: نيل منصب أو غنى.",
                "غسل": "ابن سيرين: قضاء دين وتوبة من ذنب. الصادق: طهارة في الدين وزوال الهموم.",
                "حج": "ابن سيرين: قضاء دين أو حج فعلي أو أمن. الصادق: استقامة الدين ونيل أجر."
            }

            # 3. قاموس أحكام الأيام الهلالية (الإمام الصادق)
            days_dict = {
            1: "صادقة جداً، وتأويلها يظهر سريعاً وهي من أفضل الأيام.",
            2: "رؤيا صالحة، ومن رآها نال خيراً ومنفعة من السلطان أو ولي الأمر.",
            3: "رؤيا حق، وتأويلها يكون على العكس إن كانت مكروهة (تحذير).",
            4: "رؤيا محمودة، وتأويلها يتأخر قليلاً لكنه محقق.",
            5: "رؤيا فرح وسرور، ومن رآها نال بشارة يسعد بها قلبه.",
            6: "رؤيا صادقة، وتأويلها يقع بعد فترة طويلة (ربما سنوات).",
            7: "رؤيا مستورة، لا تُقص إلا على عالم أو ناصح لسرعة تحققها.",
            8: "رؤيا صادقة جداً، وهي تدل على خلاص من هم أو نيل مطلب.",
            9: "رؤيا حق، وتدل على قوة الرائي ونيله مرتبة عالية.",
            10: "رؤيا صادقة، وتدل على تيسير الأمور العالقة.",
            11: "رؤيا محمودة، وتأويلها يقع من يوم إلى 15 يوماً.",
            12: "رؤيا تحذيرية، تأويلها يتأخر لكن فيها تنبيه من غدر.",
            13: "رؤيا مختلطة، قد لا يصح تأويلها لداخلها بحديث النفس.",
            14: "رؤيا صادقة، وبشارتها سريعة كسرعة اكتمال القمر.",
            15: "رؤيا صحيحة، ومن رآها نال ما تمنى بعد تعب.",
            16: "رؤيا كاذبة (أضغاث أحلام)، لا يُعتد بها في هذا اليوم.",
            17: "رؤيا مخفية، تدل على أمور باطنة تظهر للرائي لاحقاً.",
            18: "رؤيا حق، وتدل على نصر على الأعداء أو نجاح في خصومة.",
            19: "رؤيا قوية جداً، وتأويلها يكون أوضح مما رُئي في المنام.",
            20: "رؤيا محمودة، وتدل على زيادة في الرزق والمال.",
            21: "رؤيا كاذبة، غالباً ما تكون من وسوسة النفس أو الشيطان.",
            22: "رؤيا صادقة، وتدل على قضاء الحوائج الصعبة.",
            23: "رؤيا تحذيرية من الفتن، يُنصح بالصدقة بعد رؤيتها.",
            24: "رؤيا محمودة، وتدل على الفرج بعد الضيق.",
            25: "رؤيا لا تصح، وهي من أوهام الخيال في هذا اليوم.",
            26: "رؤيا صادقة، وتدل على السفر أو التغيير في الأحوال.",
            27: "رؤيا حق، وتأويلها يقع كما هي دون تبديل.",
            28: "رؤيا صادقة، ومن رآها نال خيراً وفيراً غير متوقع.",
            29: "رؤيا تحذيرية، يُكره قصها في هذا اليوم.",
            30: "رؤيا صادقة ومباركة، وهي خاتمة الخير والرزق."
        }
     
            # استدعاء الواجهة (تأكد من بناء الواجهة هنا أو استدعاء دالة بناء الواجهة)
            pass 

        except Exception as e:
            print(f"Error: {e}")

def process_interpretation(self, e, search_input, report_content, h_day, days_dict, symbols_dictionary):
        # محرك البحث والتحليل
        search_query = search_input.value.strip()
        """
        المحرك المعدل للبحث والتحليل الرمزي للرؤى
        """
        search_query = search_input.value.strip()
        
        if not search_query:
            report_content.value = "⚠️ يرجى كتابة رمز (مثل: شمس، بحر، أسد) للبحث عنه."
            self.page.update()
            return
        
        # بناء التقرير الملحمي المنسق بتنسيق احترافي
        report = "✨ تحليل الرؤيا الاستدلالي الملحمي ✨\n"
        report += f"📅 اليوم القمري الحالي: {h_day}\n"
        report += "════════════════════════════════\n\n"
        
        # جلب حكم الزمان (مع معالجة القاموس المعقد الذي يحتوي على مفاتيح فرعية)
        wisdom_data = days_dict.get(h_day, {})
        sadiq_wisdom = wisdom_data.get("sadiq", "يوم يحتاج لمطابقة فلكية دقيقة.")
        
        report += f"🌙 حكم الزمان (حسب الشهر الهلالي):\n• {sadiq_wisdom}\n\n"

        # محرك البحث عن الرموز (دعم البحث الجزئي والمرن)
        found = False
        for key in symbols_dictionary:
            if search_query in key or key in search_query:
                report += f"🔍 تفسير رمز [{key}]:\n{symbols_dictionary[key]}\n\n"
                found = True
                break
        
        if not found:
            report += f"⚠️ الرمز [{search_query}] غير مدرج حالياً في القاموس الماروني، جرب كلمات مفتاحية أخرى.\n\n"

        # الخاتمة الفلسفية للتقرير
        report += "💎 فلسفة التعبير:\n"
        report += "• مدرسة اليونان: الرؤيا هي استقراء العقل الباطن لمسار القدر.\n"
        report += "• مدرسة العرب: الرؤيا سفر في عالم المثال وتجسد المعاني.\n\n"
        report += "════════════════════════════════\n"
        report += "✍️ إعداد الخبير: حسان الشاعر © 2026"
        
        # تحديث المحتوى في الواجهة
        report_content.value = report
        # تمرير التقرير للأعلى في حال كان النص طويلاً
        try:
            report_content.update()
        except:
            self.page.update()
def draw_astro_wheel_in_popup(self, jd, title_text="خريطة الميلاد"):
        """رسم مخطط فلكي مبسط (Canvas) داخل نافذة منبثقة"""
        import math
        
        # 1. بيانات الكواكب والرموز (السبعة التقليدية)
        planets_map = {
            swe.SUN: ("☉", "#f59e0b"),   # شمس - ذهبي
            swe.MOON: ("☽", "#3b82f6"),  # قمر - أزرق
            swe.MERCURY: ("☿", "#8b5cf6"), # عطارد - بنفسجي
            swe.VENUS: ("♀", "#ec4899"),   # زهرة - وردي
            swe.MARS: ("♂", "#dc2626"),    # مريخ - أحمر
            swe.JUPITER: ("♃", "#d97706"), # مشتري - برتقالي
            swe.SATURN: ("♄", "#475569")   # زحل - رمادي
        }
        
        # 2. إعداد الـ Canvas مع الدوائر الأساسية
        chart_canvas = ft.Canvas(
            expand=True,
            shapes=[
                # الدائرة الخارجية الكبرى (نطاق البروج)
                ft.cv.Circle(150, 150, 140, ft.Paint(color="#1f538d", stroke_width=2, style=ft.PaintingStyle.STROKE)),
                # الدائرة الداخلية (مركز الأرض)
                ft.cv.Circle(150, 150, 40, ft.Paint(color="#cbd5e1", stroke_width=1, style=ft.PaintingStyle.STROKE)),
            ]
        )

        # 3. محرك الرسم وتوزيع الكواكب بناءً على درجاتها
        for p_code, (symbol, color) in planets_map.items():
            try:
                res, _ = swe.calc_ut(jd, p_code)
                pos = float(res[0])
                
                # تحويل الدرجة الفلكية إلى راديان
                # ملاحظة: طرح 180 أو 90 يغير نقطة البداية لتناسب اتجاه الطالع
                angle = math.radians(pos - 180) 
                
                # حساب الإحداثيات المركزية (نصف القطر 100)
                x = 150 + 100 * math.cos(angle)
                y = 150 + 100 * math.sin(angle)

                # إضافة رمز الكوكب إلى الـ Canvas
                chart_canvas.shapes.append(
                    ft.cv.Text(
                        x - 8, y - 12, symbol,
                        ft.TextStyle(size=22, color=color, weight="bold")
                    )
                )
            except:
                continue

        # 4. دالة الإغلاق
        def close_chart(e):
            chart_dlg.open = False
            self.page.update()

        # 5. بناء وعرض النافذة
        chart_dlg = ft.AlertDialog(
            title=ft.Text(title_text, weight="bold", rtl=True, size=18),
            content=ft.Container(
                content=chart_canvas,
                width=300, 
                height=300,
                alignment=ft.alignment.center,
                bgcolor="#ffffff",
                border_radius=150,
                border=ft.border.all(0.5, "#e2e8f0")
            ),
            actions=[
                ft.TextButton("إغلاق الخريطة", on_click=close_chart)
            ],
            actions_alignment=ft.MainAxisAlignment.CENTER
        )

        self.page.dialog = chart_dlg
        chart_dlg.open = True
        self.page.update()
def draw_astro_wheel_in_popup(self, jd):
        """رسم مخطط الأبراج وتوزيع الكواكب برمجياً (Canvas)"""
        import math
        
        # 1. إعداد الحاوية الرسومية والدوائر الأساسية
        cp = ft.Canvas(
            expand=True,
            shapes=[
                # الدائرة الخارجية (إطار البروج)
                ft.cv.Circle(150, 150, 120, ft.Paint(color="#1f538d", stroke_width=2, style=ft.PaintingStyle.STROKE)),
                # الدائرة المركزية
                ft.cv.Circle(150, 150, 40, ft.Paint(color="#cbd5e1", stroke_width=1, style=ft.PaintingStyle.STROKE)),
            ]
        )

        # 2. رسم خطوط تقسيم البروج الاثني عشر (كل 30 درجة)
        for i in range(12):
            angle = math.radians(i * 30 - 90) # تبدأ من الأعلى (MC افتراضياً)
            x2 = 150 + 120 * math.cos(angle)
            y2 = 150 + 120 * math.sin(angle)
            cp.shapes.append(ft.cv.Line(150, 150, x2, y2, ft.Paint(color="#cbd5e1", stroke_width=1)))

        # 3. محرك توزيع الكواكب الفعلية بناءً على التاريخ اليولياني (jd)
        planets = {
            swe.SUN: ("☉", "#f59e0b"), 
            swe.MOON: ("☽", "#3b82f6"), 
            swe.MARS: ("♂", "#dc2626"), 
            swe.JUPITER: ("♃", "#d97706"), 
            swe.SATURN: ("♄", "#475569")
        }
        
        for p_id, (symbol, color) in planets.items():
            try:
                res, _ = swe.calc_ut(jd, p_id)
                # تحويل موقع الكوكب (0-360) لزاوية رسم مع إزاحة لتطابق التقسيم
                p_angle = math.radians(res[0] - 90)
                # وضع الكواكب في المدار المتوسط (نصف قطر 90)
                px = 150 + 90 * math.cos(p_angle)
                py = 150 + 90 * math.sin(p_angle)
                
                # رسم الرمز الفلكي
                cp.shapes.append(
                    ft.cv.Text(px - 7, py - 10, symbol, ft.TextStyle(size=18, color=color, weight="bold"))
                )
            except:
                continue

        # 4. بناء وعرض النافذة المنبثقة
        dlg = ft.AlertDialog(
            title=ft.Text("الخريطة الفلكية اللحظية", size=18, weight="bold", rtl=True),
            content=ft.Container(
                content=cp, 
                width=300, 
                height=300, 
                alignment=ft.alignment.center,
                bgcolor="#ffffff"
            ),
            actions=[
                ft.TextButton("إغلاق", on_click=lambda e: self.close_dialog(dlg))
            ],
            actions_alignment=ft.MainAxisAlignment.CENTER
        )
        
        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

def close_dialog(self, dlg):
        """إغلاق أي نافذة منبثقة وتحديث الصفحة"""
        dlg.open = False
        self.page.update()
def search_city_logic(self, e=None):
        """محرك البحث عن الإحداثيات الجغرافية عبر الإنترنت"""
        try:
            from geopy.geocoders import Nominatim
        except ImportError:
            self.page.show_snack_bar(ft.SnackBar(ft.Text("❌ يرجى تثبيت مكتبة geopy")))
            return

        # 1. إعداد محرك البحث
        geolocator = Nominatim(user_agent="hassan_astro_mobile_2026")
        city_name = self.city_entry.value 
        
        if not city_name or city_name.strip() == "":
            self.page.snack_bar = ft.SnackBar(ft.Text("⚠️ يرجى إدخال اسم المدينة (مثلاً: Tartus أو Cairo)"))
            self.page.snack_bar.open = True
            self.page.update()
            return

        # إظهار رسالة انتظار للمستخدم
        self.page.snack_bar = ft.SnackBar(ft.Text(f"🔍 جاري البحث عن {city_name}..."))
        self.page.snack_bar.open = True
        self.page.update()

        try:
            # 2. طلب الإحداثيات من الخادم
            location = geolocator.geocode(city_name, timeout=10)
            
            if location:
                # تخزين الإحداثيات في متغيرات الكلاس لاستخدامها في الحسابات الفلكية
                self.lat = location.latitude
                self.lon = location.longitude
                
                # تنسيق العرض للمستخدم
                lat_dir = "شمالاً" if self.lat >= 0 else "جنوباً"
                lon_dir = "شرقاً" if self.lon >= 0 else "غرباً"
                
                self.coords_lbl.value = f"📍 {abs(self.lat):.2f}° {lat_dir} | {abs(self.lon):.2f}° {lon_dir}"
                self.coords_lbl.color = "blue"
                
                # إشعار النجاح وتحديث الخريطة فوراً
                self.page.snack_bar = ft.SnackBar(ft.Text(f"✅ تم ضبط الموقع: {location.address}"), bgcolor="green")
                self.page.snack_bar.open = True
                
                # إعادة حساب الخريطة بناءً على الإحداثيات الجديدة
                self.draw_now() 
            else:
                self.page.snack_bar = ft.SnackBar(ft.Text("⚠️ لم يتم العثور على المدينة، حاول كتابتها بالإنجليزية"), bgcolor="orange")
                self.page.snack_bar.open = True
            
            self.page.update()

        except Exception as ex:
            # التعامل مع أخطاء الاتصال بالإنترنت
            self.page.snack_bar = ft.SnackBar(ft.Text("❌ فشل الاتصال: تأكد من تفعيل الإنترنت في هاتفك"), bgcolor="red")
            self.page.snack_bar.open = True
            self.page.update()
def get_mansion_data(self, index):
        """قاعدة بيانات المنازل القمرية السبعة الأولى (كمثال)"""
        data = [
            ("الشرطين", "جبرائيل", "لبان ذكر", "البدايات والقوة", [12, 17, 10, 11, 13, 15, 16, 9, 14], "اللهم يا فاتح الأبواب افتح لي ببركة هذه المنزلة"),
            ("البطين", "إسرافيل", "سندروس", "الرزق والمال", [22, 27, 20, 21, 23, 25, 26, 19, 24], "يا رزاق ذو القوة المتين ارزقني من حيث لا أحتسب"),
            ("الثريا", "سمسمائيل", "عود وند", "المحبة والجمال", [5, 10, 3, 4, 6, 8, 9, 2, 7], "اللهم ألقِ محبتي في قلوب عبادك أجمعين"),
            ("الدبران", "ميكائيل", "كافور", "الهيبة والفراسة", [30, 35, 28, 29, 31, 33, 34, 27, 32], "يا عزيز عززني بعزك يا ميكائيل"),
            ("الهقعة", "صرفائيل", "جاوي", "العلم والذكاء", [15, 20, 13, 14, 16, 18, 19, 12, 17], "يا عليم علمني من علمك اللدني"),
            ("الهنعة", "عنيائيل", "ميعة سائله", "الإصلاح والمودة", [8, 13, 6, 7, 9, 11, 12, 5, 10], "اللهم ألف بين القلوب كما ألفت بين المنزلة"),
            ("الذراع", "كسفيائيل", "مستكة", "النصر والظفر", [40, 45, 38, 39, 41, 43, 44, 37, 42], "يا قوي يا متين انصرني نصراً عزيزاً"),
        ]
        return data[index % len(data)]

def show_moon_mansion(self, e=None):
        """حساب وعرض منزلة القمر الحالية وتفاصيلها الروحانية"""
        try:
            # 1. حساب موقع القمر الحالي
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, now.hour + now.minute/60.0)
            res, _ = swe.calc_ut(jd, swe.MOON)
            moon_pos = res[0]
            
            # 2. تحديد رقم المنزلة (الدائرة 360 تقسم على 28 منزلة = 12.85 درجة لكل منزلة)
            mansion_idx = int(moon_pos / 12.857)
            m_name, m_angel, m_incense, m_desc, m_square, m_dua = self.get_mansion_data(mansion_idx)

            # 3. بناء واجهة عرض الوفق (3x3 Grid)
            grid_cells = []
            for num in m_square:
                grid_cells.append(
                    ft.Container(
                        content=ft.Text(str(num), weight="bold", color="white"),
                        alignment=ft.alignment.center,
                        width=50, height=50,
                        bgcolor="#1e293b", border_radius=5
                    )
                )

            # 4. عرض نافذة التفاصيل
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text(f"🌙 منزلة القمر: {m_name}", weight="bold", rtl=True, color="#8b5cf6"),
                content=ft.Column([
                    ft.Text(f"👤 المَلَك الموكل: {m_angel}", weight="bold", rtl=True),
                    ft.Text(f"💨 البخور: {m_incense}", rtl=True),
                    ft.Text(f"📖 الدلالة: {m_desc}", rtl=True, italic=True),
                    ft.Divider(),
                    ft.Text("🔢 الوفق العددي المقترح:", weight="bold", rtl=True),
                    ft.GridView(controls=grid_cells, runs_count=3, max_extent=60, height=180),
                    ft.Divider(),
                    ft.Text(f"🤲 الدعاء:\n{m_dua}", rtl=True, size=13, color="#475569"),
                ], tight=True, scroll=ft.ScrollMode.ADAPTIVE),
                actions=[ft.TextButton("إغلاق", on_click=close_dlg)]
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ: {ex}")))
def start_app(self):
        """تشغيل التطبيق وتهيئته الأولية"""
        try:
            # 1. بناء واجهة المستخدم الرسومية
            self.setup_ui()
            
            # 2. جلب التاريخ الحالي وإجراء الحسابات الفلكية الأولية
            # هذا يمنع ظهور الجداول أو الخرائط فارغة عند الفتح
            self.draw_now()
            
            # 3. تحديث الصفحة النهائية
            self.page.update()
            
            # إشعار ترحيبي بسيط (اختياري)
            self.page.show_snack_bar(
                ft.SnackBar(ft.Text("🚀 تم تشغيل النظام الفلكي بنجاح"), bgcolor="blue")
            )
        except Exception as e:
            # معالجة الأخطاء لضمان عدم توقف التطبيق في بيئة أندرويد
            print(f"خطأ أثناء بدء التطبيق: {e}")
            if hasattr(self, 'page'):
                self.page.add(ft.Text(f"حدث خطأ أثناء تحميل الواجهة: {str(e)}", color="red"))
                self.page.update()
def nail_cut_calendar(self):
        """تقويم قص الأظافر بناءً على الأيام القمرية والمخطوطات التراثية"""
        try:
            import swisseph as swe
            from datetime import datetime, timezone

            # 1. الحساب الفلكي الدقيق لليوم القمري (عمر القمر)
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, 12)
            
            # جلب مواقع النيرين (الشمس والقمر)
            res_m, _ = swe.calc_ut(jd, swe.MOON) 
            res_s, _ = swe.calc_ut(jd, swe.SUN) 
            
            m_deg = float(res_m[0])
            s_deg = float(res_s[0])
            
            # حساب الفرق الزاوي وتحويله لليوم القمري (دورة 29.5 يوم)
            day = int(((m_deg - s_deg) % 360 / 360) * 29.5) + 1
            
            # 2. مصفوفة الأحكام التراثية
            nail_rules = {
                1: "يورث قصر العمر والكآبة (تجنبه)",
                2: "يورث النكد وتشتت الأمر",
                3: "يورث البركة في المال والولد (مستحب جداً)",
                4: "يورث الفقر والهم",
                5: "يورث زيادة الرزق وتيسير الأمور (مبارك)",
                6: "يورث الغم والوسواس",
                7: "يورث الهيبة والوقار والقبول عند الناس (يوم ملكي)",
                8: "يورث الضعف البدني وتراجع القوة",
                9: "يورث سوء الخلق وضيق الصدر",
                10: "يورث الرفعة والكرامة والمنزلة",
                14: "يورث الشفاء من الأوجاع الجسدية (مستحب للصحّة)",
                15: "يورث السرور والفرح وقضاء الحوائج (مبارك)",
                20: "يورث الأمن من الخوف والعز (مستحب)",
                24: "يورث الظفر بالأعداء ونيل المقاصد",
                25: "يورث الخلاص من الديون والضيق"
            }

            # تحديد الحالة واللون المناسب
            current_status = nail_rules.get(day, "يوم اعتيادي لقص الأظافر، والأفضل تحري يوم الجمعة.")
            is_suitable = "✅ مناسب ومستحب" if day in [3, 5, 7, 10, 14, 15, 20, 24, 25] else "❌ غير مناسب (يُفضل التأجيل)"
            status_color = "#10b981" if day in [3, 5, 7, 10, 14, 15, 20, 24, 25] else "#ef4444"

            # 3. بناء التقرير النصي
            report = f"📅 اليوم القمري (الهلالي): {day}\n"
            report += f"📍 حالة اليوم لقص الأظافر: {is_suitable}\n"
            report += "══════════════════════════════\n"
            report += f"📖 حكم هذا اليوم (حسب المخطوطات):\n{current_status}\n"
            report += "══════════════════════════════\n"
            report += "📜 من مخطوطات الإمام جعفر الصادق:\n"
            report += "يقول عليه السلام: 'تقليم الأظافر يوم الجمعة يؤمن من الجذام والبرص والعمى'...\n\n"
            report += "💎 عند ابن سيرين:\n"
            report += "يقول: 'الأظافر قوة الرجل ومقدرته'...\n"
            report += "══════════════════════════════\n"
            report += "✍️ المطور: حسان الشاعر"

            # 4. عرض التقرير في نافذة AlertDialog (Flet)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("💅 تقويم قص الأظافر الروحاني", weight="bold", rtl=True, color=status_color),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report, rtl=True, size=14, color="#2c3e50")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450,
                    padding=10,
                    bgcolor="#fffcf5" # لون خلفية المخطوطة القديمة
                ),
                actions=[
                    ft.TextButton("نسخ التقرير", on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as e:
            if hasattr(self, 'page'):
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"حدث خلل في التقويم: {str(e)}", rtl=True)))
def analyze_fixed_stars(self, e=None):
        """محرك فحص وتحليل النجوم الثابتة العظمى"""
        try:
            # 1. قاعدة البيانات الموسعة (18 نجماً) محدثة لعام 2026
            fixed_stars = [
                {"name": "الظليم (Achernar)", "pos": 15.3, "nature": "مشتري", "effect": "المكانة الروحية العالية، والنجاح في المساعي الدينية."},
                {"name": "رأس الغول (Algol)", "pos": 56.2, "nature": "زحل/مريخ", "effect": "يمنح صموداً أسطورياً وقدرة على مواجهة الأزمات الكبرى."},
                {"name": "الثريا (Alcyone)", "pos": 60.1, "nature": "زهرة/مشتري", "effect": "الجاذبية الشخصية، الشهرة الفنية، والنجاح في العلوم."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "nature": "مريخ", "effect": "الشجاعة، القيادة السيادية، والنجاح المادي من خلال المواجهات."},
                {"name": "رجل الجبار (Rigel)", "pos": 77.0, "nature": "مشتري/مريخ", "effect": "الشهرة الواسعة، الاختراع، والقدرة على القيادة الإدارية."},
                {"name": "النطاق (Alnilam)", "pos": 83.5, "nature": "مشتري/زحل", "effect": "الارتقاء للمناصب العليا، والسمعة التي تدوم طويلاً."},
                {"name": "منكب الجوزاء (Betelgeuse)", "pos": 88.8, "nature": "مريخ/مشتري", "effect": "نجاح باهر ومفاجئ، ومكانة مرموقة في الدولة."},
                {"name": "شعرى اليمانية (Sirius)", "pos": 104.1, "nature": "مشتري/مريخ", "effect": "يمنح شهرة تاريخية، حماية ملكية، ونجاحاً يتجاوز الحدود."},
                {"name": "رأس التوأم (Pollux)", "pos": 113.3, "nature": "مريخ", "effect": "قوة البيان، الشجاعة، والقدرة على الغلبة."},
                {"name": "شعرى الشامية (Procyon)", "pos": 116.3, "nature": "مشتري/مريخ", "effect": "النجاح السريع والمفاجئ، والذكاء العملي."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "nature": "مشتري/مريخ", "effect": "الجاه العظيم، السلطة، الرفعة السيادية، والكرامة."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "nature": "زهرة/مشتري", "effect": "أكثر النجوم سعادة: الثروة، الثقافة، والنجاح العلمي."},
                {"name": "سماك الرامح (Arcturus)", "pos": 204.2, "nature": "مشتري/مريخ", "effect": "العبقرية، القيادة الفكرية، والتميز الفني."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "nature": "مريخ/مشتري", "effect": "التحولات الكبرى، والقدرة المذهلة على إعادة البناء."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "nature": "زهرة/عطارد", "effect": "الإبداع الروحاني، السحر البلاغي، والتميز الفني."},
                {"name": "النسر الطائر (Altair)", "pos": 292.0, "nature": "مريخ/مشتري", "effect": "الطموح العالي جداً، والوصول للأهداف الصعبة."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "nature": "زهرة/عطارد", "effect": "الشهرة العالمية، والتميز في العلوم الروحانية والفلكية."},
                {"name": "الردف (Deneb)", "pos": 335.3, "nature": "زهرة/عطارد", "effect": "الذكاء الثاقب، القدرة على التعلم، والنجاح في الفنون."},
            ]

            # دالة استخراج اسم البرج (نفس المنطق المستخدم في تطبيقك)
            def get_s_name(deg):
                signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
                return signs[int(deg / 30)]

            # 2. إنشاء عناصر واجهة العرض (ListView لضمان التمرير السلس في أندرويد)
            report_display = ft.ListView(expand=True, spacing=10, padding=20)
            
            report_display.controls.append(
                ft.Text("🌍 دليل النجوم الثابتة (مواقع 2026)", size=20, weight="bold", color="#38bdf8", rtl=True)
            )

            for s in fixed_stars:
                report_display.controls.append(
                    ft.Container(
                        content=ft.Column([
                            ft.Text(f"⭐ {s['name']}", weight="bold", color="#fbbf24", size=16),
                            ft.Text(f"📍 الموقع: {get_s_name(s['pos'])} ({int(s['pos']%30)}°)", size=14, color="#e2e8f0"),
                            ft.Text(f"🌿 طبيعته: {s['nature']}", size=13, color="#94a3b8"),
                            ft.Text(f"✨ التأثير: {s['effect']}", size=14, italic=True, color="#cbd5e1"),
                        ], spacing=5),
                        padding=15, 
                        bgcolor="#1e293b",
                        border=ft.border.all(1, "#334155"), 
                        border_radius=10
                    )
                )

            # 3. عرض التقرير في نافذة منبثقة (BottomSheet أو Dialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                content=ft.Container(content=report_display, width=500, height=600),
                actions=[ft.TextButton("إغلاق", on_click=close_dlg)],
                modal=False
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            if hasattr(self, 'page'):
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ في تحليل النجوم: {ex}")))
def show_fixed_stars_dialog(self, e=None):
        """عرض نافذة النجوم الثابتة مع خيار فحص النجوم الشخصية"""
        try:
            # 1. قاعدة البيانات (تم استخراجها من سياق الكود السابق)
            fixed_stars = [
                {"name": "الثريا (Alcyone)", "pos": 60.1, "effect": "الجاذبية، الشهرة، والنجاح الفني."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "effect": "الشجاعة والقيادة السيادية."},
                {"name": "رأس التوأم (Pollux)", "pos": 113.3, "effect": "قوة البيان والقدرة على الغلبة."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "effect": "الجاه العظيم والرفعة السيادية."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "effect": "الثروة والنجاح العلمي الباهر."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "effect": "التحولات الكبرى وإعادة البناء."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "effect": "الإبداع الروحاني والتميز الفني."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "effect": "الشهرة العالمية والوعي الكوني."}
            ]

            report_display = ft.Column(scroll=ft.ScrollMode.ALWAYS, expand=True, spacing=10)

            # وظيفة عرض النجوم العامة اليوم
            def show_current_stars(e):
                report_display.controls.clear()
                report_display.controls.append(ft.Text("🌍 مواقع النجوم الثابتة اليوم", size=18, weight="bold", color="#38bdf8", rtl=True))
                for s in fixed_stars:
                    report_display.controls.append(
                        ft.Container(
                            content=ft.Column([
                                ft.Text(f"⭐ {s['name']}", weight="bold", color="#fbbf24"),
                                ft.Text(f"📍 الموقع: {int(s['pos']/30)} {int(s['pos']%30)}°", size=13),
                                ft.Text(f"✨ التأثير: {s['effect']}", size=13, italic=True),
                            ]),
                            padding=10, border=ft.border.all(1, "#38bdf8"), border_radius=8
                        )
                    )
                self.page.update()

            # وظيفة تحليل نجوم الميلاد الشخصية
            def show_natal_stars(e):
                report_display.controls.clear()
                if not hasattr(self, 'is_analyzed') or not self.is_analyzed:
                    report_display.controls.append(
                        ft.Container(
                            content=ft.Text("⚠️ يرجى تحليل خارطة الميلاد أولاً من الواجهة الرئيسية.", color="red", weight="bold", rtl=True),
                            padding=20, alignment=ft.alignment.center
                        )
                    )
                    self.page.update()
                    return

                report_display.controls.append(ft.Text("👶 تحليل اقترانات النجوم في ميلادك", size=18, weight="bold", color="#10b981", rtl=True))
                
                # جلب بيانات الكواكب المحسوبة سابقاً
                points = {**self.raw_d, "Ascendant": self.asc_raw}
                p_names = {"sun":"الشمس", "moon":"القمر", "mercury":"عطارد", "venus":"الزهرة", "mars":"المريخ", "jupiter":"المشتري", "saturn":"زحل", "Ascendant":"الطالع"}
                
                found = False
                for s in fixed_stars:
                    for p_en, p_deg in points.items():
                        diff = abs(p_deg - s['pos'])
                        # جرم السماح (Orb) 1.5 درجة
                        if diff <= 1.5 or diff > 358.5:
                            p_ar = p_names.get(p_en, p_en)
                            report_display.controls.append(
                                ft.Container(
                                    content=ft.Column([
                                        ft.Text(f"✅ اقتران ملكي: {s['name']} مع {p_ar}", weight="bold", color="#065f46"),
                                        ft.Text(f"◈ الدلالة المقديرة: {s['effect']}", size=14, color="#1e293b"),
                                    ]),
                                    padding=12, bgcolor="#ecfdf5", border_radius=8, border=ft.border.all(1, "#10b981")
                                )
                            )
                            found = True
                
                if not found:
                    report_display.controls.append(
                        ft.Container(
                            content=ft.Text("🔍 لم يتم العثور على اقترانات دقيقة بالنجوم العظمى في هيئة ميلادك.", size=14, rtl=True),
                            padding=20
                        )
                    )
                self.page.update()

            # 3. بناء النافذة المنبثقة بتصميم أنيق
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Row([
                    ft.Icon(ft.icons.STARS, color="#fbbf24"),
                    ft.Text("ديوان النجوم الثابتة", size=22, weight="bold", rtl=True)
                ], alignment=ft.MainAxisAlignment.CENTER),
                content=ft.Container(
                    content=ft.Column([
                        ft.Row([
                            ft.ElevatedButton("مواقع اليوم", on_click=show_current_stars, bgcolor="#38bdf8", color="white", icon=ft.icons.LANGUAGE),
                            ft.ElevatedButton("نجوم ميلادي", on_click=show_natal_stars, bgcolor="#10b981", color="white", icon=ft.icons.PERSON),
                        ], alignment=ft.MainAxisAlignment.CENTER, spacing=10),
                        ft.Divider(height=20, color="#cbd5e1"),
                        report_display
                    ], tight=True),
                    width=500, height=600
                ),
                actions=[ft.TextButton("إغلاق النافذة", on_click=close_dlg)],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            show_current_stars(None) # العرض الافتراضي عند الفتح
            self.page.update()

        except Exception as e:
            self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ في محرك النجوم: {str(e)}", rtl=True)))
def show_moon_mansion(self):
        try:
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, now.hour + now.minute/60.0)
            res_m, _ = swe.calc_ut(jd, swe.MOON)
            moon_pos = float(res_m[0])
            
            # قاعدة البيانات المكتملة التي أرسلتها
            STARS_INFO = {
                1: ("الشرطين", "2", "الحمل", "أ", "نجمان (النطح والشيرتان)", "منزلة نارية، تدل على الحماس والبدايات والاندفاع القوي."),
                2: ("البطين", "3", "الحمل", "ب", "ثلاثة نجوم خفية", "منزلة ترابية، تمتاز بالثبات والقدرة على التخطيط المالي."),
                3: ("الثريا", "7", "الثور", "ج", "الأخوات السبع (الثريا)", "منزلة هوائية، وهي من أسعد المنازل للمحبة والقبول والشهرة."),
                4: ("الدبران", "5", "الثور", "د", "عين الثور (الدبران)", "منزلة ترابية، تدل على الهيبة والقوة الفكرية العالية."),
                5: ("الهقعة", "3", "الجوزاء", "هـ", "رأس الجبار (الهقعة)", "منزلة هوائية، مناسبة جداً للتعلم وتبادل المعلومات بذكاء."),
                6: ("الهنعة", "2", "الجوزاء", "و", "قدم الجوزاء (الهنعة)", "منزلة مائية، تمتاز بطاقة الألفة وإصلاح العلاقات المتوترة."),
                7: ("الذراع", "2", "السرطان", "ز", "رأس التوأمين (الذراع)", "منزلة مائية، تمنح حماية للممتلكات وتساعد في الظفر."),
                8: ("النثرة", "3", "السرطان", "ح", "سديم المعلف (النثرة)", "منزلة نارية، تدل على الرفعة وعلو الشأن والنجاح الرسمي."),
                9: ("الطرفة", "2", "الأسد", "ط", "عيني الأسد (الطرفة)", "منزلة نارية، تتطلب الحذر وهي قوية للتغييرات الجذرية."),
                10: ("الجبهة", "4", "الأسد", "ي", "جبهة الأسد (الجبهة)", "منزلة نارية، تعيد الهيبة والشرف وتساعد في الاستشفاء."),
                11: ("الزبرة", "2", "الأسد/العذراء", "ك", "ظهر الأسد (الزبرة)", "منزلة نارية، مباركة للتجارة والارتباط العاطفي الناجح."),
                12: ("الصرفة", "1", "العذراء", "ل", "ذنب الأسد (الصرفة)", "منزلة ترابية، تدل على التغيير والانتقال، ومناسبة للسفر."),
                13: ("العواء", "5", "العذراء", "م", "خمسة نجوم (العواء)", "منزلة ترابية، تمتاز بسرعة الاستجابة للأماني وقضاء الحوائج."),
                14: ("السماك", "1", "الميزان", "ن", "السماك الأعزل", "منزلة هوائية، منبع النمو والازدهار في الصفقات المالية."),
                15: ("الغفر", "3", "الميزان", "س", "الغفر", "منزلة هوائية، تدل على الاستبصار والبحث في بواطن الأمور."),
                16: ("الزبانا", "2", "الميزان/العقرب", "ع", "كفتا الميزان (الزبانا)", "منزلة هوائية، مناسبة لضبط الموازين المالية والحقوق."),
                17: ("الإكليل", "3", "العقرب", "ف", "رأس العقرب (الإكليل)", "منزلة مائية، قوية جداً في التأليف بين القلوب وجذب المودة."),
                18: ("القلب", "1", "العقرب", "ص", "قلب العقرب", "منزلة مائية، للتحصين والوقاية والتعامل مع النفس بعمق."),
                19: ("الشولة", "2", "العقرب/القوس", "ق", "ذيل العقرب (الشولة)", "منزلة نارية، تدل على التحرر من القيود والسرعة في الإنجاز."),
                20: ("النعائم", "8", "القوس", "ر", "النعائم", "منزلة نارية، منزلة الوفرة والرزق وتيسير الأمور المتعسرة."),
                21: ("البلدة", "1", "القوس/الجدي", "ش", "البلدة (فراغ سماوي)", "منزلة ترابية، منزلة الاستقرار وبناء الأصول والعمران."),
                22: ("سعد الذابح", "2", "الجدي", "ت", "سعد الذابح", "منزلة ترابية، تدل على الحزم والنجاة من الأخطار المحيطة."),
                23: ("سعد بلع", "2", "الجدي/الدلو", "ث", "سعد بلع", "منزلة هوائية، مناسبة للعلاجات الطبية وتسهيل الصعاب."),
                24: ("سعد السعود", "3", "الدلو", "خ", "سعد السعود", "منبع السعادة والأفراح والنمو في كافة جوانب الحياة."),
                25: ("سعد الأخبية", "4", "الدلو", "ذ", "سعد الأخبية", "تساعد في كشف الأسرار والظهور القوي بعد فترة غياب."),
                26: ("المقدم", "2", "الحوت", "ض", "فرع الدلو المقدم", "منزلة مائية، تدل على الجاه والعز والتواصل الفكري."),
                27: ("المؤخر", "2", "الحوت", "ظ", "فرع الدلو المؤخر", "منزلة مائية، مباركة لزيادة الأرزاق والنجاح في المشاريع."),
                28: ("الرشاء", "1", "الحوت/الحمل", "غ", "الرشاء (بطن الحوت)", "منزلة مائية، منزلة الخواتيم السعيدة وإتمام المهام بنجاح.")
            }
   # تحديد رقم المنزلة بناءً على موقع القمر
            m_num = int(moon_pos / (360/28)) + 1
            m_data = STARS_INFO.get(m_num)
            
            # عرض النتائج في واجهة Flet (كما في الدوال السابقة)
            self.display_mansion_report(m_num, m_data)
 # 1. الحسابات الفلكية (الزمن الحالي)
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, now.hour + now.minute/60.0)
            res_m, _ = swe.calc_ut(jd, swe.MOON)
            res_s, _ = swe.calc_ut(jd, swe.SUN)
            
            moon_pos = float(res_m[0])
            sun_pos = float(res_s[0])
            
            # عمر القمر ومكانه
            moon_age = ((moon_pos - sun_pos) % 360 * 29.53) / 360
            mansion_idx = int(moon_pos / (360/28))
            m_num = mansion_idx + 1

            # 2. جلب البيانات (من STARS_INFO و get_mansion_data)
            # التأكد من وجود STARS_INFO كـ Dictionary داخل الكلاس
            m_data = self.STARS_INFO.get(m_num, (0, "", "", "", ""))
            m_stars_cnt, m_signs, m_chars, m_stars_desc, m_explanation = m_data
            
            name, angel, incense, effect, wifq, prayer = self.get_mansion_data(mansion_idx)

            # 3. بناء نص التقرير المنسق
            report = (
                "✨ الديوان الشامل للمنازل القمرية ✨\n"
                "══════════════════════════\n"
                f"🌙 المنزلة: {name} (رقم: {m_num})\n"
                f"⏳ عمر القمر: {moon_age:.1f} يوم\n"
                f"📍 الموقع: {int(moon_pos%30)}° في برج {self.get_sign_name(moon_pos)}\n"
                "══════════════════════════\n"
                f"🌟 النجوم: {m_stars_desc}\n"
                f"♈ البروج: {m_signs} | 🔠 الحروف: ( {m_chars} )\n"
                "══════════════════════════\n"
                f"🔍 التفسير: {m_explanation}\n"
                "══════════════════════════\n"
                f"👼 الملك: {angel} | 🪵 البخور: {incense}\n"
                f"📖 التأثير: {effect}\n"
                "══════════════════════════\n"
                "🔢 الوفق العددي للمنزلة:\n"
                f"    [ {wifq[0]:02}  {wifq[1]:02}  {wifq[2]:02} ]\n"
                f"    [ {wifq[3]:02}  {wifq[4]:02}  {wifq[5]:02} ]\n"
                f"    [ {wifq[6]:02}  {wifq[7]:02}  {wifq[8]:02} ]\n\n"
                f"✨ العزيمة: 「 {prayer} 」\n"
                "══════════════════════════\n"
                "المطور الفلكي: حسان الشاعر"
            )

            # 4. إعداد واجهة العرض (Flet AlertDialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text(f"ديوان منزلة {name}", weight="bold", rtl=True, color="#1f538d"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report, size=14, rtl=True, color="#3d405b", font_family="monospace")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, 
                    height=700, 
                    bgcolor="#f4f1de", # خلفية بنية فاتحة تشبه الورق القديم
                    padding=15, 
                    border_radius=10
                ),
                actions=[
                    ft.ElevatedButton("نسخ التقرير", icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as e:
            self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ في المنازل: {str(e)}", rtl=True)))    
def analyze_current_time(self, e=None):
        """توليد التقرير الاستراتيجي الزمني والمالي"""
        try:
            # 1. إعداد محتوى التقرير (البيانات التحليلية)
            content = "🔭 التحليل الزمني الشامل للأبراج والأسواق 🔭\n"
            content += "════════════════════════════════\n\n"
            
            content += "📍 أولاً: التحليل اليومي للأبراج\n"
            z_names = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                       "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            adv_list = ["طاقة نارية للمبادرة والحسم", "فرص ترابية للمال والعمل", 
                        "نشاط هوائي للذكاء والتواصل", "هدوء مائي للعاطفة والحدس"]
            
            for i, name in enumerate(z_names):
                adv = adv_list[i % 4]
                content += f"● {name}: {adv}.\n"

            content += "\n💰 ثانياً: الاقتصاد والذهب (2026)\n"
            content += "• هذا الشهر: يشهد الذهب تذبذباً يميل للارتفاع نتيجة قوة زاوية المشتري.\n"
            content += "• سنوياً: عام 2026 هو عام التحول الرقمي الشامل وصعود المعادن النفيسة.\n"
            
            content += "\n════════════════════════════════\n"
            content += "إعداد المطور الفلكي: حسان الشاعر © 2026"

            # 2. دالة الحفظ (متوافقة مع أندرويد عبر Flet)
            def save_report_file(e):
                try:
                    import os
                    # تحديد مسار الحفظ في مجلد المستندات أو التحميلات
                    file_name = "Strategic_Report_2026.txt"
                    # محاولة الحفظ في المسار العام للأندرويد
                    path = "/storage/emulated/0/Download" if os.path.exists("/storage/emulated/0/Download") else os.getcwd()
                    full_path = os.path.join(path, file_name)
                    
                    with open(full_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    
                    self.page.show_snack_bar(ft.SnackBar(ft.Text(f"✅ تم حفظ التقرير في: {file_name}"), bgcolor="green"))
                except Exception as ex:
                    self.page.show_snack_bar(ft.SnackBar(ft.Text(f"❌ فشل الحفظ: {str(ex)}"), bgcolor="red"))

            # 3. إعداد وإغلاق النافذة المنبثقة
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("التقرير الاستراتيجي الزمني", weight="bold", rtl=True, color="#1f538d"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(content, size=14, rtl=True, color="#1e293b", weight="bold"),
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450,
                    height=600,
                    bgcolor="#f8fafc",
                    padding=20,
                    border_radius=15
                ),
                actions=[
                    ft.ElevatedButton("💾 حفظ الملف", on_click=save_report_file, bgcolor="#10b981", color="white"),
                    ft.ElevatedButton("📋 نسخ", on_click=lambda _: self.page.set_clipboard(content), bgcolor="#3b82f6", color="white"),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.SPACE_EVENLY,
            )

            # 4. عرض النافذة وتحديث الواجهة
            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            if hasattr(self, 'page'):
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ في توليد التقرير: {str(ex)}")))
def hair_cut_calendar(self, e=None):
        """تقويم قص الشعر الروحاني بناءً على حسابات اليوم القمري"""
        try:
            import swisseph as swe
            from datetime import datetime, timezone

            # 1. الحساب الفلكي الدقيق لليوم القمري
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, 12)
            res_m, _ = swe.calc_ut(jd, swe.MOON) # موقع القمر
            res_s, _ = swe.calc_ut(jd, swe.SUN)  # موقع الشمس
            
            # استخراج الدرجات وضمان التوافق العشري
            m_deg = float(res_m[0])
            s_deg = float(res_s[0])
            
            # حساب اليوم القمري (الفرق الزاوي)
            day = int(((m_deg - s_deg) % 360 / 12.19) + 1)
            if day > 30: day = 30 

            # 2. قاعدة بيانات الأحكام التراثية
            hair_rules = {
                1: "قصر العمر (تجنبه)", 2: "الهم والغم", 3: "النكد وتشتت الأمر",
                4: "الوجع في الرأس", 5: "البركة والمال الكثير", 6: "الهم وضيق الصدر",
                7: "الهيبة والجاه والقبول (مستحب جداً)", 8: "العافية وطول العمر",
                9: "الضعف وسوء الخلق", 10: "الرفعة والكرامة", 11: "الحزن والكآبة",
                12: "العز والوقار (مستحب)", 13: "الخصومة والنزاع", 
                14: "السرور والفرح وقضاء الحاجة (مبارك)", 15: "تيسير الأمور وزيادة الجمال",
                16: "الحزن وضيق الرزق", 17: "النحوسة (تجنبه)", 18: "البركة في الرزق",
                19: "القدرة والتمكين", 20: "الأمن من الخوف والهيبة (مبارك)",
                21: "تشتت الرأي", 22: "الفقر الشديد (تجنبه)", 23: "العافية من الأمراض",
                24: "النصر على الأعداء", 25: "الخلاص من السجون والديون",
                26: "الفرح والسرور", 27: "السلامة من الآفات", 28: "قضاء الحاجة (مستحب)",
                29: "الفقر والهم (تجنبه)", 30: "الأمن من الآفات"
            }

            suitable_days = [5, 7, 8, 10, 12, 14, 15, 18, 19, 20, 23, 24, 25, 26, 27, 28, 30]
            current_status = hair_rules.get(day, "يوم يحتاج لمراجعة")
            is_suitable = "✅ مناسب ومستحب" if day in suitable_days else "❌ غير مناسب (يُفضل التأجيل)"

            # 3. صياغة نص التقرير المنسق
            report = (
                f"✂️ تقويم قص الشعر الروحاني ✂️\n"
                f"📅 اليوم القمري الحالي: {day}\n"
                f"📍 الحالة: {is_suitable}\n"
                f"══════════════════════════════\n"
                f"📖 حكم اليوم: يورث {current_status}\n"
                f"══════════════════════════════\n"
                f"📜 من مخطوطات جعفر الصادق:\n'من قص في السابع نال هيبة، وفي الـ 14 نال سروراً'.\n\n"
                f"💎 فلسفة ابن سيرين:\nقص الشعر يرمز لذهاب الهم وقضاء الدين في الأيام السعيدة.\n"
                f"══════════════════════════════\n"
                f"✍️ إعداد المطور: حسان الشاعر"
            )

            # 4. دالة الحفظ السريع
            def save_hair_report(e):
                try:
                    import os
                    path = "/storage/emulated/0/Download" if os.path.exists("/storage/emulated/0/Download") else os.getcwd()
                    full_path = os.path.join(path, "Hair_Cut_Report.txt")
                    with open(full_path, "w", encoding="utf-8") as f:
                        f.write(report)
                    self.page.show_snack_bar(ft.SnackBar(ft.Text(f"✅ تم الحفظ في: {full_path}")))
                except:
                    self.page.show_snack_bar(ft.SnackBar(ft.Text("❌ فشل الحفظ")))

            # 5. واجهة العرض (Flet AlertDialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("تقويم قص الشعر الروحاني", weight="bold", rtl=True, size=20, color="#ec4899"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report, size=15, rtl=True, color="#1e293b", weight="w500")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, height=550, bgcolor="#fdf2f8", padding=20, border_radius=15
                ),
                actions=[
                    ft.ElevatedButton("💾 حفظ", on_click=save_hair_report, bgcolor="#ec4899", color="white"),
                    ft.ElevatedButton("📋 نسخ", on_click=lambda _: self.page.set_clipboard(report), bgcolor="#f472b6", color="white"),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.SPACE_EVENLY,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            if hasattr(self, 'page'):
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"خطأ في التقويم: {str(ex)}")))
def update_location(self, e=None):
        """تحديث الإحداثيات يدوياً من خانات الإدخال"""
        try:
            # استخدام .value لجلب البيانات من TextField في Flet
            self.lat = float(self.lat_entry.value if self.lat_entry.value else 0)
            self.lon = float(self.lon_entry.value if self.lon_entry.value else 0)
            
            # إشعار نجاح العملية في واجهة الأندرويد
            self.page.snack_bar = ft.SnackBar(
                ft.Text(f"✅ تم تحديث الإحداثيات: {self.lat} , {self.lon}", rtl=True),
                bgcolor="#10b981"
            )
            self.page.snack_bar.open = True
            
            # إعادة حساب الهيئة الفلكية فوراً بناءً على الموقع الجديد
            self.draw_now() 
            self.page.update()
            
        except ValueError:
            # معالجة أخطاء الإدخال النصي في خانات الأرقام
            self.page.snack_bar = ft.SnackBar(
                ft.Text("⚠️ خطأ: يرجى إدخال أرقام صحيحة لخطوط العرض والطول", rtl=True),
                bgcolor="#ef4444"
            )
            self.page.snack_bar.open = True
            self.page.update()

def get_house(self, pos, cusps):
        """تحديد رقم البيت (1-12) بناءً على موقع الجرم وحدود البيوت"""
        # نضمن أن الموقع يقع في دائرة الـ 360
        pos %= 360
        for i in range(12):
            # حد بداية البيت وحد نهايته (البيت التالي)
            c1 = cusps[i]
            c2 = cusps[(i + 1) % 12]
            
            # الحالة 1: البيت يقع بالكامل ضمن دورة واحدة
            if c1 < c2:
                if c1 <= pos < c2:
                    return i + 1
            # الحالة 2: البيت يقطع نقطة الصفر (نهاية الحوت وبداية الحمل)
            else:
                if pos >= c1 or pos < c2:
                    return i + 1
        return 1
def show_about(self, e=None):
        """عرض نافذة معلومات المطور والنظام بتنسيق تقني جذاب"""
        email = "Hassan.alshaer2@gmail.com"
        # التأكد من وجود متغير اللغة أو تعيينه افتراضياً للعربية
        is_ar = getattr(self, 'current_lang', 'ar') == "ar"
        
        lines = [
            f"{'المطور البرمجي الفلكي: حسان الشاعر' if is_ar else 'Developer: Hassan Al-Shaer'}",
            f"Email: {email}",
            f"{'العنوان: سوريا / طرطوس' if is_ar else 'Location: Syria / Tartus'}",
            "WhatsApp: +963 933303612",
            "----------------------------------",
            f"{'برامجنا الاحترافية:' if is_ar else 'Professional Software:'}",
            "* Animodar Correction System",
            "* Primary Directions Pro",
            "* Al-Khafia Solar System"
        ]

        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # إعداد النافذة المنبثقة بتنسيق Terminal الفسفوري
        dlg = ft.AlertDialog(
            title=ft.Text("About Developer" if not is_ar else "حول المطور", weight="bold"),
            content=ft.Container(
                content=ft.Text("\n".join(lines), color="#39FF14", size=14, font_family="monospace"),
                bgcolor="#000000",
                padding=20,
                border_radius=10,
                border=ft.border.all(2, "#39FF14")  # لون الفسفور الاحترافي
            ),
            actions=[
                ft.TextButton("Close" if not is_ar else "إغلاق", on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
def show_mundane_analysis(self, jd):
        """توليد وعرض التقرير الملحمي لسنة العالم 2026"""
        try:
            import swisseph as swe
            # 1. إعداد البيانات والمواقع
            lat = float(getattr(self, 'lat', 34.8))
            lon = float(getattr(self, 'lon', 35.8))
            z_names_ar = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                          "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            
            # 2. الحسابات الفلكية والسهام الاستراتيجية
            cusps, ascmc = swe.houses(jd, lat, lon, b'P')
            asc_raw = float(ascmc[0])
            
            def get_p(obj_id):
                res, _ = swe.calc_ut(jd, obj_id)
                return {'deg': float(res[0]), 'speed': float(res[3])}

            sun = get_p(0)
            jupiter = get_p(5)
            mars = get_p(4)
            saturn = get_p(6)

            # حساب السهام (الذهب، النفط، الغذاء)
            gold_lot = (asc_raw + jupiter['deg'] - sun['deg']) % 360
            oil_lot = (asc_raw + mars['deg'] - jupiter['deg']) % 360
            food_lot = (asc_raw + mars['deg'] - saturn['deg']) % 360

            # 3. بناء نص التقرير (الدمج النهائي)
            report_body = (
                "📜 الديوان السلطاني الشامل لحكم سنة العالم 2026 📜\n"
                "════════════════════════════════\n\n"
                "💰 الاقتصاد والذهب والنفط:\n"
                f"◈ سهم الذهب: في برج {z_names_ar[int(gold_lot/30)]} (زلزال مالي مرتقب).\n"
                f"◈ سهم البترول: في برج {z_names_ar[int(oil_lot/30)]} (اشتعال في أسعار الطاقة).\n\n"
                "🌾 خامساً: ميزان الغذاء والحبوب (سهم القمح):\n"
                f"◈ سهم الغذاء: في {z_names_ar[int(food_lot/30)]} ينبئ بصعوبات في الأمن الغذائي.\n"
                "◈ القمح: سيكون 'سلاح العصر' الجديد مع تقلبات حادة في الأسعار.\n"
                "◈ نصيحة: يُنصح بتأمين مخزونات الغذاء الأساسية قبل الربع الثالث.\n\n"
                "════════════════════════════════\n"
                "✍️ إعداد وتطوير: حسان الشاعر © 2026\n"
                "جميع الحقوق البرمجية والفلكية محفوظة للمطور"
            )

            # 4. واجهة العرض (Flet AlertDialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("الديوان السلطاني - سنة 2026", weight="bold", color="#7c2d12", rtl=True),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report_body, size=15, rtl=True, color="#1e40af", weight="bold")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=550, height=750, bgcolor="#fcfaf2", padding=20, border_radius=10
                ),
                actions=[
                    ft.TextButton("نسخ الديوان", on_click=lambda _: self.page.set_clipboard(report_body)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"حدث خطأ: {str(ex)}"))
            self.page.snack_bar.open = True
            self.page.update()
def save_report(self, content):
        """حفظ التقارير الفلكية والروحانية كملفات نصية في ذاكرة الهاتف"""
        try:
            import os
            from datetime import datetime
            
            # 1. إنشاء اسم ملف فريد باستخدام الطابع الزمني
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Astro_Report_{timestamp}.txt"
            
            # 2. تحديد مسار الحفظ (متوافق مع أندرويد)
            # يحاول الحفظ في مجلد Documents أو Downloads إذا توفرت الصلاحيات
            base_path = "/storage/emulated/0/Download" if os.path.exists("/storage/emulated/0/Download") else os.getcwd()
            full_path = os.path.join(base_path, filename)
            
            # 3. عملية الكتابة
            with open(full_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            # 4. تنبيه النجاح للمستخدم عبر واجهة Flet
            self.page.snack_bar = ft.SnackBar(
                ft.Text(f"✅ تم حفظ التقرير في التحميلات باسم: {filename}", rtl=True),
                bgcolor="#10b981"
            )
            self.page.snack_bar.open = True
            self.page.update()
            
        except PermissionError:
            self.page.snack_bar = ft.SnackBar(ft.Text("❌ خطأ: يرجى منح التطبيق صلاحية الوصول للملفات"))
            self.page.snack_bar.open = True
            self.page.update()
        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"❌ تعذر الحفظ: {str(e)}"))
            self.page.snack_bar.open = True
            self.page.update()
def display_comprehensive_report(self, data):

        # 1. إعداد لوحة الألوان الملحمية
        colors = {
            "gold": "#f59e0b",   # اللون الذهبي للعناوين
            "sky": "#38bdf8",    # الأزرق السماوي للبيانات
            "pink": "#fb7185",   # الوردي للتحليل
            "danger": "#ef4444", # الأحمر للقواطع
            "info": "#10b981",   # الأخضر للشرح
            "border": "#334155"  # لون الإطارات
        }

        # 2. إنشاء قائمة العناصر (Widgets) للتقرير
        report_widgets = []

        # --- القسم الأول: المقدمة ---
        report_widgets.append(
            ft.Text("📜 التقرير الفلكي التحليلي الشامل المطور", 
                    size=22, weight="bold", color=colors["gold"], text_align="center")
        )
        report_widgets.append(ft.Divider(color=colors["gold"], height=20))

        # --- القسم الثاني: تحليل الكواكب والعقد ---
        planets_config = [
            ("الشمس", "sun_sign", "sun_house", "☀️"), ("القمر", "moon_sign", "moon_house", "🌙"),
            ("عطارد", "mercury_sign", "mercury_house", "☿"), ("الزهرة", "venus_sign", "venus_house", "♀️"),
            ("المريخ", "mars_sign", "mars_house", "♂️"), ("المشتري", "jupiter_sign", "jupiter_house", "♃"),
            ("زحل", "saturn_sign", "saturn_house", "♄"), ("أورانوس", "uranus_sign", "uranus_house", "⛢"),
            ("نبتون", "neptune_sign", "neptune_house", "♆"), ("بلوتو", "pluto_sign", "pluto_house", "♇"),
            ("العقدة الشمالية", "nn_sign", "nn_house", "🐲"),
            ("العقدة الجنوبية", "south_node", "sn_house", "🐉")
        ]

        for name, s_key, h_key, icon in planets_config:
            sign = data.get(s_key, "---")
            house = data.get(h_key, "1")
            meaning = PLANET_MEANINGS.get(name, "تأثير كوكبي فريد يشكل ملامح شخصيتك.")
            flavor = SIGN_FLAVORS.get(sign, "بطابع ممتزج.")
            field = HOUSE_FIELDS.get(str(house), "في هذا المجال الحيوي.")

            report_widgets.append(
                ft.Container(
                    content=ft.Column([
                        ft.Text(f"{icon} {name} في برج {sign} | البيت {house}", 
                                color=colors["sky"], weight="bold", size=16),
                        ft.Text(f"◈ التحليل: {meaning} {flavor} ويتجلى أثره في {field}", 
                                color=colors["pink"], size=14),
                    ], spacing=5),
                    padding=12,
                    border=ft.border.all(1, colors["sky"]),
                    border_radius=10,
                    margin=ft.margin.only(bottom=10)
                )
            )

        # --- القسم الثالث: السهام الفلكية ---
        report_widgets.append(ft.Text("🌟 ثانياً: السهام الفلكية السرية", size=19, weight="bold", color=colors["gold"]))
        parts = [
            ("سهم السعادة", data.get('fortune_part'), "نقطة الرزق المادي والقبول الاجتماعي والبهجة النفسية."),
            ("سهم الغيب", data.get('spirit_part'), "نقطة الإرادة الباطنة والنوايا الروحية والذكاء العميق.")
        ]
        for p_title, p_val, p_desc in parts:
            report_widgets.append(
                ft.Container(
                    content=ft.Column([
                        ft.Text(f"📍 {p_title}: {p_val}", color=colors["sky"], weight="bold", size=16),
                        ft.Text(f"💡 الشرح: {p_desc}", color=colors["info"], size=14),
                    ], spacing=5),
                    padding=12, border=ft.border.all(1, colors["info"]), border_radius=10, margin=ft.margin.only(bottom=10)
                )
            )

        # --- القسم الرابع: الفردارية (حاكم العمر) ---
        report_widgets.append(ft.Text("⏳ ثالثاً: حاكم المرحلة الحالية (الفردارية)", size=19, weight="bold", color=colors["gold"]))
        report_widgets.append(
            ft.Container(
                content=ft.Column([
                    ft.Text(f"👑 الحاكم المسيطر: {data.get('firdar_ruler')} - {data.get('firdar_title', '')}", 
                            color=colors["sky"], weight="bold", size=16),
                    ft.Text(f"📝 شرح المرحلة: {data.get('firdar_desc', 'مرحلة هامة لتشكيل المصير.')}", 
                            color=colors["info"], size=14, italic=True),
                ], spacing=5),
                padding=12, border=ft.border.all(1, colors["gold"]), border_radius=10, margin=ft.margin.only(bottom=10)
            )
        )

        # --- القسم الخامس: تقرير النجوم الثابتة ---
        report_widgets.append(ft.Text("🔭 رابعاً: تقرير النجوم الثابتة (أحكام الولادة)", size=19, weight="bold", color=colors["gold"]))
        stars_rep = data.get('stars_report', "")
        star_items = [ft.Text("💡 المنهجية: تحليل مواقع النجوم لحظة ولادتك مقارنة بكواكب ميلادك الشخصية.", size=12, italic=True, color=colors["info"])]
        
        if "لا يوجد" in stars_rep or not stars_rep:
            star_items.append(ft.Text("🔍 النتيجة: لم يتم العثور على اقترانات دقيقة بالنجوم العظمى.", color="#94a3b8"))
        else:
            for line in stars_rep.split('\n'):
                if line.strip():
                    star_items.append(ft.Text(line, color=colors["sky"] if "⭐" in line else "#f8fafc", size=14))

        report_widgets.append(
            ft.Container(content=ft.Column(star_items, spacing=5), padding=12, border=ft.border.all(1, colors["border"]), border_radius=10, margin=ft.margin.only(bottom=10))
        )

        # --- القسم السادس: القواطع والتحذيرات ---
        report_widgets.append(ft.Text("⚠️ خامساً: القواطع والتحذيرات الفلكية", size=19, weight="bold", color=colors["gold"]))
        warnings = data.get('warnings', "لا يوجد قواطع شديدة.")
        deep_interpretations = {
            "تراجع": "يشير التراجع إلى طاقة داخلية؛ مراجعة للذات وتأخر في الشؤون الخارجية.",
            "احتراق": "الكوكب تحت شعاع الشمس؛ دلالاته ضعيفة ظاهرياً وقوية في الخفاء.",
            "درجة حرجة": "وضعية 'المنقلب'؛ تشير إلى تحولات جذرية في دلالة الكوكب.",
            "اقتران النحسين": "ضغط عالٍ يتطلب الحذر الشديد لتجنب القرارات المتهورة."
        }

        warn_list = []
        if "لا يوجد" in warnings:
            warn_list.append(ft.Text("✅ لا يوجد قواطع شديدة في هذه الهيئة.", color=colors["info"], weight="bold"))
        else:
            for line in warnings.split('\n'):
                if line.strip():
                    warn_list.append(ft.Text(f"🛑 {line}", color=colors["danger"], weight="bold", size=14))
                    interp = next((v for k, v in deep_interpretations.items() if k in line), "يتطلب هذا الوضع انتباهاً لإعادة التوازن.")
                    warn_list.append(ft.Text(f"   🔍 التفسير: {interp}", color=colors["info"], size=13))

        report_widgets.append(
            ft.Container(content=ft.Column(warn_list, spacing=5), padding=12, border=ft.border.all(1, colors["danger"]), border_radius=10, margin=ft.margin.only(bottom=10))
        )

        # --- التذييل والخاتمة ---
        report_widgets.append(ft.Divider(color=colors["gold"]))
        report_widgets.append(ft.Text("إعداد وتطوير: حسان الشاعر © 2026", color=colors["gold"], size=14, text_align="center", weight="bold"))

        # 3. عرض التقرير في نافذة AlertDialog (متوافقة مع أندرويد)
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # دالة تجميع النص للنسخ
        def get_full_text():
            return "تقرير فلكي شامل - حسان الشاعر 2026" # يمكن تطويرها لتجمع كل النصوص

        dlg = ft.AlertDialog(
            title=ft.Text("الديوان التحليلي السلطاني", color=colors["gold"], weight="bold", rtl=True),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ALWAYS, spacing=15),
                width=650, height=1200, bgcolor="#0f172a", padding=20
            ),
            actions=[
                ft.ElevatedButton("📋 نسخ النص", on_click=lambda _: self.page.set_clipboard("تم نسخ التقرير!")),
                ft.TextButton("إغلاق", on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

def display_final_report(self, report_widgets, full_text_content):
        """عرض التقرير النهائي مع خيارات الحفظ والمشاركة"""
        import os
        import time

        # تعريف الألوان المستخدمة في التنسيق
        colors = {"header": "#7c2d12", "bg": "#0f172a"}

        # 1. دالة الحفظ الداخلي بصيغة Word
        def save_to_word_internal(e):
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('تقرير الميقاتي برو الموسع', 0)
                doc.add_paragraph(full_text_content)
                
                file_name = f"Meeqat_Report_{int(time.time())}.docx"
                # تحديد المسار (يفضل استخدام Download في أندرويد)
                base_path = "/storage/emulated/0/Download" if os.path.exists("/storage/emulated/0/Download") else os.getcwd()
                path = os.path.join(base_path, file_name)
                
                doc.save(path)
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"✅ تم الحفظ في التحميلات: {file_name}")))
            except Exception as ex:
                self.page.show_snack_bar(ft.SnackBar(ft.Text(f"❌ خطأ في الحفظ: {ex}")))
            self.page.update()

        # 2. إغلاق النافذة
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # 3. بناء النافذة المنبثقة (AlertDialog)
        dlg = ft.AlertDialog(
            modal=False,
            title=ft.Text("الديوان التحليلي السلطاني", weight="bold", color=colors["header"], rtl=True),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ADAPTIVE, spacing=15),
                width=550, 
                height=850, 
                bgcolor=colors["bg"], 
                padding=20,
                border_radius=15
            ),
            actions=[
                ft.ElevatedButton("💾 حفظ Word", on_click=save_to_word_internal, bgcolor="#10b981", color="white"),
                ft.ElevatedButton("📋 نسخ للواتساب", icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(full_text_content)),
                ft.TextButton("إغلاق", on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
def get_sign_name(self, total_deg):
        # الكود الخاص بك صحيح تماماً هنا
        signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                 "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
        return signs[int((total_deg % 360) / 30)]

def my_astronomy_logic(user_input):
    # هنا تضع منطق الحساب الخاص بك مستخدماً الـ 2500 معطى
    # مثال: نتيجة = المدخل + أحد المعطيات
     result = user_input * 0.123  # عدل هذه المعادلة لتناسب برنامجك
     return result
# --- إعدادات الحماية والأمان ---
SECRET_SALT = "MEEQAT_PRO_2026_ULTIMATE" 
TRIAL_DAYS = 3

# 1. دالة جلب بصمة الجهاز (HWID) المتوافقة مع أندرويد
def get_android_hwid(page: ft.Page):
    """إنشاء وتخزين بصمة فريدة للجهاز لمنع النسخ غير المشروع"""
    if not page.client_storage.contains_key("hwid"):
        # دمج معلومات النظام مع معرف فريد لضمان عدم تكرار البصمة
        system_info = f"{platform.processor()}-{platform.node()}-{platform.machine() or 'MobileDevice'}"
        unique_id = str(uuid.uuid4())[:8]
        seed = f"{system_info}-{unique_id}-{SECRET_SALT}"
        
        # تحويل البيانات إلى هاش (Hash) قصير وسهل القراءة
        hwid = hashlib.md5(seed.encode()).hexdigest()[:16].upper()
        page.client_storage.set("hwid", hwid)
    
    return page.client_storage.get("hwid")

# 2. دالة توليد مفتاح التفعيل (للأدمن فقط)
def generate_license_key(hwid):
    """هذه الدالة تستخدمها أنت (المطور) لتوليد المفتاح للمستخدم"""
    raw_key = f"{hwid}-{SECRET_SALT}"
    return hashlib.sha256(raw_key.encode()).hexdigest()[:20].upper()

# 3. دالة التحقق من التفعيل
def check_license(page: ft.Page):
    """التحقق مما إذا كان التطبيق مفعلاً أم لا"""
    user_hwid = get_android_hwid(page)
    stored_key = page.client_storage.get("license_key")
    
    if stored_key == generate_license_key(user_hwid):
        return True
    return False

# --- منطق الحسابات الفلكية الخاص بك ---
def get_sign_name(self, total_deg):
    signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
             "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
    return signs[int((total_deg % 360) / 30)]
def get_remaining_days(page: ft.Page):
    """حساب الأيام المتبقية من الفترة التجريبية"""
    today = datetime.now()
    if not page.client_storage.contains_key("first_run"):
        # تسجيل أول تشغيل للتطبيق
        page.client_storage.set("first_run", today.strftime("%Y-%m-%d"))
        return TRIAL_DAYS
    
    try:
        first_run_str = page.client_storage.get("first_run")
        first_date = datetime.strptime(first_run_str, "%Y-%m-%d")
        elapsed = (today - first_date).days
        remaining = TRIAL_DAYS - elapsed
        return max(0, remaining)
    except:
        return 0

def is_full_version(page: ft.Page):
    """التحقق من أن النسخة مفعلة بالكامل"""
    # 1. التحقق من وجود ملف الترخيص في مجلد التطبيق
    if os.path.exists(KEY_FILE):
        try:
            with open(KEY_FILE, "r") as f:
                saved_key = f.read().strip()
            
            # 2. توليد البصمة ومقارنتها بالمفتاح المخزن
            hwid = get_android_hwid(page) 
            # نستخدم SHA256 لزيادة الأمان في النسخة الكاملة
            correct_key = hashlib.sha256((hwid + SECRET_SALT).encode()).hexdigest()[:12].upper()
            
            return saved_key == correct_key
        except Exception:
            return False
    
    # 3. التحقق أيضاً من التخزين الداخلي كنسخة احتياطية
    if page.client_storage.contains_key("license_key"):
        stored_key = page.client_storage.get("license_key")
        hwid = get_android_hwid(page)
        correct_key = hashlib.sha256((hwid + SECRET_SALT).encode()).hexdigest()[:12].upper()
        return stored_key == correct_key
        
    return False

def get_android_hwid(page: ft.Page):
    """دالة البصمة (يجب أن تكون معرفة مسبقاً)"""
    if not page.client_storage.contains_key("hwid"):
        import uuid, platform
        seed = f"{platform.machine() or 'AndroidDevice'}-{str(uuid.uuid4())[:8]}"
        page.client_storage.set("hwid", hashlib.md5(seed.encode()).hexdigest()[:16].upper())
    return page.client_storage.get("hwid")
def main(page: ft.Page):
    # إعدادات الصفحة الأساسية
    page.title = "تفعيل الميقاتي برو 2026"
    page.rtl = True
    page.theme_mode = ft.ThemeMode.DARK
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    # 1. جلب بيانات الحماية
    hwid = get_android_hwid(page)
    remaining = get_remaining_days(page)
    is_active = is_full_version(page)

    # 2. دالة الانتقال للمحرك الفلكي (عند التفعيل أو خلال الفترة التجريبية)
    def open_astro_logic(e=None):
        page.clean()
        page.vertical_alignment = ft.MainAxisAlignment.START
        page.horizontal_alignment = ft.CrossAxisAlignment.START
        
        # هنا يتم استدعاء واجهة التطبيق الرئيسية (setup_ui)
        # كمثال: سنعرض قائمة المعطيات التي ذكرتها
        astro_list = ft.ListView(expand=True, spacing=10, padding=20)
        for i in range(1, 2501):
            astro_list.controls.append(
                ft.Text(f"💠 المعطى الفلكي {i}: نتيجة الحساب الدقيقة...", size=14)
            )

        page.add(
            ft.AppBar(
                title=ft.Text("الميقاتي الفلكي - المحرك الرئيسي"), 
                bgcolor="#1e293b",
                actions=[ft.IconButton(ft.icons.SETTINGS)]
            ),
            ft.Container(
                content=ft.Column([
                    ft.Text("📊 نتائج الحسابات الموسعة:", size=18, weight="bold", color="#38bdf8"),
                    ft.Divider(),
                    astro_list
                ], expand=True),
                padding=20,
                expand=True
            )
        )
        page.update()

    # 3. دالة معالجة التفعيل
    key_input = ft.TextField(
        label="أدخل كود التفعيل", 
        width=300, 
        text_align=ft.TextAlign.CENTER,
        password=True,
        can_reveal_password=True
    )

    def on_activate(e):
        # حساب المفتاح الصحيح بناءً على بصمة الجهاز والملح السري
        correct_key = hashlib.sha256((hwid + SECRET_SALT).encode()).hexdigest()[:12].upper()
        
        if key_input.value.strip().upper() == correct_key:
            # تخزين المفتاح في الذاكرة الدائمة للجهاز
            page.client_storage.set("license_key", correct_key)
            
            # إشعار النجاح
            page.show_snack_bar(ft.SnackBar(ft.Text("🚀 تم تفعيل النسخة الكاملة بنجاح!"), bgcolor="green"))
            
            # الدخول للتطبيق
            open_astro_logic()
        else:
            key_input.error_text = "كود التفعيل غير صحيح، يرجى التواصل مع المطور"
            page.update()

    # 4. بناء واجهة التفعيل/التجربة
    if is_active:
        # إذا كان مفعلاً، يدخل مباشرة
        open_astro_logic()
    else:
        # واجهة التفعيل
        page.add(
            ft.Icon(ft.icons.LOCK_PERSON, size=80, color="#fbbf24"),
            ft.Text("نظام حماية الميقاتي برو", size=24, weight="bold"),
            ft.Container(
                content=ft.Column([
                    ft.Text(f"🆔 بصمة جهازك (HWID):\n{hwid}", text_align="center", weight="bold", color="#38bdf8"),
                    ft.Text(f"⏳ الفترة التجريبية المتبقية: {remaining} أيام", color="#ef4444"),
                    ft.Divider(),
                    key_input,
                    ft.ElevatedButton("تفعيل النسخة الكاملة", icon=ft.icons.VPI_OUTLINED, on_click=on_activate, bgcolor="#10b981", color="white"),
                    ft.TextButton("استمرار بالنسخة التجريبية", on_click=open_astro_logic if remaining > 0 else None, visible=remaining > 0),
                    ft.Text("للحصول على الكود، أرسل الـ HWID للمطور حسان الشاعر", size=12, italic=True)
                ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=15),
                padding=30,
                border=ft.border.all(1, "#334155"),
                border_radius=15,
                bgcolor="#1e293b"
            )
        )
    # 4. القرار: ماذا نعرض للمستخدم عند فتح التطبيق؟
    if is_active:
        # إذا كان البرنامج مفعلاً، انتقل مباشرة لمحرك الحسابات
        open_astro_logic(None)
    else:
        # إذا كان غير مفعل، اعرض واجهة الحماية والتفعيل
        page.add(
            ft.Column([
                ft.Icon(ft.icons.SECURITY, size=50, color="blue"),
                ft.Text("نظام حماية الميقاتي برو", size=22, weight="bold"),
                ft.Divider(),
                #selectable=True تسمح للمستخدم بنسخ الهاش لإرساله لك
                ft.Text("قم بنسخ البصمة وإرسالها للمطور:", size=14),
                ft.Text(f"{hwid}", selectable=True, color="blue", weight="bold", size=16),
                ft.Text(f"الأيام المتبقية للفترة التجريبية: {remaining}", 
                        color="orange" if remaining > 0 else "red", weight="bold"),
                ft.Divider(),
                key_input,
                ft.ElevatedButton("تفعيل النسخة الكاملة", 
                                 on_click=on_activate, 
                                 bgcolor="green", 
                                 color="white",
                                 width=300),
                # زر الاستمرار يظهر فقط إذا كان هناك أيام متبقية
                ft.TextButton("الاستمرار بالنسخة التجريبية", 
                             on_click=open_astro_logic,
                             visible=remaining > 0),
                ft.Text("جميع الحقوق محفوظة للمطور حسان الشاعر © 2026", size=10, italic=True)
            ], 
            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
            spacing=15)
        )

# ملاحظة: في Flet نستخدم ft.app بدلاً من ft.run
if __name__ == "__main__":
 ft.run(main) # أو ft.app(target=main, view=ft.AppView.WEB_BROWSER)
