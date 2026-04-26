# 1. المكتبات الأساسية
import flet as ft
import os
import platform
import math
from datetime import datetime, timezone

# 2. المكتبات الفلكية المتوافقة مع أندرويد (تأكد من وجودها في requirements.txt)
# سنستخدم skyfield بدلاً من swisseph لأنها لا تسبب أخطاء في البناء
import skyfield.api as sf_api
from skyfield import almanac
import skyfield
from geopy.geocoders import Nominatim
from hijri_converter import Gregorian
# 3. مكتبات النصوص
import arabic_reshaper
from bidi.algorithm import get_display

# 3. مكتبات معالجة اللغة العربية وتنسيق النصوص
import arabic_reshaper
from bidi.algorithm import get_display

try:
    from docx import Document
except ImportError:
    # هذا السطر يحمي التطبيق من الانهيار إذا لم يجد المكتبة
    Document = None 
def fix_ar(text):
    if not text: return ""
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)
# ---------------------------------------------------------
# إعدادات الحماية والتشغيل (المتغيرات العامة)
# ---------------------------------------------------------
SECRET_SALT = "MEEQAT_PRO_2026_ULTIMATE"
TRIAL_DAYS = 3
# ملاحظة للأندرويد: لا تستخدم os.getcwd() كمسار حفظ نهائي، اعتمد على page.client_storage
# بدلاً من المسار القديم، نستخدم هذا المسار المتوافق مع الموبايل والويندوز معاً
if platform.system() == "Android":
    # هذا المسار يعمل دائماً في تطبيقات Flet على أندرويد
    data_dir = os.path.dirname(os.path.abspath(__file__))
    DATA_FILE = os.path.join(data_dir, ".hassan_v15_final.dat")
else:
    DATA_FILE = ".hassan_v15_final.dat"
from skyfield.api import load, Topos
from skyfield.almanac import find_discrete, sunrise_sunset

# 1. تحميل بيانات الكواكب (سيتم تحميلها مرة واحدة فقط)
eph = load('de421.bsp')
planets_map = {
    "الشمس": eph['sun'],
    "القمر": eph['moon'],
    "عطارد": eph['mercury'],
    "الزهرة": eph['venus'],
    "المريخ": eph['mars'],
    "المشتري": eph['jupiter barycenter'],
    "زحل": eph['saturn barycenter'],
    "أورانوس": eph['uranus barycenter'],
    "نبتون": eph['neptune barycenter'],
    "بلوتو": eph['pluto barycenter'],
}

ZODIAC_SIGNS = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

def get_planet_sign(planet_name, year, month, day, hour=12, minute=0):
    ts = load.timescale()
    t = ts.utc(year, month, day, hour, minute)
    
    earth = eph['earth']
    target_planet = planets_map[planet_name]
    
    # حساب الموقع الظاهري للكوكب بالنسبة للأرض
    astrometric = earth.at(t).observe(target_planet)
    lat, lon, distance = astrometric.ecliptic_latlon()
    
    # تحويل خط الطول السماوي إلى برج (كل برج 30 درجة)
    degrees = lon.degrees
    sign_index = int(degrees / 30) % 12
    return ZODIAC_SIGNS[sign_index]

# مثال للاستخدام:
# sign = get_planet_sign("الشمس", 1990, 5, 15)
# print(f"الشمس في برج: {sign}")
# print(f"التفسير: {PLANET_MEANINGS['الشمس']} {SIGN_FLAVORS[sign]}")
PLANET_MEANINGS = {
    "الشمس": "جوهر الذات، الإرادة الحيوية، والقدرة على القيادة والتألق،",
    "القمر": "الاحتياجات العاطفية، العقل الباطن، وكيفية الشعور بالأمان،",
    "عطارد": "الذكاء التحليلي، أسلوب التواصل، وقدرة العقل على معالجة البيانات،",
    "الزهرة": "التناغم الاجتماعي، القيم المادية، وطريقة التعبير عن الحب والجمال،",
    "المريخ": "طاقة المبادرة، الشجاعة، وكيفية توجيه القوة والاندفاع نحو الأهداف،",
    "المشتري": "فرص التوسع، الحظ، الفلسفة العالية وقدرة الشخص على النمو،",
    "زحل": "هيكلية الانضباط، الدروس الكرمية، والمسؤوليات التي تتطلب صبراً،",
    "أورانوس": "شرارة العبقرية، التغيير المفاجئ، والميل نحو التحرر والابتكار،",
    "نبتون": "الخيال الروحي، الحدس العميق، والقدرة على التماهي مع الفنون والروحانيات،",
    "بلوتو": "قوة التحول الجذري، تدمير القديم لبناء الجديد، والقدرة على النهوض،",
    "العقدة الشمالية": "بوصلة الروح والغاية القصوى التي يجب على الإنسان تعلمها في حياته،",
    "العقدة الجنوبية": "الرصيد الفطري والمهارات التي ولدت بها ولكنها قد تعيقك إذا تمسكت بها،"
}
ASPECT_MEANINGS = {
    "0": "اقتران (Conjunction): امتزاج طاقة الكوكبين، بداية مرحلة جديدة وقوة مركزة.",
    "180": "مقابلة (Opposition): توتر، مواجهة مع الآخرين، وحاجة للتوازن بين طرفين.",
    "90": "تربيع (Square): صراع داخلي، عقبات تتطلب مجهوداً جباراً لتجاوزها وتحويلها لنجاح.",
    "120": "تثليث (Trine): تدفق سلس للطاقة، مواهب فطرية، وحظ يأتي بدون عناء.",
    "60": "تسديس (Sextile): فرص جيدة تتطلب استغلالاً واعياً، تواصل مثمر وانسجام."
}

# تأثير البروج (الصبغة أو الأسلوب)
SIGN_FLAVORS = {
    "الحمل": "بأسلوب ناري، مبادر، ويميل للاندفاع والشجاعة المطلقة.",
    "الثور": "بأسلوب ترابي ثابت، يميل للواقعية، الصبر، والبحث عن الاستقرار المادي.",
    "الجوزاء": "بأسلوب هوائي مرن، فضولي، ويميل للتواصل المتعدد والذكاء الاجتماعي.",
    "السرطان": "بأسلوب مائي عاطفي، يميل للحماية، الرعاية، والارتباط العميق بالجذور.",
    "الأسد": "بأسلوب ناري قيادي، يميل للظهور، الثقة بالنفس، وطلب التقدير والتميز.",
    "العذراء": "بأسلوب ترابي تحليلي، يميل للدقة، التنظيم، والبحث عن الكمال والخدمة.",
    "الميزان": "بأسلوب هوائي دبلوماسي، يميل للتوازن، العدل، والبحث عن الجمال والشراكة.",
    "العقرب": "بأسلوب مائي مكثف، يميل للعمق، التحول، والقوة الخفية والإرادة الصلبة.",
    "القوس": "بأسلوب ناري استكشافي، يميل للتفاؤل، الفلسفة، وحب الحرية والمعرفة العالية.",
    "الجدي": "بأسلوب ترابي طموح، يميل للجدية، الانضباط، وبناء المكانة الاجتماعية المرموقة.",
    "الدلو": "بأسلوب هوائي ابتكاري، يميل للتحرر، التفكير الإنساني، والتمرد على القوالب التقليدية.",
    "الحوت": "بأسلوب مائي روحاني، يميل للخيال، التضحية، والتماهي العميق مع المشاعر والكون."
}
# قائمة البروج للترتيب الفلكي (30 درجة لكل برج)
ZODIAC_SIGNS = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

def get_full_interpretation(planet_name, sign_name):
    """
    تدمج معنى الكوكب مع صبغة البرج لتكوين جملة مفيدة
    """
    meaning = PLANET_MEANINGS.get(planet_name, "")
    flavor = SIGN_FLAVORS.get(sign_name, "")
    
    if meaning and flavor:
        return f"{planet_name} في {sign_name}: {meaning} {flavor}"
    return "تعذر إيجاد التفسير."

# مثال لكيفية الاستخدام في التطبيق لاحقاً:
# إذا كانت الشمس في برج الحمل:
# result = get_full_interpretation("الشمس", "الحمل")
# print(fix_ar(result)) 
# دلالات البيوت (مجال التجربة الحياتية)
HOUSE_FIELDS = {
    "1": "في بناء الهوية الشخصية، المظهر الخارجي، وكيفية مواجهة العالم لأول مرة.",
    "2": "في ميدان الموارد المالية، القيم الذاتية، والممتلكات والقدرة على الكسب.",
    "3": "في مجال التواصل اليومي، الإخوة، الدراسات الأولية، والذكاء العملي والأسفار القصيرة.",
    "4": "في أعماق الحياة الخاصة، المنزل، الجذور العائلية، والارتباط بالأرض والوالدين.",
    "5": "في عالم الإبداع، التعبير عن الذات، المواهب الفنية، العواطف، والأطفال.",
    "6": "في مجال العمل الروتيني، الصحة الجسدية، تطوير المهارات، والخدمة والواجبات.",
    "7": "في ميدان الشراكات الرسمية، الزواج، العلاقات المتوازنة، والأعداء الظاهرين.",
    "8": "في مناطق التحول الجذري، المواريث، المال المشترك، الأسرار، والعلوم الباطنية.",
    "9": "في مجال الوعي العالي، السفر البعيد، التعليم الجامعي، الفلسفة، والقانون.",
    "10": "في قمة الطموح المهني، السمعة العامة، السلطة، والنجاح في العالم الخارجي.",
    "11": "في دائرة الأصدقاء، المجموعات، الآمال الكبيرة، والمشاريع الجماعية الإنسانية.",
    "12": "في عالم الخفاء، الروحانيات، العزلة، والتعامل مع الأعداء المخفيين والماضي الكرمي."
}
ASTRO_INTERPRETATIONS = {
    "fortune_part": {
        "الشمس": "وفرة مالية عبر المناصب والظهور الاجتماعي.",
        "القمر": "رزق يأتي من التجارة، السفر، أو الدعم العائلي.",
        "المريخ": "مكاسب تأتي من الجهد البدني، الشجاعة، أو المبادرة.",
        # أضف باقي الكواكب...
    },
    "firdar_ruler": {
        "الشمس": "فترة للبروز، تحقيق الذات، وريادة الأعمال.",
        "الزهرة": "فترة ذهبية للزواج، الفنون، وتحسين العلاقات الاجتماعية.",
        "زحل": "فترة تتطلب الصبر، العمل الجاد، وبناء الأسس المتينة.",
        # أضف باقي الكواكب...
    }
}
def get_planet_house(planet_lon, ascendant_lon):
    """
    تحسب رقم البيت الذي يقع فيه الكوكب بناءً على الطالع (نظام البيوت المتساوية)
    بما أنك تستخدم Flet للأندرويد، هذا النظام هو الأكثر استقراراً برمجياً.
    """
    # فرق الدرجات بين الكوكب والطالع
    diff = (planet_lon - ascendant_lon + 360) % 360
    house_number = int(diff / 30) + 1
    return str(house_number)

def get_house_interpretation(planet_name, house_num):
    """
    تدمج معنى الكوكب مع مجال البيت
    """
    planet_mean = PLANET_MEANINGS.get(planet_name, "")
    house_field = HOUSE_FIELDS.get(house_num, "")
    return f"تتجلى طاقة {planet_name} {house_field}"
PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]

PLANETARY_HOURS_INFO = {
    "الشمس": "ساعة للظهور، طلب المناصب، التعامل مع المسؤولين، والبدايات التي تحتاج شهرة.",
    "الزهرة": "ساعة للجمال، الزواج، التصالح، الفنون، وشراء الأشياء الثمينة.",
    "عطارد": "ساعة للتجارة، الكتابة، الدراسة، المراسلات، والعمليات الحسابية والذكاء.",
    "القمر": "ساعة للتغيرات، التنقل، التعامل مع الجمهور، الأمور العائلية، والعواطف.",
    "زحل": "ساعة للعمل الشاق، البناء، العقارات، التعامل مع كبار السن، وتتطلب الصبر.",
    "المشتري": "ساعة الحظ، التوسع، طلب العلم، الأمور الدينية، والنمو المالي.",
    "المريخ": "ساعة القوة، الرياضة، الجراحة، المواجهة، ويُنصح فيها بالحذر من الاندفاع."
}
# الترتيب الكلداني للكواكب
PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]
# حكام الأيام (يبدأ اليوم الفلكي من الشروق)
DAY_RULERS = {
    "Sunday": "الشمس", "Monday": "القمر", "Tuesday": "المريخ", 
    "Wednesday": "عطارد", "Thursday": "المشتري", "Friday": "الزهرة", "Saturday": "زحل"
}
from skyfield.api import load, wgs84
from skyfield import almanac
import datetime

def get_day_night_hours(lat, lon, date_obj):
    ts = load.timescale()
    eph = load('de421.bsp')
    location = wgs84.latlon(lat, lon)
    
    # تحديد بداية ونهاية اليوم (من شروق اليوم إلى شروق الغد)
    t0 = ts.utc(date_obj.year, date_obj.month, date_obj.day, 0, 0)
    t1 = ts.utc(date_obj.year, date_obj.month, date_obj.day + 1, 12, 0)
    
    # حساب الشروق والغروب
    t, y = almanac.find_discrete(t0, t1, almanac.sunrise_sunset(eph, location))
    
    # t[0] شروق اليوم، t[1] غروب اليوم، t[2] شروق الغد
    sunrise_today = t[0].utc_datetime()
    sunset_today = t[1].utc_datetime()
    sunrise_tomorrow = t[2].utc_datetime()
    
    # حساب طول الساعة النهارية (طول النهار / 12) والليلية (طول الليل / 12)
    day_hour_len = (sunset_today - sunrise_today) / 12
    night_hour_len = (sunrise_tomorrow - sunset_today) / 12
    
    return sunrise_today, day_hour_len, sunset_today, night_hour_len

def get_current_planetary_hour(lat, lon):
    now = datetime.datetime.now(datetime.timezone.utc)
    sunrise, day_len, sunset, night_len = get_day_night_hours(lat, lon, now)
    
    # تحديد حاكم اليوم
    day_name = now.strftime("%A")
    first_hour_ruler = DAY_RULERS[day_name]
    start_index = PLANETARY_ORDER.index(first_hour_ruler)
    
    # هل نحن في النهار أم الليل؟
    if sunrise <= now < sunset:
        elapsed = (now - sunrise) / day_len
    else:
        # إذا كان الوقت قبل شروق اليوم أو بعد غروبه (حساب ساعات الليل)
        elapsed = 12 + ((now - sunset) / night_len)
        
    hour_num = int(elapsed) % 24
    current_ruler = PLANETARY_ORDER[(start_index + hour_num) % 7]
    return current_ruler
PLANETARY_WARNINGS_DB = {
    # --- التراجعات (Retrogrades) ---
    "retro_mercury": "⚠️ تراجع عطارد: فساد في الفكر والرسائل، تعطل الأسفار القريبة، وخلل في البيع والشراء.",
    "retro_venus": "⚠️ تراجع الزهرة: فساد في النكاح، نفور بين المحبين، ومراجعة في أمور الزينة والمال.",
    "retro_mars": "⚠️ تراجع المريخ: حدة في الغضب، تعطل المبادرات، ووهن في العزيمة البدنية.",
    "retro_jupiter": "⚠️ تراجع المشتري: نقص في البركة، ضيق في الرزق الواسع، ومراجعة للقضايا القانونية.",
    "retro_saturn": "⚠️ تراجع زحل: ثقل في المسؤوليات، تأخير في البناء والعقارات، وعودة ديون قديمة.",

    # --- القواطع والاقترانات الصعبة ---
    "malefic_conjunction": "🚨 اقتران النحسين (زحل والمريخ): اجتماع 'القاطع' و'المانع'. شدة وضيق تتطلب صبراً عظيماً.",
    "saturn_vital": "⚠️ حصار زحل لكوكب حيوي: تقييد، حزن، أو تأخير قسري يطال شؤون {planet}.",
    "mars_vital": "⚠️ احتكاك المريخ بكوكب حيوي: تهور، اندفاع، أو خطر التهاب ونزاع يخص {planet}.",
    "mars_opposition": "⚡ مقابلة المريخ: مواجهة مباشرة مع الأعداء، أو تهور يؤدي لنتائج عكسية.",

    # --- الحالات الفنية (احتراق، خلو مسار، درجات حرجة) ---
    "combust": "☄️ احتراق {planet}: الكوكب في 'حيز الشمس'؛ ضعفت قوته واختفت دلالته تحت الشعاع.",
    "void_of_course": "🌑 القمر خالي المسار: سعي بلا طائل. تجنب البدء بأي أمر ذي بال حتى ينتقل القمر.",
    "critical_degree": "📐 درجة حرجة (0 أو 29): الكوكب في حالة 'المنقلب'؛ نهاية مرحلة وبداية تحول مفاجئ.",
    "saturn_return": "🪐 عودة زحل: ساعة المحاسبة الكبرى. اختبار للنضج وإعادة هيكلة أساسات الحياة.",

    # --- النجوم الثابتة (Fixed Stars) ---
    "star_algol": "💀 اقتران رأس الغول (Algol): أشر النجوم؛ خطر فقدان السيطرة، أو قرارات مدمرة للعقل.",
    "star_aldebaran": "🌟 اقتران الدبران: نجاح عظيم محفوف بمخاطر السقوط الأخلاقي أو النزاعات القانونية.",
    "star_regulus": "🦁 اقتران قلب الأسد (Regulus): ارتقاء للمناصب العالية، مع خطر الانتقام من الأنداد.",
    "star_fomalhaut": "🌟 اقتران فم الحوت: نجاح في العلوم الروحية والفنون، مع خطر الوقوع في الأوهام.",
    "star_deneb_kaitos": "🌊 اقتران ذنب قيطس: دلالة على العزلة الإجبارية، التعب الجسدي، أو الحزن المفاجئ.",
    "star_antares": "🌟 اقتران قلب العقرب (Antares): تهور قتالي، نجاح سريع متبوع بنهاية مفاجئة.",

    # --- قواطع كتب السابقين ---
    "besieged": "🔒 كوكب محاصر بين النحسين: الكوكب {planet} في ضيق شديد، دلالته معطلة تماماً.",
    "hayz_violation": "🚫 فساد الهيئة (خروج عن الحيز): الكوكب يعمل عكس طبيعته، مما يسبب نتائج غير متوقعة."
}
def check_planetary_conditions(planet_positions, sun_lon):
    warnings = []
    
    for planet, lon in planet_positions.items():
        # 1. فحص الاحتراق (Combust) - إذا كان بين الكوكب والشمس أقل من 8 درجات
        diff_sun = abs(lon - sun_lon)
        if diff_sun < 8.5 or diff_sun > 351.5:
            msg = PLANETARY_WARNINGS_DB["combust"].format(planet=planet)
            warnings.append(msg)
            
        # 2. فحص الدرجات الحرجة (0 أو 29)
        deg_in_sign = lon % 30
        if deg_in_sign < 1 or deg_in_sign > 29:
            warnings.append(f"{planet}: " + PLANETARY_WARNINGS_DB["critical_degree"])

    return warnings

def is_retrograde(planet_name, t):
    """
    يفحص إذا كان الكوكب في حالة تراجع (بناءً على السرعة الزاوية)
    """
    # في Skyfield نحسب الموقع في لحظتين متقاربتين لمعرفة اتجاه الحركة
    ts = load.timescale()
    t2 = ts.utc(t.utc_datetime() + datetime.timedelta(hours=1))
    
    earth = eph['earth']
    p = planets_map[planet_name]
    
    lon1 = earth.at(t).observe(p).ecliptic_latlon()[1].degrees
    lon2 = earth.at(t2).observe(p).ecliptic_latlon()[1].degrees
    
    # إذا كان خط الطول ينقص، فهو يتراجع
    return lon2 < lon1
FIRDAR_INTERPRETATIONS = {
    "الشمس": {
        "title": "فترة الظهور والسيادة (10 سنوات)",
        "desc": "مرحلة تتركز حول تحقيق الذات، الطموح المهني، والتعامل مع أصحاب السلطة. فترة جيدة للبروز الاجتماعي."
    },
    "الزهرة": {
        "title": "فترة الجمال والمال (8 سنوات)",
        "desc": "مرحلة ذهبية للعلاقات العاطفية، الزواج، الفنون، وتحسن الدخل المادي والرفاهية."
    },
    "عطارد": {
        "title": "فترة العلم والتجارة (13 سنة)",
        "desc": "نشاط ذهني مكثف، كثرة الأسفار، نجاح في الدراسة أو التجارة والكتابة وتوسيع شبكة التواصل."
    },
    "القمر": {
        "title": "فترة المتغيرات والعائلة (9 سنوات)",
        "desc": "تركز على الاستقرار المنزلي، المشاعر، والتغيرات في السكن أو الحالة العائلية."
    },
    "زحل": {
        "title": "فترة الحصاد والمسؤولية (11 سنة)",
        "desc": "مرحلة تتطلب الصبر والعمل الجاد. بناء أسس طويلة الأمد أو مواجهة تحديات تتطلب نضجاً."
    },
    "المشتري": {
        "title": "فترة الوفرة والحظ (12 سنة)",
        "desc": "توسع في الرزق، تفاؤل، نجاحات قانونية أو دينية، وفترة مناسبة للسفر البعيد والنمو."
    },
    "المريخ": {
        "title": "فترة القوة والصراع (7 سنوات)",
        "desc": "طاقة عالية، اندفاع، مبادرات عملية شجاعة، لكنها تتطلب الحذر من الخلافات أو التهور الجسدي."
    },
    "الرأس": {
        "title": "فترة الصعود المفاجئ (3 سنوات)",
        "desc": "تغييرات قدرية، طموحات لا محدودة، وغالباً ما تشهد قفزة نوعية في مكانة الشخص."
    },
    "الذنب": {
        "title": "فترة التطهير والروحانية (سنتان)",
        "desc": "فترة زهد أو مراجعة للماضي، تخلص من علاقات أو أشياء لم تعد تخدمك، والتركيز على الباطن."
    }
}
def get_current_firdar(birth_date, is_day_birth):
    # الترتيب الفرداري (نهاري يبدأ بالشمس، ليلي يبدأ بالقمر)
    if is_day_birth:
        order = ["الشمس", "الزهرة", "عطارد", "القمر", "زحل", "المشتري", "المريخ", "الرأس", "الذنب"]
    else:
        order = ["القمر", "زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "الرأس", "الذنب"]
    
    # مدد السنوات لكل كوكب/نقطة
    durations = {
        "الشمس": 10, "الزهرة": 8, "عطارد": 13, "القمر": 9, 
        "زحل": 11, "المشتري": 12, "المريخ": 7, "الرأس": 3, "الذنب": 2
    }
    
    # حساب العمر الحالي بالأيام وتحويله لسنوات
    age_in_years = (datetime.now() - birth_date).days / 365.25
    
    accumulated_years = 0
    for planet in order:
        accumulated_years += durations[planet]
        if age_in_years <= accumulated_years:
            return planet, FIRDAR_INTERPRETATIONS[planet]
    
    return "انتهت الدورة", None
LOTS_INTERPRETATIONS = {
    "fortune": "معدن الرزق، الحظ الجسدي، والمال الحلال وصحة البدن.",
    "spirit": "معدن القوة الروحية، النوايا، الطموح، وما يخطط له العقل.",
    "love": "سهم الزهرة: يتعلق بالعواطف، الصداقات، والمتعة والحياة الاجتماعية.",
    "necessity": "سهم الشدة: يتعلق بالقيود، المتاعب، وما يُفرض على الإنسان قسراً.",
    "victory": "سهم النصر: يتعلق بالنجاح في المنازعات، القوة، والوصول للأهداف.",
    "marriage": "سهم الزواج: يحدد جودة الشراكة العاطفية واستقرار الحياة الزوجية.",
    "enemies": "سهم الأعداء: يكشف عن الخصوم المخفيين والمكائد التي قد تحاك."
}
def calculate_lots(asc_lon, sun_lon, moon_lon, is_day):
    """
    حساب السهام الأساسية بناءً على الهيئة (نهارية أم ليلية)
    """
    lots = {}
    
    if is_day:
        # حساب السهام للمواليد النهارية
        lots["fortune"] = (asc_lon + moon_lon - sun_lon) % 360
        lots["spirit"] = (asc_lon + sun_lon - moon_lon) % 360
    else:
        # حساب السهام للمواليد الليلية
        lots["fortune"] = (asc_lon + sun_lon - moon_lon) % 360
        lots["spirit"] = (asc_lon + moon_lon - sun_lon) % 360
        
    return lots
# --- دالة حساب البيانات الفلكية المطلوبة للفحص ---
def prepare_astro_data():
    """
    تقوم هذه الدالة بحساب المواقع الحالية للكواكب وتجهيز قاموس 'data'
    ليتمكن كود 'check_comprehensive_warnings' من معالجته.
    """
    ts = load.timescale()
    t = ts.now()
    eph = load('de421.bsp')
    earth = eph['earth']
    
    planets_map = {
        'الشمس': eph['sun'], 'القمر': eph['moon'], 'عطارد': eph['mercury'],
        'الزهرة': eph['venus'], 'المريخ': eph['mars'], 'المشتري': eph['jupiter barycenter'],
        'زحل': eph['saturn barycenter']
    }
    
    positions = {}
    sun_distances = {}
    
    # حساب موقع الشمس أولاً
    sun_lon = earth.at(t).observe(eph['sun']).ecliptic_latlon()[1].degrees
    
    for name, p in planets_map.items():
        lon = earth.at(t).observe(p).ecliptic_latlon()[1].degrees
        positions[name] = lon
        if name != "الشمس":
            # حساب البعد عن الشمس للاحتراق
            dist = abs(lon - sun_lon)
            sun_distances[name] = dist if dist <= 180 else 360 - dist

    # تجهيز قاموس البيانات النهائي
    data = {
        'positions': positions,
        'sun_distances': sun_distances,
        'mars_saturn_dist': abs(positions['المريخ'] - positions['زحل']),
        'is_mercury_retro': False, # يمكن حسابها بدالة السرعة لاحقاً
        'is_moon_voc': False,      # تحتاج فحص الاتصالات الأخيرة
        # أضف هنا باقي الفحوصات (التراجعات والنجوم)
    }
    return data

# --- دالتك التي أرسلتها بعد التعديل لتعمل برمجياً ---
def check_comprehensive_warnings(data):
    alerts = []
    
    # 1. التراجعات
    retro_map = {'عطارد': 'mercury', 'الزهرة': 'venus', 'المريخ': 'mars', 'المشتري': 'jupiter', 'زحل': 'saturn'}
    for ar_p, en_p in retro_map.items():
        if data.get(f'is_{en_p}_retro'):
            # استخدام get لتجنب الخطأ إذا كانت القاعدة غير مكتملة
            alerts.append(f"⚠️ تراجع {ar_p}: " + PLANETARY_WARNINGS_DB.get(f"retro_{en_p}", ""))

    # 2. النحسين (زحل والمريخ)
    if data.get('mars_saturn_dist', 100) < 6:
        alerts.append("🔴 " + PLANETARY_WARNINGS_DB.get("malefic_conjunction", ""))

    # 3. الاحتراق
    sun_distances = data.get('sun_distances', {})
    for p, dist in sun_distances.items():
        if dist < 8.5:
            msg = PLANETARY_WARNINGS_DB.get("combust", "احتراق كوكب {planet}")
            alerts.append(f"🔥 {msg.format(planet=p)}")

    # 4. الدرجات الحرجة (أول وآخر درجة في البرج)
    for p, deg in data.get('positions', {}).items():
        pos_in_sign = deg % 30
        if pos_in_sign < 1.0 or pos_in_sign > 29.0:
            msg = PLANETARY_WARNINGS_DB.get('critical_degree', "درجة حرجة")
            alerts.append(f"📍 {p}: {msg}")

    # النتيجة النهائية
    if not alerts:
        return fix_ar("✅ الهيئة متزنة ولا توجد قواطع سلبية حالياً.")
    
    return fix_ar("\n".join(alerts))
def is_moon_void_of_course(ts, eph, t):
    """
    يفحص إذا كان القمر في حالة خلو مسار.
    يكون القمر خالي المسار إذا لم يصنع أي اتصال (Aspect) رئيسي 
    مع الشمس أو الكواكب حتى ينتقل للبرج التالي.
    """
    earth = eph['earth']
    moon = eph['moon']
    
    # 1. حساب موقع القمر الحالي والبرج الحالي
    moon_lon = earth.at(t).observe(moon).ecliptic_latlon().degrees
    current_sign = int(moon_lon / 30)
    
    # 2. فحص الساعات القادمة (حتى 24 ساعة) لرؤية متى ينتقل للبرج التالي
    # وما إذا كان سيصنع اتصالاً قبل ذلك
    for hours in range(1, 24):
        future_t = ts.utc(t.utc_datetime() + datetime.timedelta(hours=hours))
        future_moon_lon = earth.at(future_t).observe(moon).ecliptic_latlon().degrees
        future_sign = int(future_moon_lon / 30)
        
        # إذا انتقل لبرج جديد، نتوقف عن البحث
        if future_sign != current_sign:
            break
            
        # هنا يتم فحص الاتصالات (التربيع، التثليث، الاقتران...) 
        # إذا وجد اتصال مع كوكب رئيسي قبل الانتقال، فليس خالي المسار.
        # للتبسيط في أندرويد: إذا لم نجد اتصالات قريبة جداً في الدرجات، نعتبره خالي مسار.
        
    # ملاحظة: برمجياً، خلو المسار حالة معقدة، ولكن الأهم هو إضافتها في التنبيهات
    return False # تعاد true إذا انطبقت الشروط
def open_astro_logic(e):
    page.clean()
    
    # 1. جلب البيانات الحقيقية (باستخدام Skyfield المتوافقة مع أندرويد)
    # ملاحظة: prepare_astro_data هي الدالة التي جهزناها سابقاً
    try:
        my_data = prepare_astro_data() 
        result_text = check_comprehensive_warnings(my_data)
    except Exception as ex:
        result_text = f"خطأ في الحساب: {ex}"
    
    # 2. تنسيق النص للعرض الصحيح
    display_text = fx(result_text)

    page.add(
        ft.AppBar(
            title=ft.Text(fx("نتائج فحص الهيئة"), color="white"),
            bgcolor=ft.colors.BLUE_GREY_900
        ),
        ft.Container(
            content=ft.Column([
                ft.Card(
                    content=ft.Container(
                        content=ft.Text(
                            display_text, 
                            rtl=True, 
                            size=16,
                            color=ft.colors.YELLOW_ACCENT,
                            text_align=ft.TextAlign.RIGHT
                        ),
                        padding=15
                    ),
                    color=ft.colors.BLACK54
                ),
                ft.ElevatedButton(
                    text=fx("العودة للرئيسية"), 
                    on_click=lambda _: main_menu(page)
                )
            ], scroll=ft.ScrollMode.ADAPTIVE, spacing=20),
            padding=20,
            expand=True,
            bgcolor=ft.colors.BLUE_GREY_900
        )
    )
    page.update()

# دالة التنسيق المختصرة التي استخدمتها أنت
def fx(text):
    if not text: return ""
    return get_display(arabic_reshaper.reshape(str(text)))
def get_hwid(page: ft.Page):
    # 1. التحقق من وجود البصمة المخزنة
    stored_hwid = page.client_storage.get("user_hwid")
    
    if stored_hwid is None:
        # 2. إنشاء بصمة تعتمد على خصائص الجهاز الفعلية (إن وجدت) + UUID
        # نستخدم uuid4 مع Salt لتعقيد البصمة
        raw_id = f"{platform.processor()}_{platform.node()}_{uuid.uuid4()}"
        
        # 3. تشفير البصمة باستخدام SHA-256 لجعلها ثابتة الطول (مثلاً 16 حرفاً)
        hashed_id = hashlib.sha256((raw_id + SECRET_SALT).encode()).hexdigest().upper()
        short_id = hashed_id[:16] # نأخذ أول 16 حرف لتسهيل التعامل معها
        
        # 4. حفظ البصمة في ذاكرة الجهاز الدائمة
        page.client_storage.set("user_hwid", short_id)
        return short_id
    
    return stored_hwid
def generate_activation_key(hwid):
    """توليد كود التفعيل بناءً على بصمة الجهاز"""
    raw_key = hwid + SECRET_SALT
    return hashlib.sha256(raw_key.encode()).hexdigest()[:12].upper()

def get_trial_info(page: ft.Page):
    """حساب الأيام المتبقية باستخدام ذاكرة التطبيق (أضمن للأندرويد)"""
    today = datetime.now().date()
    
    # التحقق من تاريخ أول تشغيل
    if not page.client_storage.contains_key("first_run_date"):
        page.client_storage.set("first_run_date", today.strftime("%Y-%m-%d"))
        return TRIAL_DAYS
    
    try:
        first_date_str = page.client_storage.get("first_run_date")
        first_date = datetime.strptime(first_date_str, "%Y-%m-%d").date()
        elapsed = (today - first_date).days
        return max(0, TRIAL_DAYS - elapsed)
    except:
        return 0

def show_activation_dialog(page: ft.Page, hwid):
    page.clean()
    code_input = ft.TextField(label=fx("أدخل كود التفعيل"), password=True, can_reveal_password=True)
    
    def verify_click(e):
        correct_key = generate_activation_key(hwid)
        if code_input.value == correct_key:
            page.client_storage.set("is_activated", True)
            page.snack_bar = ft.SnackBar(ft.Text(fx("تم التفعيل! أعد تشغيل التطبيق")))
            page.snack_bar.open = True
            page.update()
            time.sleep(2)
            main_dashboard(page)
        else:
            code_input.error_text = fx("الكود غير صحيح!")
            page.update()

    page.add(
        ft.Text(fx("انتهت الفترة التجريبية"), size=25, color="red"),
        ft.Text(fx(f"معرف الجهاز: {hwid}"), selectable=True),
        code_input,
        ft.ElevatedButton(text=fx("تفعيل الآن"), on_click=verify_click)
    )
def show_splash(page: ft.Page):
    page.clean()
    page.bgcolor = "black"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    
    logo = ft.Text("✧ ASC ✧", size=70, weight="bold", color="#FFD700")
    title = ft.Text("MEEQAT ULTIMATE PRO", size=30, weight="bold", color="#39FF14")
    progress = ft.ProgressBar(width=400, color="#FFD700", value=0)
    
    page.add(logo, title, ft.Text("By Hassan Al-Shaer", color="#BB86FC"), progress)
    
    # حلقة التحميل
    for i in range(1, 11):
        progress.value = i / 10
        page.update()
        time.sleep(0.15)
    
    # الانتقال لنظام التفعيل الذي صممته أنت
    run_auth_system(page) 
def run_auth_system(page: ft.Page):
    hwid = get_hwid(page)
    valid_key = generate_activation_key(hwid)
    days_left = get_trial_info(page)

    page.clean()
    page.bgcolor = ft.colors.BLACK
    
    # حقل النص
    entry = ft.TextField(
        label="أدخل كود التفعيل", 
        width=300, 
        password=True, 
        can_reveal_password=True,
        text_align=ft.TextAlign.CENTER
    )

    # دالة اللصق للأندرويد
def do_paste(e):
    # الطريقة الصحيحة لجلب النص من الحافظة في Flet
    def handle_clipboard(res):
        if res:
            entry.value = res
            page.update()
    page.get_clipboard(handle_clipboard)
    def on_verify(e):
        if entry.value.upper().strip() == valid_key:
            page.client_storage.set("license_key", valid_key)
            open_astro_logic(page) # دالة فتح البرنامج
        else:
            page.snack_bar = ft.SnackBar(ft.Text("كود تفعيل خاطئ"))
            page.snack_bar.open = True
            page.update()

    def on_trial(e):
        if days_left > 0:
            open_astro_logic(page)
        else:
            page.snack_bar = ft.SnackBar(ft.Text("انتهت الفترة التجريبية"))
            page.snack_bar.open = True
            page.update()

    # بناء الواجهة
    page.add(
        ft.Text("تفعيل الميقاتي برو", size=25, weight="bold", color="white"),
        ft.Text(f"ID: {hwid}", color="gray", size=12),
        ft.Row([
            entry, 
            ft.IconButton(icon=ft.icons.PASTE, on_click=do_paste, tooltip="لصق الكود")
        ], alignment=ft.MainAxisAlignment.CENTER),
        ft.ElevatedButton("تفعيل النسخة الأصلية", on_click=on_verify, bgcolor="green", color="white"),
        ft.TextButton(f"الدخول كنسخة تجريبية ({days_left} يوم)", on_click=on_trial),
        ft.Text(f"الأيام المتبقية: {days_left}", color="orange")
    )
    page.update()
def main(page: ft.Page):
    # إعدادات الصفحة الأساسية
    page.window_width = 400
    page.window_height = 800
    # ابدأ بشاشة التحميل
    show_splash(page)

ft.app(target=main)
# ==========================================
# 3. محرك البرنامج الرئيسي (Master Engine)
# ==========================================
class HassanAstroMobile:
    def __init__(self, page: ft.Page, mode):
        self.page = page
        
        # 1. إعدادات العنوان (المتوافقة مع أندرويد)
        version_type = "الكاملة" if mode == "full" else "التجريبية"
        self.page.title = f"Hassan Astro Pro - {version_type}"
        self.page.bgcolor = "#ffffff"
        self.page.rtl = True # دعم العربية
        
        # 2. تعريف المتغيرات (بدل tk.StringVar نستخدم متغيرات عادية)
        self.house_system = "Placidus"
 # 1. أرباب البيوت
        self.domiciles = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", 
                          "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
        
        # 2. أرباب الشرف
        self.exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}

        # 3. أرباب المثلثات (نهار/ليل)
        self.triplicities = {
            "ناري":  {"day": "الشمس", "night": "المشتري"},
            "ترابي": {"day": "الزهرة", "night": "القمر"},
            "هوائي": {"day": "زحل", "night": "عطارد"},
            "مائي":  {"day": "الزهرة", "night": "المريخ"}
        }

        # 4. الحدود المصرية (Egyptian Terms) - الدقة القصوى
        # كل قائمة تحتوي على (الدرجة القصوى للحد، الكوكب صاحب الحد)
        self.egyptian_terms = {
            0:  [(6, "المشتري"), (12, "الزهرة"), (20, "عطارد"), (25, "المريخ"), (30, "زحل")],  # الحمل
            1:  [(8, "الزهرة"), (14, "عطارد"), (22, "المشتري"), (27, "زحل"), (30, "المريخ")],  # الثور
            2:  [(6, "عطارد"), (12, "المشتري"), (17, "الزهرة"), (24, "المريخ"), (30, "زحل")],  # الجوزاء
            3:  [(7, "المريخ"), (13, "الزهرة"), (19, "عطارد"), (26, "المشتري"), (30, "زحل")],  # السرطان
            4:  [(6, "المشتري"), (11, "الزهرة"), (18, "زحل"), (24, "عطارد"), (30, "المريخ")],  # الأسد
            5:  [(7, "عطارد"), (17, "الزهرة"), (21, "المشتري"), (28, "المريخ"), (30, "زحل")],  # العذراء
            6:  [(6, "زحل"), (14, "الزهرة"), (21, "المشتري"), (28, "عطارد"), (30, "المريخ")],  # الميزان
            7:  [(7, "المريخ"), (11, "الزهرة"), (19, "عطارد"), (24, "المشتري"), (30, "زحل")],  # العقرب
            8:  [(12, "المشتري"), (17, "الزهرة"), (21, "عطارد"), (26, "زحل"), (30, "المريخ")],  # القوس
            9:  [(7, "عطارد"), (14, "المشتري"), (22, "الزهرة"), (26, "زحل"), (30, "المريخ")],  # الجدي
            10: [(7, "عطارد"), (13, "الزهرة"), (20, "المشتري"), (25, "المريخ"), (30, "زحل")],  # الدلو
            11: [(12, "الزهرة"), (16, "المشتري"), (19, "عطارد"), (28, "المريخ"), (30, "زحل")]   # الحوت
        }
                    # قاعدة البيانات المكتملة التي أرسلتها
        STARS_INFO = {
                1: ("الشرطين", "2", "الحمل", "أ", "نجمان (النطح والشيرتان)", "منزلة نارية، تدل على الحماس والبدايات والاندفاع القوي."),
                2: ("البطين", "3", "الحمل", "ب", "ثلاثة نجوم خفية", "منزلة ترابية، تمتاز بالثبات والقدرة على التخطيط المالي."),
                3: ("الثريا", "7", "الثور", "ج", "الأخوات السبع (الثريا)", "منزلة هوائية، وهي من أسعد المنازل للمحبة والقبول والشهرة."),
                4: ("الدبران", "5", "الثور", "د", "عين الثور (الدبران)", "منزلة ترابية، تدل على الهيبة والقوة الفكرية العالية."),
                5: ("الهقعة", "3", "الجوزاء", "هـ", "رأس الجبار (الهقعة)", "منزلة هوائية، مناسبة جداً للتعلم وتبادل المعلومات بذكاء."),
                6: ("الهنعة", "2", "الجوزاء", "و", "قدم الجوزاء (الهنعة)", "منزلة مائية، تمتاز بطاقة الألفة وإصلاح العلاقات المتوترة."),
                7: ("الذراع", "2", "السرطان", "ز", "رأس التوأمين (الذراع)", "منزلة مائية، تمنح حماية للممتلكات وتساعد في الظفر."),
                8: ("النثرة", "3", "السرطان", "ح", "سديم المعلف (النثرة)", "منزلة نارية، تدل على الرفعة وعلو الشأن والنجاح الرسمي."),
                9: ("الطرفة", "2", "الأسد", "ط", "عيني الأسد (الطرفة)", "منزلة نارية، تتطلب الحذر وهي قوية للتغييرات الجذرية."),
                10: ("الجبهة", "4", "الأسد", "ي", "جبهة الأسد (الجبهة)", "منزلة نارية، تعيد الهيبة والشرف وتساعد في الاستشفاء."),
                11: ("الزبرة", "2", "الأسد/العذراء", "ك", "ظهر الأسد (الزبرة)", "منزلة نارية، مباركة للتجارة والارتباط العاطفي الناجح."),
                12: ("الصرفة", "1", "العذراء", "ل", "ذنب الأسد (الصرفة)", "منزلة ترابية، تدل على التغيير والانتقال، ومناسبة للسفر."),
                13: ("العواء", "5", "العذراء", "م", "خمسة نجوم (العواء)", "منزلة ترابية، تمتاز بسرعة الاستجابة للأماني وقضاء الحوائج."),
                14: ("السماك", "1", "الميزان", "ن", "السماك الأعزل", "منزلة هوائية، منبع النمو والازدهار في الصفقات المالية."),
                15: ("الغفر", "3", "الميزان", "س", "الغفر", "منزلة هوائية، تدل على الاستبصار والبحث في بواطن الأمور."),
                16: ("الزبانا", "2", "الميزان/العقرب", "ع", "كفتا الميزان (الزبانا)", "منزلة هوائية، مناسبة لضبط الموازين المالية والحقوق."),
                17: ("الإكليل", "3", "العقرب", "ف", "رأس العقرب (الإكليل)", "منزلة مائية، قوية جداً في التأليف بين القلوب وجذب المودة."),
                18: ("القلب", "1", "العقرب", "ص", "قلب العقرب", "منزلة مائية، للتحصين والوقاية والتعامل مع النفس بعمق."),
                19: ("الشولة", "2", "العقرب/القوس", "ق", "ذيل العقرب (الشولة)", "منزلة نارية، تدل على التحرر من القيود والسرعة في الإنجاز."),
                20: ("النعائم", "8", "القوس", "ر", "النعائم", "منزلة نارية، منزلة الوفرة والرزق وتيسير الأمور المتعسرة."),
                21: ("البلدة", "1", "القوس/الجدي", "ش", "البلدة (فراغ سماوي)", "منزلة ترابية، منزلة الاستقرار وبناء الأصول والعمران."),
                22: ("سعد الذابح", "2", "الجدي", "ت", "سعد الذابح", "منزلة ترابية، تدل على الحزم والنجاة من الأخطار المحيطة."),
                23: ("سعد بلع", "2", "الجدي/الدلو", "ث", "سعد بلع", "منزلة هوائية، مناسبة للعلاجات الطبية وتسهيل الصعاب."),
                24: ("سعد السعود", "3", "الدلو", "خ", "سعد السعود", "منبع السعادة والأفراح والنمو في كافة جوانب الحياة."),
                25: ("سعد الأخبية", "4", "الدلو", "ذ", "سعد الأخبية", "تساعد في كشف الأسرار والظهور القوي بعد فترة غياب."),
                26: ("المقدم", "2", "الحوت", "ض", "فرع الدلو المقدم", "منزلة مائية، تدل على الجاه والعز والتواصل الفكري."),
                27: ("المؤخر", "2", "الحوت", "ظ", "فرع الدلو المؤخر", "منزلة مائية، مباركة لزيادة الأرزاق والنجاح في المشاريع."),
                28: ("الرشاء", "1", "الحوت/الحمل", "غ", "الرشاء (بطن الحوت)", "منزلة مائية، منزلة الخواتيم السعيدة وإتمام المهام بنجاح.")
            }
        
    def __init__(self, page: ft.Page, mode):
        self.page = page
        version_type = "الكاملة" if mode == "full" else "التجريبية"
        self.page.title = f"Hassan Astro Pro - {version_type}"
        self.page.rtl = True 
        
        # رموز وأسماء البروج
        self.z_syms = ["♈", "♉", "♊", "♋", "♌", "♍", "♎", "♏", "♐", "♑", "♒", "♓"]
        self.z_names = ["الحمل ♈", "الثور ♉", "الجوزاء ♊", "السرطان ♋", "الأسد ♌", "العذراء ♍", "الميزان ♎", "العقرب ♏", "القوس ♐", "الجدي ♑", "الدلو ♒", "الحوت ♓"]

        # --- حساب التاريخ الهجري (متوافق مع أندرويد) ---
        try:
            from hijri_converter import Gregorian
            now = datetime.now()
            hijri_now = Gregorian(now.year, now.month, now.day).to_hijri()
            self.hijri_day = hijri_now.day
            months_ar = ["محرم", "صفر", "ربيع الأول", "ربيع الثاني", "جمادى الأولى", "جمادى الآخرة", "رجب", "شعبان", "رمضان", "شوال", "ذو القعدة", "ذو الحجة"]
            self.hijri_month_name = months_ar[hijri_now.month - 1]
        except Exception as e:
            import math
            diff = datetime.now() - datetime(2000, 1, 6)
            self.hijri_day = math.floor(((diff.days / 29.5305) - math.floor(diff.days / 29.5305)) * 29.53) + 1
            self.hijri_month_name = "قمري (حسابي)"

        # تعريف الفلسفة
        self.dream_philosophy = {
            "ibn_arabi": "الرؤيا هي اتصال الروح بعالم المثال (الخيال المنفصل).",
            "chaldean": "الأحلام هي انعكاس لنفوذ الكواكب السبعة على النفس البشرية."
        }

        # بناء الواجهة
        self.build_flet_ui()

    # --- دالة استخراج صاحب الحد (خارج الـ __init__) ---
    def get_term_ruler(self, sign_index, degree_in_sign):
        """تستخرج صاحب الحد بناءً على الدرجة وقائمة الحدود المصرية"""
        # egyptian_terms يجب أن تكون معرفة في self
        terms = self.egyptian_terms.get(sign_index, [])
        for max_deg, planet in terms:
            if degree_in_sign <= max_deg:
                return planet
        return "غير معروف"
    def build_flet_ui(self):
        # تنظيف الصفحة وبناء الواجهة الأساسية
        self.page.clean()
        
        # حاوية النتائج (التي ستستقبل الـ 104 أسطر)
        self.report_container = ft.Column(
            controls=[], 
            scroll=ft.ScrollMode.ALWAYS,
            expand=True
        )

        self.page.add(
            ft.AppBar(
                title=ft.Text("الميقاتي الفلكي - Master Edition"), 
                bgcolor=ft.colors.SURFACE_VARIANT,
                actions=[
                    ft.IconButton(ft.icons.COPY, on_click=self.copy_report) # زر لنسخ التقرير كاملاً
                ]
            ),
            ft.Text(f"📅 اليوم القمري: {self.hijri_day}", size=18, weight="bold", rtl=True),
            ft.Divider(),
            self.report_container,
            ft.FloatingActionButton(
                icon=ft.icons.PLAY_ARROW, 
                text="بدء الحسابات الملحمية",
                on_click=lambda _: self.start_calculations()
            )
        )

    # هذه الدالة يجب أن تكون خارج build_flet_ui ولكن داخل الـ Class
    def log_to_flet(self, message, color="black"):
        """البديل الرسمي لـ txt.insert - متوافق مع أندرويد"""
        self.report_container.controls.append(
            ft.Text(value=message, color=color, size=16, rtl=True, selectable=True)
        )
        self.page.update()
# الدالة هنا (داخل الكلاس)
    def copy_report(self, e):
        """دالة لنسخ كافة النتائج المسجلة في الحاوية"""
        full_text = "\n".join([c.value for c in self.report_container.controls if isinstance(c, ft.Text)])
        self.page.set_clipboard(full_text)
        
        self.page.snack_bar = ft.SnackBar(
            content=ft.Text(fx("تم نسخ التقرير الفلكي بنجاح!")),
        )
        self.page.snack_bar.open = True
        self.page.update()
    def get_egyptian_bounds(self, sign_idx, deg): 
        """استخراج صاحب الحد بناءً على البرج والدرجة"""
        # نعتمد على القاموس المعرف مسبقاً في self.egyptian_terms لتوفير الذاكرة
        terms = self.egyptian_terms 
        
        # البحث عن الكوكب الذي يقع ضمن نطاقه الدرجة المعطاة
        for limit, planet in terms.get(sign_idx, []):
            if deg <= limit: 
                return planet
                
        # في حال عدم مطابقة الدرجات (حالة نادرة)، يعود بزحل كونه صاحب الدرجات الأخيرة غالباً
        return "زحل"
    def calculate_real_almuten(self, target_deg, t, is_day):
        """حساب المبتز (المستولي) بناءً على الحظوظ الخمسة"""
        # (استخدمنا القواميس المعرفة سابقاً في self لتوفير الذاكرة)
        domiciles = self.domiciles
        exaltations = self.exaltations
        
        sign_idx = int(target_deg / 30) % 12
        deg_in_sign = target_deg % 30
        scores = {k: 0 for k in ["الشمس", "القمر", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل"]}
        
        # 1. رب البيت (5 نقاط)
        scores[domiciles[sign_idx]] += 5
        
        # 2. رب الشرف (4 نقاط)
        if sign_idx in exaltations: 
            scores[exaltations[sign_idx]] += 4
            
        # 3. رب المثلثة (3 نقاط)
        elements = ["Fire", "Earth", "Air", "Water"]
        element_name = elements[sign_idx % 4]
        # triplicities_map يجب أن يكون معرفاً في الكلاس
        tri_ruler = self.triplicities[element_name]["day" if is_day else "night"]
        scores[tri_ruler] += 3
        
        # 4. رب الحد (نقطتان)
        scores[self.get_egyptian_bounds(sign_idx, deg_in_sign)] += 2
        
        # 5. رب الوجه (نقطة واحدة)
        faces = ["المريخ", "الشمس", "الزهرة", "عطارد", "القمر", "زحل", "المشتري"]
        scores[faces[int(target_deg / 10) % 7]] += 1

        # تحديد الكوكب الفائز (المبتز)
        winner_name = max(scores, key=scores.get)
        
        # --- التعديل الجوهري للأندرويد: استخدام Skyfield بدلاً من Swe ---
        eph = load('de421.bsp')
        earth = eph['earth']
        
        # خريطة الكواكب لـ Skyfield
        p_sf_map = {
            "الشمس": eph['sun'], "القمر": eph['moon'], "عطارد": eph['mercury'], 
            "الزهرة": eph['venus'], "المريخ": eph['mars'], 
            "المشتري": eph['jupiter barycenter'], "زحل": eph['saturn barycenter']
        }
        
        # حساب موقع الكوكب المبتز في اللحظة المطلوبة t
        astrometric = earth.at(t).observe(p_sf_map[winner_name])
        lat, lon, distance = astrometric.ecliptic_latlon()
        
        return winner_name, lon.degrees # إرجاع اسم الكوكب ودرجته
    def draw_astro_wheel(self, t, title):
        """حساب المواقع الفلكية وعرضها في واجهة Flet (نسخة الأندرويد)"""
        # 1. تجهيز المحرك (Skyfield)
        ts = load.timescale()
        eph = load('de421.bsp')
        earth = eph['earth']
        
        # 2. حساب الطالع (ASC) ووسط السماء (MC)
        # ملاحظة: حساب الأوتاد بدقة يتطلب خط الطول والعرض والوقت المحلي
        # للتبسيط هنا نستخدم حسابات تقريبية أو نعتمد على نظام Whole Sign المتوافق مع أندرويد
        asc = 0 # سيتم حسابه بناءً على الموقع والزمن الفلكي
        mc = (asc + 270) % 360

        # 3. تعريف بيانات الكواكب
        p_data = {
            'sun': ("الشمس", "☉", "#f59e0b"), 'moon': ("القمر", "☽", "#3b82f6"), 
            'mercury': ("عطارد", "☿", "#8b5cf6"), 'venus': ("الزهرة", "♀", "#ec4899"), 
            'mars': ("المريخ", "♂", "#dc2626"), 'jupiter barycenter': ("المشتري", "♃", "#d97706"), 
            'saturn barycenter': ("زحل", "♄", "#475569")
        }

        # 4. عرض النتائج في Flet
        self.log_to_flet(f"✨ {title}", color="blue")
        self.log_to_flet("📐 الأوتاد الرئيسية:", color="red")
        
        # عرض الطالع
        self.log_to_flet(f"الطالع (ASC): {self.z_names[int(asc/30)]} {int(asc%30)}°", color="#ef4444")

        # 5. حساب وعرض الكواكب (باستخدام Skyfield)
        self.log_to_flet("\n🪐 وضعية الكواكب:", color="blue")
        
        planet_positions = {}
        for pid, (n, s, c) in p_data.items():
            astrometric = earth.at(t).observe(eph[pid])
            _, lon, _ = astrometric.ecliptic_latlon()
            pos = lon.degrees
            planet_positions[n] = pos
            
            sign_idx = int(pos / 30)
            deg = int(pos % 30)
            mnt = int((pos - int(pos)) * 60)
            
            # طباعة النتائج في واجهة Flet
            self.log_to_flet(f"{s} {n}: {self.z_names[sign_idx]} {deg}°{mnt:02d}'", color=c)

        # 6. حساب السهام السرية (السعادة والغيب)
        sun_pos = planet_positions["الشمس"]
        moon_pos = planet_positions["القمر"]
        is_day = True # تحدد بناءً على موقع الشمس فوق أو تحت الأفق
        
        fortuna = (asc + moon_pos - sun_pos) % 360 if is_day else (asc + sun_pos - moon_pos) % 360
        
        self.log_to_flet("\n💎 السهام السرية:", color="orange")
        self.log_to_flet(f"⊗ سهم السعادة: {self.z_names[int(fortuna/30)]} {int(fortuna%30)}°", color="#f97316")

        self.page.update()
    def draw_now(self):
        # الحصول على الوقت الحالي (Skyfield يستخدم كائن الوقت مباشرة)
        ts = load.timescale()
        t = ts.now() # هذا البديل لـ jd في Skyfield
        
        # استدعاء دالة الرسم التي قمنا بتحويلها سابقاً لـ Skyfield
        self.draw_astro_wheel(t, "خريطة اللحظة (Time Now)")

    def draw_2026(self):
        # إعداد تاريخ دخول سنة 2026
        ts = load.timescale()
        # 20 مارس 2026 الساعة 09:24 UTC
        t_2026 = ts.utc(2026, 3, 20, 9, 24)
        
        # 1. رسم الخريطة الفلكية لعام 2026
        self.draw_astro_wheel(t_2026, "خريطة طالع سنة 2026")
        
        # 2. التحليل العالمي
        if hasattr(self, 'show_mundane_analysis'):
            # تأكد أن هذه الدالة أيضاً تم تحويلها لتستقبل t بدلاً من jd
            self.show_mundane_analysis(t_2026)
    def setup_ui(self):
        # إعدادات عامة للصفحة
        self.page.padding = 10
        self.page.spacing = 10
        self.page.rtl = True
        self.page.theme_mode = ft.ThemeMode.LIGHT
        self.page.scroll = ft.ScrollMode.ADAPTIVE

        # --- دالة مساعدة للعناوين الفرعية ---
        def section_title(text, color="#f1f5f9"):
            return ft.Container(
                content=ft.Text(text, weight="bold", size=14, color="#1e293b"),
                bgcolor=color,
                padding=8,
                alignment=ft.alignment.center,
                width=float("inf"),
                border_radius=8
            )

        # --- 1. القائمة اليمين (جدول الكواكب) ---
        self.tree = ft.DataTable(
            columns=[
                ft.DataColumn(ft.Text("الجرم")),
                ft.DataColumn(ft.Text("البرج")),
                ft.DataColumn(ft.Text("الدرجة")),
                ft.DataColumn(ft.Text("بيت")),
            ],
            rows=[],
            column_spacing=10,
            heading_row_height=40,
            data_row_min_height=35,
        )

        left_sidebar = ft.Container(
            content=ft.Column([
                section_title("📊 بيانات الهيئة", color="#e2e8f0"),
                ft.Container(content=self.tree, overflow=ft.PropertyChangeProvider)
            ], scroll=ft.ScrollMode.ADAPTIVE),
            bgcolor="#ffffff",
            border=ft.border.all(1, "#cbd5e1"),
            border_radius=10,
            padding=10,
        )

        # --- 2. إعدادات حقول الإدخال (الموقع والوقت) ---
        self.city_entry = ft.TextField(label="المدينة", value="Tartus", text_align="center", dense=True)
        self.house_menu = ft.Dropdown(
            label="نظام البيوت",
            options=[ft.dropdown.Option("Equal"), ft.dropdown.Option("Placidus"), ft.dropdown.Option("Whole Sign")],
            value="Equal",
            dense=True
        )
        self.coords_lbl = ft.Text("عرض: 00° 00' | طول: 00° 00'", size=12, text_align="center", color="#64748b")

        # حقول الوقت مع كيبورد الأرقام للأندرويد
        self.day_ent = ft.TextField(label="يوم", width=60, text_align="center", keyboard_type=ft.KeyboardType.NUMBER)
        self.month_ent = ft.TextField(label="شهر", width=60, text_align="center", keyboard_type=ft.KeyboardType.NUMBER)
        self.year_ent = ft.TextField(label="سنة", width=80, text_align="center", keyboard_type=ft.KeyboardType.NUMBER)
        self.hour_ent = ft.TextField(label="س", width=60, text_align="center", keyboard_type=ft.KeyboardType.NUMBER)
        self.min_ent = ft.TextField(label="د", width=60, text_align="center", keyboard_type=ft.KeyboardType.NUMBER)

        # أزرار الروحانيات
        spiritual_btns = ft.Column([
            ft.ElevatedButton("🌙 منازل القمر", bgcolor="#8b5cf6", color="white", on_click=self.show_moon_mansion, width=250),
            ft.ElevatedButton("💭 تفسير الأحلام", bgcolor="#4b5563", color="white", on_click=self.show_dream_interpreter, width=250),
            ft.ElevatedButton("⏳ ساعات الكواكب", bgcolor="#0f172a", color="white", on_click=self.show_planetary_hours, width=250),
        ], horizontal_alignment="center", spacing=5)

        # --- 3. القائمة اليسرى (لوحة التحكم) ---
        right_sidebar_content = ft.Column([
            ft.Text("الميقاتي الفلكي - حسان الشاعر", color="#a11616", size=18, weight="bold", text_align="center"),
            section_title("📍 الموقع والبيوت"),
            self.city_entry,
            ft.ElevatedButton("🌍 جلب الإحداثيات", bgcolor="#10b981", color="white", on_click=self.search_city_logic, width=250),
            self.coords_lbl,
            self.house_menu,
            
            section_title("🔮 المحرك الرئيسي"),
            ft.ElevatedButton("🕒 الوقت الآن", bgcolor="#2563eb", color="white", on_click=lambda _: self.draw_now(), width=250),
            ft.ElevatedButton("🌍 طالع سنة العالم", bgcolor="#dc2626", color="white", on_click=lambda _: self.draw_2026(), width=250),
            
            section_title("👶 بيانات الميلاد"),
            ft.Row([self.day_ent, self.month_ent, self.year_ent], alignment="center", spacing=5),
            ft.Row([self.hour_ent, self.min_ent], alignment="center", spacing=5),
            ft.ElevatedButton("👶 تحليل الولادة", bgcolor="#059669", color="white", on_click=self.analyze_birth_chart, width=250),
            
            section_title("🌙 الروحانيات والأدوات"),
            spiritual_btns,
        ], scroll=ft.ScrollMode.ALWAYS, horizontal_alignment="center", spacing=10)

        right_sidebar = ft.Container(
            content=right_sidebar_content,
            bgcolor="#f8fafc",
            border=ft.border.all(1, "#cbd5e1"),
            border_radius=10,
            padding=15
        )

        # --- 4. منطقة النتائج المركزية ---
        self.report_container = ft.Column(scroll=ft.ScrollMode.ALWAYS, expand=True)
        self.chart_container = ft.Container(
            content=self.report_container,
            bgcolor="#ffffff",
            border=ft.border.all(1, "#cbd5e1"),
            border_radius=10,
            padding=15,
            min_height=300
        )

        # --- 5. التجميع النهائي بنظام ResponsiveRow للجوال ---
        main_layout = ft.ResponsiveRow([
            ft.Column([right_sidebar], col={"sm": 12, "md": 4}),
            ft.Column([self.chart_container], col={"sm": 12, "md": 8}),
            ft.Column([left_sidebar], col={"sm": 12, "md": 12}),
        ], spacing=10)

        self.page.add(
            ft.AppBar(title=ft.Text("Hassan Astro Ultimate"), bgcolor=ft.colors.BLUE_GREY_900, color="white"),
            main_layout
        )
        self.page.update()
def auto_skip(self, e, next_control, limit):
    if len(e.control.value) >= limit:
        next_control.focus()
        self.page.update()
def show_planetary_hours(self):
    try:
        # 1. جلب البيانات الجغرافية
        lat = float(getattr(self, 'lat', 34.88))
        lon = float(getattr(self, 'lon', 35.88))
        
        # 2. حساب الشروق والغروب باستخدام المحرك الذي جهزناه
        ts = load.timescale()
        eph = load('de421.bsp')
        from skyfield import almanac
        from skyfield.api import wgs84
        
        location = wgs84.latlon(lat, lon)
        now = datetime.now(timezone.utc)
        t0 = ts.utc(now.year, now.month, now.day, 0)
        t1 = ts.utc(now.year, now.month, now.day, 24)
        
        # حساب لحظة الشروق الفعلي لليوم
        t, y = almanac.find_discrete(t0, t1, almanac.sunrise_sunset(eph, location))
        sunrise_time = t[0].utc_datetime() if len(t) > 0 else now
        
        # 3. تحديد حاكم اليوم (يبدأ من الشروق)
        day_name = now.strftime("%A")
        # DAY_RULERS و PLANETARY_ORDER معرفة في القواميس السابقة
        first_hour_ruler = DAY_RULERS.get(day_name, "الشمس")
        start_idx = PLANETARY_ORDER.index(first_hour_ruler)
        
        # 4. حساب الساعة الفلكية (كم ساعة مرت منذ الشروق)
        elapsed_hours = int((now - sunrise_time).total_seconds() / 3600)
        ruler = PLANETARY_ORDER[(start_idx + elapsed_hours) % 7]

        # 5. عرض النتيجة في AlertDialog (كما في كودك)
        analysis_text = f"🕒 الوقت الحالي: {now.strftime('%H:%M')}\n"
        analysis_text += f"🌅 شروق الشمس: {sunrise_time.strftime('%H:%M')}\n\n"
        analysis_text += f"📜 التفسير:\n{PLANETARY_HOURS_INFO.get(ruler)}"

        dlg = ft.AlertDialog(
            title=ft.Text(f"⏳ حاكم الساعة: {ruler}", rtl=True),
            content=ft.Text(analysis_text, rtl=True),
            actions=[ft.TextButton("إغلاق", on_click=lambda _: self.close_dialog(dlg))]
        )
        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

    except Exception as e:
        self.log_to_flet(f"خطأ في حساب الساعات: {e}", color="red")
    def show_animodar_correction(self):
        try:
            # 1. جلب البيانات من واجهة Flet بأمان
            try:
                d = int(self.day_ent.value) if self.day_ent.value else 1
                m = int(self.month_ent.value) if self.month_ent.value else 1
                y = int(self.year_ent.value) if self.year_ent.value else 2000
                h = int(self.hour_ent.value) if self.hour_ent.value else 12
                mn = int(self.min_ent.value) if self.min_ent.value else 0
                offset = 3.0 # التوقيت الافتراضي (يمكنك ربطه بحقل لاحقاً)
            except ValueError:
                raise Exception("يرجى إدخال أرقام صحيحة في خانات الميلاد")

            # 2. المحرك الفلكي (Skyfield)
            ts = load.timescale()
            t = ts.utc(y, m, d, h - offset, mn)
            eph = load('de421.bsp')
            earth = eph['earth']

            # 3. حساب الكواكب (الشمس والقمر)
            sun_lon = earth.at(t).observe(eph['sun']).ecliptic_latlon().degrees
            moon_lon = earth.at(t).observe(eph['moon']).ecliptic_latlon().degrees
            
            # حساب الطالع (ASC) - باستخدام نظام Equal للسرعة في أندرويد
            # الطالع يحتاج لزمن فلكي، هنا سنستخدم قيمة تقريبية أو نربطه بدالة الأوتاد
            abs_asc = 0 # سيتم استبداله بدالة حساب ASC الحقيقية لديك

            # 4. قاعدة الأنمودار: اختيار الكوكب الهدف (المبتز على السيجي)
            is_day = not (180 <= (sun_lon - abs_asc) % 360 <= 360)
            target_deg = sun_lon if is_day else moon_lon
            
            # تحديد صاحب البرج (المبتز المبسط)
            sign_idx = int(target_deg / 30) % 12
            winner_name = self.domiciles[sign_idx] # استخدام القائمة التي عرفتها في الكلاس
            
            # حساب درجة الكوكب المبتز حالياً
            p_sf_map = {'الشمس': 'sun', 'القمر': 'moon', 'عطارد': 'mercury', 'الزهرة': 'venus', 'المريخ': 'mars', 'المشتري': 'jupiter barycenter', 'زحل': 'saturn barycenter'}
            winner_p = eph[p_sf_map[winner_name]]
            winner_lon = earth.at(t).observe(winner_p).ecliptic_latlon().degrees
            
            # 5. استخراج الدرجة والبرج
            winner_deg_in_sign = int(winner_lon % 30)
            winner_sign_name = self.z_names[int(winner_lon / 30) % 12]

            # 6. النتيجة النهائية
            result_msg = f"✅ وفقاً للأنمودار: يجب أن تكون درجة الطالع قريبة من {winner_deg_in_sign}° في برج {winner_sign_name} (درجة كوكب {winner_name})."
            self.log_to_flet(result_msg, color="purple")

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"⚠️ خطأ في التصحيح: {str(e)}", rtl=True))
            self.page.snack_bar.open = True
            self.page.update()
    def display_detailed_animodar(self, y, m, d, h, mn, s_sign, s_deg, a_sign, a_curr, a_corr, planet_name):
        # 1. بناء نص التقرير (تم تنظيفه ليدعم العربية بشكل أفضل)
        full_text = (
            f"تقرير تصحيح وقت الولادة (النمودار التقليدي المتقدم)\n"
            f"تاريخ الميلاد: {y}/{m}/{d} | الوقت: {h:02d}:{mn:02d}\n\n"
            f"۞ البعد الفلسفي للنمودار:\n"
            f"يرتبط الطالع بدرجة الكوكب المبتز المهيمن فلكياً.\n\n"
            f"النتائج النهائية:\n"
            f"• الكوكب المبتز: {planet_name}\n"
            f"• درجة الكوكب: {s_deg}° في برج {s_sign}\n"
            f"• الطالع الحالي: {a_curr}° في برج {a_sign}\n"
            f"• الدرجة المصححة: [{a_corr}°] تماماً.\n\n"
            f"المطور: حسان الشاعر"
        )

        # 2. دالة حفظ ملف Word (معدلة للأندرويد)
        def save_to_word(e):
            try:
                from docx import Document
                import os
                
                doc = Document()
                # ملاحظة: النصوص العربية في docx تحتاج أحياناً لتنسيق RTL
                p = doc.add_heading('تقرير تصحيح الطالع', 0)
                doc.add_paragraph(full_text)
                
                file_name = f"Animodar_{y}_{m}_{d}.docx"
                
                # تحديد المسار المتوافق مع أندرويد وويندوز
                if platform.system() == "Android":
                    # الحفظ في مجلد ملفات التطبيق العام لضمان الصلاحيات
                    save_dir = "/storage/emulated/0/Download"
                    if not os.path.exists(save_dir):
                        save_dir = os.path.expanduser("~")
                else:
                    save_dir = os.path.expanduser("~/Downloads")
                
                save_path = os.path.join(save_dir, file_name)
                doc.save(save_path)
                
                # استخدام SnackBar الحديث في Flet
                self.page.snack_bar = ft.SnackBar(ft.Text(f"✅ تم الحفظ في: {save_path}"), bgcolor="green")
                self.page.snack_bar.open = True
            except Exception as ex:
                self.page.snack_bar = ft.SnackBar(ft.Text(f"❌ خطأ: {str(ex)}"), bgcolor="red")
                self.page.snack_bar.open = True
            self.page.update()

        # 3. إغلاق النافذة
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # 4. واجهة التقرير
        dlg = ft.AlertDialog(
            modal=True,
            title=ft.Text(fx("نتائج التصحيح الفلكي")),
            content=ft.Container(
                content=ft.Column([
                    ft.Text(fx(full_text), size=14, color=ft.colors.ON_SURFACE_VARIANT, rtl=True),
                ], scroll=ft.ScrollMode.ALWAYS, tight=True),
                width=400, height=500
            ),
            actions=[
                ft.ElevatedButton(fx("حفظ Word"), icon=ft.icons.SAVE, on_click=save_to_word),
                ft.TextButton(fx("إغلاق"), on_click=close_dlg),
            ],
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
    def get_accurate_planets(self, year, month, day, hour_decimal, lat, lon):
        try:
            # 1. إعداد المحرك الفلكي (Skyfield)
            ts = load.timescale()
            # تصحيح التوقيت العالمي (UTC) -3 ساعات
            t = ts.utc(year, month, day, hour_decimal - 3)
            eph = load('de421.bsp')
            earth = eph['earth']
            
            p_sf_map = {
                "الشمس": eph['sun'], "القمر": eph['moon'], "عطارد": eph['mercury'],
                "الزهرة": eph['venus'], "المريخ": eph['mars'], 
                "المشتري": eph['jupiter barycenter'], "زحل": eph['saturn barycenter']
            }
            
            results = {"positions_raw": {}}
            signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                     "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

            # 2. حساب مواقع الكواكب
            for name, p_obj in p_sf_map.items():
                astrometric = earth.at(t).observe(p_obj)
                _, lon_pos_obj, _ = astrometric.ecliptic_latlon()
                lon_pos = lon_pos_obj.degrees
                
                results[name] = {
                    "sign": signs[int(lon_pos / 30) % 12], 
                    "full_deg": lon_pos,
                    "deg_in_sign": lon_pos % 30
                }
                results["positions_raw"][name] = lon_pos

            # 3. حساب الطالع والبيوت (نظام Equal لضمان السرعة والأمان في أندرويد)
            # ملاحظة: Placidus يتطلب خوارزميات معقدة جداً غير متوفرة بـ Skyfield مباشرة
            # لذا سنعتمد نظام البيوت المتساوية كبديل احترافي ومستقر
            asc = 0 # هنا تضع دالة حساب الطالع بناءً على الزمن الفلكي
            results["الطالع"] = asc
            results["وسط السماء"] = (asc + 270) % 360
            results["البيوت"] = [(asc + (i * 30)) % 360 for i in range(12)]
            
            return results

        except Exception as e:
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(
                    ft.Text(fx(f"خطأ في الحساب: {e}"), bgcolor="red")
                )
                self.page.snack_bar.open = True
                self.page.update()
            return None
def calculate_almuten(self, planets_data):
        """
        حساب الكوكب المبتز (المسيطر) بناءً على نقاط القوة في الأبراج.
        """
        # 1. تعريف أرباب البيوت (بالترتيب من الحمل إلى الحوت)
        domiciles = [
            "المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", 
            "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"
        ]
        
        # 2. تعريف أرباب الشرف (البرج: الكوكب)
        # الحمل: شمس، الثور: قمر، السرطان: مشتري، العذراء: عطارد، الميزان: زحل، الجدي: مريخ، الحوت: زهرة
        exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}

        # تهيئة سجل النقاط للكواكب السبعة فقط
        scores = {k: 0 for k in ["الشمس", "القمر", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل"]}
        
        try:
            if not planets_data:
                return "لا توجد بيانات", 0

            for p_name in scores.keys():
                p_info = planets_data.get(p_name)
                if not p_info: 
                    continue
                
                pos = p_info["full_deg"]
                sign_idx = int(pos / 30) % 12
                
                # أ- قوة بيت الكوكب (5 نقاط)
                # إذا كان الكوكب يتواجد في برج هو ربه الأصلي
                if domiciles[sign_idx] == p_name:
                    scores[p_name] += 5
                    
                # ب- قوة الشرف (4 نقاط)
                # إذا كان الكوكب يتواجد في برج شرفه
                if exaltations.get(sign_idx) == p_name:
                    scores[p_name] += 4
                
                # إمكانية التطوير المستقبلي:
                # المثلثات (+3)، الحدود (+2)، الوجوه (+1)

            # 3. تحديد الكوكب صاحب أعلى نقاط
            winner = max(scores, key=scores.get)
            
            # إذا تعادلت النقاط عند الصفر
            if scores[winner] == 0:
                return "غير محدد", 0
                
            return winner, scores[winner]
            
        except Exception as e:
            # تنبيه المستخدم في حال حدوث خطأ أثناء المعالجة
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(ft.Text(fx(f"خطأ في حساب المبتز: {e}")))
                self.page.snack_bar.open = True
                self.page.update()
def analyze_birth_chart(self):
        try:
            # 1. جلب المدخلات بأمان من واجهة Flet
            def g_v(attr, default):
                element = getattr(self, attr, None)
                return element.value if element and element.value else default

            try:
                y, m, d = int(g_v('year_ent', 1990)), int(g_v('month_ent', 1)), int(g_v('day_ent', 1))
                h, mi = int(g_v('hour_ent', 12)), int(g_v('min_ent', 0))
                offset = float(g_v('gmt_ent', 3.0))
            except ValueError:
                raise Exception("يرجى إدخال أرقام صحيحة في تاريخ ووقت الميلاد")

            lat = float(getattr(self, 'lat', 34.88))
            lon = float(getattr(self, 'lon', 35.88))
            
            # --- المحرك الفلكي (Skyfield) البديل لـ SwissEph ---
            ts = load.timescale()
            t = ts.utc(y, m, d, h - offset, mi)
            eph = load('de421.bsp')
            earth = eph['earth']
            
            # 2. حساب الطالع (ASC) والبيوت بنظام Equal لضمان السرعة
            asc_raw = 0 # هنا نضع قيمة الطالع المحسوبة (ASC)
            zodiacs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                       "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            
            # حساب الكواكب
            p_sf_map = [
                ("الشمس", "sun"), ("القمر", "moon"), ("عطارد", "mercury"), 
                ("الزهرة", "venus"), ("المريخ", "mars"), ("المشتري", "jupiter barycenter"), 
                ("زحل", "saturn barycenter")
            ]

            results = {}
            raw_d = {}
            raw_speed = {}

            for ar, en_id in p_sf_map:
                astrometric = earth.at(t).observe(eph[en_id])
                _, lon_pos, _ = astrometric.ecliptic_latlon()
                deg = lon_pos.degrees
                raw_d[en_id] = deg
                
                # حساب السرعة للتراجع (Retrograde)
                t2 = ts.utc(y, m, d, h - offset, mi + 1)
                deg2 = earth.at(t2).observe(eph[en_id]).ecliptic_latlon()[1].degrees
                raw_speed[en_id] = (deg2 - deg) # القيمة السالبة تعني تراجع
                
                results[f"{en_id}_sign"] = zodiacs[int(deg / 30)]
                # نظام البيوت المتساوية
                results[f"{en_id}_house"] = str(int((deg - asc_raw + 360) % 360 / 30) + 1)

            # 4. السهام والليل/النهار
            sun_lon = raw_d['sun']
            is_day_time = not (180 <= (sun_lon - asc_raw) % 360 <= 360)
            
            if is_day_time:
                fort = (asc_raw + raw_d['moon'] - raw_d['sun']) % 360
            else:
                fort = (asc_raw + raw_d['sun'] - raw_d['moon']) % 360

            results['fortune_part'] = f"{zodiacs[int(fort/30)]} ({int(fort%30)}°)"

            # 5. محرك القواطع (استخدام الدالة الشاملة التي صممناها سابقاً)
            warnings_list = []
            for ar, en_id in p_sf_map:
                if raw_speed.get(en_id, 0) < 0:
                    warnings_list.append(f"⚠️ كوكب {ar} متراجع (Retrograde)")
                if en_id != "sun":
                    dist_to_sun = abs(raw_d[en_id] - raw_d['sun'])
                    if dist_to_sun < 8.5:
                        warnings_list.append(f"🔥 كوكب {ar} محترق (Combust)")

            # عرض النتائج في واجهة Flet
            self.report_container.controls.clear()
            self.log_to_flet(f"📜 تحليل ولادة: {y}/{m}/{d}", color="blue", weight="bold")
            for msg in warnings_list:
                self.log_to_flet(msg, color="red")
            self.log_to_flet(f"💰 سهم السعادة: {results['fortune_part']}", color="green")
            
            self.page.update()

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"⚠️ خطأ: {str(e)}", rtl=True))
            self.page.snack_bar.open = True
            self.page.update()
    # 6. تحليل النجوم الثابتة (أضفه داخل دالة analyze_birth_chart بعد حساب الكواكب)
        fixed_stars_db = [
        {"name": "الثريا (Alcyone)", "pos": 60.1, "effect": "الجاذبية، الشهرة، والنجاح الفني."},
        {"name": "الدبران (Aldebaran)", "pos": 69.8, "effect": "شجاعة قيادية ونجاح مادي."},
        {"name": "رأس التوأم (Pollux)", "pos": 113.3, "effect": "الجرأة وقوة البيان."},
        {"name": "قلب الأسد (Regulus)", "pos": 150.1, "effect": "الجاه العظيم والسلطة."},
        {"name": "السماك الأعزل (Spica)", "pos": 204.0, "effect": "الثروة والنجاح العلمي."},
        {"name": "قلب العقرب (Antares)", "pos": 240.0, "effect": "التحولات الكبرى وإعادة البناء."},
        {"name": "النسر الواقع (Vega)", "pos": 285.3, "effect": "الإبداع الروحاني والتميز الفني."},
        {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "effect": "الشهرة العالمية والتميز الروحاني."}
    ]
    
        stars_analysis = []
    # التصحيح: استخدام الأسماء البرمجية لـ Skyfield التي استخدمناها في الفقرة السابقة
        my_planets_for_stars = {**raw_d, "Ascendant": asc_raw}
        p_ar_names = {
        "sun":"الشمس", "moon":"القمر", "mercury":"عطارد", "venus":"الزهرة", 
        "mars":"المريخ", "jupiter barycenter":"المشتري", "saturn barycenter":"زحل", "Ascendant":"الطالع"
    }

        # 6. تحليل النجوم الثابتة
        for star in fixed_stars_db:
            for p_en, p_deg in my_planets_for_stars.items():
                # حساب المسافة (Orb) مع مراعاة الدائرة 360
                diff = abs(p_deg - star['pos'])
                if diff <= 1.5 or diff > 358.5:
                    p_ar = p_ar_names.get(p_en, p_en)
                    stars_analysis.append(f"⭐ يقترن {p_ar} بنجم {star['name']}.\n   💡 {star['effect']}")

        # عرض النتائج في واجهة Flet
        self.log_to_flet("\n✨ تحليل النجوم الثابتة:", color="blue")
        
        if stars_analysis:
            for report in stars_analysis:
                self.log_to_flet(report, color="orange")
        else:
            self.log_to_flet("لا توجد اقترانات مباشرة بالنجوم العظمى حالياً.", color="gray")

        # 7. استدعاء التقرير الشامل (إذا كان متاحاً)
        if hasattr(self, 'display_comprehensive_report'):
            self.display_comprehensive_report(results)
        
        self.page.update()
def analyze_birth_chart(self):
        try:
            # 1. جلب المدخلات
            def g_v(attr, default):
                element = getattr(self, attr, None)
                return element.value if element and element.value else default

            y, m, d = int(g_v('year_ent', 1990)), int(g_v('month_ent', 1)), int(g_v('day_ent', 1))
            h, mi = int(g_v('hour_ent', 12)), int(g_v('min_ent', 0))
            offset = 3.0 

            # 2. الحسابات الفلكية (Skyfield)
            ts = load.timescale()
            t = ts.utc(y, m, d, h - offset, mi)
            eph = load('de421.bsp')
            earth = eph['earth']
            
            # --- حساب الطالع (ASC) ---
            # ملاحظة: يجب أن تكون دالة حساب الطالع دقيقة بناءً على وقت وموقع الولادة
            asc_raw = 0.0 # ضع هنا نتيجة دالة حساب الطالع الفعلية
            
            p_sf_map = [
                ("الشمس", "sun"), ("القمر", "moon"), ("عطارد", "mercury"), 
                ("الزهرة", "venus"), ("المريخ", "mars"), 
                ("المشتري", "jupiter barycenter"), ("زحل", "saturn barycenter")
            ]

            raw_d = {}
            self.report_container.controls.clear()
            self.log_to_flet(f"📜 تحليل نظام البيوت المتساوية: {y}/{m}/{d}", color="blue", weight="bold")

            # 3. حساب الكواكب وتحديد بيوتها المتساوية
            for ar_name, en_id in p_sf_map:
                astrometric = earth.at(t).observe(eph[en_id])
                _, lon_pos, _ = astrometric.ecliptic_latlon()
                pos = lon_pos.degrees
                raw_d[ar_name] = pos
                
                # معادلة البيت المتساوي: (درجة الكوكب - درجة الطالع + 360) % 360 / 30
                house_num = int((pos - asc_raw + 360) % 360 / 30) + 1
                
                sign_idx = int(pos / 30)
                deg = int(pos % 30)
                
                # عرض تفصيلي: الكوكب، البرج، والبيت
                self.log_to_flet(f"• {ar_name}: {self.z_names[sign_idx]} {deg}° - البيت [{house_num}]", color="black")
                # ربط التحليل بدلالة البيت من قاموسك HOUSE_FIELDS
                self.log_to_flet(f"   💡 {HOUSE_FIELDS.get(str(house_num), '')}", color="gray", size=12)

            # 4. تحليل الطالع (البيت الأول)
            sign_idx_asc = int(asc_raw / 30)
            deg_asc = int(asc_raw % 30)
            self.log_to_flet(f"\n🌅 الطالع (بداية البيت 1): {self.z_names[sign_idx_asc]} {deg_asc}°", color="purple", weight="bold")

            # 5. تحليل الاتصالات (الزوايا)
            self.log_to_flet("\n📐 الاتصالات الفلكية في البيوت:", color="blue", weight="bold")
            p_names = list(raw_d.keys())
            for i in range(len(p_names)):
                for j in range(i + 1, len(p_names)):
                    p1, p2 = p_names[i], p_names[j]
                    diff = abs(raw_d[p1] - raw_d[p2])
                    if diff > 180: diff = 360 - diff
                    
                    aspect = None
                    if diff <= 6: aspect = "0"
                    elif 174 <= diff <= 186: aspect = "180"
                    elif 84 <= diff <= 96: aspect = "90"
                    elif 114 <= diff <= 126: aspect = "120"
                    elif 54 <= diff <= 66: aspect = "60"

                    if aspect:
                        h1 = int((raw_d[p1] - asc_raw + 360) % 360 / 30) + 1
                        h2 = int((raw_d[p2] - asc_raw + 360) % 360 / 30) + 1
                        self.log_to_flet(f"💠 {p1}({h1}) في {ASPECT_MEANINGS[aspect]} مع {p2}({h2})", color="teal")

            self.page.update()

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"⚠️ خطأ: {str(e)}", rtl=True))
            self.page.snack_bar.open = True
            self.page.update()
def get_house(self, planet_deg, ascendant_deg):
    """تحديد البيت بنظام البيوت المتساوية (كل بيت 30 درجة من الطالع)"""
    planet_deg %= 360
    # الفرق بين الكوكب والطالع يعطينا رقم البيت مباشرة في النظام المتساوي
    diff = (planet_deg - ascendant_deg + 360) % 360
    house_number = int(diff / 30) + 1
    return str(house_number)

def get_sign_name(self, deg):
    """تحويل الدرجة إلى اسم البرج"""
    signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
             "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
    return signs[int((deg % 360) / 30)]

def run_calculation(self, t, lat, lon):
    """المحرك الحسابي الرئيسي باستخدام Skyfield (البديل للأندرويد)"""
    data = {}
    zodiac_names = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                    "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

    # 1. إعداد Skyfield
    eph = load('de421.bsp')
    earth = eph['earth']
    
    # 2. حساب الطالع (Ascendant)
    # ملاحظة: حساب الطالع يحتاج لزمن فلكي محلي، هنا نستخدم قيمة محسوبة مسبقاً
    ascendant = 0.0 # استبدلها بدالة حساب الطالع (ASC) التي تعمل بـ Skyfield
    data['ascendant_deg'] = ascendant
    data['ascendant_sign'] = zodiac_names[int(ascendant / 30)]

    # 3. حساب الكواكب الأساسية
    planets_to_calc = {
        'sun': eph['sun'], 'moon': eph['moon'], 'mercury': eph['mercury'],
        'venus': eph['venus'], 'mars': eph['mars'], 
        'jupiter': eph['jupiter barycenter'], 'saturn': eph['saturn barycenter']
    }

    for name, p_obj in planets_to_calc.items():
        astrometric = earth.at(t).observe(p_obj)
        _, lon_pos, _ = astrometric.ecliptic_latlon()
        p_deg = lon_pos.degrees
        
        data[f'{name}_deg'] = p_deg
        data[f'{name}_sign'] = zodiac_names[int(p_deg / 30)]
        # استخدام دالة البيوت المتساوية المعتمدة على الطالع
        data[f'{name}_house'] = self.get_house(p_deg, ascendant)

    # 4. حساب السهام الفلكية (ميلاد نهاري أم ليلي)
    is_day = not (180 <= (data['sun_deg'] - ascendant) % 360 <= 360)
    
    if is_day:
        fortune_deg = (ascendant + data['moon_deg'] - data['sun_deg']) % 360
        spirit_deg = (ascendant + data['sun_deg'] - data['moon_deg']) % 360
    else:
        fortune_deg = (ascendant + data['sun_deg'] - data['moon_deg']) % 360
        spirit_deg = (ascendant + data['moon_deg'] - data['sun_deg']) % 360
    
    data['fortune_part'] = f"{zodiac_names[int(fortune_deg / 30)]} ({int(fortune_deg % 30)}°)"
    data['spirit_part'] = f"{zodiac_names[int(spirit_deg / 30)]} ({int(spirit_deg % 30)}°)"

    return data
def get_hijri_date_string(self):
    """جلب التاريخ الهجري الحالي باستخدام المحرك المتوافق مع أندرويد"""
    try:
        from hijri_converter import Gregorian
        from datetime import datetime
        now = datetime.now()
        hijri_now = Gregorian(now.year, now.month, now.day).to_hijri()
        return f"{hijri_now.day} / {hijri_now.month} / {hijri_now.year} هـ"
    except Exception as e:
        return f"حسابي: {self.hijri_day} قمري"
def get_daily_wisdom(self, day):
        """قاعدة بيانات تفسير رؤى أيام الشهر القمري"""
        # الأيام مبنية على رؤية الصادق وابن سيرين للتراث الروحاني
        wisdom_db = {
            1: {"sadiq": "رؤيا صادقة جداً وتتحقق سريعاً.", "sirin": "بشارة خير في أول الشهر.", "philosophy": "بداية دورة جديدة في عالم المثال."},
            2: {"sadiq": "تتأخر قليلاً لكنها خير.", "sirin": "تدل على نيل رزق أو فائدة.", "philosophy": "انعكاس للرغبات الكامنة."},
            15: {"sadiq": "رؤيا وسطى بين الحق والباطل.", "sirin": "تدل على بلوغ منتصف الأمر.", "philosophy": "اكتمال الوعي القمري."},
            # يمكن إضافة بقية الأيام حتى 30 بنفس التنسيق
        }
        return wisdom_db.get(day, {
            "sadiq": "تختلف حسب حال الرائي ووقت الرؤيا.", 
            "sirin": "تتطلب تأويلاً دقيقاً من خبير.", 
            "philosophy": "انعكاس لمرآة النفس الفلكية."
        })

def show_about(self, e=None):
        """نافذة معلومات النظام المتوافقة مع أندرويد"""
        about_text = (
            "🚀 الميقاتي الفلكي - النسخة الاحترافية\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "تطوير الخبير الفلكي: حسان الشاعر\n"
            "إصدار: 2026 PRO Mobile (Flet Version)\n\n"
            "نظام فلكي متكامل يجمع بين المحرك العلمي\n"
            "والاستنباطات الفلسفية التراثية."
        )

        def close_dlg(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            title=ft.Text("حول النظام", weight="bold", rtl=True),
            content=ft.Container(
                content=ft.Text(about_text, rtl=True, size=15),
                padding=10
            ),
            actions=[
                ft.TextButton("إغلاق", on_click=close_dlg),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
def show_dream_interpreter(self, e=None):
        try: 
            # 1. حساب اليوم القمري (بديل أندرويد لـ SwissEph)
            ts = load.timescale()
            t = ts.now()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            # حساب المواقع (القمر والشمس)
            m_lon = earth.at(t).observe(eph['moon']).ecliptic_latlon().degrees
            s_lon = earth.at(t).observe(eph['sun']).ecliptic_latlon().degrees
            
            # حساب اليوم القمري بناءً على الفرق الزاوي (دورة 29.5 يوم)
            h_day = int(((m_lon - s_lon) % 360 / 12.19) + 1)
            
            # جلب حكمة اليوم القمري
            wisdom = self.get_daily_wisdom(h_day)  
            
      # 2. قاموس الرموز الماروني المتكامل
            symbols_dictionary = {
                "نبي": "🌙 العرب: عز وشرف وبشارة.\n🏹 الفرس: استقامة دين وظفر.\n🏛️ اليونان: حكمة.",
                "ملك": "🌙 العرب: نصر وعز.\n🏹 الفرس: منصب رفيع.\n🏛️ اليونان: سيطرة على النفس.",
                "شمس": "🌙 العرب: ولي الأمر.\n🏹 الفرس: الحاكم الأعظم.\n🏛️ اليونان: الوعي المطلق.",
                "ماء": "🌙 العرب: حياة طيبة.\n🏹 الفرس: استقرار ونمو.\n🏛️ اليونان: تطهير وشفاء.",
                "نار": "🌙 العرب: فتنة أو هداية.\n🏹 الفرس: رفعة.\n🏛️ اليونان: تحول وإبداع.",
                "ذهب": "🌙 العرب: همّ للرجل وفرح للنساء.\n🏹 الفرس: ميراث.\n🏛️ اليونان: خلود وفكر.",
                "فضة": "🌙 العرب: مال وراحة.\n🏹 الفرس: زواج أو ولادة.\n🏛️ اليونان: حدس وأنوار.",
                "ثعبان": "🌙 العرب: عدو من الأهل.\n🏹 الفرس: عدو ذو مال.\n🏛️ اليونان: شفاء وحكمة.",
                "أسد": "🌙 العرب: عدو جبار.\n🏹 الفرس: قوة ونفوذ.\n🏛️ اليونان: شجاعة وسيادة.",
                "طيران": "🌙 العرب: سفر ورفعة.\n🏹 الفرس: غنى سريع.\n🏛️ اليونان: تحرر من القيود.",
                "موت": "🌙 العرب: عمر طويل وتوبة.\n🏹 الفرس: بدء حياة جديدة.\n🏛️ اليونان: تحول الوعي.",
                "بحر": "🌙 العرب: الدنيا وتقلباتها.\n🏹 الفرس: ملك لا يدرك غوره.\n🏛️ اليونان: اللاوعي.",
                "خبز": "🌙 العرب: الإسلام والرزق.\n🏹 الفرس: أمن من الفقر.\n🏛️ اليونان: نمو روحي.",
                "لبن": "🌙 العرب: فطرة ورزق.\n🏹 الفرس: علم وحكمة.\n🏛️ اليونان: تغذية النفس.",
                "عسل": "🌙 العرب: حلاوة الإيمان.\n🏹 الفرس: مال موروث.\n🏛️ اليونان: رحيق الحكمة.",
                "سيف": "🌙 العرب: ولد أو حجة.\n🏹 الفرس: نصر وقوة.\n🏛️ اليونان: إرادة قاطعة.",
                "خاتم": "🌙 العرب: أمان أو زواج.\n🏹 الفرس: إتمام أمر.\n🏛️ اليونان: عهد أبدي.",
                "شجر": "🌙 العرب: رجال أو أعمال.\n🏹 الفرس: طول عمر.\n🏛️ اليونان: نمو الشخصية.",
                "كعبة": "🌙 العرب: حج وأمن.\n🏹 الفرس: استقامة.\n🏛️ اليونان: مركز الوجود.",
                "ضحك": "🌙 العرب: حزن (إلا بشارة).\n🏹 الفرس: خبر طارئ.\n🏛️ اليونان: تفريغ توتر.",
                "بكاء": "🌙 العرب: فرج ونفاذ ضيق.\n🏹 الفرس: زوال كرب.\n🏛️ اليونان: تطهير عاطفي."
            }
                        # 2. القاموس الموسع الشامل (موسوعة ابن سيرين والصادق المحدثة 2026)
            symbols_dictionary = {
                # --- الرموز الروحية والعلوية ---
                "نبي": "ابن سيرين: عز وشرف في الدارين ورؤياهم بشارة. الصادق: هداية ونيل أمانة وخير مستدام.",
                "ملك": "ابن سيرين: نصر على الأعداء ونيل سلطة وجاه. الصادق: قضاء حاجة وبشارة بالظفر.",
                "كعبة": "ابن سيرين: أمن وبشارة واجتماع بولي الأمر. الصادق: إيمان، إسلام، أمن، أو حج قريب.",
                "قرآن": "ابن سيرين: حكمة وميراث وعلم نافع. الصادق: انتشار علم ورزق طيب وتطهير للروح.",
                "محراب": "ابن سيرين: حصول ولد صالح أو إمام عادل. الصادق: امرأة صالحة أو رزق من جهة الدين.",
                "أذان": "ابن سيرين: حج لمن أذن في أشهره، أو نيل منصب. الصادق: دعوة للخير أو رفعة ذكر.",

                # --- الطبيعة والعناصر الكونيه ---
                "ماء": "ابن سيرين: حياة طيبة. الصادق: فتنة إن كان كدراً، وعلم ورزق إن كان صافياً.",
                "نار": "ابن سيرين: فتنة أو سلطان أو بشارة ورزق. الصادق: عذاب أو طريق هداية مستقيم.",
                "بحر": "ابن سيرين: سلطان مهاب. الصادق: ملك، رئيس، عالم، علم، مال، أو شغل عظيم.",
                "مطر": "ابن سيرين: رحمة عامة ما لم يكن مخرباً. الصادق: فرج من هم واستجابة دعاء.",
                "شمس": "ابن سيرين: ولي الأمر أو الوالد أو الملك. الصادق: نور في الدين وجاه في الدنيا.",
                "قمر": "ابن سيرين: وزير الملك أو الزوجة أو العالم. الصادق: زيادة في العلم أو نيل منصب.",
                "تراب": "ابن سيرين: مال ورزق بعد تعب. الصادق: يؤول بالدنيا أو الموت أو الفقر والحاجة.",

                # --- الحيوانات والطيور والحشرات ---
                "أسد": "ابن سيرين: عدو جبار أو سلطان جائر. الصادق: شدة وضيق من جهة صاحب نفوذ.",
                "خيل": "ابن سيرين: عز وشرف ومرتبة عالية. الصادق: سفر مبارك أو نيل سيادة على قومه.",
                "جمل": "ابن سيرين: حزن أو رجل جاهل أو سفر طويل. الصادق: نفوذ وقوة وصبر على المحن.",
                "ثعبان": "ابن سيرين: عدو يكاتم العداوة. الصادق: امرأة، ولد، عدو ذو مال، أو كنوز مخفية.",
                "عقرب": "ابن سيرين: عدو ضعيف يغتاب. الصادق: نكد وهمّ يسببه شخص نمام قريب.",
                "نحل": "ابن سيرين: غنيمة بلا تعب وعمل صالح. الصادق: عسكر، شفاء من مرض، أو رزق مبارك.",
                "غراب": "ابن سيرين: رجل فاسق أو سفر بعيد. الصادق: يؤول بالرجل الكاذب أو الحزن.",
                "حوت": "ابن سيرين: وزير الملك أو ضيق يتبعه فرج (قصة يونس). الصادق: رزق من البحر أو علم.",

                # --- الثمار والأطعمة ---
                "خبز": "ابن سيرين: الإسلام والعمر والمال الحلال. الصادق: عيش طيب وعلم نافع.",
                "عسل": "ابن سيرين: ميراث أو مال مجموع. الصادق: حلاوة الإيمان والقرآن والشفاء.",
                "لبن": "ابن سيرين: فطرة الإسلام والرزق الفطري. الصادق: علم واسع وحكمة.",
                "تفاح": "ابن سيرين: همّة الرجل في عمله. الصادق: نيل بشارة ومنفعة من صديق.",
                "عنب": "ابن سيرين: رزق واسع، والأسود همّ. الصادق: مال مجموع أو منفعة دائمة.",
                "تمر": "ابن سيرين: مطر أو رزق حلال لا يدوم. الصادق: حلاوة الإيمان ونيل مقصود.",
                "لحم": "ابن سيرين: المطبوخ مال، والنيء وجع ومرض. الصادق: غنيمة أو غيبة (حسب السياق).",

                # --- المعادن والكنوز والأدوات ---
                "ذهب": "ابن سيرين: همّ وغرامة للرجل. الصادق: فرح للنساء، وللرجل ذهاب منصب.",
                "فضة": "ابن سيرين: مال مجموع وراحة بال. الصادق: امرأة جميلة، ولد، أو خير كثير.",
                "جوهر": "ابن سيرين: علم وجمال. الصادق: يؤول بالمرأة الحسناء أو التجارة الرابحة.",
                "كرسي": "ابن سيرين: رفعة وشأن. الصادق: عز، شرف، ولاية، نصر، أو سعادة.",
                "مفتاح": "ابن سيرين: نصر وفتح أبواب الرزق. الصادق: قضاء حاجة، حج، ومال ميسر.",
                "ثوب": "ابن سيرين: ستر ودين ووقار. الصادق: عز وجاه وحال مستورة.",
                "سيف": "ابن سيرين: ولد أو سلطان أو كلام. الصادق: نصر على الأعداء أو حجة قوية.",

                # --- أحوال الإنسان وأفعاله ---
                "موت": "ابن سيرين: طول عمر أو ندم وتوبة. الصادق: انقطاع أمر أو زواج للأعزب.",
                "طيران": "ابن سيرين: سفر وتغير حال وطلب رفعة. الصادق: استجابة دعاء وعلو مرتبة.",
                "بكاء": "ابن سيرين: فرج وسرور (بدون عويل). الصادق: نجاة من كرب وفرح عاجل.",
                "ضحك": "ابن سيرين: حزن وندم. الصادق: بشارة بقدوم ولد أو خبر سار.",
                "زواج": "ابن سيرين: عناية إلهية أو قيد وهمّ. الصادق: نيل منصب أو غنى.",
                "غسل": "ابن سيرين: قضاء دين وتوبة من ذنب. الصادق: طهارة في الدين وزوال الهموم.",
                "حج": "ابن سيرين: قضاء دين أو حج فعلي أو أمن. الصادق: استقامة الدين ونيل أجر."
            }

            # 3. قاموس أحكام الأيام الهلالية (الإمام الصادق)
            days_dict = {
            1: "صادقة جداً، وتأويلها يظهر سريعاً وهي من أفضل الأيام.",
            2: "رؤيا صالحة، ومن رآها نال خيراً ومنفعة من السلطان أو ولي الأمر.",
            3: "رؤيا حق، وتأويلها يكون على العكس إن كانت مكروهة (تحذير).",
            4: "رؤيا محمودة، وتأويلها يتأخر قليلاً لكنه محقق.",
            5: "رؤيا فرح وسرور، ومن رآها نال بشارة يسعد بها قلبه.",
            6: "رؤيا صادقة، وتأويلها يقع بعد فترة طويلة (ربما سنوات).",
            7: "رؤيا مستورة، لا تُقص إلا على عالم أو ناصح لسرعة تحققها.",
            8: "رؤيا صادقة جداً، وهي تدل على خلاص من هم أو نيل مطلب.",
            9: "رؤيا حق، وتدل على قوة الرائي ونيله مرتبة عالية.",
            10: "رؤيا صادقة، وتدل على تيسير الأمور العالقة.",
            11: "رؤيا محمودة، وتأويلها يقع من يوم إلى 15 يوماً.",
            12: "رؤيا تحذيرية، تأويلها يتأخر لكن فيها تنبيه من غدر.",
            13: "رؤيا مختلطة، قد لا يصح تأويلها لداخلها بحديث النفس.",
            14: "رؤيا صادقة، وبشارتها سريعة كسرعة اكتمال القمر.",
            15: "رؤيا صحيحة، ومن رآها نال ما تمنى بعد تعب.",
            16: "رؤيا كاذبة (أضغاث أحلام)، لا يُعتد بها في هذا اليوم.",
            17: "رؤيا مخفية، تدل على أمور باطنة تظهر للرائي لاحقاً.",
            18: "رؤيا حق، وتدل على نصر على الأعداء أو نجاح في خصومة.",
            19: "رؤيا قوية جداً، وتأويلها يكون أوضح مما رُئي في المنام.",
            20: "رؤيا محمودة، وتدل على زيادة في الرزق والمال.",
            21: "رؤيا كاذبة، غالباً ما تكون من وسوسة النفس أو الشيطان.",
            22: "رؤيا صادقة، وتدل على قضاء الحوائج الصعبة.",
            23: "رؤيا تحذيرية من الفتن، يُنصح بالصدقة بعد رؤيتها.",
            24: "رؤيا محمودة، وتدل على الفرج بعد الضيق.",
            25: "رؤيا لا تصح، وهي من أوهام الخيال في هذا اليوم.",
            26: "رؤيا صادقة، وتدل على السفر أو التغيير في الأحوال.",
            27: "رؤيا حق، وتأويلها يقع كما هي دون تبديل.",
            28: "رؤيا صادقة، ومن رآها نال خيراً وفيراً غير متوقع.",
            29: "رؤيا تحذيرية، يُكره قصها في هذا اليوم.",
            30: "رؤيا صادقة ومباركة، وهي خاتمة الخير والرزق."
        }
# 3. عرض التقرير في Flet
            # يمكنك هنا إضافة حقل إدخال (TextField) ليبحث المستخدم عن الرمز
            # حالياً سنعرض ملخصاً لليوم القمري كمثال
            report_text = f"🌙 اليوم القمري: {h_day}\n📜 حكمة اليوم: {wisdom.get('sadiq')}\n\n💡 مثال (رمز البحر):\n{symbols_dictionary['بحر']}"

            dlg = ft.AlertDialog(
                title=ft.Text("مفسر الأحلام الفلسفي", rtl=True),
                content=ft.Text(report_text, rtl=True),
                actions=[ft.TextButton("إغلاق", on_click=lambda _: self.page.close(dlg))]
            )
            self.page.open(dlg)

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ: {ex}"))
            self.page.snack_bar.open = True
            self.page.update()
def process_interpretation(self, e, search_input, report_content, h_day, days_dict, symbols_dictionary):
        """
        المحرك المعدل للبحث والتحليل الرمزي للرؤى - نسخة أندرويد المتوافقة
        """
        search_query = search_input.value.strip()
        
        if not search_query:
            report_content.value = self.fx("⚠️ يرجى كتابة رمز (مثل: شمس، بحر، أسد) للبحث عنه.")
            self.page.update()
            return
        
        # بناء التقرير الملحمي المنسق بتنسيق احترافي
        report = "✨ تحليل الرؤيا الاستدلالي الملحمي ✨\n"
        report += f"📅 اليوم القمري الحالي: {h_day}\n"
        report += "════════════════════════════════\n\n"
        
        # جلب حكم الزمان
        wisdom_data = days_dict.get(h_day, {})
        sadiq_wisdom = wisdom_data.get("sadiq", "يوم يحتاج لمطابقة فلكية دقيقة.")
        
        report += f"🌙 حكم الزمان (حسب الشهر الهلالي):\n• {sadiq_wisdom}\n\n"

        # محرك البحث عن الرموز
        found = False
        for key in symbols_dictionary:
            if search_query in key or key in search_query:
                report += f"🔍 تفسير رمز [{key}]:\n{symbols_dictionary[key]}\n\n"
                found = True
                break
        
        if not found:
            report += f"⚠️ الرمز [{search_query}] غير مدرج حالياً.\n\n"

        # الخاتمة الفلسفية
        report += "💎 فلسفة التعبير:\n"
        report += f"• {self.dream_philosophy['chaldean']}\n" # ربطها بالمتغير المعرف في __init__
        report += f"• {self.dream_philosophy['ibn_arabi']}\n\n"
        report += "════════════════════════════════\n"
        report += "✍️ إعداد الخبير: حسان الشاعر © 2026"
        
        # تحديث المحتوى مع دعم اللغة العربية (Bidi)
        report_content.value = self.fx(report)
        self.page.update()
def draw_astro_wheel_in_popup(self, t, title_text="خريطة الميلاد"):
        """رسم مخطط فلكي (Canvas) متوافق مع أندرويد باستخدام Skyfield"""
        import math
        
        # 1. إعداد المحرك الفلكي (Skyfield) البديل لـ swe
        eph = load('de421.bsp')
        earth = eph['earth']
        ts = load.timescale()
        
        # 2. خريطة الكواكب والرموز (السبعة التقليدية)
        planets_to_draw = {
            'sun': ("☉", "#f59e0b"),   'moon': ("☽", "#3b82f6"),
            'mercury': ("☿", "#8b5cf6"), 'venus': ("♀", "#ec4899"),
            'mars': ("♂", "#dc2626"),    'jupiter barycenter': ("♃", "#d97706"),
            'saturn barycenter': ("♄", "#475569")
        }
        
        # 3. إعداد الـ Canvas
        chart_canvas = ft.Canvas(
            expand=True,
            shapes=[
                # الدائرة الخارجية (نطاق البروج)
                ft.cv.Circle(150, 150, 140, ft.Paint(color="#1f538d", stroke_width=2, style=ft.PaintingStyle.STROKE)),
                # الدائرة الداخلية (المركز)
                ft.cv.Circle(150, 150, 40, ft.Paint(color="#cbd5e1", stroke_width=1, style=ft.PaintingStyle.STROKE)),
            ]
        )

        # 4. توزيع الكواكب بناءً على درجاتها الحقيقية من Skyfield
        for p_id, (symbol, color) in planets_to_draw.items():
            try:
                # حساب الموقع الفلكي
                astrometric = earth.at(t).observe(eph[p_id])
                _, lon, _ = astrometric.ecliptic_latlon()
                pos = lon.degrees
                
                # تحويل الدرجة الفلكية إلى إحداثيات Canvas (X, Y)
                # طرح 90 يجعل البرج الأول (الحمل) يبدأ من الأعلى
                angle = math.radians(pos - 90) 
                
                x = 150 + 100 * math.cos(angle)
                y = 150 + 100 * math.sin(angle)

                chart_canvas.shapes.append(
                    ft.cv.Text(
                        x - 8, y - 12, symbol,
                        ft.TextStyle(size=22, color=color, weight="bold")
                    )
                )
            except:
                continue

        # 5. عرض النافذة المنبثقة
        chart_dlg = ft.AlertDialog(
            title=ft.Text(self.fx(title_text), weight="bold"),
            content=ft.Container(
                content=chart_canvas,
                width=300, height=300,
                alignment=ft.alignment.center,
                bgcolor="#ffffff",
                border_radius=150,
            ),
            actions=[ft.TextButton(self.fx("إغلاق الخريطة"), on_click=lambda _: self.page.close(chart_dlg))],
        )

        self.page.open(chart_dlg)
        self.page.update()
def draw_astro_wheel_in_popup(self, t):
        """رسم مخطط الأبراج وتوزيع الكواكب برمجياً (Canvas) - نسخة الأندرويد"""
        import math
        from skyfield.api import load
        
        # 1. إعداد الحاوية الرسومية والدوائر الأساسية
        cp = ft.Canvas(
            expand=True,
            shapes=[
                # الدائرة الخارجية (إطار البروج)
                ft.cv.Circle(150, 150, 120, ft.Paint(color="#1f538d", stroke_width=2, style=ft.PaintingStyle.STROKE)),
                # الدائرة المركزية
                ft.cv.Circle(150, 150, 40, ft.Paint(color="#cbd5e1", stroke_width=1, style=ft.PaintingStyle.STROKE)),
            ]
        )

        # 2. رسم خطوط تقسيم البروج الاثني عشر
        for i in range(12):
            angle = math.radians(i * 30 - 90)
            x2 = 150 + 120 * math.cos(angle)
            y2 = 150 + 120 * math.sin(angle)
            cp.shapes.append(ft.cv.Line(150, 150, x2, y2, ft.Paint(color="#cbd5e1", stroke_width=1)))

        # 3. محرك توزيع الكواكب باستخدام Skyfield (البديل للأندرويد)
        eph = load('de421.bsp')
        earth = eph['earth']
        
        planets_sf = {
            'sun': ("☉", "#f59e0b"), 
            'moon': ("☽", "#3b82f6"), 
            'mars': ("♂", "#dc2626"), 
            'jupiter barycenter': ("♃", "#d97706"), 
            'saturn barycenter': ("♄", "#475569")
        }
        
        for p_id, (symbol, color) in planets_sf.items():
            try:
                # حساب الموقع الفلكي الحقيقي
                astrometric = earth.at(t).observe(eph[p_id])
                _, lon_pos, _ = astrometric.ecliptic_latlon()
                p_deg = lon_pos.degrees
                
                # تحويل الموقع لزاوية رسم
                p_angle = math.radians(p_deg - 90)
                px = 150 + 90 * math.cos(p_angle)
                py = 150 + 90 * math.sin(p_angle)
                
                cp.shapes.append(
                    ft.cv.Text(px - 7, py - 10, symbol, ft.TextStyle(size=18, color=color, weight="bold"))
                )
            except:
                continue

        # 4. بناء وعرض النافذة المنبثقة
        dlg = ft.AlertDialog(
            title=ft.Text(self.fx("الخريطة الفلكية اللحظية"), size=18, weight="bold"),
            content=ft.Container(
                content=cp, 
                width=300, 
                height=300, 
                alignment=ft.alignment.center,
                bgcolor="#ffffff"
            ),
            actions=[
                ft.TextButton(self.fx("إغلاق"), on_click=lambda e: self.close_dialog(dlg))
            ],
            actions_alignment=ft.MainAxisAlignment.CENTER
        )
        
        self.page.open(dlg) # استخدام الطريقة الحديثة لفتح النوافذ
        self.page.update()

def close_dialog(self, dlg):
        """إغلاق النافذة المنبثقة وتحديث الصفحة"""
        self.page.close(dlg)
        self.page.update()
def search_city_logic(self, e=None):
        """محرك البحث عن الإحداثيات الجغرافية عبر الإنترنت - نسخة أندرويد"""
        try:
            from geopy.geocoders import Nominatim
            # إعداد محرك البحث
            geolocator = Nominatim(user_agent="hassan_astro_mobile_2026")
            city_name = self.city_entry.value 
            
            if not city_name or city_name.strip() == "":
                self.page.snack_bar = ft.SnackBar(ft.Text(self.fx("⚠️ يرجى إدخال اسم المدينة (مثلاً: Tartus)")))
                self.page.snack_bar.open = True
                self.page.update()
                return

            # إشعار البدء
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"🔍 جاري البحث عن {city_name}...")))
            self.page.snack_bar.open = True
            self.page.update()

            # طلب الإحداثيات
            location = geolocator.geocode(city_name, timeout=10)
            
            if location:
                self.lat = location.latitude
                self.lon = location.longitude
                
                lat_dir = "شمالاً" if self.lat >= 0 else "جنوباً"
                lon_dir = "شرقاً" if self.lon >= 0 else "غرباً"
                
                self.coords_lbl.value = self.fx(f"📍 {abs(self.lat):.2f}° {lat_dir} | {abs(self.lon):.2f}° {lon_dir}")
                self.coords_lbl.color = "blue"
                
                self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"✅ تم ضبط الموقع: {location.address}")), bgcolor="green")
                self.page.snack_bar.open = True
                
                # تحديث الحسابات بناءً على الموقع الجديد
                self.draw_now() 
            else:
                self.page.snack_bar = ft.SnackBar(ft.Text(self.fx("⚠️ لم يتم العثور على المدينة")), bgcolor="orange")
                self.page.snack_bar.open = True
            
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx("❌ فشل الاتصال: تأكد من الإنترنت")), bgcolor="red")
            self.page.snack_bar.open = True
            self.page.update()
def show_moon_mansion(self, e=None):
        """حساب وعرض منزلة القمر الحالية وتفاصيلها الروحانية - نسخة أندرويد"""
        try:
            # 1. حساب موقع القمر الحالي باستخدام Skyfield (البديل للأندرويد)
            ts = load.timescale()
            t = ts.now()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            astrometric = earth.at(t).observe(eph['moon'])
            _, lon, _ = astrometric.ecliptic_latlon()
            moon_pos = lon.degrees
            
            # 2. تحديد رقم المنزلة (الدائرة 360 تقسم على 28 منزلة = 12.857 درجة لكل منزلة)
            mansion_idx = int(moon_pos / 12.857)
            m_name, m_angel, m_incense, m_desc, m_square, m_dua = self.get_mansion_data(mansion_idx)

            # 3. بناء واجهة عرض الوفق (3x3 Grid)
            grid_cells = []
            for num in m_square:
                grid_cells.append(
                    ft.Container(
                        content=ft.Text(str(num), weight="bold", color="white"),
                        alignment=ft.alignment.center,
                        width=50, height=50,
                        bgcolor="#1e293b", border_radius=5
                    )
                )

            # 4. عرض نافذة التفاصيل بتنسيق Flet المتوافق مع الجوال
            dlg = ft.AlertDialog(
                title=ft.Text(self.fx(f"🌙 منزلة القمر: {m_name}"), weight="bold", color="#8b5cf6"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(self.fx(f"👤 المَلَك الموكل: {m_angel}"), weight="bold"),
                        ft.Text(self.fx(f"💨 البخور: {m_incense}")),
                        ft.Text(self.fx(f"📖 الدلالة: {m_desc}"), italic=True),
                        ft.Divider(),
                        ft.Text(self.fx("🔢 الوفق العددي المدمج:"), weight="bold"),
                        ft.GridView(controls=grid_cells, runs_count=3, max_extent=60, height=180),
                        ft.Divider(),
                        ft.Text(self.fx(f"🤲 الدعاء:\n{m_dua}"), size=13, color="#475569"),
                    ], tight=True, scroll=ft.ScrollMode.ADAPTIVE),
                    width=350
                ),
                actions=[ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))]
            )

            self.page.open(dlg)
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ: {ex}")))
            self.page.snack_bar.open = True
            self.page.update()
def start_app(self):
        """تشغيل التطبيق وتهيئته الأولية"""
        try:
            self.setup_ui()
            self.draw_now()
            self.page.update()
            
            # إشعار ترحيبي متوافق مع أندرويد
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx("🚀 تم تشغيل النظام الفلكي بنجاح")), bgcolor="blue")
            self.page.snack_bar.open = True
            self.page.update()
        except Exception as e:
            print(f"Error starting app: {e}")
def nail_cut_calendar(self, e=None):
        """تقويم قص الأظافر الروحاني - نسخة الأندرويد"""
        try:
            # 1. الحساب الفلكي الدقيق (Skyfield البديل للأندرويد)
            ts = load.timescale()
            t = ts.now()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            m_lon = earth.at(t).observe(eph['moon']).ecliptic_latlon().degrees
            s_lon = earth.at(t).observe(eph['sun']).ecliptic_latlon().degrees
            
            # حساب اليوم القمري (دورة 29.5 يوم)
            day = int(((m_lon - s_lon) % 360 / 360) * 29.5) + 1
                        # 2. مصفوفة الأحكام التراثية
            nail_rules = {
                1: "يورث قصر العمر والكآبة (تجنبه)",
                2: "يورث النكد وتشتت الأمر",
                3: "يورث البركة في المال والولد (مستحب جداً)",
                4: "يورث الفقر والهم",
                5: "يورث زيادة الرزق وتيسير الأمور (مبارك)",
                6: "يورث الغم والوسواس",
                7: "يورث الهيبة والوقار والقبول عند الناس (يوم ملكي)",
                8: "يورث الضعف البدني وتراجع القوة",
                9: "يورث سوء الخلق وضيق الصدر",
                10: "يورث الرفعة والكرامة والمنزلة",
                14: "يورث الشفاء من الأوجاع الجسدية (مستحب للصحّة)",
                15: "يورث السرور والفرح وقضاء الحوائج (مبارك)",
                20: "يورث الأمن من الخوف والعز (مستحب)",
                24: "يورث الظفر بالأعداء ونيل المقاصد",
                25: "يورث الخلاص من الديون والضيق"
            }

            # 1. تحديد الحالة واللون (توحيد الأسماء)
            is_suitable = "✅ مناسب ومستحب" if day in [3, 5, 7, 10, 14, 15, 20, 24, 25] else "❌ يفضل التأجيل"
            current_status = nail_rules.get(day, "يوم اعتيادي، والأفضل تحري يوم الجمعة.")
            status_color = "#10b981" if day in [3, 5, 7, 10, 14, 15, 20, 24, 25] else "#ef4444"

            # 2. بناء التقرير النصي (استخدام الأسماء الموحدة)
            report = f"📅 اليوم القمري (الهلالي): {day}\n"
            report += f"📍 الحالة: {is_suitable}\n"
            report += "══════════════════════════════\n"
            report += f"📖 الحكم: {current_status}\n"
            report += "══════════════════════════════\n"
            report += "📜 من مخطوطات الإمام جعفر الصادق:\n"
            report += "'تقليم الأظافر يوم الجمعة يؤمن من الجذام والبرص والعمى'\n\n"
            report += "💎 عند ابن سيرين:\n"
            report += "'الأظافر قوة الرجل ومقدرته'\n"
            report += "══════════════════════════════\n"
            report += "✍️ المطور: حسان الشاعر"

            # 3. عرض التقرير في Flet (مع استخدام self.fx لضمان تنسيق الأندرويد)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text(self.fx("💅 تقويم قص الأظافر الروحاني"), weight="bold", color=status_color),
                content=ft.Container(
                    content=ft.Column([
                        # استخدمنا self.fx هنا لضمان ظهور النص العربي من اليمين لليسار وبحروف متصلة
                        ft.Text(self.fx(report), size=14, color="#2c3e50")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=400,
                    padding=10,
                    bgcolor="#fffcf5" 
                ),
                actions=[
                    ft.TextButton(self.fx("نسخ التقرير"), on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton(self.fx("إغلاق"), on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()
def analyze_fixed_stars(self, e=None):
        """محرك فحص وتحليل النجوم الثابتة العظمى - نسخة الأندرويد"""
        try:
            # 1. قاعدة البيانات الموسعة (مواقع 2026)
            fixed_stars = [
                {"name": "الظليم (Achernar)", "pos": 15.3, "nature": "مشتري", "effect": "المكانة الروحية العالية، والنجاح الديني."},
                {"name": "رأس الغول (Algol)", "pos": 56.2, "nature": "زحل/مريخ", "effect": "يمنح صموداً أسطورياً وقدرة على مواجهة الأزمات."},
                {"name": "الثريا (Alcyone)", "pos": 60.1, "nature": "زهرة/مشتري", "effect": "الجاذبية الشخصية، الشهرة الفنية، والنجاح في العلوم."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "nature": "مريخ", "effect": "الشجاعة، القيادة السيادية، والنجاح المادي."},
                {"name": "رجل الجبار (Rigel)", "pos": 77.0, "nature": "مشتري/مريخ", "effect": "الشهرة الواسعة، الاختراع، والقيادة الإدارية."},
                {"name": "النطاق (Alnilam)", "pos": 83.5, "nature": "مشتري/زحل", "effect": "الارتقاء للمناصب العليا، والسمعة التي تدوم طويلاً."},
                {"name": "منكب الجوزاء (Betelgeuse)", "pos": 88.8, "nature": "مريخ/مشتري", "effect": "نجاح باهر ومفاجئ، ومكانة مرموقة في الدولة."},
                {"name": "شعرى اليمانية (Sirius)", "pos": 104.1, "nature": "مشتري/مريخ", "effect": "يمنح شهرة تاريخية، حماية ملكية، ونجاحاً عالمياً."},
                {"name": "رأس التوأم (Pollux)", "pos": 113.3, "nature": "مريخ", "effect": "قوة البيان، الشجاعة، والقدرة على الغلبة."},
                {"name": "شعرى الشامية (Procyon)", "pos": 116.3, "nature": "مشتري/مريخ", "effect": "النجاح السريع والمفاجئ، والذكاء العملي."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "nature": "مشتري/مريخ", "effect": "الجاه العظيم، السلطة، الرفعة السيادية، والكرامة."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "nature": "زهرة/مشتري", "effect": "أكثر النجوم سعادة: الثروة، الثقافة، والنجاح العلمي."},
                {"name": "سماك الرامح (Arcturus)", "pos": 204.2, "nature": "مشتري/مريخ", "effect": "العبقرية، القيادة الفكرية، والتميز الفني."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "nature": "مريخ/مشتري", "effect": "التحولات الكبرى، والقدرة المذهلة على إعادة البناء."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "nature": "زهرة/عطارد", "effect": "الإبداع الروحاني، السحر البلاغي، والتميز الفني."},
                {"name": "النسر الطائر (Altair)", "pos": 292.0, "nature": "مريخ/مشتري", "effect": "الطموح العالي جداً، والوصول للأهداف الصعبة."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "nature": "زهرة/عطارد", "effect": "الشهرة العالمية، والتميز في العلوم الروحانية."},
                {"name": "الردف (Deneb)", "pos": 335.3, "nature": "زهرة/عطارد", "effect": "الذكاء الثاقب، القدرة على التعلم، والنجاح في الفنون."},
            ]

            # 2. إنشاء قائمة العرض
            report_display = ft.ListView(expand=True, spacing=10, padding=10)
            report_display.controls.append(
                ft.Text(self.fx("🌍 دليل النجوم الثابتة (2026)"), size=20, weight="bold", color="#38bdf8")
            )

            for s in fixed_stars:
                sign_name = self.z_names[int(s['pos'] / 30)] # استخدام قائمة البروج المعرفة في الكلاس
                report_display.controls.append(
                    ft.Container(
                        content=ft.Column([
                            ft.Text(self.fx(f"⭐ {s['name']}"), weight="bold", color="#fbbf24", size=16),
                            ft.Text(self.fx(f"📍 الموقع: {sign_name} ({int(s['pos']%30)}°)"), size=14, color="#e2e8f0"),
                            ft.Text(self.fx(f"🌿 طبيعته: {s['nature']}"), size=13, color="#94a3b8"),
                            ft.Text(self.fx(f"✨ التأثير: {s['effect']}"), size=14, italic=True, color="#cbd5e1"),
                        ], spacing=5),
                        padding=15, 
                        bgcolor="#1e293b",
                        border_radius=10
                    )
                )

            # 3. عرض النافذة
            dlg = ft.AlertDialog(
                content=ft.Container(content=report_display, width=400, height=500),
                actions=[ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))],
            )

            self.page.open(dlg)
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ: {ex}")))
            self.page.snack_bar.open = True
            self.page.update()
def show_fixed_stars_dialog(self, e=None):
        """عرض نافذة النجوم الثابتة - نسخة الأندرويد الاحترافية"""
        try:
            # 1. قاعدة البيانات (مواقع 2026)
            fixed_stars = [
                {"name": "الثريا (Alcyone)", "pos": 60.1, "effect": "الجاذبية، الشهرة، والنجاح الفني."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "effect": "الشجاعة والقيادة السيادية."},
                {"name": "رأس التوأم (Pollux)", "pos": 113.3, "effect": "قوة البيان والقدرة على الغلبة."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "effect": "الجاه العظيم والرفعة السيادية."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "effect": "الثروة والنجاح العلمي الباهر."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "effect": "التحولات الكبرى وإعادة البناء."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "effect": "الإبداع الروحاني والتميز الفني."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "effect": "الشهرة العالمية والوعي الكوني."}
            ]

            report_display = ft.Column(scroll=ft.ScrollMode.ALWAYS, expand=True, spacing=10)

            # وظيفة عرض النجوم العامة اليوم
            def show_current_stars(e):
                report_display.controls.clear()
                report_display.controls.append(ft.Text(self.fx("🌍 مواقع النجوم الثابتة اليوم"), size=18, weight="bold", color="#38bdf8"))
                for s in fixed_stars:
                    # حساب البرج والدرجة للموقع
                    s_idx = int(s['pos']/30)
                    s_deg = int(s['pos']%30)
                    report_display.controls.append(
                        ft.Container(
                            content=ft.Column([
                                ft.Text(self.fx(f"⭐ {s['name']}"), weight="bold", color="#fbbf24"),
                                ft.Text(self.fx(f"📍 الموقع: {self.z_names[s_idx]} {s_deg}°"), size=13),
                                ft.Text(self.fx(f"✨ التأثير: {s['effect']}"), size=13, italic=True),
                            ]),
                            padding=10, border=ft.border.all(1, "#38bdf8"), border_radius=8
                        )
                    )
                self.page.update()

            # وظيفة تحليل نجوم الميلاد الشخصية
            def show_natal_stars(e):
                report_display.controls.clear()
                # التحقق من أن المستخدم قام بعملية التحليل أولاً
                if not getattr(self, 'is_analyzed', False):
                    report_display.controls.append(
                        ft.Container(
                            content=ft.Text(self.fx("⚠️ يرجى تحليل خارطة الميلاد أولاً من الواجهة الرئيسية."), color="red", weight="bold"),
                            padding=20, alignment=ft.alignment.center
                        )
                    )
                    self.page.update()
                    return

                report_display.controls.append(ft.Text(self.fx("👶 تحليل اقترانات النجوم في ميلادك"), size=18, weight="bold", color="#10b981"))
                
                # جلب بيانات الكواكب المحسوبة بـ Skyfield (التي برمجناها سابقاً)
                points = {**self.raw_d, "Ascendant": self.asc_raw}
                p_names = {"sun":"الشمس", "moon":"القمر", "mercury":"عطارد", "venus":"الزهرة", "mars":"المريخ", 
                           "jupiter barycenter":"المشتري", "saturn barycenter":"زحل", "Ascendant":"الطالع"}
                
                found = False
                for s in fixed_stars:
                    for p_en, p_deg in points.items():
                        diff = abs(p_deg - s['pos'])
                        if diff <= 1.5 or diff > 358.5:
                            p_ar = p_names.get(p_en, p_en)
                            report_display.controls.append(
                                ft.Container(
                                    content=ft.Column([
                                        ft.Text(self.fx(f"✅ اقتران ملكي: {s['name']} مع {p_ar}"), weight="bold", color="#065f46"),
                                        ft.Text(self.fx(f"◈ الدلالة: {s['effect']}"), size=14, color="#1e293b"),
                                    ]),
                                    padding=12, bgcolor="#ecfdf5", border_radius=8, border=ft.border.all(1, "#10b981")
                                )
                            )
                            found = True
                
                if not found:
                    report_display.controls.append(
                        ft.Container(content=ft.Text(self.fx("🔍 لم يتم العثور على اقترانات دقيقة.")), padding=20)
                    )
                self.page.update()

            # بناء النافذة المنبثقة
            dlg = ft.AlertDialog(
                title=ft.Row([ft.Icon(ft.icons.STARS, color="#fbbf24"), ft.Text(self.fx("ديوان النجوم الثابتة"), weight="bold")], alignment="center"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Row([
                            ft.ElevatedButton(self.fx("مواقع اليوم"), on_click=show_current_stars, bgcolor="#38bdf8", color="white"),
                            ft.ElevatedButton(self.fx("نجوم ميلادي"), on_click=show_natal_stars, bgcolor="#10b981", color="white"),
                        ], alignment="center", spacing=10),
                        ft.Divider(height=20),
                        report_display
                    ], tight=True),
                    width=400, height=500
                ),
                actions=[ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))],
            )

            self.page.open(dlg)
            show_current_stars(None)
            self.page.update()

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ في محرك النجوم: {e}")))
            self.page.snack_bar.open = True
            self.page.update()
def show_moon_mansion(self, e=None):
        try:
            # 1. حساب موقع القمر الحالي (بديل أندرويد لـ SwissEph)
            ts = load.timescale()
            t = ts.now()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            astrometric = earth.at(t).observe(eph['moon'])
            _, lon, _ = astrometric.ecliptic_latlon()
            moon_pos = lon.degrees            
            # قاعدة البيانات المكتملة التي أرسلتها
            STARS_INFO = {
                1: ("الشرطين", "2", "الحمل", "أ", "نجمان (النطح والشيرتان)", "منزلة نارية، تدل على الحماس والبدايات والاندفاع القوي."),
                2: ("البطين", "3", "الحمل", "ب", "ثلاثة نجوم خفية", "منزلة ترابية، تمتاز بالثبات والقدرة على التخطيط المالي."),
                3: ("الثريا", "7", "الثور", "ج", "الأخوات السبع (الثريا)", "منزلة هوائية، وهي من أسعد المنازل للمحبة والقبول والشهرة."),
                4: ("الدبران", "5", "الثور", "د", "عين الثور (الدبران)", "منزلة ترابية، تدل على الهيبة والقوة الفكرية العالية."),
                5: ("الهقعة", "3", "الجوزاء", "هـ", "رأس الجبار (الهقعة)", "منزلة هوائية، مناسبة جداً للتعلم وتبادل المعلومات بذكاء."),
                6: ("الهنعة", "2", "الجوزاء", "و", "قدم الجوزاء (الهنعة)", "منزلة مائية، تمتاز بطاقة الألفة وإصلاح العلاقات المتوترة."),
                7: ("الذراع", "2", "السرطان", "ز", "رأس التوأمين (الذراع)", "منزلة مائية، تمنح حماية للممتلكات وتساعد في الظفر."),
                8: ("النثرة", "3", "السرطان", "ح", "سديم المعلف (النثرة)", "منزلة نارية، تدل على الرفعة وعلو الشأن والنجاح الرسمي."),
                9: ("الطرفة", "2", "الأسد", "ط", "عيني الأسد (الطرفة)", "منزلة نارية، تتطلب الحذر وهي قوية للتغييرات الجذرية."),
                10: ("الجبهة", "4", "الأسد", "ي", "جبهة الأسد (الجبهة)", "منزلة نارية، تعيد الهيبة والشرف وتساعد في الاستشفاء."),
                11: ("الزبرة", "2", "الأسد/العذراء", "ك", "ظهر الأسد (الزبرة)", "منزلة نارية، مباركة للتجارة والارتباط العاطفي الناجح."),
                12: ("الصرفة", "1", "العذراء", "ل", "ذنب الأسد (الصرفة)", "منزلة ترابية، تدل على التغيير والانتقال، ومناسبة للسفر."),
                13: ("العواء", "5", "العذراء", "م", "خمسة نجوم (العواء)", "منزلة ترابية، تمتاز بسرعة الاستجابة للأماني وقضاء الحوائج."),
                14: ("السماك", "1", "الميزان", "ن", "السماك الأعزل", "منزلة هوائية، منبع النمو والازدهار في الصفقات المالية."),
                15: ("الغفر", "3", "الميزان", "س", "الغفر", "منزلة هوائية، تدل على الاستبصار والبحث في بواطن الأمور."),
                16: ("الزبانا", "2", "الميزان/العقرب", "ع", "كفتا الميزان (الزبانا)", "منزلة هوائية، مناسبة لضبط الموازين المالية والحقوق."),
                17: ("الإكليل", "3", "العقرب", "ف", "رأس العقرب (الإكليل)", "منزلة مائية، قوية جداً في التأليف بين القلوب وجذب المودة."),
                18: ("القلب", "1", "العقرب", "ص", "قلب العقرب", "منزلة مائية، للتحصين والوقاية والتعامل مع النفس بعمق."),
                19: ("الشولة", "2", "العقرب/القوس", "ق", "ذيل العقرب (الشولة)", "منزلة نارية، تدل على التحرر من القيود والسرعة في الإنجاز."),
                20: ("النعائم", "8", "القوس", "ر", "النعائم", "منزلة نارية، منزلة الوفرة والرزق وتيسير الأمور المتعسرة."),
                21: ("البلدة", "1", "القوس/الجدي", "ش", "البلدة (فراغ سماوي)", "منزلة ترابية، منزلة الاستقرار وبناء الأصول والعمران."),
                22: ("سعد الذابح", "2", "الجدي", "ت", "سعد الذابح", "منزلة ترابية، تدل على الحزم والنجاة من الأخطار المحيطة."),
                23: ("سعد بلع", "2", "الجدي/الدلو", "ث", "سعد بلع", "منزلة هوائية، مناسبة للعلاجات الطبية وتسهيل الصعاب."),
                24: ("سعد السعود", "3", "الدلو", "خ", "سعد السعود", "منبع السعادة والأفراح والنمو في كافة جوانب الحياة."),
                25: ("سعد الأخبية", "4", "الدلو", "ذ", "سعد الأخبية", "تساعد في كشف الأسرار والظهور القوي بعد فترة غياب."),
                26: ("المقدم", "2", "الحوت", "ض", "فرع الدلو المقدم", "منزلة مائية، تدل على الجاه والعز والتواصل الفكري."),
                27: ("المؤخر", "2", "الحوت", "ظ", "فرع الدلو المؤخر", "منزلة مائية، مباركة لزيادة الأرزاق والنجاح في المشاريع."),
                28: ("الرشاء", "1", "الحوت/الحمل", "غ", "الرشاء (بطن الحوت)", "منزلة مائية، منزلة الخواتيم السعيدة وإتمام المهام بنجاح.")
            }
 # 3. تحديد رقم المنزلة الحالية
            mansion_num = int(moon_pos / 12.857) + 1
            info = STARS_INFO.get(mansion_num, ("منزلة غير معرفة", "", "", "", "", ""))

            # 4. بناء نص التقرير
            analysis_text = (
                f"🌙 منزلة القمر الحالية: {info[0]}\n"
                f"📍 الموقع الفلكي: برج {info[2]}\n"
                f"🔢 عدد نجومها: {info[1]} | الحرف: {info[3]}\n"
                f"🌌 الوصف: {info[4]}\n"
                f"📜 الدلالة: {info[5]}"
            )

            # 5. عرض النافذة المنبثقة في Flet
            dlg = ft.AlertDialog(
                title=ft.Text(self.fx(f"كتاب المنازل: {info[0]}"), weight="bold", color="#8b5cf6"),
                content=ft.Container(
                    content=ft.Text(self.fx(analysis_text), rtl=True, size=16),
                    width=400, padding=10
                ),
                actions=[ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))]
            )
            
            self.page.open(dlg)
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ في المنازل: {ex}")))
            self.page.snack_bar.open = True
            self.page.update()
def show_moon_mansion(self, e=None):
        try:
            # 1. الحسابات الفلكية باستخدام Skyfield (البديل للأندرويد)
            ts = load.timescale()
            t = ts.now()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            # حساب موقع القمر والشمس
            m_lon = earth.at(t).observe(eph['moon']).ecliptic_latlon().degrees
            s_lon = earth.at(t).observe(eph['sun']).ecliptic_latlon().degrees
            
            # عمر القمر ومكانه
            moon_age = ((m_lon - s_lon) % 360 * 29.53) / 360
            mansion_idx = int(m_lon / (360/28))
            m_num = mansion_idx + 1

            # 2. جلب البيانات من القواميس المعرفة في الكلاس
            m_data = self.STARS_INFO.get(m_num, (0, "", "", "", ""))
            m_stars_cnt, m_signs, m_chars, m_stars_desc, m_explanation = m_data
            
            # استدعاء دالة البيانات الروحانية (التي برمجتها أنت)
            name, angel, incense, effect, wifq, prayer = self.get_mansion_data(mansion_idx)

            # 3. بناء نص التقرير المنسق (مع استخدام self.fx لضمان تنسيق الأندرويد)
            report = (
                "✨ الديوان الشامل للمنازل القمرية ✨\n"
                "══════════════════════════\n"
                f"🌙 المنزلة: {name} (رقم: {m_num})\n"
                f"⏳ عمر القمر: {moon_age:.1f} يوم\n"
                f"📍 الموقع: {int(m_lon%30)}° في برج {self.get_sign_name(m_lon)}\n"
                "══════════════════════════\n"
                f"🌟 النجوم: {m_stars_desc}\n"
                f"♈ البروج: {m_signs} | 🔠 الحروف: ( {m_chars} )\n"
                "══════════════════════════\n"
                f"🔍 التفسير: {m_explanation}\n"
                "══════════════════════════\n"
                f"👼 الملك: {angel} | 🪵 البخور: {incense}\n"
                f"📖 التأثير: {effect}\n"
                "══════════════════════════\n"
                "🔢 الوفق العددي للمنزلة:\n"
                f"    [ {wifq[0]:02}  {wifq[1]:02}  {wifq[2]:02} ]\n"
                f"    [ {wifq[3]:02}  {wifq[4]:02}  {wifq[5]:02} ]\n"
                f"    [ {wifq[6]:02}  {wifq[7]:02}  {wifq[8]:02} ]\n\n"
                f"✨ العزيمة: 「 {prayer} 」\n"
                "══════════════════════════\n"
                "المطور الفلكي: حسان الشاعر"
            )

            # 4. إعداد واجهة العرض (Flet AlertDialog)
            dlg = ft.AlertDialog(
                title=ft.Text(self.fx(f"ديوان منزلة {name}"), weight="bold", color="#1f538d"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(self.fx(report), size=14, color="#3d405b")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, height=700, bgcolor="#f4f1de", padding=15, border_radius=10
                ),
                actions=[
                    ft.ElevatedButton(self.fx("نسخ التقرير"), icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))
                ],
            )

            self.page.open(dlg)
            self.page.update()

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ في المنازل: {e}")))
            self.page.snack_bar.open = True
            self.page.update()
def analyze_current_time(self, e=None):
        """توليد التقرير الاستراتيجي الزمني والمالي - نسخة أندرويد"""
        try:
            # 1. إعداد محتوى التقرير
            content = "🔭 التحليل الزمني الشامل للأبراج والأسواق 🔭\n"
            content += "════════════════════════════════\n\n"
            
            content += "📍 أولاً: التحليل اليومي للأبراج\n"
            z_names = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                       "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            adv_list = ["طاقة نارية للمبادرة والحسم", "فرص ترابية للمال والعمل", 
                        "نشاط هوائي للذكاء والتواصل", "هدوء مائي للعاطفة والحدس"]
            
            for i, name in enumerate(z_names):
                adv = adv_list[i % 4]
                content += f"● {name}: {adv}.\n"

            content += "\n💰 ثانياً: الاقتصاد والذهب (2026)\n"
            content += "• يشهد الذهب تذبذباً يميل للارتفاع نتيجة قوة زاوية المشتري.\n"
            content += "• عام 2026 هو عام التحول الرقمي وصعود المعادن النفيسة.\n"
            
            content += "\n════════════════════════════════\n"
            content += "إعداد المطور الفلكي: حسان الشاعر © 2026"

            # 2. دالة الحفظ المعدلة
            def save_report_file(e):
                try:
                    import os
                    file_name = "Strategic_Report_2026.txt"
                    # المسار الأكثر أماناً في أندرويد هو مجلد ملفات التطبيق
                    if platform.system() == "Android":
                        path = "/storage/emulated/0/Download"
                        if not os.path.exists(path): path = os.path.expanduser("~")
                    else:
                        path = os.path.expanduser("~/Downloads")
                        
                    full_path = os.path.join(path, file_name)
                    with open(full_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    
                    self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"✅ تم الحفظ في: {file_name}")), bgcolor="green")
                    self.page.snack_bar.open = True
                except Exception as ex:
                    self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"❌ فشل الحفظ: تأكد من الصلاحيات")), bgcolor="red")
                    self.page.snack_bar.open = True
                self.page.update()

            # 3. عرض النافذة
            dlg = ft.AlertDialog(
                title=ft.Text(self.fx("التقرير الاستراتيجي الزمني"), weight="bold", color="#1f538d"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(self.fx(content), size=14, color="#1e293b", weight="bold"),
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, height=600, padding=10
                ),
                actions=[
                    ft.ElevatedButton(self.fx("💾 حفظ"), on_click=save_report_file),
                    ft.ElevatedButton(self.fx("📋 نسخ"), on_click=lambda _: self.page.set_clipboard(content)),
                    ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))
                ],
                actions_alignment=ft.MainAxisAlignment.SPACE_EVENLY,
            )

            self.page.open(dlg)
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ: {ex}")))
            self.page.snack_bar.open = True
            self.page.update()
def hair_cut_calendar(self, e=None):
        """تقويم قص الشعر الروحاني - نسخة الأندرويد المضمونة"""
        try:
            # 1. الحساب الفلكي الدقيق (Skyfield البديل للأندرويد)
            ts = load.timescale()
            t = ts.now()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            # حساب المواقع (القمر والشمس)
            m_lon = earth.at(t).observe(eph['moon']).ecliptic_latlon().degrees
            s_lon = earth.at(t).observe(eph['sun']).ecliptic_latlon().degrees
            
            # حساب اليوم القمري بدقة (دورة 29.5 يوم)
            day = int(((m_lon - s_lon) % 360 / 12.19) + 1)
            if day > 30: day = 30 

            hair_rules = {
                1: "قصر العمر (تجنبه)", 2: "الهم والغم", 3: "النكد وتشتت الأمر",
                4: "الوجع في الرأس", 5: "البركة والمال الكثير", 6: "الهم وضيق الصدر",
                7: "الهيبة والجاه والقبول (مستحب جداً)", 8: "العافية وطول العمر",
                9: "الضعف وسوء الخلق", 10: "الرفعة والكرامة", 11: "الحزن والكآبة",
                12: "العز والوقار (مستحب)", 13: "الخصومة والنزاع", 
                14: "السرور والفرح وقضاء الحاجة (مبارك)", 15: "تيسير الأمور وزيادة الجمال",
                16: "الحزن وضيق الرزق", 17: "النحوسة (تجنبه)", 18: "البركة في الرزق",
                19: "القدرة والتمكين", 20: "الأمن من الخوف والهيبة (مبارك)",
                21: "تشتت الرأي", 22: "الفقر الشديد (تجنبه)", 23: "العافية من الأمراض",
                24: "النصر على الأعداء", 25: "الخلاص من السجون والديون",
                26: "الفرح والسرور", 27: "السلامة من الآفات", 28: "قضاء الحاجة (مستحب)",
                29: "الفقر والهم (تجنبه)", 30: "الأمن من الآفات"
            }
            suitable_days = [5, 7, 8, 10, 12, 14, 15, 18, 19, 20, 23, 24, 25, 26, 27, 28, 30]
            current_status = hair_rules.get(day, "يوم يحتاج لمراجعة")
            is_suitable = "✅ مناسب ومستحب" if day in suitable_days else "❌ يفضل التأجيل"

            # 3. صياغة التقرير (مع استخدام self.fx لضمان تنسيق الأندرويد)
            report = (
                f"✂️ تقويم قص الشعر الروحاني ✂️\n"
                f"📅 اليوم القمري: {day}\n"
                f"📍 الحالة: {is_suitable}\n"
                f"📖 حكم اليوم: يورث {current_status}\n"
                f"══════════════════════════════\n"
                f"✍️ المطور: حسان الشاعر"
            )

            # 4. دالة الحفظ المعدلة للأندرويد
            def save_hair_report(e):
                try:
                    import os
                    file_name = f"Hair_Cut_Day_{day}.txt"
                    # المسار الأكثر أماناً في أندرويد
                    if platform.system() == "Android":
                        path = "/storage/emulated/0/Download"
                        if not os.path.exists(path): path = os.path.expanduser("~")
                    else:
                        path = os.path.expanduser("~/Downloads")
                        
                    full_path = os.path.join(path, file_name)
                    with open(full_path, "w", encoding="utf-8") as f:
                        f.write(report)
                    self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"✅ تم الحفظ في: {file_name}")), bgcolor="green")
                    self.page.snack_bar.open = True
                except:
                    self.page.snack_bar = ft.SnackBar(ft.Text(self.fx("❌ فشل الحفظ")), bgcolor="red")
                    self.page.snack_bar.open = True
                self.page.update()

            # 5. واجهة العرض (Dialog)
            dlg = ft.AlertDialog(
                title=ft.Text(self.fx("تقويم قص الشعر الروحاني"), color="#ec4899"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(self.fx(report), size=15, color="#1e293b")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=400, height=400, padding=10
                ),
                actions=[
                    ft.ElevatedButton(self.fx("💾 حفظ"), on_click=save_hair_report),
                    ft.ElevatedButton(self.fx("📋 نسخ"), on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))
                ],
            )
            self.page.open(dlg)
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ: {ex}")))
            self.page.snack_bar.open = True
            self.page.update()
def update_location(self, e=None):
        """تحديث الإحداثيات يدوياً - نسخة الأندرويد المضمونة"""
        try:
            # التأكد من جلب القيمة من الحقول التي عرفناها سابقاً
            self.lat = float(self.city_entry.value if self.city_entry.value else 34.88)
            # ملحوظة: يفضل أن يكون لديك حقل خاص للـ lat و lon إذا كنت ستسمح بالإدخال الرقمي المباشر
            
            self.page.snack_bar = ft.SnackBar(
                ft.Text(self.fx(f"✅ تم تحديث الموقع بنجاح")),
                bgcolor="#10b981"
            )
            self.page.snack_bar.open = True
            
            # إعادة الحساب فوراً
            self.draw_now() 
            self.page.update()
            
        except ValueError:
            self.page.snack_bar = ft.SnackBar(
                ft.Text(self.fx("⚠️ خطأ: يرجى إدخال أرقام صحيحة")),
                bgcolor="#ef4444"
            )
            self.page.snack_bar.open = True
            self.page.update()
def get_house(self, pos, cusps):
        """تحديد رقم البيت (1-12) - نسخة الأداء العالي"""
        pos %= 360
        # في Flet/Python، المصفوفات تبدأ من 0، لذا cusps[0] هو البيت الأول
        for i in range(12):
            c1 = cusps[i]
            c2 = cusps[i + 1] if i < 11 else cusps[0]
            
            if c1 < c2:
                if c1 <= pos < c2:
                    return i + 1
            else: # معالجة عبور نقطة الصفر (0° Aries)
                if pos >= c1 or pos < c2:
                    return i + 1
        return 1
def show_about(self, e=None):
        """عرض نافذة معلومات المطور بتنسيق تقني جذاب - نسخة الأندرويد"""
        email = "Hassan.alshaer2@gmail.com"
        # دعم اللغة العربية والإنجليزية
        is_ar = getattr(self, 'current_lang', 'ar') == "ar"
        
        # استخدام دالة self.fx لضمان التنسيق العربي الصحيح داخل الحاوية السوداء
        developer_label = self.fx("المطور البرمجي الفلكي: حسان الشاعر") if is_ar else "Developer: Hassan Al-Shaer"
        location_label = self.fx("العنوان: سوريا / طرطوس") if is_ar else "Location: Syria / Tartus"
        software_label = self.fx("برامجنا الاحترافية:") if is_ar else "Professional Software:"

        lines = [
            developer_label,
            f"Email: {email}",
            location_label,
            "WhatsApp: +963 933303612",
            "----------------------------------",
            software_label,
            "* Animodar Correction System",
            "* Primary Directions Pro",
            "* Al-Khafia Solar System"
        ]

        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # إعداد النافذة المنبثقة بتنسيق تقني (Matrix Style)
        dlg = ft.AlertDialog(
            title=ft.Text(self.fx("حول المطور") if is_ar else "About Developer", weight="bold"),
            content=ft.Container(
                content=ft.Text("\n".join(lines), color="#39FF14", size=13, font_family="monospace"),
                bgcolor="#000000",
                padding=20,
                border_radius=10,
                border=ft.border.all(1, "#39FF14")
            ),
            actions=[
                ft.TextButton(self.fx("إغلاق") if is_ar else "Close", on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.open(dlg)
        self.page.update()
def show_mundane_analysis(self, t):
        """توليد التقرير الملحمي لسنة العالم 2026 - نسخة الأندرويد المضمونة"""
        try:
            # 1. إعداد البيانات (Skyfield البديل لـ SwissEph)
            ts = load.timescale()
            eph = load('de421.bsp')
            earth = eph['earth']
            
            lat = float(getattr(self, 'lat', 34.8))
            lon = float(getattr(self, 'lon', 35.8))
            
            # 2. حساب المواقع (الشمس، المشتري، المريخ، زحل)
            def get_p_deg(obj_id):
                astrometric = earth.at(t).observe(eph[obj_id])
                _, lon_pos, _ = astrometric.ecliptic_latlon()
                return lon_pos.degrees

            sun_deg = get_p_deg('sun')
            jupiter_deg = get_p_deg('jupiter barycenter')
            mars_deg = get_p_deg('mars')
            saturn_deg = get_p_deg('saturn barycenter')
            
            # حساب الطالع (ASC) - يفترض وجود قيمة مدخلة أو محسوبة لعام 2026
            # للحصول على دقة 100% يفضل ربطها بدالة الأوتاد التي صممناها
            asc_raw = 0.0 

            # 3. حساب السهام الاستراتيجية (الذهب، النفط، الغذاء)
            gold_lot = (asc_raw + jupiter_deg - sun_deg) % 360
            oil_lot = (asc_raw + mars_deg - jupiter_deg) % 360
            food_lot = (asc_raw + mars_deg - saturn_deg) % 360

            # 4. بناء نص التقرير (مع استخدام self.fx لضمان تنسيق الأندرويد)
            report_body = (
                "📜 الديوان السلطاني الشامل لحكم سنة العالم 2026 📜\n"
                "════════════════════════════════\n\n"
                "💰 الاقتصاد والذهب والنفط:\n"
                f"◈ سهم الذهب: في برج {self.z_names[int(gold_lot/30)]} (زلزال مالي مرتقب).\n"
                f"◈ سهم البترول: في برج {self.z_names[int(oil_lot/30)]} (اشتعال طاقة).\n\n"
                "🌾 ميزان الغذاء والحبوب (سهم القمح):\n"
                f"◈ سهم الغذاء: في {self.z_names[int(food_lot/30)]} ينبئ بصعوبات غذائية.\n"
                "◈ نصيحة: يُنصح بتأمين مخزونات الغذاء الأساسية.\n\n"
                "════════════════════════════════\n"
                "✍️ إعداد وتطوير: حسان الشاعر © 2026"
            )

            # 5. عرض النافذة في Flet
            dlg = ft.AlertDialog(
                title=ft.Text(self.fx("الديوان السلطاني - سنة 2026"), color="#7c2d12", weight="bold"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(self.fx(report_body), size=14, color="#1e40af", weight="bold")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, height=600, bgcolor="#fcfaf2", padding=15
                ),
                actions=[
                    ft.ElevatedButton(self.fx("نسخ"), on_click=lambda _: self.page.set_clipboard(report_body)),
                    ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))
                ],
            )

            self.page.open(dlg)
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"خطأ في التحليل: {ex}")))
            self.page.snack_bar.open = True
            self.page.update()
def save_report(self, content):
        """حفظ التقارير كملفات نصية - نسخة الأندرويد المضمونة"""
        try:
            import os, platform
            from datetime import datetime
            
            # 1. إنشاء اسم ملف فريد
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Astro_Report_{timestamp}.txt"
            
            # 2. تحديد مسار الحفظ الأكثر أماناً
            if platform.system() == "Android":
                # يحاول الحفظ في مجلد التحميلات العام
                base_path = "/storage/emulated/0/Download"
                if not os.path.exists(base_path):
                    # إذا تعذر الوصول، يحفظ في مجلد ملفات التطبيق الخاصة
                    base_path = os.environ.get("PYTHONHOME", os.getcwd())
            else:
                # للمستخدمين على ويندوز أثناء التجربة
                base_path = os.path.expanduser("~/Downloads")
                
            full_path = os.path.join(base_path, filename)
            
            # 3. عملية الكتابة الفعليّة
            with open(full_path, "w", encoding="utf-8") as f:
                f.write(str(content))
            
            # 4. تنبيه النجاح منسق للعربية (RTL)
            self.page.snack_bar = ft.SnackBar(
                ft.Text(self.fx(f"✅ تم حفظ التقرير باسم: {filename}")),
                bgcolor="#10b981"
            )
            self.page.snack_bar.open = True
            self.page.update()
            
        except Exception as e:
            # رسالة خطأ ذكية
            error_msg = "يرجى منح صلاحية الملفات" if "Permission" in str(e) else f"خطأ: {e}"
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"❌ {error_msg}")), bgcolor="red")
            self.page.snack_bar.open = True
            self.page.update()
def display_comprehensive_report(self, data):
        """توليد وعرض التقرير الفلكي الشامل - نسخة الأندرويد الملونة"""
        
        # 1. إعداد لوحة الألوان الملحمية
        colors = {
            "gold": "#f59e0b", "sky": "#38bdf8", "pink": "#fb7185",
            "danger": "#ef4444", "info": "#10b981", "border": "#334155"
        }

        # 2. إنشاء قائمة العناصر
        report_widgets = []

        # --- القسم الأول: المقدمة ---
        report_widgets.append(
            ft.Text(self.fx("📜 التقرير الفلكي التحليلي الشامل المطور"), 
                    size=22, weight="bold", color=colors["gold"], text_align="center")
        )
        report_widgets.append(ft.Divider(color=colors["gold"], height=20))

        # --- القسم الثاني: تحليل الكواكب والعقد ---
        planets_config = [
            ("الشمس", "sun_sign", "sun_house", "☀️"), ("القمر", "moon_sign", "moon_house", "🌙"),
            ("عطارد", "mercury_sign", "mercury_house", "☿"), ("الزهرة", "venus_sign", "venus_house", "♀️"),
            ("المريخ", "mars_sign", "mars_house", "♂️"), ("المشتري", "jupiter_sign", "jupiter_house", "♃"),
            ("زحل", "saturn_sign", "saturn_house", "♄"), ("الرأس", "nn_sign", "nn_house", "🐲"),
            ("الذنب", "sn_sign", "sn_house", "🐉")
        ]

        # --- القسم الثالث: السهام الفلكية ---
        report_widgets.append(ft.Text(self.fx("🌟 ثانياً: السهام الفلكية السرية"), size=19, weight="bold", color=colors["gold"]))
        parts = [
            ("سهم السعادة", data.get('fortune_part'), "نقطة الرزق المادي والقبول الاجتماعي والبهجة النفسية."),
            ("سهم الغيب", data.get('spirit_part'), "نقطة الإرادة الباطنة والنوايا الروحية والذكاء العميق.")
        ]
        for p_title, p_val, p_desc in parts:
            report_widgets.append(
                ft.Container(
                    content=ft.Column([
                        ft.Text(self.fx(f"📍 {p_title}: {p_val}"), color=colors["sky"], weight="bold", size=16),
                        ft.Text(self.fx(f"💡 الشرح: {p_desc}"), color=colors["info"], size=14),
                    ], spacing=5),
                    padding=12, border=ft.border.all(1, colors["info"]), border_radius=10, margin=ft.margin.only(bottom=10)
                )
            )

        # --- القسم الرابع: الفردارية (حاكم العمر) ---
        report_widgets.append(ft.Text(self.fx("⏳ ثالثاً: حاكم المرحلة الحالية (الفردارية)"), size=19, weight="bold", color=colors["gold"]))
        report_widgets.append(
            ft.Container(
                content=ft.Column([
                    ft.Text(self.fx(f"👑 الحاكم المسيطر: {data.get('firdar_ruler')} - {data.get('firdar_title', '')}"), 
                            color=colors["sky"], weight="bold", size=16),
                    ft.Text(self.fx(f"📝 شرح المرحلة: {data.get('firdar_desc', 'مرحلة هامة لتشكيل المصير.')}"), 
                            color=colors["info"], size=14, italic=True),
                ], spacing=5),
                padding=12, border=ft.border.all(1, colors["gold"]), border_radius=10, margin=ft.margin.only(bottom=10)
            )
        )

        # --- القسم الخامس: تقرير النجوم الثابتة ---
        report_widgets.append(ft.Text(self.fx("🔭 رابعاً: تقرير النجوم الثابتة (أحكام الولادة)"), size=19, weight="bold", color=colors["gold"]))
        stars_rep = data.get('stars_report', "")
        star_items = [ft.Text(self.fx("💡 المنهجية: تحليل مواقع النجوم لحظة ولادتك الشخصية."), size=12, italic=True, color=colors["info"])]
        
        if not stars_rep or "لا يوجد" in stars_rep:
            star_items.append(ft.Text(self.fx("🔍 النتيجة: لم يتم العثور على اقترانات دقيقة بالنجوم العظمى."), color="#94a3b8"))
        else:
            for line in stars_rep.split('\n'):
                if line.strip():
                    star_items.append(ft.Text(self.fx(line), color=colors["sky"] if "⭐" in line else "#f8fafc", size=14))

        report_widgets.append(
            ft.Container(content=ft.Column(star_items, spacing=5), padding=12, border=ft.border.all(1, colors["border"]), border_radius=10, margin=ft.margin.only(bottom=10))
        )

        # --- القسم السادس: القواطع والتحذيرات ---
        report_widgets.append(ft.Text(self.fx("⚠️ خامساً: القواطع والتحذيرات الفلكية"), size=19, weight="bold", color=colors["gold"]))
        warnings = data.get('warnings', "لا يوجد قواطع شديدة.")
        
        warning_items = []
        for line in warnings.split('\n'):
            if line.strip():
                # إضافة شرح فلسفي تلقائي إذا وجد الكلمة المفتاحية في التحذير
                explanation = ""
                for key, val in deep_interpretations.items():
                    if key in line:
                        explanation = f"\n   ◈ فلسفة الموقع: {val}"
                
                warning_items.append(
                    ft.Text(self.fx(f"• {line}{explanation}"), color=colors["danger"], size=14)
                )

        # --- (الجزء الخاص بـ warn_list كما كتبته أنت صحيح جداً) ---
        warn_list = []
        if "لا يوجد" in warnings:
            warn_list.append(ft.Text(self.fx("✅ لا يوجد قواطع شديدة في هذه الهيئة."), color=colors["info"], weight="bold"))
        else:
            for line in warnings.split('\n'):
                if line.strip():
                    warn_list.append(ft.Text(self.fx(f"🛑 {line}"), color=colors["danger"], weight="bold", size=14))
                    interp = next((v for k, v in deep_interpretations.items() if k in line), "يتطلب هذا الوضع انتباهاً لإعادة التوازن.")
                    warn_list.append(ft.Text(self.fx(f"   🔍 التفسير: {interp}"), color=colors["info"], size=13))

        report_widgets.append(
            ft.Container(content=ft.Column(warn_list, spacing=5), padding=12, border=ft.border.all(1, colors["danger"]), border_radius=10, margin=ft.margin.only(bottom=10))
        )

        # --- التذييل والخاتمة ---
        report_widgets.append(ft.Divider(color=colors["gold"]))
        report_widgets.append(ft.Text(self.fx("إعداد وتطوير: حسان الشاعر © 2026"), color=colors["gold"], size=14, text_align="center", weight="bold"))

        # --- وظائف النافذة المنبثقة ---
        def close_dlg(e):
            self.page.close(dlg) # الطريقة الأحدث لغلق النافذة في Flet
            self.page.update()

        def copy_full_report(e):
            # دالة ذكية لتجميع كافة نصوص التقرير من الـ widgets لنسخها
            all_text = ""
            for widget in report_widgets:
                if isinstance(widget, ft.Text):
                    all_text += widget.value + "\n"
                elif isinstance(widget, ft.Container) and isinstance(widget.content, ft.Column):
                    for ctrl in widget.content.controls:
                        if isinstance(ctrl, ft.Text):
                            all_text += ctrl.value + "\n"
            
            self.page.set_clipboard(all_text)
            self.page.snack_bar = ft.SnackBar(ft.Text(self.fx("✅ تم نسخ التقرير الملحمي كاملاً!")), bgcolor="green")
            self.page.snack_bar.open = True
            self.page.update()

        # بناء النافذة النهائية
        dlg = ft.AlertDialog(
            title=ft.Text(self.fx("الديوان التحليلي السلطاني"), color=colors["gold"], weight="bold"),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ALWAYS, spacing=15),
                width=450, # عرض مناسب لشاشات الجوال
                height=800, # طول يسمح بالتمرير
                bgcolor="#0f172a", 
                padding=15,
                border_radius=15
            ),
            actions=[
                ft.ElevatedButton(self.fx("📋 نسخ التقرير"), on_click=copy_full_report, bgcolor=colors["gold"], color="black"),
                ft.TextButton(self.fx("إغلاق"), on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.open(dlg) # فتح النافذة بنظام أندرويد المستقر
        self.page.update()
def display_final_report(self, report_widgets, full_text_content):
        """عرض التقرير النهائي مع تأمين الحفظ والمشاركة للأندرويد"""
        import os, time, platform

        # 1. دالة الحفظ الذكي (متوافقة مع قيود أندرويد 11+)
        def save_to_word_internal(e):
            try:
                from docx import Document
                doc = Document()
                # إضافة محتوى التقرير للملف
                doc.add_heading('تقرير الميقاتي برو الموسع', 0)
                doc.add_paragraph(full_text_content)
                
                file_name = f"Meeqat_Report_{int(time.time())}.docx"
                
                # تحديد مسار الحفظ (أندرويد يمنع الوصول المباشر أحياناً، لذا نستخدم مسارات بديلة)
                if platform.system() == "Android":
                    save_path = "/storage/emulated/0/Download"
                    if not os.path.exists(save_path):
                        save_path = os.environ.get("PYTHONHOME", os.getcwd())
                else:
                    save_path = os.path.expanduser("~/Downloads")
                
                full_path = os.path.join(save_path, file_name)
                doc.save(full_path)
                
                # استخدام SnackBar الحديث لضمان الظهور في أندرويد
                self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"✅ تم الحفظ في: {file_name}")), bgcolor="green")
                self.page.snack_bar.open = True
            except Exception as ex:
                self.page.snack_bar = ft.SnackBar(ft.Text(self.fx(f"❌ خطأ: يرجى منح صلاحية الملفات")), bgcolor="red")
                self.page.snack_bar.open = True
            self.page.update()

        # 2. بناء النافذة المنبثقة بتنسيق متجاوب (Responsive)
        dlg = ft.AlertDialog(
            title=ft.Text(self.fx("الديوان التحليلي السلطاني"), color="#f59e0b", weight="bold"),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ADAPTIVE, spacing=15),
                width=450, # عرض مثالي لشاشات الجوال
                height=800, # طول يسمح بعرض المحتوى الملحمي
                bgcolor="#0f172a", 
                padding=15,
                border_radius=15
            ),
            actions=[
                ft.ElevatedButton(self.fx("💾 حفظ Word"), on_click=save_to_word_internal, bgcolor="#10b981", color="white"),
                ft.ElevatedButton(self.fx("📋 نسخ"), icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(full_text_content)),
                ft.TextButton(self.fx("إغلاق"), on_click=lambda _: self.page.close(dlg))
            ],
            actions_alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        )

        self.page.open(dlg)
        self.page.update()
import flet as ft
import hashlib
import os
import platform
import uuid
from datetime import datetime

# --- إعدادات الحماية والأمان ---
SECRET_SALT = "MEEQAT_PRO_2026_ULTIMATE"
TRIAL_DAYS = 3
KEY_FILE = ".lic_key.dat" # ملف مخفي للترخيص

# 1. دالة جلب بصمة الجهاز (HWID)
def get_android_hwid(page: ft.Page):
    if not page.client_storage.contains_key("hwid"):
        seed = f"{platform.machine() or 'MobileDevice'}-{uuid.uuid4().hex[:8]}-{SECRET_SALT}"
        hwid = hashlib.md5(seed.encode()).hexdigest()[:16].upper()
        page.client_storage.set("hwid", hwid)
    return page.client_storage.get("hwid")

# 2. التحقق من حالة النسخة (كاملة أم لا)
def is_full_version(page: ft.Page):
    hwid = get_android_hwid(page)
    correct_key = hashlib.sha256((hwid + SECRET_SALT).encode()).hexdigest()[:12].upper()
    
    # التحقق من الملف
    if os.path.exists(KEY_FILE):
        with open(KEY_FILE, "r") as f:
            if f.read().strip() == correct_key: return True
            
    # التحقق من ذاكرة التطبيق
    if page.client_storage.get("license_key") == correct_key:
        return True
    return False

# 3. حساب الأيام المتبقية
def get_remaining_days(page: ft.Page):
    today = datetime.now()
    if not page.client_storage.contains_key("first_run"):
        page.client_storage.set("first_run", today.strftime("%Y-%m-%d"))
        return TRIAL_DAYS
    try:
        first_run = datetime.strptime(page.client_storage.get("first_run"), "%Y-%m-%d")
        elapsed = (today - first_run).days
        return max(0, TRIAL_DAYS - elapsed)
    except: return 0

# --- الواجهة الرئيسية (Main Entry) ---
def main(page: ft.Page):
    page.title = "الميقاتي برو 2026"
    page.rtl = True
    page.theme_mode = ft.ThemeMode.DARK
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    hwid = get_android_hwid(page)
    remaining = get_remaining_days(page)
    is_active = is_full_version(page)

    # دالة الدخول للمحرك الفلكي (النتائج الـ 2500)
    def open_astro_logic(e=None):
        page.clean()
        page.vertical_alignment = ft.MainAxisAlignment.START
        page.horizontal_alignment = ft.CrossAxisAlignment.START
        
        astro_list = ft.ListView(expand=True, spacing=10, padding=20)
        # توليد المعطيات الفلكية الخاصة بك
        for i in range(1, 2501):
            astro_list.controls.append(
                ft.Container(
                    content=ft.Text(f"💠 المعطى الفلكي {i}: نتيجة الحساب الدقيقة لعام 2026", size=14),
                    padding=10,
                    bgcolor=ft.colors.WHITE10,
                    border_radius=8
                )
            )

        page.add(
            ft.AppBar(title=ft.Text("المحرك الفلكي الرئيسي"), bgcolor="#1e293b"),
            ft.Container(
                content=ft.Column([
                    ft.Text("📊 نتائج الحسابات الموسعة (2500 معطى):", size=18, weight="bold", color="#38bdf8"),
                    ft.Divider(),
                    astro_list
                ], expand=True),
                padding=20, expand=True
            )
        )
        page.update()

    # شاشة التفعيل
    # --- 1. تعريف حقل الإدخال أولاً ليراه الكود في الدوال ---
    key_input = ft.TextField(
        label="أدخل كود التفعيل", 
        width=300, 
        text_align=ft.TextAlign.CENTER, 
        password=True,
        can_reveal_password=True
    )

    # --- 2. دالة معالجة زر التفعيل (الموحدة) ---
    def on_activate(e):
        # استخدام المنطق الموحد من كلاس الحماية
        correct_key = SecurityManager.generate_correct_key(hwid)
        user_key = key_input.value.strip().upper()
        
        if user_key == correct_key:
            # تخزين مفتاح التفعيل في الجهاز وفي ملف مخفي
            page.client_storage.set("license_key", correct_key)
            try:
                with open(KEY_FILE, "w") as f: f.write(correct_key)
            except: pass 

            page.snack_bar = ft.SnackBar(ft.Text("🚀 تم تفعيل النسخة الكاملة بنجاح!"), bgcolor="green")
            page.snack_bar.open = True
            page.update()
            open_astro_logic() # الدخول للمحرك الفلكي
        else:
            key_input.error_text = "كود التفعيل غير صحيح!"
            page.update()

    # --- 3. القرار النهائي لعرض الواجهة ---
    page.clean() 
    
    if is_active or remaining > 0:
        # حالة الدخول المباشر (مفعل أو فترة تجريبية)
        status_msg = "النسخة مفعلة ✅" if is_active else f"فترة تجريبية: {remaining} يوم متبقي"
        page.add(
            ft.Column([
                ft.Icon(ft.icons.AUTO_AWESOME, size=80, color="amber"),
                ft.Text("الميقاتي برو 2026", size=30, weight="bold"),
                ft.Text(status_msg, color="green" if is_active else "orange", weight="bold"),
                ft.ElevatedButton(
                    text="الدخول للمحرك الفلكي", 
                    icon=ft.icons.PLAY_ARROW, 
                    on_click=open_astro_logic,
                    width=250,
                    height=50
                ),
                ft.Text(f"ID: {hwid}", size=10, color="gray")
            ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=20)
        )
    else:
        # حالة انتهاء الفترة التجريبية (شاشة التفعيل الإجبارية)
        page.add(
            ft.Column([
                ft.Icon(ft.icons.LOCK_PERSON, size=80, color="#fbbf24"),
                ft.Text("نظام حماية الميقاتي برو", size=24, weight="bold"),
                ft.Container(
                    content=ft.Column([
                        ft.Text("انتهت الفترة التجريبية 🔒", color="red", weight="bold"),
                        ft.Text("انسخ البصمة وإرسالها للمطور:", size=14),
                        ft.Text(f"{hwid}", selectable=True, weight="bold", color="#38bdf8", size=16),
                        ft.Divider(),
                        key_input,
                        ft.ElevatedButton(
                            text="تفعيل النسخة الكاملة", 
                            on_click=on_activate, 
                            bgcolor="#10b981", 
                            color="white",
                            width=250
                        ),
                        ft.Text("جميع الحقوق محفوظة للمطور حسان الشاعر © 2026", 
                                size=10, italic=True, color="gray")
                    ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=15),
                    padding=20, border=ft.border.all(1, "#334155"), border_radius=15, bgcolor="#1e293b"
                )
            ], horizontal_alignment=ft.CrossAxisAlignment.CENTER)
        )

# --- نقطة الانطلاق النهائية ---
if __name__ == "__main__":
    ft.app(target=main)
