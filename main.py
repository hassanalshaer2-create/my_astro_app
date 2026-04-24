# 1. المكتبات الأساسية للنظام وواجهة المستخدم
import flet as ft
import hashlib
import os
import platform
import math
import uuid
import time
from datetime import datetime, timezone

# 2. المكتبات الفلكية والجغرافية (تأكد من وجودها في requirements.txt)
import swisseph as swe
from ummalqura.hijri_date import HijriDate
from geopy.geocoders import Nominatim

# 3. مكتبات معالجة اللغة العربية وتنسيق النصوص
import arabic_reshaper
from bidi.algorithm import get_display

try:
    from docx import Document
except ImportError:
    # هذا السطر يحمي التطبيق من الانهيار إذا لم يجد المكتبة
    Document = None 

# ---------------------------------------------------------
# إعدادات الحماية والتشغيل (المتغيرات العامة)
# ---------------------------------------------------------
SECRET_SALT = "MEEQAT_PRO_2026_ULTIMATE"
TRIAL_DAYS = 3
# ملاحظة للأندرويد: لا تستخدم os.getcwd() كمسار حفظ نهائي، اعتمد على page.client_storage
# بدلاً من المسار القديم، نستخدم هذا المسار المتوافق مع الموبايل والويندوز معاً
if platform.system() == "Android":
    DATA_FILE = os.path.join(os.getcwd(), ".hassan_v15_final.dat")
else:
    DATA_FILE = os.path.join(os.path.expanduser("~"), ".hassan_v15_final.dat")
PLANET_MEANINGS = {
    "الشمس": "جوهر الذات، الإرادة الحيوية، والقدرة على القيادة والتألق،",
    "القمر": "الاحتياجات العاطفية، العقل الباطن، وكيفية الشعور بالأمان،",
    "عطارد": "الذكاء التحليلي، أسلوب التواصل، وقدرة العقل على معالجة البيانات،",
    "الزهرة": "التناغم الاجتماعي، القيم المادية، وطريقة التعبير عن الحب والجمال،",
    "المريخ": "طاقة المبادرة، الشجاعة، وكيفية توجيه القوة والاندفاع نحو الأهداف،",
    "المشتري": "فرص التوسع، الحظ، الفلسفة العالية وقدرة الشخص على النمو،",
    "زحل": "هيكلية الانضباط، الدروس الكرمية، والمسؤوليات التي تتطلب صبراً،",
    "أورانوس": "شرارة العبقرية، التغيير المفاجئ، والميل نحو التحرر والابتكار،",
    "نبتون": "الخيال الروحي، الحدس العميق، والقدرة على التماهي مع الفنون والروحانيات،",
    "بلوتو": "قوة التحول الجذري، تدمير القديم لبناء الجديد، والقدرة على النهوض،",
    "العقدة الشمالية": "بوصلة الروح والغاية القصوى التي يجب على الإنسان تعلمها في حياته،",
    "العقدة الجنوبية": "الرصيد الفطري والمهارات التي ولدت بها ولكنها قد تعيقك إذا تمسكت بها،"
}
# تأثير البروج (الصبغة أو الأسلوب)
SIGN_FLAVORS = {
    "الحمل": "بأسلوب ناري، مبادر، ويميل للاندفاع والشجاعة المطلقة.",
    "الثور": "بأسلوب ترابي ثابت، يميل للواقعية، الصبر، والبحث عن الاستقرار المادي.",
    "الجوزاء": "بأسلوب هوائي مرن، فضولي، ويميل للتواصل المتعدد والذكاء الاجتماعي.",
    "السرطان": "بأسلوب مائي عاطفي، يميل للحماية، الرعاية، والارتباط العميق بالجذور.",
    "الأسد": "بأسلوب ناري قيادي، يميل للظهور، الثقة بالنفس، وطلب التقدير والتميز.",
    "العذراء": "بأسلوب ترابي تحليلي، يميل للدقة، التنظيم، والبحث عن الكمال والخدمة.",
    "الميزان": "بأسلوب هوائي دبلوماسي، يميل للتوازن، العدل، والبحث عن الجمال والشراكة.",
    "العقرب": "بأسلوب مائي مكثف، يميل للعمق، التحول، والقوة الخفية والإرادة الصلبة.",
    "القوس": "بأسلوب ناري استكشافي، يميل للتفاؤل، الفلسفة، وحب الحرية والمعرفة العالية.",
    "الجدي": "بأسلوب ترابي طموح، يميل للجدية، الانضباط، وبناء المكانة الاجتماعية المرموقة.",
    "الدلو": "بأسلوب هوائي ابتكاري، يميل للتحرر، التفكير الإنساني، والتمرد على القوالب التقليدية.",
    "الحوت": "بأسلوب مائي روحاني، يميل للخيال، التضحية، والتماهي العميق مع المشاعر والكون."
}

# دلالات البيوت (مجال التجربة الحياتية)
HOUSE_FIELDS = {
    "1": "في بناء الهوية الشخصية، المظهر الخارجي، وكيفية مواجهة العالم لأول مرة.",
    "2": "في ميدان الموارد المالية، القيم الذاتية، والممتلكات والقدرة على الكسب.",
    "3": "في مجال التواصل اليومي، الإخوة، الدراسات الأولية، والذكاء العملي والأسفار القصيرة.",
    "4": "في أعماق الحياة الخاصة، المنزل، الجذور العائلية، والارتباط بالأرض والوالدين.",
    "5": "في عالم الإبداع، التعبير عن الذات، المواهب الفنية، العواطف، والأطفال.",
    "6": "في مجال العمل الروتيني، الصحة الجسدية، تطوير المهارات، والخدمة والواجبات.",
    "7": "في ميدان الشراكات الرسمية، الزواج، العلاقات المتوازنة، والأعداء الظاهرين.",
    "8": "في مناطق التحول الجذري، المواريث، المال المشترك، الأسرار، والعلوم الباطنية.",
    "9": "في مجال الوعي العالي، السفر البعيد، التعليم الجامعي، الفلسفة، والقانون.",
    "10": "في قمة الطموح المهني، السمعة العامة، السلطة، والنجاح في العالم الخارجي.",
    "11": "في دائرة الأصدقاء، المجموعات، الآمال الكبيرة، والمشاريع الجماعية الإنسانية.",
    "12": "في عالم الخفاء، الروحانيات، العزلة، والتعامل مع الأعداء المخفيين والماضي الكرمي."
}
ASTRO_INTERPRETATIONS = {
    "fortune_part": {
        "الشمس": "وفرة مالية عبر المناصب والظهور الاجتماعي.",
        "القمر": "رزق يأتي من التجارة، السفر، أو الدعم العائلي.",
        "المريخ": "مكاسب تأتي من الجهد البدني، الشجاعة، أو المبادرة.",
        # أضف باقي الكواكب...
    },
    "firdar_ruler": {
        "الشمس": "فترة للبروز، تحقيق الذات، وريادة الأعمال.",
        "الزهرة": "فترة ذهبية للزواج، الفنون، وتحسين العلاقات الاجتماعية.",
        "زحل": "فترة تتطلب الصبر، العمل الجاد، وبناء الأسس المتينة.",
        # أضف باقي الكواكب...
    }
}
PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]

PLANETARY_HOURS_INFO = {
    "الشمس": "ساعة للظهور، طلب المناصب، التعامل مع المسؤولين، والبدايات التي تحتاج شهرة.",
    "الزهرة": "ساعة للجمال، الزواج، التصالح، الفنون، وشراء الأشياء الثمينة.",
    "عطارد": "ساعة للتجارة، الكتابة، الدراسة، المراسلات، والعمليات الحسابية والذكاء.",
    "القمر": "ساعة للتغيرات، التنقل، التعامل مع الجمهور، الأمور العائلية، والعواطف.",
    "زحل": "ساعة للعمل الشاق، البناء، العقارات، التعامل مع كبار السن، وتتطلب الصبر.",
    "المشتري": "ساعة الحظ، التوسع، طلب العلم، الأمور الدينية، والنمو المالي.",
    "المريخ": "ساعة القوة، الرياضة، الجراحة، المواجهة، ويُنصح فيها بالحذر من الاندفاع."
}
# الترتيب الكلداني للكواكب
PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]
# حكام الأيام (يبدأ اليوم الفلكي من الشروق)
DAY_RULERS = {
    "Sunday": "الشمس", "Monday": "القمر", "Tuesday": "المريخ", 
    "Wednesday": "عطارد", "Thursday": "المشتري", "Friday": "الزهرة", "Saturday": "زحل"
}

PLANETARY_WARNINGS_DB = {
    # --- التراجعات (Retrogrades) ---
    "retro_mercury": "⚠️ تراجع عطارد: فساد في الفكر والرسائل، تعطل الأسفار القريبة، وخلل في البيع والشراء.",
    "retro_venus": "⚠️ تراجع الزهرة: فساد في النكاح، نفور بين المحبين، ومراجعة في أمور الزينة والمال.",
    "retro_mars": "⚠️ تراجع المريخ: حدة في الغضب، تعطل المبادرات، ووهن في العزيمة البدنية.",
    "retro_jupiter": "⚠️ تراجع المشتري: نقص في البركة، ضيق في الرزق الواسع، ومراجعة للقضايا القانونية.",
    "retro_saturn": "⚠️ تراجع زحل: ثقل في المسؤوليات، تأخير في البناء والعقارات، وعودة ديون قديمة.",

    # --- القواطع والاقترانات الصعبة ---
    "malefic_conjunction": "🚨 اقتران النحسين (زحل والمريخ): اجتماع 'القاطع' و'المانع'. شدة وضيق تتطلب صبراً عظيماً.",
    "saturn_vital": "⚠️ حصار زحل لكوكب حيوي: تقييد، حزن، أو تأخير قسري يطال شؤون {planet}.",
    "mars_vital": "⚠️ احتكاك المريخ بكوكب حيوي: تهور، اندفاع، أو خطر التهاب ونزاع يخص {planet}.",
    "mars_opposition": "⚡ مقابلة المريخ: مواجهة مباشرة مع الأعداء، أو تهور يؤدي لنتائج عكسية.",

    # --- الحالات الفنية (احتراق، خلو مسار، درجات حرجة) ---
    "combust": "☄️ احتراق {planet}: الكوكب في 'حيز الشمس'؛ ضعفت قوته واختفت دلالته تحت الشعاع.",
    "void_of_course": "🌑 القمر خالي المسار: سعي بلا طائل. تجنب البدء بأي أمر ذي بال حتى ينتقل القمر.",
    "critical_degree": "📐 درجة حرجة (0 أو 29): الكوكب في حالة 'المنقلب'؛ نهاية مرحلة وبداية تحول مفاجئ.",
    "saturn_return": "🪐 عودة زحل: ساعة المحاسبة الكبرى. اختبار للنضج وإعادة هيكلة أساسات الحياة.",

    # --- النجوم الثابتة (Fixed Stars) ---
    "star_algol": "💀 اقتران رأس الغول (Algol): أشر النجوم؛ خطر فقدان السيطرة، أو قرارات مدمرة للعقل.",
    "star_aldebaran": "🌟 اقتران الدبران: نجاح عظيم محفوف بمخاطر السقوط الأخلاقي أو النزاعات القانونية.",
    "star_regulus": "🦁 اقتران قلب الأسد (Regulus): ارتقاء للمناصب العالية، مع خطر الانتقام من الأنداد.",
    "star_fomalhaut": "🌟 اقتران فم الحوت: نجاح في العلوم الروحية والفنون، مع خطر الوقوع في الأوهام.",
    "star_deneb_kaitos": "🌊 اقتران ذنب قيطس: دلالة على العزلة الإجبارية، التعب الجسدي، أو الحزن المفاجئ.",
    "star_antares": "🌟 اقتران قلب العقرب (Antares): تهور قتالي، نجاح سريع متبوع بنهاية مفاجئة.",

    # --- قواطع كتب السابقين ---
    "besieged": "🔒 كوكب محاصر بين النحسين: الكوكب {planet} في ضيق شديد، دلالته معطلة تماماً.",
    "hayz_violation": "🚫 فساد الهيئة (خروج عن الحيز): الكوكب يعمل عكس طبيعته، مما يسبب نتائج غير متوقعة."
}
FIRDAR_INTERPRETATIONS = {
    "الشمس": {
        "title": "فترة الظهور والسيادة (10 سنوات)",
        "desc": "مرحلة تتركز حول تحقيق الذات، الطموح المهني، والتعامل مع أصحاب السلطة. فترة جيدة للبروز الاجتماعي."
    },
    "الزهرة": {
        "title": "فترة الجمال والمال (8 سنوات)",
        "desc": "مرحلة ذهبية للعلاقات العاطفية، الزواج، الفنون، وتحسن الدخل المادي والرفاهية."
    },
    "عطارد": {
        "title": "فترة العلم والتجارة (13 سنة)",
        "desc": "نشاط ذهني مكثف، كثرة الأسفار، نجاح في الدراسة أو التجارة والكتابة وتوسيع شبكة التواصل."
    },
    "القمر": {
        "title": "فترة المتغيرات والعائلة (9 سنوات)",
        "desc": "تركز على الاستقرار المنزلي، المشاعر، والتغيرات في السكن أو الحالة العائلية."
    },
    "زحل": {
        "title": "فترة الحصاد والمسؤولية (11 سنة)",
        "desc": "مرحلة تتطلب الصبر والعمل الجاد. بناء أسس طويلة الأمد أو مواجهة تحديات تتطلب نضجاً."
    },
    "المشتري": {
        "title": "فترة الوفرة والحظ (12 سنة)",
        "desc": "توسع في الرزق، تفاؤل، نجاحات قانونية أو دينية، وفترة مناسبة للسفر البعيد والنمو."
    },
    "المريخ": {
        "title": "فترة القوة والصراع (7 سنوات)",
        "desc": "طاقة عالية، اندفاع، مبادرات عملية شجاعة، لكنها تتطلب الحذر من الخلافات أو التهور الجسدي."
    },
    "الرأس": {
        "title": "فترة الصعود المفاجئ (3 سنوات)",
        "desc": "تغييرات قدرية، طموحات لا محدودة، وغالباً ما تشهد قفزة نوعية في مكانة الشخص."
    },
    "الذنب": {
        "title": "فترة التطهير والروحانية (سنتان)",
        "desc": "فترة زهد أو مراجعة للماضي، تخلص من علاقات أو أشياء لم تعد تخدمك، والتركيز على الباطن."
    }
}
LOTS_INTERPRETATIONS = {
    "fortune": "معدن الرزق، الحظ الجسدي، والمال الحلال وصحة البدن.",
    "spirit": "معدن القوة الروحية، النوايا، الطموح، وما يخطط له العقل.",
    "love": "سهم الزهرة: يتعلق بالعواطف، الصداقات، والمتعة والحياة الاجتماعية.",
    "necessity": "سهم الشدة: يتعلق بالقيود، المتاعب، وما يُفرض على الإنسان قسراً.",
    "victory": "سهم النصر: يتعلق بالنجاح في المنازعات، القوة، والوصول للأهداف.",
    "marriage": "سهم الزواج: يحدد جودة الشراكة العاطفية واستقرار الحياة الزوجية.",
    "enemies": "سهم الأعداء: يكشف عن الخصوم المخفيين والمكائد التي قد تحاك."
}

def check_comprehensive_warnings(data):
    alerts = []
    
    # 1. التراجعات
    retro_map = {'عطارد': 'mercury', 'الزهرة': 'venus', 'المريخ': 'mars', 'المشتري': 'jupiter', 'زحل': 'saturn'}
    for ar_p, en_p in retro_map.items():
        if data.get(f'is_{en_p}_retro'):
            alerts.append(PLANETARY_WARNINGS_DB[f"retro_{en_p}"])

    # 2. النحسين (زحل والمريخ)
    if data.get('mars_saturn_dist', 100) < 6:
        alerts.append(PLANETARY_WARNINGS_DB["malefic_conjunction"])

    # 3. الاحتراق (أقل من 8 درجات من الشمس)
    for p, dist in data.get('sun_distances', {}).items():
        if dist < 8.5:
            alerts.append(PLANETARY_WARNINGS_DB["combust"].format(planet=p))

    # 4. النجوم الثابتة (أورب 1 درجة)
    for star in ['algol', 'aldebaran', 'regulus', 'fomalhaut', 'deneb_kaitos', 'antares']:
        if data.get(f'conjunct_{star}'):
            alerts.append(PLANETARY_WARNINGS_DB[f"star_{star}"])

    # 5. خلو المسار والدرجات الحرجة
    if data.get('is_moon_voc'):
        alerts.append(PLANETARY_WARNINGS_DB["void_of_course"])
    
    for p, deg in data.get('positions', {}).items():
        if deg % 30 in [0, 29]:
            alerts.append(f"📍 {p}: {PLANETARY_WARNINGS_DB['critical_degree']}")

    return "\n".join(alerts) if alerts else "✅ الهيئة متزنة ولا توجد قواطع سلبية حالياً."
def open_astro_logic(e):
    page.clean()
    
    # 1. تجهيز البيانات (مثال)
    my_data = {
        "is_mercury_retro": True,
        "sun_distances": {"الزهرة": 5.2},
        "positions": {"القمر": 29.5}
    }
    
    # 2. استدعاء دالة التحذيرات التي كتبتها أنت
    result_text = check_comprehensive_warnings(my_data)
    
    # 3. عرض النتيجة للمستخدم (هذا هو السطر المقصود)
    page.add(
        ft.AppBar(title=ft.Text("نتائج فحص الهيئة")),
        ft.Container(
            content=ft.Text(
                result_text, 
                rtl=True, 
                size=18,
                text_align=ft.TextAlign.RIGHT
            ),
            padding=20
        )
    )
    page.update()

def fx(text):
    """إرجاع النص مباشرة مع دعم Flet الأصلي للعربية"""
    if not text: 
        return ""
    return str(text)

def get_hwid(page: ft.Page):
    # الفحص إذا كان هناك بصمة مسجلة مسبقاً في ذاكرة الجهاز
    if not page.client_storage.contains_key("user_hwid"):
        # إنشاء بصمة عشوائية فريدة لأول مرة فقط
        import uuid
        new_id = str(uuid.uuid4())[:12].upper()
        page.client_storage.set("user_hwid", new_id)
    
    return page.client_storage.get("user_hwid")

def generate_activation_key(hwid):
    """توليد كود التفعيل بناءً على بصمة الجهاز"""
    raw_key = hwid + SECRET_SALT
    return hashlib.sha256(raw_key.encode()).hexdigest()[:12].upper()

def get_trial_info(page: ft.Page):
    """حساب الأيام المتبقية باستخدام ذاكرة التطبيق (أضمن للأندرويد)"""
    today = datetime.now().date()
    
    # التحقق من تاريخ أول تشغيل
    if not page.client_storage.contains_key("first_run_date"):
        page.client_storage.set("first_run_date", today.strftime("%Y-%m-%d"))
        return TRIAL_DAYS
    
    try:
        first_date_str = page.client_storage.get("first_run_date")
        first_date = datetime.strptime(first_date_str, "%Y-%m-%d").date()
        elapsed = (today - first_date).days
        return max(0, TRIAL_DAYS - elapsed)
    except:
        return 0

import flet as ft
import time
import hashlib

# 1. شاشة البداية
def show_splash(page: ft.Page):
    page.clean()
    page.bgcolor = "black"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    
    logo = ft.Text("✧ ASC ✧", size=70, weight="bold", color="#FFD700")
    title = ft.Text("MEEQAT ULTIMATE PRO", size=30, weight="bold", color="#39FF14")
    progress = ft.ProgressBar(width=400, color="#FFD700", value=0)
    
    page.add(logo, title, ft.Text("By Hassan Al-Shaer", color="#BB86FC"), progress)
    
    for i in range(1, 11):
        progress.value = i / 10
        page.update()
        time.sleep(0.15)
    
    # بعد انتهاء التحميل ننتقل لنظام التفعيل
    run_auth_system(page)

# 2. نظام التفعيل (متوافق تماماً مع أندرويد)
def run_auth_system(page: ft.Page):
    hwid = get_hwid(page)
    valid_key = generate_activation_key(hwid)
    days_left = get_trial_info(page)

    page.clean()
    page.bgcolor = ft.colors.BLACK
    
    # حقل النص
    entry = ft.TextField(
        label="أدخل كود التفعيل", 
        width=300, 
        password=True, 
        can_reveal_password=True,
        text_align=ft.TextAlign.CENTER
    )

    # دالة اللصق للأندرويد
    def do_paste(e):
        # جلب النص من حافظة الموبايل ووضعه في الخانة
        entry.value = page.get_clipboard()
        page.update()

    def on_verify(e):
        if entry.value.upper().strip() == valid_key:
            page.client_storage.set("license_key", valid_key)
            open_astro_logic(page) # دالة فتح البرنامج
        else:
            page.snack_bar = ft.SnackBar(ft.Text("كود تفعيل خاطئ"))
            page.snack_bar.open = True
            page.update()

    def on_trial(e):
        if days_left > 0:
            open_astro_logic(page)
        else:
            page.snack_bar = ft.SnackBar(ft.Text("انتهت الفترة التجريبية"))
            page.snack_bar.open = True
            page.update()

    # بناء الواجهة
    page.add(
        ft.Text("تفعيل الميقاتي برو", size=25, weight="bold", color="white"),
        ft.Text(f"ID: {hwid}", color="gray", size=12),
        ft.Row([
            entry, 
            ft.IconButton(icon=ft.icons.PASTE, on_click=do_paste, tooltip="لصق الكود")
        ], alignment=ft.MainAxisAlignment.CENTER),
        ft.ElevatedButton("تفعيل النسخة الأصلية", on_click=on_verify, bgcolor="green", color="white"),
        ft.TextButton(f"الدخول كنسخة تجريبية ({days_left} يوم)", on_click=on_trial),
        ft.Text(f"الأيام المتبقية: {days_left}", color="orange")
    )
    page.update()
# ==========================================
# 3. محرك البرنامج الرئيسي (Master Engine)
# ==========================================
class HassanAstroMobile:
    def __init__(self, page: ft.Page, mode):
        self.page = page
        
        # 1. إعدادات العنوان (المتوافقة مع أندرويد)
        version_type = "الكاملة" if mode == "full" else "التجريبية"
        self.page.title = f"Hassan Astro Pro - {version_type}"
        self.page.bgcolor = "#ffffff"
        self.page.rtl = True # دعم العربية
        
        # 2. تعريف المتغيرات (بدل tk.StringVar نستخدم متغيرات عادية)
        self.house_system = "Placidus"
 # 1. أرباب البيوت
        self.domiciles = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", 
                          "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
        
        # 2. أرباب الشرف
        self.exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}

        # 3. أرباب المثلثات (نهار/ليل)
        self.triplicities = {
            "ناري":  {"day": "الشمس", "night": "المشتري"},
            "ترابي": {"day": "الزهرة", "night": "القمر"},
            "هوائي": {"day": "زحل", "night": "عطارد"},
            "مائي":  {"day": "الزهرة", "night": "المريخ"}
        }

        # 4. الحدود المصرية (Egyptian Terms) - الدقة القصوى
        # كل قائمة تحتوي على (الدرجة القصوى للحد، الكوكب صاحب الحد)
        self.egyptian_terms = {
            0:  [(6, "المشتري"), (12, "الزهرة"), (20, "عطارد"), (25, "المريخ"), (30, "زحل")],  # الحمل
            1:  [(8, "الزهرة"), (14, "عطارد"), (22, "المشتري"), (27, "زحل"), (30, "المريخ")],  # الثور
            2:  [(6, "عطارد"), (12, "المشتري"), (17, "الزهرة"), (24, "المريخ"), (30, "زحل")],  # الجوزاء
            3:  [(7, "المريخ"), (13, "الزهرة"), (19, "عطارد"), (26, "المشتري"), (30, "زحل")],  # السرطان
            4:  [(6, "المشتري"), (11, "الزهرة"), (18, "زحل"), (24, "عطارد"), (30, "المريخ")],  # الأسد
            5:  [(7, "عطارد"), (17, "الزهرة"), (21, "المشتري"), (28, "المريخ"), (30, "زحل")],  # العذراء
            6:  [(6, "زحل"), (14, "الزهرة"), (21, "المشتري"), (28, "عطارد"), (30, "المريخ")],  # الميزان
            7:  [(7, "المريخ"), (11, "الزهرة"), (19, "عطارد"), (24, "المشتري"), (30, "زحل")],  # العقرب
            8:  [(12, "المشتري"), (17, "الزهرة"), (21, "عطارد"), (26, "زحل"), (30, "المريخ")],  # القوس
            9:  [(7, "عطارد"), (14, "المشتري"), (22, "الزهرة"), (26, "زحل"), (30, "المريخ")],  # الجدي
            10: [(7, "عطارد"), (13, "الزهرة"), (20, "المشتري"), (25, "المريخ"), (30, "زحل")],  # الدلو
            11: [(12, "الزهرة"), (16, "المشتري"), (19, "عطارد"), (28, "المريخ"), (30, "زحل")]   # الحوت
        }        
        self.z_syms = ["♈", "♉", "♊", "♋", "♌", "♍", "♎", "♏", "♐", "♑", "♒", "♓"]
        self.z_names = ["الحمل ♈", "الثور ♉", "الجوزاء ♊", "السرطان ♋", "الأسد ♌", "العذراء ♍", "الميزان ♎", "العقرب ♏", "القوس ♐", "الجدي ♑", "الدلو ♒", "الحوت ♓"]

        # --- الإضافة الجديدة: إعدادات محرك تفسير الأحلام الفلسفي ---
        # حساب اليوم القمري تقريبياً لربطه بحكمة الإمام الصادق
        import math
        diff = datetime.now() - datetime(2000, 1, 6)
        self.hijri_day = math.floor(((diff.days / 29.5305) - math.floor(diff.days / 29.5305)) * 29.53) + 1
        
        # تعريف فلسفة ابن عربي والكلدانيين كمرجع ثابت
        self.dream_philosophy = {
            "ibn_arabi": "الرؤيا هي اتصال الروح بعالم المثال (الخيال المنفصل).",
            "chaldean": "الأحلام هي انعكاس لنفوذ الكواكب السبعة على النفس البشرية."
        }
        # -------------------------------------------------------

        self.build_flet_ui() # دالة سنقوم بتعريفها لبناء واجهة Flet

    def build_flet_ui(self):
        # تنظيف الصفحة وبناء الواجهة الأساسية
        self.page.clean()
        
        # حاوية النتائج (التي ستستقبل الـ 104 أسطر)
        self.report_container = ft.Column(
            controls=[], 
            scroll=ft.ScrollMode.ALWAYS,
            expand=True
        )

        self.page.add(
            ft.AppBar(
                title=ft.Text("الميقاتي الفلكي - Master Edition"), 
                bgcolor=ft.colors.SURFACE_VARIANT,
                actions=[
                    ft.IconButton(ft.icons.COPY, on_click=self.copy_report) # زر لنسخ التقرير كاملاً
                ]
            ),
            ft.Text(f"📅 اليوم القمري: {self.hijri_day}", size=18, weight="bold", rtl=True),
            ft.Divider(),
            self.report_container,
            ft.FloatingActionButton(
                icon=ft.icons.PLAY_ARROW, 
                text="بدء الحسابات الملحمية",
                on_click=lambda _: self.start_calculations()
            )
        )

    # هذه الدالة يجب أن تكون خارج build_flet_ui ولكن داخل الـ Class
    def log_to_flet(self, message, color="black"):
        """البديل الرسمي لـ txt.insert - متوافق مع أندرويد"""
        self.report_container.controls.append(
            ft.Text(value=message, color=color, size=16, rtl=True, selectable=True)
        )
        self.page.update()

    def copy_report(self, e):
        """دالة لنسخ كافة النتائج المسجلة في الحاوية"""
        full_text = "\n".join([c.value for c in self.report_container.controls if isinstance(c, ft.Text)])
        self.page.set_clipboard(full_text)
        self.page.show_snack_bar(ft.SnackBar(ft.Text("تم نسخ التقرير الفلكي بنجاح!")))

def get_egyptian_bounds(self): # السطر 459 (مثلاً)
    """جدول الحدود المصرية الكامل (أدق تقسيم فلكي تقليدي)""" # السطر 460 (يجب أن يبدأ بـ 4 مسافات)
    bounds_data = { ... } # بقية الكود المزاح أيضاً
    terms = {
            0: [(6, "المشتري"), (12, "الزهرة"), (20, "عطارد"), (25, "المريخ"), (30, "زحل")],
            1: [(8, "الزهرة"), (14, "عطارد"), (22, "المشتري"), (27, "زحل"), (30, "المريخ")],
            2: [(6, "عطارد"), (12, "المشتري"), (17, "الزهرة"), (24, "المريخ"), (30, "زحل")],
            3: [(7, "المريخ"), (13, "الزهرة"), (19, "عطارد"), (26, "المشتري"), (30, "زحل")],
            4: [(6, "المشتري"), (11, "الزهرة"), (18, "زحل"), (24, "عطارد"), (30, "المريخ")],
            5: [(7, "عطارد"), (17, "الزهرة"), (21, "المشتري"), (28, "المريخ"), (30, "زحل")],
            6: [(6, "زحل"), (14, "الزهرة"), (21, "المشتري"), (28, "عطارد"), (30, "المريخ")],
            7: [(7, "المريخ"), (11, "الزهرة"), (19, "عطارد"), (24, "المشتري"), (30, "زحل")],
            8: [(12, "المشتري"), (17, "الزهرة"), (21, "عطارد"), (26, "زحل"), (30, "المريخ")],
            9: [(7, "عطارد"), (14, "المشتري"), (22, "الزهرة"), (26, "زحل"), (30, "المريخ")],
            10:[(7, "عطارد"), (13, "الزهرة"), (20, "المشتري"), (25, "المريخ"), (30, "زحل")],
            11:[(12, "الزهرة"), (16, "المشتري"), (19, "عطارد"), (28, "المريخ"), (30, "زحل")]
        }
    for limit, planet in terms.get(sign_idx, []):
            if deg <= limit: return planet
    return "زحل"

    def calculate_real_almuten(self, target_deg, jd, is_day):
        """حساب المبتز (المستولي) بناءً على الحظوظ الخمسة"""
        domiciles = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
        exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}
        triplicities = {
            "Fire": ["الشمس", "المشتري"], "Earth": ["الزهرة", "القمر"],
            "Air": ["زحل", "عطارد"], "Water": ["الزهرة", "المريخ"]
        }
        
        sign_idx = int(target_deg / 30) % 12
        deg_in_sign = target_deg % 30
        scores = {k: 0 for k in ["الشمس", "القمر", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل"]}
        
        # 1. رب البيت (5 نقاط)
        scores[domiciles[sign_idx]] += 5
        # 2. رب الشرف (4 نقاط)
        if sign_idx in exaltations: scores[exaltations[sign_idx]] += 4
        # 3. رب المثلثة (3 نقاط)
        element = ["Fire", "Earth", "Air", "Water"][sign_idx % 4]
        scores[triplicities[element][0 if is_day else 1]] += 3
        # 4. رب الحد (نقطتان)
        scores[self.get_egyptian_term(sign_idx, deg_in_sign)] += 2
        # 5. رب الوجه (نقطة واحدة)
        faces = ["المريخ", "الشمس", "الزهرة", "عطارد", "القمر", "زحل", "المشتري"]
        scores[faces[int(target_deg / 10) % 7]] += 1

        winner = max(scores, key=scores.get)
        p_map = {"الشمس":0, "القمر":1, "عطارد":2, "الزهرة":3, "المريخ":4, "المشتري":5, "زحل":6}
        p_res = swe.calc_ut(jd, p_map[winner])
        return winner, p_res[0] if isinstance(p_res, (tuple, list)) else p_res     
    def draw_astro_wheel(self, jd, title):
        # تنظيف الإطارات والجداول قبل إعادة الرسم
        for widget in self.chart_frame.winfo_children(): 
            widget.destroy()
        self.tree.delete(*self.tree.get_children())
# استدعاء النظام المختار من القائمة
        selected_system = self.systems_map.get(self.house_system.get(), b'P')
        cusps, ascmc = swe.houses(jd, self.lat, self.lon, selected_system)

        # --- التصحيح: ربط نظام البيوت المختار بالمحرك السويسري ---
        selected_house_system = self.systems_map.get(self.house_system.get(), b'P')
        cusps, ascmc = swe.houses(jd, self.lat, self.lon, selected_house_system)
        
        asc = float(ascmc[0])
        mc = float(ascmc[1])
        dsc, ic = (asc + 180) % 360, (mc + 180) % 360

        # إعداد الرسم بأقصى تكبير
        fig = plt.Figure(figsize=(12, 12), facecolor='white')
        fig.subplots_adjust(left=0, right=1, top=1, bottom=0)
        ax = fig.add_subplot(111, projection='polar')
        ax.set_theta_direction(1)
        ax.set_theta_offset(np.deg2rad(180 - asc))

        # رسم الأوتاد (ASC, DSC, MC, IC) بخطوط عريضة وتفاصيل الدقائق
        angles = {asc: ("ASC", "#ef4444"), dsc: ("DSC", "#ef4444"), mc: ("MC", "#10b981"), ic: ("IC", "#10b981")}
        for deg, (lbl, clr) in angles.items():
            rad = np.deg2rad(deg)
            ax.plot([rad, rad], [3.2, 4.6], color=clr, lw=6, zorder=10)
            
            # حساب الدرجة والدقيقة
            d, m = int(deg % 30), int((deg * 60) % 60)
            ax.text(rad, 4.9, f"{lbl}\n{d}°{m:02d}'", fontsize=11, fontweight='bold', color=clr, ha='center')
            
            # إدراج الدرجة والدقيقة في الجدول الجانبي (اللست)
            self.tree.insert("", "end", values=(lbl, self.z_names[int(deg/30)], f"{d}°{m:02d}'", "-"), tags=(lbl,))
            self.tree.tag_configure(lbl, foreground=clr)

        # رسم حزام البروج الـ 12
        for i in range(12):
            theta = np.deg2rad(i * 30)
            ax.bar(theta, 0.8, width=np.deg2rad(30), bottom=3.2, color=['#fee2e2','#fef3c7','#ecfdf5','#eff6ff'][i%4], edgecolor='#cbd5e1', align='edge')
            ax.text(theta + np.deg2rad(15), 3.6, self.z_syms[i], fontsize=32, ha='center', va='center')

        # الكواكب والسهام
        p_data = {0:("الشمس","☉","#f59e0b"), 1:("القمر","☽","#3b82f6"), 2:("عطارد","☿","#8b5cf6"), 
                  3:("الزهرة","♀","#ec4899"), 4:("المريخ","♂","#dc2626"), 5:("المشتري","♃","#d97706"), 
                  6:("زحل","♄","#475569"), 7:("أورانس","♅","#06b6d4"), 8:("نبتون","♆","#6366f1"), 
                  9:("بلوتو","♇","#a855f7"), 10:("الرأس","☊","#059669")}
  # أضف هذه الرموز داخل المصفوفات في __init__
        self.p_names_extra = {11: ("الذنب", "☋", "#059669"), 12: ("ليلث", "⚸", "#000000"), 13: ("كايرون", " Chiron", "#5b21b6")}
      
        plot_list = []
        for pid, (n, s, c) in p_data.items():
            res, _ = swe.calc_ut(jd, pid)
            pos = float(res[0])
            h = self.get_house(pos, cusps) # حساب رقم البيت لكل كوكب
            plot_list.append((pos, s, c))
            self.tree.insert("", "end", values=(f"{s} {n}", self.z_names[int(pos/30)], f"{int(pos%30)}°{int((pos*60)%60):02d}'", h), tags=(n,))
            self.tree.tag_configure(n, foreground=c)
            if pid == 10: # إضافة العقدة الجنوبية تلقائياً
                plot_list.append(((pos+180)%360, "☋", "#059669"))

        # --- الحسابات الفلكية (سليمة تماماً وتبقى كما هي) ---
        sun_pos_raw, _ = swe.calc_ut(jd, 0)
        moon_res, _ = swe.calc_ut(jd, 1)
        sun, moon = float(sun_pos_raw[0]), float(moon_res[0])
        is_day = not (180 <= (sun - asc) % 360 <= 360) 
        
        fortuna = (asc + moon - sun) % 360 if is_day else (asc + sun - moon) % 360
        spirit = (asc + sun - moon) % 360 if is_day else (asc + moon - sun) % 360

        # --- تحويل العرض من الرسم والجداول إلى قائمة Flet الأنيقة ---
        self.log_to_flet("🪐 وضعية الكواكب والسهام:", color="blue")
        
        # عرض الكواكب
        for pid, (n, s, c) in p_data.items():
            res, _ = swe.calc_ut(jd, pid)
            pos = float(res[0])
            h = self.get_house(pos, cusps)
            sign_name = self.z_names[int(pos/30)]
            deg, mnt = int(pos%30), int((pos*60)%60)
            
            # البديل لـ self.tree.insert
            self.log_to_flet(f"{s} {n}: {sign_name} {deg}°{mnt:02d}' - البيت ({h})", color=c)

        # عرض السهام السرية
        for s_pos, s_name, s_sym in [(fortuna, "سهم السعادة", "⊗"), (spirit, "سهم الغيب", "❂")]:
            h_s = self.get_house(s_pos, cusps)
            sign_s = self.z_names[int(s_pos/30)]
            self.log_to_flet(f"{s_sym} {s_name}: {sign_s} {int(s_pos%30)}° - البيت ({h_s})", color="#f97316")

        # عرض حدود البيوت (بدلاً من الرسم القطبي)
        self.log_to_flet("🏠 حدود البيوت الفلكية:", color="green")
        for i, c in enumerate(cusps):
            self.log_to_flet(f"البيت {i+1}: {int(c/30)} {self.z_names[int(c/30)]} {int(c%30)}°")

        # --- حذف أوامر الرسم التي تسبب الانهيار ---
        # احذف سطر FigureCanvasTkAgg وكل أسطر ax.plot و ax.text

    def draw_now(self):
        n = datetime.now(timezone.utc)
        self.draw_astro_wheel(swe.julday(n.year, n.month, n.day, n.hour + n.minute/60.0), "Time Now")
    def draw_2026(self):
        jd = swe.julday(2026, 3, 20, 9 + 24/60.0) # إذا كنت تقصد 9:24 صباحاً
        self.draw_astro_wheel(jd, "2026 Ingress")
        self.show_mundane_analysis(jd)



def setup_ui(self):
    # إعدادات عامة للصفحة
    self.page.padding = 10
    self.page.spacing = 10
    self.page.rtl = True # دعم اللغة العربية بالكامل

    # 1. القائمة اليسرى (جدول الكواكب) - DataTable
    self.tree = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("الجرم", weight="bold")),
            ft.DataColumn(ft.Text("البرج", weight="bold")),
            ft.DataColumn(ft.Text("الدرجة", weight="bold")),
            ft.DataColumn(ft.Text("بيت", weight="bold")),
        ],
        rows=[], # يتم تعبئتها ديناميكياً
        column_spacing=15,
        heading_row_height=40,
        data_row_min_height=35,
    )

    left_sidebar = ft.Container(
        content=ft.Column([self.tree], scroll=ft.ScrollMode.ADAPTIVE),
        bgcolor="#ffffff",
        border=ft.border.all(1, "#cbd5e1"),
        border_radius=8,
        expand=2,
    )

    # 2. القائمة اليمنى (أدوات التحكم)
    # تنسيق العناوين الجانبية
    def section_title(text, color="#e2e8f0"):
        return ft.Container(
            content=ft.Text(text, weight="bold", size=14),
            bgcolor=color,
            padding=5,
            alignment=ft.alignment.center,
            width=float("inf")
        )

    # حقول إدخال الموقع والبيوت
    self.city_entry = ft.TextField(label="المدينة", value="Tartus", text_align="center", dense=True)
    self.house_menu = ft.Dropdown(
        label="نظام البيوت",
        options=[ft.dropdown.Option(k) for k in self.systems_map.keys()],
        value="Placidus",
        on_change=lambda e: self.draw_now(),
        dense=True
    )
    
    self.coords_lbl = ft.Text("عرض: 00° 00' | طول: 00° 00'", size=12, text_align="center")

    # حقول إدخال الوقت والتاريخ (التنقل التلقائي)
    self.day_ent = ft.TextField(label="يوم", width=65, text_align="center", on_change=lambda e: self.auto_skip(e, self.month_ent, 2))
    self.month_ent = ft.TextField(label="شهر", width=65, text_align="center", on_change=lambda e: self.auto_skip(e, self.year_ent, 2))
    self.year_ent = ft.TextField(label="سنة", width=85, text_align="center", on_change=lambda e: self.auto_skip(e, self.hour_ent, 4))
    self.hour_ent = ft.TextField(label="س", width=65, text_align="center", on_change=lambda e: self.auto_skip(e, self.min_ent, 2))
    self.min_ent = ft.TextField(label="د", width=65, text_align="center")

    # أزرار الروحانيات والأدوات
    spiritual_btns = ft.Column([
        ft.ElevatedButton("🌙 منازل القمر", bgcolor="#8b5cf6", color="white", on_click=self.show_moon_mansion, width=250),
        ft.ElevatedButton("💭 تفسير الأحلام", bgcolor="#4b5563", color="white", on_click=self.show_dream_interpreter, width=250),
        ft.ElevatedButton("✂️ قص الشعر", bgcolor="#ec4899", color="white", on_click=self.hair_cut_calendar, width=250),
        ft.ElevatedButton("💅 قص الأظافر", bgcolor="#be185d", color="white", on_click=self.nail_cut_calendar, width=250),
        ft.ElevatedButton("⏳ ساعات الكواكب", bgcolor="#0f172a", color="white", on_click=self.show_planetary_hours, width=250),
    ], horizontal_alignment="center")

    right_sidebar_content = ft.Column(
        [
            ft.Text("الميقاتي الفلكي - حسان الشاعر", color="#a11616", size=18, weight="bold", text_align="center"),
            section_title("📍 الموقع والبيوت"),
            self.city_entry,
            self.house_menu,
            ft.ElevatedButton("🌍 جلب الإحداثيات", bgcolor="#10b981", color="white", on_click=self.search_city_logic, width=250),
            self.coords_lbl,
            
            section_title("🔮 المحرك الرئيسي"),
            ft.ElevatedButton("⚖️ تصحيح الطالع", bgcolor="#7c3aed", color="white", on_click=self.show_animodar_correction, width=250),
            ft.ElevatedButton("🕒 الوقت الآن", bgcolor="#2563eb", color="white", on_click=self.draw_now, width=250),
            ft.ElevatedButton("🌍 طالع سنة العالم", bgcolor="#dc2626", color="white", on_click=self.draw_2026, width=250),
            
            section_title("👶 بيانات الميلاد"),
            ft.Row([self.day_ent, self.month_ent, self.year_ent], alignment="center"),
            ft.Row([self.hour_ent, self.min_ent], alignment="center"),
            ft.ElevatedButton("👶 تحليل الولادة", bgcolor="#059669", color="white", on_click=self.analyze_birth_chart, width=250),
            
            section_title("🌙 الروحانيات والأدوات"),
            spiritual_btns,
            
            ft.ElevatedButton("حول البرنامج", bgcolor="#1f538d", color="white", on_click=self.show_about, width=150),
        ],
        scroll=ft.ScrollMode.ALWAYS,
        horizontal_alignment="center",
        spacing=10
    )

    right_sidebar = ft.Container(
        content=right_sidebar_content,
        width=280,
        bgcolor="#f8fafc",
        border=ft.border.all(1, "#cbd5e1"),
        padding=10
    )

    # 3. منطقة الخريطة المركزية (Canvas)
    self.chart_container = ft.Container(
        content=ft.Text("الخريطة الفلكية تظهر هنا", color="#94a3b8"),
        expand=3,
        bgcolor="#ffffff",
        alignment=ft.alignment.center
    )

    # تجميع الواجهة النهائية
    self.page.add(
        ft.Row(
            [
                left_sidebar,
                self.chart_container,
                right_sidebar,
            ],
            expand=True,
        )
    )

# وظيفة التنقل التلقائي المعدلة لـ Flet
    def auto_skip(self, e, next_control, limit):
     if len(e.control.value) >= limit:
        next_control.focus()
        self.page.update()
        # ربط الخانات بالتنقل التلقائي (بعد إكمال عدد الأرقام المطلوب)
    def show_planetary_hours(self):
        try:
            
            # 1. جلب البيانات (Flet يستخدم .value وليس .get)
            lat = getattr(self, 'lat', 34.88)
            lon = getattr(self, 'lon', 35.88)
            city_name = self.city_entry.value if hasattr(self, 'city_entry') else "غير محدد"
            now = datetime.now()
            
            # 2. حساب الحاكم (بناءً على الترتيب الكلداني)
            # PLANETARY_ORDER = ["زحل", "المشتري", "المريخ", "الشمس", "الزهرة", "عطارد", "القمر"]
            weekday = now.weekday() 
            day_rulers_start = {6: 3, 0: 6, 1: 2, 2: 5, 3: 1, 4: 4, 5: 0}
            start_idx = day_rulers_start.get(weekday, 0)
            ruler = PLANETARY_ORDER[(start_idx + now.hour) % 7]

            # 3. تجهيز نص التحليل
            analysis_text = f"""📍 الموقع: {city_name}
🕒 الوقت: {now.strftime('%H:%M')}

📜 التفسير الموسع:
{PLANETARY_HOURS_INFO.get(ruler, "المعلومات غير متوفرة حالياً")}

💡 نصيحة الساعة:
يفضل ممارسة الأنشطة المرتبطة بطبيعة كوكب {ruler} لضمان التوافق مع القوى الفلكية المهيمنة الآن."""

            # 4. عرض النافذة المنبثقة (طريقة Flet الصحيحة)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                modal=False,
                title=ft.Text(f"⏳ حاكم الساعة: {ruler}", size=20, weight="bold", color="#38bdf8", rtl=True),
                content=ft.Container(
                    content=ft.Text(analysis_text, size=16, rtl=True),
                    padding=10,
                    height=300, # تحديد ارتفاع لضمان التمرير إذا كان النص طويلاً
                ),
                actions=[
                    ft.TextButton("نسخ التحليل", on_click=lambda _: self.page.set_clipboard(analysis_text)),
                    ft.TextButton("إغلاق", on_click=close_dlg),
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as e:
            # عرض الخطأ لمستخدم أندرويد بشكل أنيق
            snack = ft.SnackBar(ft.Text(f"حدث خلل: {str(e)}"))
            self.page.snack_bar = snack
            snack.open = True
            self.page.update()
    def show_animodar_correction(self):
        try:
            
            def clean_float(val):
                if isinstance(val, (list, tuple)): val = val[0]
                return float(val)

            # 1. جلب البيانات (تعديل .get() إلى .value وتوفير قيم افتراضية)
            lat = float(getattr(self, 'lat', 34.88))
            lon = float(getattr(self, 'lon', 35.88))
            
            # التأكد من وجود قيم في الحقول لتجنب ValueError
            try:
                d = int(self.day_ent.value or 1)
                m = int(self.month_ent.value or 1)
                y = int(self.year_ent.value or 2000)
                h = int(self.hour_ent.value or 12)
                mn = int(self.min_ent.value or 0)
                offset = float(self.gmt_ent.value if hasattr(self, 'gmt_ent') and self.gmt_ent.value else 3.0)
            except ValueError:
                raise Exception("يرجى التأكد من إدخال أرقام صحيحة في خانات الميلاد")

            jd_birth = swe.julday(y, m, d, (h + mn/60.0) - offset)
            
            # 2. حساب الطالع والزوايا
            res_houses = swe.houses_ex(jd_birth, lat, lon, b'P')
            abs_asc = clean_float(res_houses[1][0])

            # 3. حساب درجة السيجي
            sun_pos = clean_float(swe.calc_ut(jd_birth, swe.SUN)[0])
            
            # نهار أم ليل
            is_day = abs_asc < sun_pos < (abs_asc + 180) % 360
            if is_day:
                target_deg = sun_pos
            else:
                target_deg = clean_float(swe.calc_ut(jd_birth, swe.MOON)[0])

            # 4. حساب المبتز
            p_names = {0: "الشمس", 1: "القمر", 2: "عطارد", 3: "الزهرة", 4: "المريخ", 5: "المشتري", 6: "زحل"}
            domiciles = [4, 2, 2, 1, 0, 2, 3, 4, 5, 6, 6, 5] 
            
            sign_idx = int(target_deg / 30) % 12
            scores = {i: 0 for i in range(7)}
            scores[domiciles[sign_idx]] += 5 
            
            winner_id = max(scores, key=scores.get)
            winner_name = p_names[winner_id]
            abs_winner = clean_float(swe.calc_ut(jd_birth, winner_id)[0])

            # 5. استخراج الأسماء والدرجات
            z_list = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            s_sign = z_list[int(abs_winner / 30) % 12]
            a_sign = z_list[int(abs_asc / 30) % 12]
            
            # 6. عرض النتيجة باستخدام Flet SnackBar أو استدعاء دالة عرض تفصيلية
            # تأكد أن دالة display_detailed_animodar مصممة لـ Flet أيضاً
            self.display_detailed_animodar(y, m, d, h, mn, s_sign, int(abs_winner % 30), 
                                           a_sign, int(abs_asc % 30), int(abs_winner % 30), winner_name)

        except Exception as e:
            # طريقة عرض الخطأ الصحيحة في أندرويد (Flet)
            snack = ft.SnackBar(ft.Text(f"خطأ في التصحيح: {str(e)}"))
            self.page.snack_bar = snack
            snack.open = True
            self.page.update()

    def display_detailed_animodar(self, y, m, d, h, mn, s_sign, s_deg, a_sign, a_curr, a_corr, planet_name):
        # 1. بناء نص التقرير بشكل صحيح داخل متغير واحد
        full_text = f"""تقرير تصحيح وقت الولادة (النمودار التقليدي المتقدم)
تاريخ الميلاد: {y}/{m}/{d} | الوقت المحلي: {h:02d}:{mn:02d}

۞ البعد الفلسفي للنمودار:
يقوم تصحيح الطالع على مبدأ 'التناظر الكوني'...

النتائج النهائية:
• الكوكب المبتز: {planet_name}
• درجة الكوكب: {s_deg}° في برج {s_sign}
• الطالع الحالي: {a_curr}° في برج {a_sign}
• الدرجة المصححة الواجب ضبطها: [{a_corr}°] تماماً.

المطور: الفلكي البرمجي حسان الشاعر"""

        # 2. دالة الحفظ
        def save_to_word(e):
            nonlocal full_text
            try:
                from docx import Document
                import os
                
                doc = Document()
                doc.add_heading('تقرير تصحيح الطالع - النمودار', 0)
                doc.add_paragraph(full_text)
                
                file_name = f"Animodar_{y}_{m}_{d}.docx"
                # مسار متوافق مع أندرويد
                data_dir = self.page.client_storage.get("app_path") or os.getcwd()
                save_path = os.path.join(data_dir, file_name)
                
                doc.save(save_path)
                
                self.page.snack_bar = ft.SnackBar(ft.Text(f"✅ تم الحفظ: {file_name}"), bgcolor="green")
                self.page.snack_bar.open = True
                self.page.update()
            except Exception as ex:
                self.page.snack_bar = ft.SnackBar(ft.Text(f"❌ خطأ: {str(ex)}"), bgcolor="red")
                self.page.snack_bar.open = True
                self.page.update()

        # 3. إغلاق النافذة
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # 4. واجهة العرض
        report_column = ft.Column(
            scroll=ft.ScrollMode.ALWAYS,
            expand=True,
            controls=[
                ft.Text("التقرير الملكي للنمودار", size=22, weight="bold", color="#38bdf8", text_align="center"),
                ft.Divider(color="#38bdf8"),
                ft.Text(full_text, size=16, color="#e2e8f0", rtl=True),
            ]
        )

        dlg = ft.AlertDialog(
            modal=True,
            title=ft.Text("نتائج التصحيح", rtl=True),
            content=ft.Container(content=report_column, width=500, height=600, padding=10),
            actions=[
                ft.ElevatedButton("💾 حفظ Word", icon=ft.icons.SAVE, on_click=save_to_word),
                ft.TextButton("إغلاق", on_click=close_dlg),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
    def auto_skip(self, e, next_widget, limit):
    # e.control هو العنصر الذي أطلق الحدث في Flet
     if len(e.control.value) >= limit:
        if next_widget:
            next_widget.focus()
            self.page.update()
    def get_accurate_planets(self, year, month, day, hour_decimal, lat, lon):
        
        try:
            # تصحيح التوقيت العالمي
            utc_hour = hour_decimal - 3 
            jd = swe.julday(year, month, day, utc_hour)
            
            planets = {
                "الشمس": swe.SUN, "القمر": swe.MOON, "عطارد": swe.MERCURY,
                "الزهرة": swe.VENUS, "المريخ": swe.MARS, "المشتري": swe.JUPITER,
                "زحل": swe.SATURN, "الرأس": swe.MEAN_NODE
            }
            
            results = {}
            signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                     "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

            for name, p_id in planets.items():
                res, ret = swe.calc_ut(jd, p_id)
                lon_pos = res[0]
                results[name] = {
                    "sign": signs[int(lon_pos / 30) % 12], 
                    "full_deg": lon_pos,
                    "deg_in_sign": lon_pos % 30
                }
                
            # حساب البيوت والطالع (نظام بلازيدوس)
            houses, ascmc = swe.houses_ex(jd, lat, lon, b'P')
            
            results["الطالع"] = ascmc[0]
            results["وسط السماء"] = ascmc[1]
            results["البيوت"] = houses
            
            return results
        except Exception as e:
            # في أندرويد يفضل إظهار الخطأ في Snackbar بدلاً من print
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ في الحساب الفلكي: {e}"))
                self.page.snack_bar.open = True
                self.page.update()
            return None
    def calculate_almuten(self, planets_data):
        # 1. تعريف أرباب البيوت (إصلاح الإزاحة هنا)
        domiciles = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
        
        # 2. تعريف أرباب الشرف
        exaltations = {0: "الشمس", 1: "القمر", 3: "المشتري", 5: "عطارد", 6: "زحل", 9: "المريخ", 11: "الزهرة"}

        scores = {k: 0 for k in ["الشمس", "القمر", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل"]}
        
        try:
            for p_name in scores.keys():
                # استخدام .get لتجنب الخطأ إذا نقص كوكب من البيانات
                p_info = planets_data.get(p_name)
                if not p_info: continue
                
                pos = p_info["full_deg"]
                sign_idx = int(pos / 30) % 12
                
                # أ- قوة بيت الكوكب (5 نقاط)
                if domiciles[sign_idx] == p_name:
                    scores[p_name] += 5
                    
                # ب- قوة الشرف (4 نقاط)
                if exaltations.get(sign_idx) == p_name:
                    scores[p_name] += 4
                
                # ملاحظة: يمكنك هنا إضافة (المثلثات +3، الحدود +2، الوجوه +1) لاحقاً

            # 3. تحديد الفائز
            winner = max(scores, key=scores.get)
            
            if scores[winner] == 0:
                return "غير محدد", 0
                
            return winner, scores[winner]
            
        except Exception as e:
            # طباعة الخطأ في Console الأندرويد للمساعدة في التصحيح
            print(f"Almuten Error: {e}")
            return "خطأ في الحساب", 0

    def analyze_birth_chart(self):
        try:
            import swisseph as swe
            from datetime import datetime

            # 1. جلب المدخلات من واجهة Flet (دعم أندرويد)
            def g_v(attr, default):
                element = getattr(self, attr, None)
                return element.value if element and element.value else default

            y, m, d = int(g_v('year_ent', 1990)), int(g_v('month_ent', 1)), int(g_v('day_ent', 1))
            h, mi = int(g_v('hour_ent', 12)), int(g_v('min_ent', 0))
            lat = float(getattr(self, 'lat', 34.88))
            lon = float(getattr(self, 'lon', 35.88))
            
            # تصحيح التوقيت العالمي (UT)
            offset = float(g_v('gmt_ent', 3.0))
            ut_h = h - offset + (mi / 60.0)
            jd = swe.julday(y, m, d, ut_h)

            # 2. حساب البيوت والطالع (نظام بلاسيدوس)
            cusps_raw, ascmc_raw = swe.houses(jd, lat, lon, b'P')
            asc_raw = float(ascmc_raw[0])
            c = [float(cusps_raw[i]) for i in range(1, 13)]

            def find_house(deg):
                for i in range(12):
                    s, e = c[i], c[(i + 1) % 12]
                    if s < e:
                        if s <= deg < e: return str(i + 1)
                    else: # عبور نقطة الصفر
                        if s <= deg or deg < e: return str(i + 1)
                return "12"

            # 3. حساب الكواكب والسرعات (للتراجع والاحتراق)
            zodiacs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                       "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            
            p_map = [("الشمس",swe.SUN,"sun"), ("القمر",swe.MOON,"moon"), ("عطارد",swe.MERCURY,"mercury"), 
                     ("الزهرة",swe.VENUS,"venus"), ("المريخ",swe.MARS,"mars"), ("المشتري",swe.JUPITER,"jupiter"), 
                     ("زحل",swe.SATURN,"saturn"), ("أورانوس",swe.URANUS,"uranus"), ("نبتون",swe.NEPTUNE,"neptune"), 
                     ("بلوتو",swe.PLUTO,"pluto"), ("الرأس",swe.MEAN_NODE,"nn")]

            results = {}
            raw_d = {}
            raw_speed = {}

            for ar, code, en in p_map:
                res, _ = swe.calc_ut(jd, code)
                deg = float(res[0])
                speed = float(res[3]) # السرعة
                raw_d[en] = deg
                raw_speed[en] = speed
                results[f"{en}_sign"] = zodiacs[int(deg / 30)]
                results[f"{en}_house"] = find_house(deg)

            # 4. تحديد حالة النهار والليل آلياً وحساب السهام
            s_h = int(results.get('sun_house', 1))
            is_day_time = 7 <= s_h <= 12
            
            # سهم السعادة (Fortune)
            if is_day_time:
                fort = (asc_raw + raw_d['moon'] - raw_d['sun']) % 360
                spir = (asc_raw + raw_d['sun'] - raw_d['moon']) % 360
            else:
                fort = (asc_raw + raw_d['sun'] - raw_d['moon']) % 360
                spir = (asc_raw + raw_d['moon'] - raw_d['sun']) % 360

            results['fortune_part'] = f"{zodiacs[int(fort/30)]} ({int(fort%30)}°)"
            results['spirit_part'] = f"{zodiacs[int(spir/30)]} ({int(spir%30)}°)"

            # 5. حساب الفردارية والمستولي (Almuten)
            if hasattr(self, 'calculate_current_firdaria'):
                firdar_name = self.calculate_current_firdaria(datetime(y,m,d), is_day_time)
                results['firdar_ruler'] = firdar_name
            
            # 6. محرك القواطع والتحذيرات (احتراق، تراجع، نحوس)
            warnings_list = []
            for ar, code, en in p_map:
                if raw_speed.get(en, 0) < 0:
                    warnings_list.append(f"⚠️ كوكب {ar} متراجع (Retrograde)")
                if en != "sun":
                    dist_to_sun = abs(raw_d[en] - raw_d['sun'])
                    if dist_to_sun < 8 or dist_to_sun > 352:
                        warnings_list.append(f"🔥 كوكب {ar} محترق (Combust) لقربه من الشمس")
            
            if abs(raw_d['saturn'] - raw_d['mars']) < 5:
                warnings_list.append("💀 تحذير: اقتران النحسين (زحل والمريخ) في الهيئة")
            
            results['warnings'] = "\n".join(warnings_list) if warnings_list else "لا يوجد قواطع فلكية سلبية"

            # 7. تحليل النجوم الثابتة المتقدم (محرك الفحص الشامل)
            fixed_stars_db = [
                {"name": "الثريا (Alcyone)", "pos": 60.1, "effect": "الجاذبية، الشهرة، والنجاح الفني."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "effect": "شجاعة قيادية ونجاح مادي محفوف بالمواجهات."},
                {"name": "رأس التوأم (Pollux)", "pos": 113.3, "effect": "الجرأة وقوة البيان والقدرة على الغلبة."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "effect": "الجاه العظيم، السلطة، والرفعة السيادية."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "effect": "الثروة، الثقافة، والنجاح العلمي الباهر."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "effect": "التحولات الكبرى وإعادة البناء بعد الهدم."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "effect": "الإبداع الروحاني والتميز الفني."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "effect": "الشهرة العالمية والتميز الروحاني الرفيع."}
            ]
            
            stars_analysis = []
            my_planets_for_stars = {**raw_d, "Ascendant": asc_raw}
            p_ar_names = {"sun":"الشمس", "moon":"القمر", "mercury":"عطارد", "venus":"الزهرة", "mars":"المريخ", "jupiter":"المشتري", "saturn":"زحل", "Ascendant":"الطالع"}

            for star in fixed_stars_db:
                for p_en, p_deg in my_planets_for_stars.items():
                    diff = abs(p_deg - star['pos'])
                    if diff <= 1.5 or diff > 358.5:
                        p_ar = p_ar_names.get(p_en, p_en)
                        stars_analysis.append(f"⭐ يقترن {p_ar} بنجم {star['name']}.\n   💡 التأثير: {star['effect']}")

            results['stars_report'] = "\n\n".join(stars_analysis) if stars_analysis else "لا توجد اقترانات مباشرة بالنجوم العظمى."

            # 8. تخزين البيانات النهائية لربطها بالواجهة
            self.raw_d = raw_d
            self.asc_raw = asc_raw
            self.is_analyzed = True
            
            # استدعاء التقرير الشامل (Display Report)
            self.display_comprehensive_report(results)

        except Exception as e:
            # رسالة خطأ أندرويد (تمنع الشاشة البيضاء)
            if hasattr(self, 'page'):
                snack = ft.SnackBar(ft.Text(f"خطأ في التحليل الشامل: {str(e)}"))
                self.page.snack_bar = snack
                snack.open = True
                self.page.update()
    def get_house(self, planet_deg, houses):
        """تحديد البيت الذي يقع فيه الكوكب بناءً على حدود البيوت (Cusps)"""
        # نضمن أن الدرجة تقع بين 0 و 360
        planet_deg %= 360
        for i in range(12):
            s = houses[i]
            e = houses[(i + 1) % 12]
            
            if s < e:
                # الحالة العادية: البداية أقل من النهاية
                if s <= planet_deg < e:
                    return str(i + 1)
            else:
                # حالة عبور نقطة الصفر (برج الحمل)
                if s <= planet_deg or planet_deg < e:
                    return str(i + 1)
        return "1"

    def get_sign_name(self, deg):
        """تحويل الدرجة المطلقة إلى اسم البرج المقابل لها"""
        signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                 "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
        # استخدام int لضمان الحصول على فهرس صحيح من 0 إلى 11
        index = int((deg % 360) / 30)
        return signs[index]
    def run_calculation(self, jd, lat, lon):
        data = {}
        zodiac_names = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                        "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]

        # 1. حساب نظام البيوت
        houses, ascmc = swe.houses(jd, lat, lon, b'P') 
        ascendant = ascmc[0]
        data['ascendant_deg'] = ascendant
        data['ascendant_sign'] = zodiac_names[int(ascendant / 30)]

        # 2. دالة تحديد البيت (المعدلة لعلاج عبور نقطة الصفر)
        def get_house_safe(planet_deg, h_list):
            planet_deg %= 360
            for i in range(1, 13):
                s = h_list[i]
                e = h_list[i+1] if i < 12 else h_list[1]
                if s < e:
                    if s <= planet_deg < e: return i
                else: # حالة عبور نقطة الصفر
                    if s <= planet_deg or planet_deg < e: return i
            return 1

        # 3. حساب الكواكب
        planets_to_calc = {
            'sun': swe.SUN, 'moon': swe.MOON, 'mercury': swe.MERCURY,
            'venus': swe.VENUS, 'mars': swe.MARS, 'jupiter': swe.JUPITER,
            'saturn': swe.SATURN, 'uranus': swe.URANUS, 'neptune': swe.NEPTUNE,
            'pluto': swe.PLUTO
        }

        for name, code in planets_to_calc.items():
            res, _ = swe.calc_ut(jd, code)
            p_deg = res[0]
            data[f'{name}_deg'] = p_deg
            data[f'{name}_sign'] = zodiac_names[int(p_deg / 30)]
            data[f'{name}_house'] = str(get_house_safe(p_deg, houses))

        # 4. العقد القمرية والسهام
        res_node, _ = swe.calc_ut(jd, swe.MEAN_NODE)
        nn_deg = res_node[0]
        data['north_node_deg'] = nn_deg
        data['north_node_sign'] = zodiac_names[int(nn_deg / 30)]
        
        # سهم السعادة وسهم الغيب
        fortune_deg = (ascendant + data['moon_deg'] - data['sun_deg']) % 360
        spirit_deg = (ascendant + data['sun_deg'] - data['moon_deg']) % 360
        
        data['fortune_part'] = f"{zodiac_names[int(fortune_deg / 30)]} ({int(fortune_deg % 30)}°)"
        data['spirit_part'] = f"{zodiac_names[int(spirit_deg / 30)]} ({int(spirit_deg % 30)}°)"

        return data

    def get_hijri_date_string(self):
        try:
            from ummalqura.hijri_date import HijriDate
            now = HijriDate.today()
            # إرجاع اليوم والشهر والسنة بتنسيق عربي
            return f"{now.day} / {now.month} / {now.year} هـ"
        except ImportError:
            return "يرجى إضافة ummalqura للمتطلبات"
        except Exception as e:
            return f"خطأ: {str(e)}"
    def get_daily_wisdom(self, day):
        # قاعدة بيانات التفسير حسب الأيام (نموذج لليوم الحالي)
        wisdom_db = {
            1: {"sadiq": "رؤيا صادقة جداً وتتحقق سريعاً.", "sirin": "تدل على بشارة خير في أول الشهر.", "philosophy": "بداية في عالم المثال."},
            # يمكنك إكمال بقية الأيام 1-30 هنا
        }
        return wisdom_db.get(day, {"sadiq": "تختلف حسب حال الرائي.", "sirin": "تطلب تأويلاً دقيقاً.", "philosophy": "انعكاس لنفس الرائي."})
    def show_about(self, e=None):
        """إظهار معلومات النظام بشكل يتوافق مع أندرويد"""
        about_text = (
            "🚀 الميقاتي الفلكي - النسخة الاحترافية\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "تطوير الخبير: حسان الشاعر\n"
            "إصدار: 2026 PRO Mobile\n\n"
            "نظام فلكي متكامل يجمع بين الحسابات السويسرية\n"
            "والاستنباطات الفلسفية التراثية، مصمم ليعمل\n"
            "بكفاءة عالية على أنظمة أندرويد."
        )

        def close_dlg(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            title=ft.Text("حول النظام", weight="bold", rtl=True),
            content=ft.Text(about_text, rtl=True, size=14),
            actions=[
                ft.TextButton("إغلاق", on_click=close_dlg),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
    self.page.update()

    # يجب أن يكون هناك سطر فارغ وتكون def على نفس مستوى الـ def التي قبلها
    def show_dream_interpreter(self):
        try: 
            # 1. حساب اليوم القمري (المحرك الفلكي)
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, 12)
            res_m, _ = swe.calc_ut(jd, 1) # القمر
            res_s, _ = swe.calc_ut(jd, 0) # الشمس
            
            # حساب اليوم القمري (الفرق بين القمر والشمس)
            h_day = int(((float(res_m[0]) - float(res_s[0])) % 360 / 12.2) + 1)
            # ... باقي الكود
        except Exception as e:
            print(f"Error: {e}")            # جلب حكمة اليوم القمري
            wisdom = self.get_daily_wisdom(h_day)  
            symbols_dictionary = {
                "نبي": "🌙 العرب: عز وشرف وبشارة.\n🏹 الفرس: استقامة دين وظفر.\n🏛️ اليونان: اتصال بالعقل الكلي والحكمة.",
                "ملك": "🌙 العرب: نصر وعز وقضاء حاجة.\n🏹 الفرس: نيل منصب رفيع.\n🏛️ اليونان: سيطرة على النفس والأنا العليا.",
                "شمس": "🌙 العرب: ولي الأمر أو الوالد.\n🏹 الفرس: الحاكم الأعظم.\n🏛️ اليونان: الحقيقة والوعي المطلق.",
                "ماء": "🌙 العرب: حياة طيبة ورزق.\n🏹 الفرس: استقرار ونمو تجاري.\n🏛️ اليونان: تطهير وشفاء جسدي.",
                "نار": "🌙 العرب: فتنة أو هداية.\n🏹 الفرس: منصب نوراني ورفعة.\n🏛️ اليونان: تحول وإبداع أو غضب.",
                "ذهب": "🌙 العرب: همّ للرجل وفرح للنساء.\n🏹 الفرس: ميراث أو ولاية.\n🏛️ اليونان: خلود وفكر نقي.",
                "فضة": "🌙 العرب: مال وراحة.\n🏹 الفرس: زواج أو ولادة.\n🏛️ اليونان: حدس وأنوار باطنية.",
                "ثعبان": "🌙 العرب: عدو من الأهل.\n🏹 الفرس: عدو ذو مال.\n🏛️ اليونان: شفاء وحكمة الأرض.",
                "أسد": "🌙 العرب: عدو جبار.\n🏹 الفرس: قوة ونفوذ.\n🏛️ اليونان: شجاعة وسيادة.",
                "طيران": "🌙 العرب: سفر ورفعة.\n🏹 الفرس: غنى سريع.\n🏛️ اليونان: تحرر من قيود المادة.",
                "موت": "🌙 العرب: عمر طويل وتوبة.\n🏹 الفرس: بدء حياة جديدة.\n🏛️ اليونان: تحول الوعي.",
                "بحر": "🌙 العرب: الدنيا وتقلباتها.\n🏹 الفرس: ملك لا يدرك غوره.\n🏛️ اليونان: اللاوعي العميق.",
                "خبز": "🌙 العرب: الإسلام والرزق.\n🏹 الفرس: أمن من الفقر.\n🏛️ اليونان: نمو روحي.",
                "لبن": "🌙 العرب: فطرة ورزق حلال.\n🏹 الفرس: علم وحكمة.\n🏛️ اليونان: تغذية النفس.",
                "عسل": "🌙 العرب: حلاوة الإيمان.\n🏹 الفرس: مال موروث.\n🏛️ اليونان: رحيق الحكمة.",
                "سيف": "🌙 العرب: ولد أو حجة قوية.\n🏹 الفرس: نصر وقوة.\n🏛️ اليونان: إرادة قاطعة.",
                "خاتم": "🌙 العرب: أمان أو منصب أو زواج.\n🏹 الفرس: إتمام أمر عظيم.\n🏛️ اليونان: عهد أبدي ودائرة الكمال.",
                "شجر": "🌙 العرب: رجال أو أعمال.\n🏹 الفرس: طول عمر وبركة.\n🏛️ اليونان: نمو الشخصية والجذور.",
                "كعبة": "🌙 العرب: حج وأمن وبشارة.\n🏹 الفرس: استقامة وقبلة المراد.\n🏛️ اليونان: مركز الوجود والاتزان.",
                "ضحك": "🌙 العرب: حزن (إلا بشارة بمولود).\n🏹 الفرس: خبر طارئ.\n🏛️ اليونان: إفراغ توتر العقل الباطن.",
                "بكاء": "🌙 العرب: فرج ونفاذ ضيق.\n🏹 الفرس: زوال مرض أو كرب.\n🏛️ اليونان: تطهير عاطفي (Catharsis)."
            }

            # 2. القاموس الموسع الشامل (موسوعة ابن سيرين والصادق المحدثة 2026)
            symbols_dictionary = {
                # --- الرموز الروحية والعلوية ---
                "نبي": "ابن سيرين: عز وشرف في الدارين ورؤياهم بشارة. الصادق: هداية ونيل أمانة وخير مستدام.",
                "ملك": "ابن سيرين: نصر على الأعداء ونيل سلطة وجاه. الصادق: قضاء حاجة وبشارة بالظفر.",
                "كعبة": "ابن سيرين: أمن وبشارة واجتماع بولي الأمر. الصادق: إيمان، إسلام، أمن، أو حج قريب.",
                "قرآن": "ابن سيرين: حكمة وميراث وعلم نافع. الصادق: انتشار علم ورزق طيب وتطهير للروح.",
                "محراب": "ابن سيرين: حصول ولد صالح أو إمام عادل. الصادق: امرأة صالحة أو رزق من جهة الدين.",
                "أذان": "ابن سيرين: حج لمن أذن في أشهره، أو نيل منصب. الصادق: دعوة للخير أو رفعة ذكر.",

                # --- الطبيعة والعناصر الكونيه ---
                "ماء": "ابن سيرين: حياة طيبة. الصادق: فتنة إن كان كدراً، وعلم ورزق إن كان صافياً.",
                "نار": "ابن سيرين: فتنة أو سلطان أو بشارة ورزق. الصادق: عذاب أو طريق هداية مستقيم.",
                "بحر": "ابن سيرين: سلطان مهاب. الصادق: ملك، رئيس، عالم، علم، مال، أو شغل عظيم.",
                "مطر": "ابن سيرين: رحمة عامة ما لم يكن مخرباً. الصادق: فرج من هم واستجابة دعاء.",
                "شمس": "ابن سيرين: ولي الأمر أو الوالد أو الملك. الصادق: نور في الدين وجاه في الدنيا.",
                "قمر": "ابن سيرين: وزير الملك أو الزوجة أو العالم. الصادق: زيادة في العلم أو نيل منصب.",
                "تراب": "ابن سيرين: مال ورزق بعد تعب. الصادق: يؤول بالدنيا أو الموت أو الفقر والحاجة.",

                # --- الحيوانات والطيور والحشرات ---
                "أسد": "ابن سيرين: عدو جبار أو سلطان جائر. الصادق: شدة وضيق من جهة صاحب نفوذ.",
                "خيل": "ابن سيرين: عز وشرف ومرتبة عالية. الصادق: سفر مبارك أو نيل سيادة على قومه.",
                "جمل": "ابن سيرين: حزن أو رجل جاهل أو سفر طويل. الصادق: نفوذ وقوة وصبر على المحن.",
                "ثعبان": "ابن سيرين: عدو يكاتم العداوة. الصادق: امرأة، ولد، عدو ذو مال، أو كنوز مخفية.",
                "عقرب": "ابن سيرين: عدو ضعيف يغتاب. الصادق: نكد وهمّ يسببه شخص نمام قريب.",
                "نحل": "ابن سيرين: غنيمة بلا تعب وعمل صالح. الصادق: عسكر، شفاء من مرض، أو رزق مبارك.",
                "غراب": "ابن سيرين: رجل فاسق أو سفر بعيد. الصادق: يؤول بالرجل الكاذب أو الحزن.",
                "حوت": "ابن سيرين: وزير الملك أو ضيق يتبعه فرج (قصة يونس). الصادق: رزق من البحر أو علم.",

                # --- الثمار والأطعمة ---
                "خبز": "ابن سيرين: الإسلام والعمر والمال الحلال. الصادق: عيش طيب وعلم نافع.",
                "عسل": "ابن سيرين: ميراث أو مال مجموع. الصادق: حلاوة الإيمان والقرآن والشفاء.",
                "لبن": "ابن سيرين: فطرة الإسلام والرزق الفطري. الصادق: علم واسع وحكمة.",
                "تفاح": "ابن سيرين: همّة الرجل في عمله. الصادق: نيل بشارة ومنفعة من صديق.",
                "عنب": "ابن سيرين: رزق واسع، والأسود همّ. الصادق: مال مجموع أو منفعة دائمة.",
                "تمر": "ابن سيرين: مطر أو رزق حلال لا يدوم. الصادق: حلاوة الإيمان ونيل مقصود.",
                "لحم": "ابن سيرين: المطبوخ مال، والنيء وجع ومرض. الصادق: غنيمة أو غيبة (حسب السياق).",

                # --- المعادن والكنوز والأدوات ---
                "ذهب": "ابن سيرين: همّ وغرامة للرجل. الصادق: فرح للنساء، وللرجل ذهاب منصب.",
                "فضة": "ابن سيرين: مال مجموع وراحة بال. الصادق: امرأة جميلة، ولد، أو خير كثير.",
                "جوهر": "ابن سيرين: علم وجمال. الصادق: يؤول بالمرأة الحسناء أو التجارة الرابحة.",
                "كرسي": "ابن سيرين: رفعة وشأن. الصادق: عز، شرف، ولاية، نصر، أو سعادة.",
                "مفتاح": "ابن سيرين: نصر وفتح أبواب الرزق. الصادق: قضاء حاجة، حج، ومال ميسر.",
                "ثوب": "ابن سيرين: ستر ودين ووقار. الصادق: عز وجاه وحال مستورة.",
                "سيف": "ابن سيرين: ولد أو سلطان أو كلام. الصادق: نصر على الأعداء أو حجة قوية.",

                # --- أحوال الإنسان وأفعاله ---
                "موت": "ابن سيرين: طول عمر أو ندم وتوبة. الصادق: انقطاع أمر أو زواج للأعزب.",
                "طيران": "ابن سيرين: سفر وتغير حال وطلب رفعة. الصادق: استجابة دعاء وعلو مرتبة.",
                "بكاء": "ابن سيرين: فرج وسرور (بدون عويل). الصادق: نجاة من كرب وفرح عاجل.",
                "ضحك": "ابن سيرين: حزن وندم. الصادق: بشارة بقدوم ولد أو خبر سار.",
                "زواج": "ابن سيرين: عناية إلهية أو قيد وهمّ. الصادق: نيل منصب أو غنى.",
                "غسل": "ابن سيرين: قضاء دين وتوبة من ذنب. الصادق: طهارة في الدين وزوال الهموم.",
                "حج": "ابن سيرين: قضاء دين أو حج فعلي أو أمن. الصادق: استقامة الدين ونيل أجر."
            }

            # 3. قاموس أحكام الأيام الهلالية (الإمام الصادق)
            days_dict = {
            1: "صادقة جداً، وتأويلها يظهر سريعاً وهي من أفضل الأيام.",
            2: "رؤيا صالحة، ومن رآها نال خيراً ومنفعة من السلطان أو ولي الأمر.",
            3: "رؤيا حق، وتأويلها يكون على العكس إن كانت مكروهة (تحذير).",
            4: "رؤيا محمودة، وتأويلها يتأخر قليلاً لكنه محقق.",
            5: "رؤيا فرح وسرور، ومن رآها نال بشارة يسعد بها قلبه.",
            6: "رؤيا صادقة، وتأويلها يقع بعد فترة طويلة (ربما سنوات).",
            7: "رؤيا مستورة، لا تُقص إلا على عالم أو ناصح لسرعة تحققها.",
            8: "رؤيا صادقة جداً، وهي تدل على خلاص من هم أو نيل مطلب.",
            9: "رؤيا حق، وتدل على قوة الرائي ونيله مرتبة عالية.",
            10: "رؤيا صادقة، وتدل على تيسير الأمور العالقة.",
            11: "رؤيا محمودة، وتأويلها يقع من يوم إلى 15 يوماً.",
            12: "رؤيا تحذيرية، تأويلها يتأخر لكن فيها تنبيه من غدر.",
            13: "رؤيا مختلطة، قد لا يصح تأويلها لداخلها بحديث النفس.",
            14: "رؤيا صادقة، وبشارتها سريعة كسرعة اكتمال القمر.",
            15: "رؤيا صحيحة، ومن رآها نال ما تمنى بعد تعب.",
            16: "رؤيا كاذبة (أضغاث أحلام)، لا يُعتد بها في هذا اليوم.",
            17: "رؤيا مخفية، تدل على أمور باطنة تظهر للرائي لاحقاً.",
            18: "رؤيا حق، وتدل على نصر على الأعداء أو نجاح في خصومة.",
            19: "رؤيا قوية جداً، وتأويلها يكون أوضح مما رُئي في المنام.",
            20: "رؤيا محمودة، وتدل على زيادة في الرزق والمال.",
            21: "رؤيا كاذبة، غالباً ما تكون من وسوسة النفس أو الشيطان.",
            22: "رؤيا صادقة، وتدل على قضاء الحوائج الصعبة.",
            23: "رؤيا تحذيرية من الفتن، يُنصح بالصدقة بعد رؤيتها.",
            24: "رؤيا محمودة، وتدل على الفرج بعد الضيق.",
            25: "رؤيا لا تصح، وهي من أوهام الخيال في هذا اليوم.",
            26: "رؤيا صادقة، وتدل على السفر أو التغيير في الأحوال.",
            27: "رؤيا حق، وتأويلها يقع كما هي دون تبديل.",
            28: "رؤيا صادقة، ومن رآها نال خيراً وفيراً غير متوقع.",
            29: "رؤيا تحذيرية، يُكره قصها في هذا اليوم.",
            30: "رؤيا صادقة ومباركة، وهي خاتمة الخير والرزق."
        }
    def process_interpretation(self, e, search_input, report_content, h_day, days_dict, symbols_dictionary):
        # المحرك المعدل للبحث
        search_query = search_input.value.strip()
        if not search_query:
            report_content.value = "⚠️ يرجى كتابة رمز للبحث عنه."
            self.page.update()
            return
        
        # بناء التقرير الملحمي المنسق
        report = "✨ تحليل الرؤيا الاستدلالي الملحمي ✨\n"
        report += f"📅 اليوم القمري: {h_day}\n"
        report += "════════════════════════════════\n\n"
        report += f"🌙 حكم الزمان (حسب الشهر الهلالي):\n• {days_dict.get(h_day, 'يوم يحتاج لمطابقة فلكية دقيقة.')}\n\n"

        # محرك البحث عن الرموز (دعم البحث الجزئي)
        found = False
        for key in symbols_dictionary:
            if search_query in key:
                report += f"🔍 تفسير رمز [{key}]:\n{symbols_dictionary[key]}\n\n"
                found = True
                break
        
        if not found:
            report += f"⚠️ الرمز [{search_query}] غير مدرج حالياً، جرب مرادفات أخرى.\n\n"

        report += "💎 فلسفة التعبير:\n"
        report += "• مدرسة اليونان: الرؤيا هي استقراء العقل الباطن لمسار القدر.\n"
        report += "• مدرسة العرب: الرؤيا سفر في عالم المثال وتجسد المعاني.\n\n"
        report += "════════════════════════════════\n"
        report += "✍️ إعداد: حسان الشاعر © 2026"
        
        report_content.value = report
        self.page.update()

    def draw_astro_wheel_in_popup(self, jd, title_text="خريطة الميلاد"):
        # 1. بيانات الكواكب والرموز
        planets_map = {
            swe.SUN: "☉", swe.MOON: "☽", swe.MERCURY: "☿", 
            swe.VENUS: "♀", swe.MARS: "♂", swe.JUPITER: "♃", swe.SATURN: "♄"
        }
        
        # 2. إعداد الـ Canvas
        chart_canvas = ft.Canvas(
            expand=True,
            shapes=[
                # الدائرة الكبرى
                ft.cv.Circle(150, 150, 130, ft.Paint(color="#1f538d", stroke_width=2, style=ft.PaintingStyle.STROKE)),
                # الدائرة الصغرى
                ft.cv.Circle(150, 150, 35, ft.Paint(color="#cbd5e1", stroke_width=1, style=ft.PaintingStyle.STROKE)),
            ]
        )

        # 3. محرك الرسم وتوزيع الكواكب
        for p_code, symbol in planets_map.items():
            res, _ = swe.calc_ut(jd, p_code)
            pos = float(res[0]) # التأكد من القيمة العشرية
            
            # تحويل الدرجة الفلكية إلى إحداثيات (X, Y)
            # نقص 90 لأن الصفر الفلكي (الحمل) يبدأ من الشرق رياضياً
            angle = math.radians(pos - 90)
            x = 150 + 100 * math.cos(angle)
            y = 150 + 100 * math.sin(angle)

            # رسم الرمز بدلاً من الدائرة لجمالية التقرير
            chart_canvas.shapes.append(
                ft.cv.Text(
                    x - 10, y - 10, symbol,
                    ft.TextStyle(size=20, color="#dc2626", weight="bold")
                )
            )

        # 4. عرض النافذة (AlertDialog)
        def close_chart(e):
            chart_dlg.open = False
            self.page.update()

        chart_dlg = ft.AlertDialog(
            title=ft.Text(title_text, weight="bold", rtl=True),
            content=ft.Container(
                content=chart_canvas,
                width=300, height=300,
                alignment=ft.alignment.center,
                bgcolor="#f8fafc", # خلفية هادئة للرسم
                border_radius=150 # لجعل الحاوية دائرية
            ),
            actions=[ft.TextButton("إغلاق", on_click=close_chart)]
        )

        self.page.dialog = chart_dlg
        chart_dlg.open = True
        self.page.update()
    def draw_astro_wheel_in_popup(self, jd):
        # 1. إعداد الحاوية الرسومية
        cp = ft.Canvas(
            expand=True,
            shapes=[
                # الدائرة الخارجية والداخلية
                ft.cv.Circle(150, 150, 120, ft.Paint(color="#1f538d", stroke_width=2, style=ft.PaintingStyle.STROKE)),
                ft.cv.Circle(150, 150, 40, ft.Paint(color="#cbd5e1", stroke_width=1, style=ft.PaintingStyle.STROKE)),
            ]
        )

        # 2. رسم خطوط تقسيم البروج (كل 30 درجة)
        for i in range(12):
            angle = math.radians(i * 30 - 90) # تعديل الزاوية لتبدأ من الأعلى
            x2 = 150 + 120 * math.cos(angle)
            y2 = 150 + 120 * math.sin(angle)
            cp.shapes.append(ft.cv.Line(150, 150, x2, y2, ft.Paint(color="#cbd5e1", stroke_width=1)))

        # 3. محرك جلب ورسم الكواكب الفعلية (إضافة ضرورية للدقة)
        planets = {swe.SUN: "☉", swe.MOON: "☽", swe.MARS: "♂", swe.JUPITER: "♃", swe.SATURN: "♄"}
        for p_id, symbol in planets.items():
            res, _ = swe.calc_ut(jd, p_id)
            # تحويل موقع الكوكب لزاوية رسم
            p_angle = math.radians(res[0] - 90)
            px = 150 + 90 * math.cos(p_angle)
            py = 150 + 90 * math.sin(p_angle)
            
            # رسم رمز الكوكب في موقعه الصحيح
            cp.shapes.append(
                ft.cv.Text(px - 5, py - 5, symbol, ft.TextStyle(size=16, color="#dc2626"))
            )

        # 4. عرض النافذة المنبثقة
        dlg = ft.AlertDialog(
            title=ft.Text("الخريطة الفلكية اللحظية", size=18, weight="bold", rtl=True),
            content=ft.Container(content=cp, width=300, height=300, alignment=ft.alignment.center),
            actions=[ft.TextButton("إغلاق", on_click=lambda e: self.close_dialog(dlg))]
        )
        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

    def close_dialog(self, dlg):
        """إغلاق أي نافذة منبثقة في التطبيق"""
        dlg.open = False
        self.page.update()

    def search_city_logic(self, e=None):
        # إعداد محرك البحث عن المواقع
        geolocator = Nominatim(user_agent="hassan_astro_mobile_2026")
        city_name = self.city_entry.value 
        
        if not city_name:
            self.page.snack_bar = ft.SnackBar(ft.Text("⚠️ يرجى إدخال اسم المدينة أولاً"))
            self.page.snack_bar.open = True
            self.page.update()
            return

        try:
            location = geolocator.geocode(city_name, timeout=10)
            if location:
                self.lat = location.latitude
                self.lon = location.longitude
                
                # تحديث النص الظاهر للإحداثيات
                lat_d = int(abs(self.lat))
                lon_d = int(abs(self.lon))
                self.coords_lbl.value = f"عرض: {lat_d}° | طول: {lon_d}°"
                
                # إشعار النجاح
                self.page.snack_bar = ft.SnackBar(ft.Text(f"📍 تم تحديد: {location.address}"))
                self.page.snack_bar.open = True
                
                # تحديث الحسابات والخريطة فوراً
                self.draw_now() 
            else:
                self.page.snack_bar = ft.SnackBar(ft.Text("⚠️ المدينة غير موجودة، جرب كتابتها بالإنجليزية"))
                self.page.snack_bar.open = True
            
            self.page.update()
        except Exception as ex:
            # التعامل مع أخطاء الاتصال أو الخادم
            self.page.snack_bar = ft.SnackBar(ft.Text(f"❌ خطأ: تأكد من تشغيل الإنترنت"))
            self.page.snack_bar.open = True
            self.page.update()
    def get_mansion_data(self, index):
        # قاعدة البيانات الشاملة: (الاسم، الملك، البخور، الدلالة، الوفق_المقترح، الدعاء)
        data = [
            ("الشرطين", "جبرائيل", "لبان ذكر", "البدايات والقوة", [12, 17, 10, 11, 13, 15, 16, 9, 14], "اللهم يا فاتح الأبواب افتح لي ببركة هذه المنزلة وملكهما جبرائيل"),
            ("البطين", "إسرافيل", "سندروس", "الرزق والمال", [22, 27, 20, 21, 23, 25, 26, 19, 24], "يا رزاق ذو القوة المتين ارزقني من حيث لا أحتسب ببركة إسرافيل"),
            ("الثريا", "سمسمائيل", "عود وند", "المحبة والجمال", [5, 10, 3, 4, 6, 8, 9, 2, 7], "اللهم ألقِ محبتي في قلوب عبادك أجمعين ببركة سر الثريا وسمسمائيل"),
            ("الدبران", "ميكائيل", "كافور", "الهيبة والفراسة", [30, 35, 28, 29, 31, 33, 34, 27, 32], "يا عزيز عززني بعزك يا ميكائيل واكفني شر الأعداء"),
            ("الهقعة", "صرفائيل", "جاوي", "العلم والذكاء", [15, 20, 13, 14, 16, 18, 19, 12, 17], "يا عليم علمني من علمك اللدني ببركة صرفائيل الموكل بالهقعة"),
            ("الهنعة", "عنيائيل", "ميعة سائله", "الإصلاح والمودة", [8, 13, 6, 7, 9, 11, 12, 5, 10], "اللهم ألف بين القلوب كما ألفت بين المنزلة وخادمها عنيائيل"),
            ("الذراع", "كسفيائيل", "مستكة", "النصر والظفر", [40, 45, 38, 39, 41, 43, 44, 37, 42], "يا قوي يا متين انصرني نصراً عزيزاً ببركة كسفيائيل والذراع"),
        ]
        # استخدام التكرار الدائري لضمان عدم خروج المؤشر عن النطاق
        return data[index % len(data)]

    # ملاحظة: الأسطر التالية يجب أن تكون داخل دالة التشغيل __init__ أو دالة البداية
    # لضمان عدم ظهور الشاشة البيضاء عند فتح التطبيق
    def start_app(self):
        try:
            # 1. بناء واجهة المستخدم (Flet)
            self.setup_ui()
            
            # 2. الحساب الأولي للخريطة (تجنب القيم الفارغة)
            self.draw_now()
            
            # 3. تحديث الصفحة لضمان ظهور كل شيء فوراً
            self.page.update()
        except Exception as e:
            # في حال حدث أي خطأ أثناء التشغيل، يظهر تنبيه بدلاً من الشاشة البيضاء
            print(f"Error starting app: {e}")

    def nail_cut_calendar(self):
        try:

            # 1. الحساب الفلكي الدقيق لليوم القمري
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, 12)
            res_m, _ = swe.calc_ut(jd, 1) # موقع القمر
            res_s, _ = swe.calc_ut(jd, 0) # موقع الشمس
            # استخراج الدرجة الأولى من المصفوفة لمنع خطأ TypeError
            m_deg = float(res_m[0])
            s_deg = float(res_s[0])
            day = int(((m_deg - s_deg) % 360 / 360) * 29.5) + 1
            
            # 2. مصفوفة الأحكام (بياناتك الأصلية دون حذف)
            nail_rules = {
                1: "يورث قصر العمر والكآبة (تجنبه)",
                2: "يورث النكد وتشتت الأمر",
                3: "يورث البركة في المال والولد (مستحب جداً)",
                4: "يورث الفقر والهم",
                5: "يورث زيادة الرزق وتيسير الأمور (مبارك)",
                6: "يورث الغم والوسواس",
                7: "يورث الهيبة والوقار والقبول عند الناس (يوم ملكي)",
                8: "يورث الضعف البدني وتراجع القوة",
                9: "يورث سوء الخلق وضيق الصدر",
                10: "يورث الرفعة والكرامة والمنزلة",
                14: "يورث الشفاء من الأوجاع الجسدية (مستحب للصحّة)",
                15: "يورث السرور والفرح وقضاء الحوائج (مبارك)",
                20: "يورث الأمن من الخوف والعز (مستحب)",
                24: "يورث الظفر بالأعداء ونيل المقاصد",
                25: "يورث الخلاص من الديون والضيق"
            }

            current_status = nail_rules.get(day, "يوم اعتيادي لقص الأظافر، والأفضل تحري يوم الجمعة.")
            is_suitable = "✅ مناسب ومستحب" if day in [3, 5, 7, 10, 14, 15, 20, 24, 25] else "❌ غير مناسب (يُفضل التأجيل)"

            # 3. بناء التقرير (Flet)
            report = f"📅 اليوم القمري (الهلالي): {day}\n"
            report += f"📍 حالة اليوم لقص الأظافر: {is_suitable}\n"
            report += "══════════════════════════════\n"
            report += f"📖 حكم هذا اليوم (حسب المخطوطات):\n{current_status}\n"
            report += "══════════════════════════════\n"
            report += "📜 من مخطوطات الإمام جعفر الصادق:\n"
            report += "يقول عليه السلام: 'تقليم الأظافر يوم الجمعة يؤمن من الجذام والبرص والعمى'...\n\n"
            report += "💎 عند ابن سيرين:\n"
            report += "يقول: 'الأظافر قوة الرجل ومقدرته'...\n"
            report += "══════════════════════════════\n"
            report += "✍️ المطور: حسان الشاعر"

            # 4. عرض التقرير في نافذة أندرويد (AlertDialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("💅 تقويم قص الأظافر الروحاني", weight="bold", rtl=True),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report, rtl=True, size=14, color="#2c3e50")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450,
                    bgcolor="#fffcf5" # لون الخلفية الأصلي
                ),
                actions=[
                    ft.TextButton("نسخ التقرير", on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ]
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as e:
            # تنبيه الخطأ في أندرويد (SnackBar)
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ: {str(e)}"))
                self.page.snack_bar.open = True
                self.page.update()
    def analyze_fixed_stars(self):
        try:
            # قاعدة البيانات الموسعة (18 نجماً)
            fixed_stars = [
                {"name": "الظليم (Achernar)", "pos": 15.3, "nature": "مشتري", "effect": "المكانة الروحية العالية، الاحترام العام، والنجاح في المساعي الدينية."},
                {"name": "رأس الغول (Algol)", "pos": 56.2, "nature": "زحل/مريخ", "effect": "أقوى نجوم السماء طاقة؛ يمنح صموداً أسطورياً وقدرة على مواجهة الأزمات الكبرى."},
                {"name": "الثريا (Alcyone)", "pos": 60.1, "nature": "زهرة/مشتري", "effect": "الجاذبية الشخصية، الشهرة الفنية، والنجاح في التجارة والعلوم."},
                {"name": "الدبران (Aldebaran)", "pos": 69.8, "nature": "مريخ", "effect": "عين الثور: الشجاعة، القيادة السيادية، والنجاح المادي من خلال المواجهات."},
                {"name": "رجل الجبار (Rigel)", "pos": 77.0, "nature": "مشتري/مريخ", "effect": "الشهرة الواسعة، الاختراع، والقدرة على القيادة العسكرية أو الإدارية."},
                {"name": "النطاق (Alnilam)", "pos": 83.5, "nature": "مشتري/ساترن", "effect": "الارتقاء للمناصب العليا، والسمعة التي تدوم طويلاً بعد غياب الشخص."},
                {"name": "منكب الجوزاء (Betelgeuse)", "pos": 88.8, "nature": "مريخ/مشتري", "effect": "نجاح باهر ومفاجئ، ثروة عظيمة، ومكانة مرموقة في الدولة."},
                {"name": "شعرى اليمانية (Sirius)", "pos": 104.1, "nature": "مشتري/مريخ", "effect": "النجم الأسطع: يمنح شهرة تاريخية، حماية ملكية، ونجاحاً يتجاوز الحدود."},
                {"name": "رأس التوأم (Pollux)", "nature": "مريخ", "pos": 113.3, "effect": "قوة البيان، الشجاعة في الدفاع عن المعتقدات، والقدرة على الغلبة."},
                {"name": "شعرى الشامية (Procyon)", "pos": 116.3, "nature": "مشتري/مريخ", "effect": "النجاح السريع والمفاجئ، الذكاء العملي، والقدرة على الإنجاز الفوري."},
                {"name": "قلب الأسد (Regulus)", "pos": 150.1, "nature": "مشتري/مريخ", "effect": "النجم الملكي: الجاه العظيم، السلطة، الرفعة السيادية، والكرامة العالية."},
                {"name": "السماك الأعزل (Spica)", "pos": 204.0, "nature": "زهرة/مشتري", "effect": "أكثر النجوم سعادة: الثروة، الثقافة الرفيعة، والنجاح العلمي الباهر."},
                {"name": "سماك الرامح (Arcturus)", "pos": 204.2, "nature": "مشتري/مريخ", "effect": "العبقرية، القيادة الفكرية، والنجاح من خلال الابتكار والتميز."},
                {"name": "قلب العقرب (Antares)", "pos": 240.0, "nature": "مريخ/مشتري", "effect": "التحولات الكبرى، المغامرة، والقدرة المذهلة على إعادة البناء بعد الهدم."},
                {"name": "النسر الواقع (Vega)", "pos": 285.3, "nature": "زهرة/عطارد", "effect": "الإبداع الروحاني، السحر البلاغي، والتميز الفني والروحاني الفريد."},
                {"name": "النسر الطائر (Altair)", "pos": 292.0, "nature": "مريخ/مشتري", "effect": "الطموح العالي جداً، القوة البدنية، والوصول للأهداف الصعبة والمستحيلة."},
                {"name": "فم الحوت (Fomalhaut)", "pos": 334.0, "nature": "زهرة/عطارد", "effect": "الشهرة العالمية، الوعي الكوني، والتميز في العلوم الروحانية والفلكية."},
                {"name": "الردف (Deneb)", "pos": 335.3, "nature": "زهرة/عطارد", "effect": "الذكاء الثاقب، القدرة على التعلم السريع، والنجاح في الفنون والعلوم."},
            ]

# دالة استخراج البرج
            def get_sign_name(deg):
                signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
                return signs[int(deg / 30)]

            # 2. إنشاء عناصر واجهة Flet
            report_display = ft.Column(scroll=ft.ScrollMode.ALWAYS, expand=True)

            def show_current_stars(e):
                report_display.controls.clear()
                report_display.controls.append(ft.Text("🌍 مواقع النجوم الثابتة اليوم (دليل 2026)", size=18, weight="bold", color="#38bdf8"))
                for s in fixed_stars:
                    report_display.controls.append(
                        ft.Container(
                            content=ft.Column([
                                ft.Text(f"⭐ {s['name']}", weight="bold", color="#fbbf24"),
                                ft.Text(f"📍 الموقع: {get_sign_name(s['pos'])} ({int(s['pos']%30)}°)", size=13),
                                ft.Text(f"✨ التأثير: {s['effect']}", size=13, italic=True),
                            ]),
                            padding=10, border=ft.border.all(1, "#38bdf8"), border_radius=8
                        )
                    )
                self.page.update()

            def show_natal_stars(e):
                report_display.controls.clear()
                if not hasattr(self, 'is_analyzed') or not self.is_analyzed:
                    report_display.controls.append(ft.Text("⚠️ خطأ: يرجى تحليل خارطة الميلاد من الواجهة الرئيسية أولاً.", color="red"))
                    self.page.update()
                    return

                report_display.controls.append(ft.Text("👶 تحليل اقترانات النجوم في ميلادك", size=18, weight="bold", color="#10b981"))
                points = {**self.raw_d, "Ascendant": self.asc_raw}
                p_names = {"sun":"الشمس", "moon":"القمر", "mercury":"عطارد", "venus":"الزهرة", "mars":"المريخ", "jupiter":"المشتري", "saturn":"زحل", "Ascendant":"الطالع"}
                
                found = False
                for s in fixed_stars:
                    for p_en, p_deg in points.items():
                        diff = abs(p_deg - s['pos'])
                        if diff <= 1.5 or diff > 358.5:
                            p_ar = p_names.get(p_en, p_en)
                            report_display.controls.append(
                                ft.Container(
                                    content=ft.Text(f"✅ اقتران قوي: نجم {s['name']} مع {p_ar}\n◈ الدلالة: {s['effect']}", size=14),
                                    padding=10, bgcolor="#e2e8f0", border_radius=8
                                )
                            )
                            found = True
                if not found:
                    report_display.controls.append(ft.Text("🔍 النتيجة: لم يتم العثور على اقترانات دقيقة بالنجوم العظمى."))
                self.page.update()

            # 3. عرض النافذة المنبثقة (Dialog)
            dlg = ft.AlertDialog(
                title=ft.Text("✨ ديوان النجوم الثابتة", size=20, weight="bold"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Row([
                            ft.ElevatedButton("مواقع اليوم", on_click=show_current_stars, bgcolor="#38bdf8", color="white"),
                            ft.ElevatedButton("نجوم ميلادي", on_click=show_natal_stars, bgcolor="#10b981", color="white"),
                        ], alignment=ft.MainAxisAlignment.CENTER),
                        ft.Divider(),
                        report_display
                    ], tight=True),
                    width=500, height=600
                )
            )

            self.page.dialog = dlg
            dlg.open = True
            show_current_stars(None) # العرض الافتراضي
            self.page.update()

        except Exception as e:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ في محرك النجوم: {str(e)}"))
            self.page.snack_bar.open = True
            self.page.update()
    def show_moon_mansion(self):
        try:

            # 1. الحسابات الفلكية الدقيقة (UT)
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, now.hour + now.minute/60.0)
            
            res_m, _ = swe.calc_ut(jd, swe.MOON)
            res_s, _ = swe.calc_ut(jd, swe.SUN)
            moon_pos, sun_pos = float(res_m[0]), float(res_s[0])
            
            # حساب المرحلة وعمر القمر
            phase_angle = (moon_pos - sun_pos) % 360
            moon_age = (phase_angle * 29.53) / 36            
            def get_phase(angle):
                if angle < 45: return "هلال وليد"
                elif angle < 90: return "تربيع أول"
                elif angle < 135: return "أحدب متزايد"
                elif angle < 180: return "بدر مكتمل"
                elif angle < 225: return "أحدب متناقص"
                elif angle < 270: return "تربيع ثاني"
                elif angle < 315: return "هلال غارب"
                else: return "محاق"

            # 2. تحديد المنزلة (1-28)
            mansion_idx = int(moon_pos / (360/28))
            m_num = mansion_idx + 1

            # 3. قاعدة البيانات العلمية والنجوم
            # 3. قاعدة البيانات العلمية والروحانية (توزيع هرمس - نظام القوائم المرتبة)
            # الترتيب: (عدد النجوم، البروج، الحرف، وصف النجوم، الشرح الموسع)
            STARS_INFO = {
                1: ("2", "الحمل", "أ", "نجمان (النطح والشيرتان)", "منزلة نارية، تدل على الحماس والبدايات والاندفاع القوي."),
                2: ("3", "الحمل", "ب", "ثلاثة نجوم خفية", "منزلة ترابية، تمتاز بالثبات والقدرة على التخطيط المالي."),
                3: ("7", "الثور", "ج", "الأخوات السبع (الثريا)", "منزلة هوائية، وهي من أسعد المنازل للمحبة والقبول والشهرة."),
                4: ("5", "الثور", "د", "عين الثور (الدبران)", "منزلة ترابية، تدل على الهيبة والقوة الفكرية العالية."),
                5: ("3", "الجوزاء", "هـ", "رأس الجبار (الهقعة)", "منزلة هوائية، مناسبة جداً للتعلم وتبادل المعلومات بذكاء."),
                6: ("2", "الجوزاء", "و", "قدم الجوزاء (الهنعة)", "منزلة مائية، تمتاز بطاقة الألفة وإصلاح العلاقات المتوترة."),
                7: ("2", "السرطان", "ز", "رأس التوأمين (الذراع)", "منزلة مائية، تمنح حماية للممتلكات وتساعد في الظفر."),
                8: ("3", "السرطان", "ح", "سديم المعلف (النثرة)", "منزلة نارية، تدل على الرفعة وعلو الشأن والنجاح الرسمي."),
                9: ("2", "الأسد", "ط", "عيني الأسد (الطرفة)", "منزلة نارية، تتطلب الحذر وهي قوية للتغييرات الجذرية."),
                10: ("4", "الأسد", "ي", "جبهة الأسد (الجبهة)", "منزلة نارية، تعيد الهيبة والشرف وتساعد في الاستشفاء."),
                11: ("2", "الأسد/العذراء", "ك", "ظهر الأسد (الزبرة)", "منزلة نارية، مباركة للتجارة والارتباط العاطفي الناجح."),
                12: ("1", "العذراء", "ل", "ذنب الأسد (الصرفة)", "منزلة ترابية، تدل على التغيير والانتقال، ومناسبة للسفر."),
                13: ("5", "العذراء", "م", "خمسة نجوم (العواء)", "منزلة ترابية، تمتاز بسرعة الاستجابة للأماني وقضاء الحوائج."),
                14: ("1", "الميزان", "ن", "السماك الأعزل", "منزلة هوائية، منبع النمو والازدهار في الصفقات المالية."),
                15: ("3", "الميزان", "س", "الغفر", "منزلة هوائية، تدل على الاستبصار والبحث في بواطن الأمور."),
                16: ("2", "الميزان/العقرب", "ع", "كفتا الميزان (الزبانا)", "منزلة هوائية، مناسبة لضبط الموازين المالية والحقوق."),
                17: ("3", "العقرب", "ف", "رأس العقرب (الإكليل)", "منزلة مائية، قوية جداً في التأليف بين القلوب وجذب المودة."),
                18: ("1", "العقرب", "ص", "قلب العقرب", "منزلة مائية، للتحصين والوقاية والتعامل مع النفس بعمق."),
                19: ("2", "العقرب/القوس", "ق", "ذيل العقرب (الشولة)", "منزلة نارية، تدل على التحرر من القيود والسرعة في الإنجاز."),
                20: ("8", "القوس", "ر", "النعائم", "منزلة نارية، منزلة الوفرة والرزق وتيسير الأمور المتعسرة."),
                21: ("1", "القوس/الجدي", "ش", "البلدة (فراغ سماوي)", "منزلة ترابية، منزلة الاستقرار وبناء الأصول والعمران."),
                22: ("2", "الجدي", "ت", "سعد الذابح", "منزلة ترابية، تدل على الحزم والنجاة من الأخطار المحيطة."),
                23: ("2", "الجدي/الدلو", "ث", "سعد بلع", "منزلة هوائية، مناسبة للعلاجات الطبية وتسهيل الصعاب."),
                24: ("3", "الدلو", "خ", "سعد السعود", "منبع السعادة والأفراح والنمو في كافة جوانب الحياة."),
                25: ("4", "الدلو", "ذ", "سعد الأخبية", "تساعد في كشف الأسرار والظهور القوي بعد فترة غياب."),
                26: ("2", "الحوت", "ض", "فرع الدلو المقدم", "منزلة مائية، تدل على الجاه والعز والتواصل الفكري."),
                27: ("2", "الحوت", "ظ", "فرع الدلو المؤخر", "منزلة مائية، مباركة لزيادة الأرزاق والنجاح في المشاريع."),
                28: ("1", "الحوت/الحمل", "غ", "الرشاء (بطن الحوت)", "منزلة مائية، منزلة الخواتيم السعيدة وإتمام المهام بنجاح.")
            }

 # 2. جلب البيانات (تأكد من وجود STARS_INFO في كلاسك)
            # سنفترض وجود 5 عناصر كما ذكرت (عدد، بروج، حرف، وصف، شرح)
            m_stars_cnt, m_signs, m_chars, m_stars_desc, m_explanation = self.STARS_INFO.get(m_num, (0, "", "", "", ""))
            name, angel, incense, effect, wifq, prayer = self.get_mansion_data(mansion_idx)

            # 3. بناء نص التقرير المنسق (بدلاً من tk.Text)
            report = (
                "✨ الديوان الشامل للمنازل القمرية ✨\n"
                "══════════════════════════\n"
                f"🌙 المنزلة: {name} (رقم: {m_num})\n"
                f"⏳ عمر القمر: {moon_age:.1f} يوم\n"
                f"📍 الموقع: {int(moon_pos%30)}° في برج {self.get_sign_name(moon_pos)}\n"
                "══════════════════════════\n"
                f"🌟 النجوم: {m_stars_desc}\n"
                f"♈ البروج: {m_signs} | 🔠 الحروف: ( {m_chars} )\n"
                "══════════════════════════\n"
                f"🔍 التفسير: {m_explanation}\n"
                "══════════════════════════\n"
                f"👼 الملك: {angel} | 🪵 البخور: {incense}\n"
                f"📖 التأثير: {effect}\n"
                "══════════════════════════\n"
                "🔢 الوفق العددي للمنزلة:\n"
                f"    [ {wifq[0]}  {wifq[1]}  {wifq[2]} ]\n"
                f"    [ {wifq[3]}  {wifq[4]}  {wifq[5]} ]\n"
                f"    [ {wifq[6]}  {wifq[7]}  {wifq[8]} ]\n\n"
                f"✨ العزيمة: 「 {prayer} 」\n"
                "══════════════════════════\n"
                "المطور الفلكي: حسان الشاعر"
            )

            # 4. واجهة العرض (AlertDialog للأندرويد)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text(f"منزلة {name}", weight="bold", rtl=True),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report, size=14, rtl=True, color="#3d405b")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, height=700, bgcolor="#f4f1de", padding=15, border_radius=10
                ),
                actions=[
                    ft.TextButton("نسخ التقرير", on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ]
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as e:
            # منع الشاشة البيضاء عبر SnackBar
            self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ في المنازل: {str(e)}"))
            self.page.snack_bar.open = True
            self.page.update()
    def analyze_current_time(self, e=None):
        try:

            # 1. إعداد محتوى التقرير
            content = "🔭 التحليل الزمني الشامل للأبراج والأسواق 🔭\n"
            content += "════════════════════════════════\n\n"
            
            content += "📍 أولاً: التحليل اليومي للأبراج\n"
            z_names = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                       "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            adv_list = ["طاقة نارية للمبادرة والحسم", "فرص ترابية للمال والعمل", 
                        "نشاط هوائي للذكاء والتواصل", "هدوء مائي للعاطفة والحدس"]
            
            for i, name in enumerate(z_names):
                adv = adv_list[i % 4]
                content += f"● {name}: {adv}.\n"

            content += "\n💰 ثانياً: الاقتصاد والذهب (2026)\n"
            content += "هذا الشهر: يشهد الذهب تذبذباً يميل للارتفاع نتيجة قوة المشتري. "
            content += "سنوياً: عام 2026 هو عام التحول الرقمي والذكاء الاصطناعي وصعود المعادن النفيسة.\n"
            
            content += "\n════════════════════════════════\n"
            content += "إعداد المطور: حسان الشاعر"

            # 2. دالة الحفظ السريع كملف نصي (متوافقة مع أندرويد)
            def save_report_file(e):
                try:
                    # الحفظ في مسار ملفات التطبيق المؤقتة لضمان الصلاحيات
                    file_path = "Strategic_Report_2026.txt"
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    self.page.snack_bar = ft.SnackBar(ft.Text(f"✅ تم حفظ التقرير باسم: {file_path}"))
                    self.page.snack_bar.open = True
                    self.page.update()
                except Exception as ex:
                    self.page.snack_bar = ft.SnackBar(ft.Text(f"❌ فشل الحفظ: {str(ex)}"))
                    self.page.snack_bar.open = True
                    self.page.update()

            # 3. إعداد النافذة المنبثقة (AlertDialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("التقرير الاستراتيجي الزمني", weight="bold", rtl=True, color="#1f538d"),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(content, size=14, rtl=True, color="#1e293b", weight="bold"),
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450,
                    height=600,
                    bgcolor="#f8fafc",
                    padding=20,
                    border_radius=15
                ),
                actions=[
                    # أزرار العمليات مدمجة في أسفل التقرير
                    ft.ElevatedButton("💾 حفظ الملف", icon=ft.icons.SAVE, on_click=save_report_file),
                    ft.ElevatedButton("📋 نسخ", icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(content)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.SPACE_EVENLY,
            )

            # 4. عرض النافذة
            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ: {str(ex)}"))
                self.page.snack_bar.open = True
                self.page.update()
    def hair_cut_calendar(self, e=None):
        try:
            # 1. الحساب الفلكي الدقيق لليوم القمري
            now = datetime.now(timezone.utc)
            jd = swe.julday(now.year, now.month, now.day, 12)
            res_m, _ = swe.calc_ut(jd, 1) # القمر
            res_s, _ = swe.calc_ut(jd, 0) # الشمس
            
            # استخراج الدرجة الأولى لضمان توافق أندرويد
            m_deg = float(res_m[0])
            s_deg = float(res_s[0])
            day = int(((m_deg - s_deg) % 360 / 12.2) + 1)
            if day > 30: day = 30 # قيد لليوم القمري

            # 2. قاعدة بيانات الأحكام (المخطوطات)
            hair_rules = {
                1: "يورث قصر العمر (تجنبه)",
                2: "يورث الهم والغم",
                3: "يورث النكد وتشتت الأمر",
                4: "يورث الوجع في الرأس (تجنبه)",
                5: "يورث البركة والمال الكثير",
                6: "يورث الهم وضيق الصدر",
                7: "يورث الهيبة والجاه والقبول عند الناس (مستحب جداً)",
                8: "يورث العافية وطول العمر",
                9: "يورث الضعف وسوء الخلق",
                10: "يورث الرفعة والكرامة",
                11: "يورث الحزن والكآبة",
                12: "يورث العز والوقار (مستحب)",
                13: "يورث الخصومة والنزاع",
                14: "يورث السرور والفرح وقضاء الحاجة (مبارك)",
                15: "يورث تيسير الأمور وزيادة الجمال (مستحب)",
                16: "يورث الحزن وضيق الرزق",
                17: "يورث النحوسة (تجنبه)",
                18: "يورث البركة في الرزق",
                19: "يورث القدرة والتمكين",
                20: "يورث الأمن من الخوف والهيبة (مبارك)",
                21: "يورث تشتت الرأي",
                22: "يورث الفقر الشديد (تجنبه)",
                23: "يورث العافية من الأمراض",
                24: "يورث النصر على الأعداء",
                25: "يورث الخلاص من السجون والديون",
                26: "يورث الفرح والسرور",
                27: "يورث السلامة من الآفات",
                28: "يورث قضاء الحاجة (مستحب)",
                29: "يورث الفقر والهم (تجنبه)",
                30: "يورث الأمن من الآفات"
            }

            current_status = hair_rules.get(day, "يوم مجهول الحكم")
            suitable_days = [5, 7, 8, 10, 12, 14, 15, 18, 19, 20, 23, 24, 25, 26, 27, 28, 30]
            is_suitable = "✅ مناسب ومستحب" if day in suitable_days else "❌ غير مناسب (يُفضل التأجيل)"

            # 3. صياغة نص التقرير الملحمي
            report = (
                f"✂️ تقويم قص الشعر الروحاني ✂️\n"
                f"📅 اليوم القمري الحالي: {day}\n"
                f"📍 الحالة: {is_suitable}\n"
                f"══════════════════════════════\n"
                f"📖 حكم اليوم: يورث {current_status}\n"
                f"══════════════════════════════\n"
                f"📜 من مخطوطات جعفر الصادق:\n'من قص في السابع نال هيبة، وفي الـ 14 نال سروراً'.\n\n"
                f"💎 فلسفة ابن سيرين:\nقص الشعر في الأيام المباركة يرمز لذهاب الهم وقضاء الدين.\n"
                f"══════════════════════════════\n"
                f"✍️ إعداد المطور: حسان الشاعر"
            )

            # 4. عرض النتيجة في AlertDialog (أفضل للأندرويد)
            self.page.dialog = ft.AlertDialog(
                title=ft.Text("نتائج التقويم الروحاني", rtl=True),
                content=ft.Text(report, rtl=True, size=16),
                actions=[ft.TextButton("إغلاق", on_click=lambda _: self.page.close(self.page.dialog))]
            )
            self.page.dialog.open = True
            self.page.update()

        except Exception as ex:
            self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ: {str(ex)}"))
            self.page.snack_bar.open = True
            self.page.update()
            # 4. دالة الحفظ السريع
            def save_hair_report(e):
                try:
                    with open("Hair_Cut_Report.txt", "w", encoding="utf-8") as f:
                        f.write(report)
                    self.page.snack_bar = ft.SnackBar(ft.Text("✅ تم حفظ التقرير بنجاح"))
                    self.page.snack_bar.open = True
                    self.page.update()
                except: pass

            # 5. واجهة العرض (Flet AlertDialog)
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("تقويم قص الشعر الروحاني", weight="bold", rtl=True),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report, size=14, rtl=True, color="#1e293b", weight="bold")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=450, height=600, bgcolor="#f8fafc", padding=20, border_radius=10
                ),
                actions=[
                    ft.ElevatedButton("💾 حفظ", icon=ft.icons.SAVE, on_click=save_hair_report),
                    ft.ElevatedButton("📋 نسخ", icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard(report)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.SPACE_EVENLY,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            if hasattr(self, 'page'):
                self.page.snack_bar = ft.SnackBar(ft.Text(f"خطأ: {str(ex)}"))
                self.page.snack_bar.open = True
                self.page.update()
    def update_location(self, e=None):
        try:
            # استخدام .value بدلاً من .get() لـ Flet
            self.lat = float(self.lat_entry.value)
            self.lon = float(self.lon_entry.value)
            
            # تنبيه أندرويد (SnackBar) بدلاً من MessageBox
            self.page.snack_bar = ft.SnackBar(
                ft.Text(f"✅ تم التحديث: عرض {self.lat} | طول {self.lon}"),
                bgcolor="#10b981"
            )
            self.page.snack_bar.open = True
            self.draw_now() # إعادة الحساب فوراً
            self.page.update()
        except ValueError:
            self.page.snack_bar = ft.SnackBar(ft.Text("⚠️ خطأ: يرجى إدخال أرقام صحيحة"))
            self.page.snack_bar.open = True
            self.page.update()
    def get_house(self, pos, cusps):
        for i in range(12):
            c1, c2 = cusps[i], cusps[(i+1)%12]
            if (c1 < c2 and c1 <= pos < c2) or (c1 > c2 and (pos >= c1 or pos < c2)): return i + 1
        return 1
    def show_about(self, e=None):
        email = "Hassan.alshaer2@gmail.com"
        is_ar = self.current_lang == "ar"
        
        lines = [
            f"{'المطور البرمجي الفلكي: حسان الشاعر' if is_ar else 'Developer: Hassan Al-Shaer'}",
            f"Email: {email}",
            f"{'العنوان: سوريا / طرطوس' if is_ar else 'Location: Syria / Tartus'}",
            "WhatsApp: +963 933303612",
            "----------------------------------",
            f"{'برامجنا الاحترافية:' if is_ar else 'Professional Software:'}",
            "* Animodar Correction System",
            "* Primary Directions Pro",
            "* Al-Khafia Solar System"
        ]

        def close_dlg(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            title=ft.Text("About Developer" if not is_ar else "حول المطور", weight="bold"),
            content=ft.Container(
                content=ft.Text("\n".join(lines), color="#39FF14", size=14, font_family="monospace"),
                bgcolor="#000000",
                padding=20,
                border_radius=10,
                border=ft.border.all(2, "#39FF14") # لون الفسفور الذي اخترته
            ),
            actions=[ft.TextButton("Close" if not is_ar else "إغلاق", on_click=close_dlg)],
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
    def show_mundane_analysis(self, jd):
        try:

            # 1. جلب الإحداثيات والبيانات الأساسية
            lat = float(self.lat) if hasattr(self, 'lat') else 34.8
            lon = float(self.lon) if hasattr(self, 'lon') else 35.8
            city_name = self.city_entry.value if hasattr(self, 'city_entry') else "الموقع المحدد"

            # 2. المحرك الحسابي (حساب الطالع والسهام والكواكب)
            cusps, ascmc = swe.houses(jd, lat, lon, b'P')
            asc_raw = float(ascmc[0])
            
            def get_p(obj_id):
                res, _ = swe.calc_ut(jd, obj_id)
                # res[0] الموقع، res[3] السرعة لتحديد التراجع
                return {'deg': float(res[0]), 'sign': int(res[0]/30), 'speed': float(res[3])}

            p = {n: get_p(i) for n, i in [('sun', 0), ('moon', 1), ('mercury', 2), ('venus', 3), 
                                           ('mars', 4), ('jupiter', 5), ('saturn', 6)]}

            z_names_ar = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
            z_lords = ["المريخ", "الزهرة", "عطارد", "القمر", "الشمس", "عطارد", "الزهرة", "المريخ", "المشتري", "زحل", "زحل", "المشتري"]
            
            asc_idx = int(asc_raw / 30)
            ruler_name = z_lords[asc_idx]

            # حساب السهام الاستراتيجية (الذهب، النفط، الغذاء)
            gold_lot = (asc_raw + p['jupiter']['deg'] - p['sun']['deg']) % 360
            oil_lot = (asc_raw + p['mars']['deg'] - p['jupiter']['deg']) % 360
            food_lot = (asc_raw + p['mars']['deg'] - p['saturn']['deg']) % 360

            # 3. بناء نص التقرير الملحمي (التنسيق المفصل)
            
            # --- [القسم 1: المقدمة والمنهجية] ---
            report_body = (
                "📜 الديوان السلطاني الشامل لحكم سنة العالم 2026 📜\n"
                "════════════════════════════════\n\n"
                "📖 منهجية الاستخراج:\n"
                f"يُستخرج هذا الطالع لحظة دخول الشمس (0° حمل). "
                f"الطالع المحلي لمدينة {city_name} هو الذي يحدد نصيب الإقليم من الأحداث.\n\n"
                
                "📍 بيانات الطالع وصاحب السنة:\n"
                f"◈ الطالع المحلي: برج {z_names_ar[asc_idx]} بالدرجة {int(asc_raw%30)}°.\n"
                f"◈ رب السنة (المستولي): كوكب {ruler_name}.\n"
                f"◈ حالة الحاكم: {'في قوة واستقامة (دليل إنجاز)' if p[('sun' if ruler_name=='الشمس' else 'moon' if ruler_name=='القمر' else 'mercury' if ruler_name=='عطارد' else 'venus' if ruler_name=='الزهرة' else 'mars' if ruler_name=='المريخ' else 'jupiter' if ruler_name=='المشتري' else 'saturn')]['speed'] > 0 else 'في حالة تراجع (دليل تأخير وإعادة نظر)'}.\n\n"
                
                "🌍 أولاً: شؤون الأقاليم والسياسة العالمية:\n"
                "◈ الأقاليم الشمالية: تعيش حالة من 'الجمود الاستراتيجي' وإعادة تموضع القوى العسكرية.\n"
                "◈ إقليم الشرق الأوسط: صعود سيادي غير مسبوق؛ وتحول المنطقة لمركز ثقل سياسي واقتصادي يفرض شروطه.\n"
                "◈ إقليم الشرق الأقصى: قوة تكنولوجية كاسحة تغير موازين القوى في المحيطات.\n"
                "◈ الأقاليم الغربية: حالة من التراجع نتيجة اضطرابات في القوانين الداخلية.\n\n"
                
                "💰 ثانياً: الاقتصاد، الذهب، والعملات الرقمية:\n"
                f"◈ سهم الذهب: استقر في {z_names_ar[int(gold_lot/30)]}، مما ينبئ بـ 'زلزال مالي' يعيد الاعتبار للأصول الحقيقية.\n"
                "◈ الذهب والفضة: يتجهان لتحطيم قمم تاريخية نتيجة فقدان الثقة في العملات الورقية.\n"
                "◈ العملات الرقمية: تدخل 'عصر القوننة الشاملة' بظهور عملات مدعومة من الدول والذهب.\n\n"
                
                "🛢️ ثالثاً: ميزان الطاقة والبترول والغاز:\n"
                f"◈ سهم البترول: استقر في برج {z_names_ar[int(oil_lot/30)]}.\n"
                "◈ أسعار الطاقة: اشتعال مفاجئ نتيجة 'أزمات لوجستية' أو إغلاق ممرات مائية حيوية.\n"
                "◈ التحول الطاقي: الاعتماد على 'الهيدروجين والطاقة البديلة' يتسارع بشكل قسري.\n\n"
            )

            # --- [القسم 2: القواطع والمناخ والغذاء] ---
            is_danger = p['saturn']['speed'] < 0 or p['mars']['speed'] < 0
            if is_danger:
                report_body += (
                    "🚨 رابعاً: تحذير القواطع الكونية والمناخ:\n"
                    "◈ الزلازل: خطر مرتفع لهزات أرضية تضرب 'خطوط الصدع' الدولية (المتوسط وجنوب آسيا).\n"
                    "◈ المناخ: سنة 'التطرف المناخي'؛ فيضانات عارمة تضرب السواحل والموانئ.\n\n"
                )
            else:
                report_body += (
                    "✅ رابعاً: إشارة السلامة الكونية والمناخ:\n"
                    "◈ المناخ: اعتدال العناصر، أمطار في أوانها ورياح لواقح، وسكون في طبقات الأرض.\n\n"
                )

            report_body += (
                "🌾 خامساً: ميزان الغذاء والحبوب (سهم القمح):\n"
                f"◈ سهم الغذاء: في {z_names_ar[int(food_lot/30)]} ينبئ بصعوبات في الأمن الغذائي.\n"
                "◈ القمح: سيكون 'سلاح العصر' الجديد مع تقلبات حادة في الأسعار.\n"
                "◈ نصيحة: يُنصح بتأمين مخزونات الغذاء الأساسية قبل الربع الثالث من السنة.\n\n"
                "════════════════════════════════\n"
                "✍️ إعداد وتطوير: حسان الشاعر © 2026\n"
                "جميع الحقوق البرمجية والفلكية محفوظة للمطور"
            )

            # 4. واجهة العرض (Flet AlertDialog) لضمان عدم ظهور شاشة بيضاء
            def close_dlg(e):
                dlg.open = False
                self.page.update()

            dlg = ft.AlertDialog(
                title=ft.Text("الديوان السلطاني الشامل - سنة 2026", weight="bold", color="#7c2d12", rtl=True),
                content=ft.Container(
                    content=ft.Column([
                        ft.Text(report_body, size=15, rtl=True, color="#1e40af", weight="bold")
                    ], scroll=ft.ScrollMode.ADAPTIVE, tight=True),
                    width=550, height=750, bgcolor="#fcfaf2", padding=20, border_radius=10
                ),
                actions=[
                    ft.TextButton("نسخ الديوان بالكامل", on_click=lambda _: self.page.set_clipboard(report_body)),
                    ft.TextButton("إغلاق", on_click=close_dlg)
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            self.page.dialog = dlg
            dlg.open = True
            self.page.update()

        except Exception as ex:
            # رسالة خطأ متوافقة مع أندرويد لسهولة التصحيح
            self.page.snack_bar = ft.SnackBar(ft.Text(f"حدث خطأ في توليد التقرير: {str(ex)}"))
            self.page.snack_bar.open = True
            self.page.update()
    def save_report(self, content):
        try:
            import os
            from datetime import datetime
            
            # إنشاء اسم ملف فريد
            filename = f"Hassan_Master_Report_2026_{int(datetime.now().timestamp())}.txt"
            
            # الحفظ في مسار التطبيق (متوافق مع أندرويد)
            with open(filename, "w", encoding="utf-8") as f:
                f.write(content)
            
            # تنبيه النجاح في أندرويد (SnackBar)
            self.page.snack_bar = ft.SnackBar(
                ft.Text(f"✅ تم حفظ التقرير باسم: {filename}"),
                bgcolor="#10b981"
            )
            self.page.snack_bar.open = True
            self.page.update()
            
        except Exception as e:
            # تنبيه الخطأ
            self.page.snack_bar = ft.SnackBar(ft.Text(f"❌ تعذر حفظ الملف: {str(e)}"))
            self.page.snack_bar.open = True
            self.page.update()

    def display_comprehensive_report(self, data):

        # 1. إعداد لوحة الألوان الملحمية
        colors = {
            "gold": "#f59e0b",   # اللون الذهبي للعناوين
            "sky": "#38bdf8",    # الأزرق السماوي للبيانات
            "pink": "#fb7185",   # الوردي للتحليل
            "danger": "#ef4444", # الأحمر للقواطع
            "info": "#10b981",   # الأخضر للشرح
            "border": "#334155"  # لون الإطارات
        }

        # 2. إنشاء قائمة العناصر (Widgets) للتقرير
        report_widgets = []

        # --- القسم الأول: المقدمة ---
        report_widgets.append(
            ft.Text("📜 التقرير الفلكي التحليلي الشامل المطور", 
                    size=22, weight="bold", color=colors["gold"], text_align="center")
        )
        report_widgets.append(ft.Divider(color=colors["gold"], height=20))

        # --- القسم الثاني: تحليل الكواكب والعقد ---
        planets_config = [
            ("الشمس", "sun_sign", "sun_house", "☀️"), ("القمر", "moon_sign", "moon_house", "🌙"),
            ("عطارد", "mercury_sign", "mercury_house", "☿"), ("الزهرة", "venus_sign", "venus_house", "♀️"),
            ("المريخ", "mars_sign", "mars_house", "♂️"), ("المشتري", "jupiter_sign", "jupiter_house", "♃"),
            ("زحل", "saturn_sign", "saturn_house", "♄"), ("أورانوس", "uranus_sign", "uranus_house", "⛢"),
            ("نبتون", "neptune_sign", "neptune_house", "♆"), ("بلوتو", "pluto_sign", "pluto_house", "♇"),
            ("العقدة الشمالية", "nn_sign", "nn_house", "🐲"),
            ("العقدة الجنوبية", "south_node", "sn_house", "🐉")
        ]

        for name, s_key, h_key, icon in planets_config:
            sign = data.get(s_key, "---")
            house = data.get(h_key, "1")
            meaning = PLANET_MEANINGS.get(name, "تأثير كوكبي فريد يشكل ملامح شخصيتك.")
            flavor = SIGN_FLAVORS.get(sign, "بطابع ممتزج.")
            field = HOUSE_FIELDS.get(str(house), "في هذا المجال الحيوي.")

            report_widgets.append(
                ft.Container(
                    content=ft.Column([
                        ft.Text(f"{icon} {name} في برج {sign} | البيت {house}", 
                                color=colors["sky"], weight="bold", size=16),
                        ft.Text(f"◈ التحليل: {meaning} {flavor} ويتجلى أثره في {field}", 
                                color=colors["pink"], size=14),
                    ], spacing=5),
                    padding=12,
                    border=ft.border.all(1, colors["sky"]),
                    border_radius=10,
                    margin=ft.margin.only(bottom=10)
                )
            )

        # --- القسم الثالث: السهام الفلكية ---
        report_widgets.append(ft.Text("🌟 ثانياً: السهام الفلكية السرية", size=19, weight="bold", color=colors["gold"]))
        parts = [
            ("سهم السعادة", data.get('fortune_part'), "نقطة الرزق المادي والقبول الاجتماعي والبهجة النفسية."),
            ("سهم الغيب", data.get('spirit_part'), "نقطة الإرادة الباطنة والنوايا الروحية والذكاء العميق.")
        ]
        for p_title, p_val, p_desc in parts:
            report_widgets.append(
                ft.Container(
                    content=ft.Column([
                        ft.Text(f"📍 {p_title}: {p_val}", color=colors["sky"], weight="bold", size=16),
                        ft.Text(f"💡 الشرح: {p_desc}", color=colors["info"], size=14),
                    ], spacing=5),
                    padding=12, border=ft.border.all(1, colors["info"]), border_radius=10, margin=ft.margin.only(bottom=10)
                )
            )

        # --- القسم الرابع: الفردارية (حاكم العمر) ---
        report_widgets.append(ft.Text("⏳ ثالثاً: حاكم المرحلة الحالية (الفردارية)", size=19, weight="bold", color=colors["gold"]))
        report_widgets.append(
            ft.Container(
                content=ft.Column([
                    ft.Text(f"👑 الحاكم المسيطر: {data.get('firdar_ruler')} - {data.get('firdar_title', '')}", 
                            color=colors["sky"], weight="bold", size=16),
                    ft.Text(f"📝 شرح المرحلة: {data.get('firdar_desc', 'مرحلة هامة لتشكيل المصير.')}", 
                            color=colors["info"], size=14, italic=True),
                ], spacing=5),
                padding=12, border=ft.border.all(1, colors["gold"]), border_radius=10, margin=ft.margin.only(bottom=10)
            )
        )

        # --- القسم الخامس: تقرير النجوم الثابتة ---
        report_widgets.append(ft.Text("🔭 رابعاً: تقرير النجوم الثابتة (أحكام الولادة)", size=19, weight="bold", color=colors["gold"]))
        stars_rep = data.get('stars_report', "")
        star_items = [ft.Text("💡 المنهجية: تحليل مواقع النجوم لحظة ولادتك مقارنة بكواكب ميلادك الشخصية.", size=12, italic=True, color=colors["info"])]
        
        if "لا يوجد" in stars_rep or not stars_rep:
            star_items.append(ft.Text("🔍 النتيجة: لم يتم العثور على اقترانات دقيقة بالنجوم العظمى.", color="#94a3b8"))
        else:
            for line in stars_rep.split('\n'):
                if line.strip():
                    star_items.append(ft.Text(line, color=colors["sky"] if "⭐" in line else "#f8fafc", size=14))

        report_widgets.append(
            ft.Container(content=ft.Column(star_items, spacing=5), padding=12, border=ft.border.all(1, colors["border"]), border_radius=10, margin=ft.margin.only(bottom=10))
        )

        # --- القسم السادس: القواطع والتحذيرات ---
        report_widgets.append(ft.Text("⚠️ خامساً: القواطع والتحذيرات الفلكية", size=19, weight="bold", color=colors["gold"]))
        warnings = data.get('warnings', "لا يوجد قواطع شديدة.")
        deep_interpretations = {
            "تراجع": "يشير التراجع إلى طاقة داخلية؛ مراجعة للذات وتأخر في الشؤون الخارجية.",
            "احتراق": "الكوكب تحت شعاع الشمس؛ دلالاته ضعيفة ظاهرياً وقوية في الخفاء.",
            "درجة حرجة": "وضعية 'المنقلب'؛ تشير إلى تحولات جذرية في دلالة الكوكب.",
            "اقتران النحسين": "ضغط عالٍ يتطلب الحذر الشديد لتجنب القرارات المتهورة."
        }

        warn_list = []
        if "لا يوجد" in warnings:
            warn_list.append(ft.Text("✅ لا يوجد قواطع شديدة في هذه الهيئة.", color=colors["info"], weight="bold"))
        else:
            for line in warnings.split('\n'):
                if line.strip():
                    warn_list.append(ft.Text(f"🛑 {line}", color=colors["danger"], weight="bold", size=14))
                    interp = next((v for k, v in deep_interpretations.items() if k in line), "يتطلب هذا الوضع انتباهاً لإعادة التوازن.")
                    warn_list.append(ft.Text(f"   🔍 التفسير: {interp}", color=colors["info"], size=13))

        report_widgets.append(
            ft.Container(content=ft.Column(warn_list, spacing=5), padding=12, border=ft.border.all(1, colors["danger"]), border_radius=10, margin=ft.margin.only(bottom=10))
        )

        # --- التذييل والخاتمة ---
        report_widgets.append(ft.Divider(color=colors["gold"]))
        report_widgets.append(ft.Text("إعداد وتطوير: حسان الشاعر © 2026", color=colors["gold"], size=14, text_align="center", weight="bold"))

        # 3. عرض التقرير في نافذة AlertDialog (متوافقة مع أندرويد)
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        # دالة تجميع النص للنسخ
        def get_full_text():
            return "تقرير فلكي شامل - حسان الشاعر 2026" # يمكن تطويرها لتجمع كل النصوص

        dlg = ft.AlertDialog(
            title=ft.Text("الديوان التحليلي السلطاني", color=colors["gold"], weight="bold", rtl=True),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ALWAYS, spacing=15),
                width=650, height=1200, bgcolor="#0f172a", padding=20
            ),
            actions=[
                ft.ElevatedButton("📋 نسخ النص", on_click=lambda _: self.page.set_clipboard("تم نسخ التقرير!")),
                ft.TextButton("إغلاق", on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()

    def display_comprehensive_report(self, data):
        """الدالة الملحمية الكاملة - التقرير السلطاني الموسع (نسخة الأندرويد)"""
        
        # 1. إعداد الحاوية وقاموس الألوان
        report_widgets = []
        self.full_report_content = "" # النص الذي سيتم حفظه في ملف Word
        
        colors = {
            "header": "#f59e0b", "planet_title": "#38bdf8", "body": "#e2e8f0",
            "danger": "#ef4444", "border": "#475569", "gold": "#FFD700",
            "info": "#a8a29e", "sky": "#0ea5e9"
        }

        # دالة داخلية لإضافة النصوص للواجهة والملف في آن واحد
        def log_it(text, style_type):
            self.full_report_content += text
            report_widgets.append(
                ft.Text(
                    value=text,
                    color=colors.get(style_type, "#ffffff"),
                    size=22 if style_type == "header" else 17 if style_type == "planet_title" else 14,
                    weight="bold" if style_type in ["header", "planet_title", "gold"] else None,
                    italic=True if style_type == "info" else False,
                    rtl=True
                )
            )

        # 2. القواميس الموسعة (الشروحات التي أرسلتها)
        PART_DESCRIPTIONS = {
            "fortune_part": "سهم السعادة (Part of Fortune): هو نقطة تلاقي الروح والجسد والمادة. يمثل 'البوابة اليدوية' للرزق، حيث يشير موقعه إلى المجال الذي يجد فيه الإنسان سهولة في الكسب المادي والقبول الاجتماعي والشعور بالرضا والبهجة.",
            "spirit_part": "سهم الغيب (Part of Spirit): هو سهم النوايا والارادة الباطنة. يعبر عن الجانب الروحي العميق، ويحدد أين يضع الإنسان شغفه الحقيقي وتطلعاته التي لا يخبر بها أحداً، وهو مفتاح النجاح المعنوي والفكري."
        }

        deep_interpretations = {
            "تراجع عطارد": "اضطراب في التواصل الرقمي واللفظي؛ قد تجد صعوبة في إيصال فكرتك أو تتعرض أجهزتك لأعطال فنية. أعد قراءة العقود مرتين.",
            "تراجع الزهرة": "مرحلة مراجعة للمشاعر والقيم المالية؛ ليس وقتاً جيداً لتغييرات جذرية في المظهر.",
            "تراجع المريخ": "انخفاض في مستوى الطاقة البدنية والمبادرة؛ وجه هذه الطاقة للداخل لإعادة تقييم أهدافك.",
            "احتراق": "الكوكب في حالة 'غياب قسري'؛ طاقته ذائبة في سطوع الشمس، مما يجعل دلالاته تعمل خلف الكواليس.",
            "درجة حرجة": "الكوكب يقف على 'عتبة' تغيير كبير؛ حالة من عدم الاستقرار تنبئ بنهاية مرحلة وبداية أخرى.",
            "اقتران النحسين": "اجتماع طاقة المنع (زحل) مع طاقة الاندفاع (المريخ)؛ يتطلب انضباطاً حديدياً لتجنب الانفجار."
        }

        # 3. بناء المحتوى - (أ) القسم الافتتاحي
        log_it("📜 التقرير الفلكي التحليلي الشامل\n", "header")
        log_it("═" * 40 + "\n\n", "border")

        # (ب) قسم الكواكب والعقد
        planets_to_show = [
            ("الشمس", "sun_sign", "sun_house", "☀️"), ("القمر", "moon_sign", "moon_house", "🌙"),
            ("عطارد", "mercury_sign", "mercury_house", "☿"), ("الزهرة", "venus_sign", "venus_house", "♀️"),
            ("المريخ", "mars_sign", "mars_house", "♂️"), ("المشتري", "jupiter_sign", "jupiter_house", "♃"),
            ("زحل", "saturn_sign", "saturn_house", "♄"),
            ("العقدة الشمالية", "north_node", "nn_house", "🐲"),
            ("العقدة الجنوبية", "south_node", "sn_house", "🐉")
        ]

        for name, s_key, h_key, icon in planets_to_show:
            sign = data.get(s_key, "---")
            house = data.get(h_key, "1")
            log_it(f"╔{'═'*35}╗\n", "border")
            log_it(f"  {icon} {name} في برج {sign} | البيت {house}\n", "planet_title")
            
            # جلب الشرح (كوكب أو عقدة)
            explanation = data.get(f"{s_key}_desc", "تأثير فلكي فريد يشكل جزءاً من شخصيتك.")
            log_it(f"  ◈ الشرح: {explanation}\n", "body")
            log_it(f"╚{'═'*35}╝\n\n", "border")

        # (ج) قسم النجوم الثابتة
        log_it("🔭 تقرير النجوم الثابتة (الاقترانات الشخصية)\n", "gold")
        log_it("💡 المنهجية: لا يظهر النجم إلا إذا اقترن بكوكبك بمسافة لا تتعدى 1.5 درجة.\n", "info")
        stars_rep = data.get('stars_report', "لا توجد اقترانات وثيقة حالياً.")
        log_it(f"{stars_rep}\n", "sky")

        # (د) قسم السهام الفلكية
        log_it("🌟 السهام الفلكية ونقاط القوة\n", "header")
        for s_key, s_title in [("fortune_part", "سهم السعادة"), ("spirit_part", "سهم الغيب")]:
            val = data.get(s_key, "غير محسوب")
            log_it(f"📍 {s_title}: {val}\n", "planet_title")
            log_it(f"💡 {PART_DESCRIPTIONS[s_key]}\n", "body")

        # (هـ) قسم القواطع والتحذيرات الذكي
        log_it("\n⚠️ القواطع والتحذيرات الفلكية الشاملة\n", "header")
        warnings = data.get('warnings', "لا يوجد قواطع شديدة.")
        
        warn_widgets = []
        if "لا يوجد" in warnings:
            warn_widgets.append(ft.Text("✅ الهيئة خالية من القواطع، مما يشير لسريان طاقة الكواكب طبيعياً.", color=colors["info"], rtl=True))
        else:
            for line in warnings.split('\n'):
                if line.strip():
                    warn_widgets.append(ft.Text(f"🛑 {line}", color=colors["danger"], weight="bold", rtl=True))
                    found_desc = "يتطلب هذا الوضع انتباهاً خاصاً لضمان توازن المسار."
                    for key, desc in deep_interpretations.items():
                        if key in line: found_desc = desc; break
                    warn_widgets.append(ft.Text(f"   🔍 التفسير: {found_desc}", size=13, color=colors["info"], italic=True, rtl=True))
                    self.full_report_content += f"🛑 {line}\n   🔍 التفسير: {found_desc}\n"

        report_widgets.append(ft.Container(content=ft.Column(warn_widgets), padding=10, border=ft.border.all(1, colors["danger"]), border_radius=8))

        # (و) خاتمة المطور
        log_it("\n" + "─"*30 + "\n", "border")
        log_it("✍️ إعداد وتطوير: حسان الشاعر\nجميع الحقوق محفوظة للمطور © 2026", "header")

        # 4. دالة الحفظ والدالة النهائية للنافذة
        def save_to_word_internal(e):
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('تقرير الميقاتي برو الموسع', 0)
                doc.add_paragraph(self.full_report_content)
                file_name = f"Meeqat_Report_{int(time.time())}.docx"
                # مسار أندرويد
                path = os.path.join(self.page.client_storage.get("app_path") or os.getcwd(), file_name)
                doc.save(path)
                self.page.open(ft.SnackBar(ft.Text(f"✅ تم الحفظ بنجاح: {file_name}")))
            except Exception as ex:
                self.page.open(ft.SnackBar(ft.Text(f"❌ خطأ: {ex}")))

        dlg = ft.AlertDialog(
            title=ft.Text("التقرير التحليلي الملحمي", rtl=True),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ALWAYS, expand=True),
                width=650, height=900
            ),
            actions=[
                ft.ElevatedButton("💾 حفظ Word", on_click=save_to_word_internal, bgcolor=colors["header"], color="white"),
                ft.TextButton("إغلاق", on_click=lambda _: self.page.close(dlg))
            ],
            actions_alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        )
        
        self.page.open(dlg)
        self.page.update()
        def close_dlg(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            modal=False,
            title=ft.Text("الديوان التحليلي السلطاني", weight="bold", color=colors["header"], rtl=True),
            content=ft.Container(
                content=ft.Column(report_widgets, scroll=ft.ScrollMode.ADAPTIVE, spacing=15),
                width=550, height=850, bgcolor="#0f172a", padding=20
            ),
            actions=[
                ft.ElevatedButton("📋 نسخ للواتساب", icon=ft.icons.COPY, on_click=lambda _: self.page.set_clipboard("نص التقرير")),
                ft.TextButton("إغلاق", on_click=close_dlg)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dlg
        dlg.open = True
        self.page.update()
    def get_sign_name(self, total_deg):
        # الكود الخاص بك صحيح تماماً هنا
        signs = ["الحمل", "الثور", "الجوزاء", "السرطان", "الأسد", "العذراء", 
                 "الميزان", "العقرب", "القوس", "الجدي", "الدلو", "الحوت"]
        return signs[int((total_deg % 360) / 30)]

def my_astronomy_logic(user_input):
    # هنا تضع منطق الحساب الخاص بك مستخدماً الـ 2500 معطى
    # مثال: نتيجة = المدخل + أحد المعطيات
     result = user_input * 0.123  # عدل هذه المعادلة لتناسب برنامجك
     return result


# --- إعداداتك الأصلية ---
SECRET_SALT = "MEEQAT_PRO_2026_ULTIMATE" 
TRIAL_DAYS = 3
DATA_PATH = os.getcwd() 
TRIAL_FILE = os.path.join(DATA_PATH, "trial_info.dat")
KEY_FILE = os.path.join(DATA_PATH, "license.lic")

# 1. دالة جلب بصمة الموبايل (HWID)
# 1. دالة جلب البصمة (المتوافقة مع أندرويد)
def get_android_hwid(page: ft.Page):
    if not page.client_storage.contains_key("hwid"):
        # ملاحظة: platform.machine() قد تعيد نصاً فارغاً في بعض هواتف أندرويد
        # لذا سنضيف "AndroidDevice" كقيمة احتياطية لضمان عدم حدوث خطأ
        seed = f"{platform.machine() or 'AndroidDevice'}-{str(uuid.uuid4())[:8]}"
        page.client_storage.set("hwid", hashlib.md5(seed.encode()).hexdigest()[:16].upper())
    return page.client_storage.get("hwid")

def get_remaining_days(page: ft.Page):
    today = datetime.now()
    if not page.client_storage.contains_key("first_run"):
        page.client_storage.set("first_run", today.strftime("%Y-%m-%d"))
        return TRIAL_DAYS
    
    try:
        first_run_str = page.client_storage.get("first_run")
        first_date = datetime.strptime(first_run_str, "%Y-%m-%d")
        elapsed = (today - first_date).days
        return max(0, TRIAL_DAYS - elapsed)
    except:
        return 0 # في حال حدث خطأ في تنسيق التاريخ المخزن# 3. دالة التحقق من التفعيل الأصلي
# 1. دالة التحقق من النسخة الكاملة
def is_full_version(page: ft.Page, secret_salt: str):
    # نتحقق أولاً هل الملف موجود (مع مراعاة مسار الأندرويد)
    if os.path.exists(KEY_FILE):
        try:
            with open(KEY_FILE, "r") as f:
                saved_key = f.read().strip()
            
            # تمرير الـ page هنا ضروري جداً
            hwid = get_android_hwid(page) 
            
            # حساب المفتاح الصحيح للمقارنة
            correct_key = hashlib.sha256((hwid + secret_salt).encode()).hexdigest()[:12].upper()
            
            return saved_key == correct_key
        except Exception:
            return False
    return False

# --- داخل دالة main ---
def main(page: ft.Page):
    page.title = "تفعيل الميقاتي برو"
    page.rtl = True
    page.theme_mode = ft.ThemeMode.DARK
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    # 1. التحقق من التفعيل عند التشغيل
    is_active = is_full_version(page, SECRET_SALT)
    hwid = get_android_hwid(page)
    remaining = get_remaining_days(page)

    # 2. تعريف الدوال الداخلية (Logic)
    def open_astro_logic(e):
        page.clean()
        page.vertical_alignment = ft.MainAxisAlignment.START
        astro_list = ft.ListView(expand=1, spacing=10, padding=20)

        # إضافة المعطيات الـ 2500
        for i in range(2500):
            astro_list.controls.append(ft.Text(f"المعطى الفلكي {i}: نتيجة الحساب..."))

        page.add(
            ft.AppBar(title=ft.Text("لوحة الحسابات"), bgcolor=ft.colors.SURFACE_VARIANT),
            ft.Text("نتائج الحسابات:", size=18, weight="bold"),
            astro_list
        )
        page.update()

    def on_activate(e):
        correct_key = hashlib.sha256((hwid + SECRET_SALT).encode()).hexdigest()[:12].upper()
        if key_input.value.strip().upper() == correct_key:
            page.client_storage.set("license_key", correct_key)
            page.snack_bar = ft.SnackBar(ft.Text("تم التفعيل بنجاح!"))
            page.snack_bar.open = True
            open_astro_logic(None)
        else:
            key_input.error_text = "كود التفعيل غير صحيح"
            page.update()

    # 3. تعريف عناصر واجهة التفعيل
    key_input = ft.TextField(label="أدخل كود التفعيل هنا", password=True, width=300)

    # 4. القرار: ماذا نعرض للمستخدم الآن؟
    if is_active:
        open_astro_logic(None)
    else:
        # هذه الخطوة كانت ناقصة في كودك (إضافة العناصر للشاشة)
        page.add(
            ft.Text("نظام حماية الميقاتي برو", size=22, weight="bold"),
            ft.Text(f"بصمة جهازك: {hwid}", selectable=True, color="blue"),
            ft.Text(f"الأيام المتبقية: {remaining}", color="orange"),
            key_input,
            ft.ElevatedButton("تفعيل النسخة الكاملة", on_click=on_activate, bgcolor="green", color="white"),
            ft.TextButton("الاستمرار بالنسخة التجريبية", on_click=open_astro_logic)
        )

if __name__ == "__main__":
 ft.run(main)
